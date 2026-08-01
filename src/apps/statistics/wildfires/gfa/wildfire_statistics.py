#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Burnt-area statistics for the Global Fire Atlas wildfires.

Reports, per country and year, how many fires there were and the smallest,
largest and total area burnt, in hectares::

    Country          Year    Fires     Minimum      Maximum        Total
    World            2021   482119       25.31    229104.55  38104772.10
    World            2020   461003       25.31    198441.02  33970118.44
    World            Total  943122       25.31    229104.55  72074890.54

    Spain            2021     1204       25.34     12904.11    481203.77
    Spain            2020      876       25.34      8110.02    377411.20
    Spain            Total    2080       25.34     12904.11    858614.97

    France           2021      331       25.31      6002.45    120884.10
    ...

The **World** block comes first and summarises every country in scope, year by
year and then over all of them. It is not a total of the columns below it any
more than a country's ``Total`` is a total of the years above it: only ``Fires``
and ``Total (ha)`` are sums, while ``Minimum`` and ``Maximum`` are the smallest
and largest single fire anywhere in scope.

The ``Total`` row of a country works the same way over that country's years.

.. note::

   The World block is **omitted when ``--country`` is given**, where it could
   only repeat that country's rows word for word.

Run it over everything, or narrow it to one country, one year, or both::

    python3 -m src.apps.statistics.wildfires.gfa.wildfire_statistics --csv burnt.csv
    python3 -m src.apps.statistics.wildfires.gfa.wildfire_statistics \\
        --country Spain --year 2021 --csv spain_2021.csv --docx spain_2021.docx

At least one of ``--csv`` and ``--docx`` is required: an application that
computed a report and then printed nothing would be a strange thing.

The application only reads. Database settings come from the environment
(``.env``, see :mod:`src.settings`); every one of them can be overridden with a
command-line argument.

This is the GWIS report over a different dataset
------------------------------------------------

Deliberately so: it is the same four figures, grouped the same way, written the
same way, so that a GFA report and a
:mod:`GWIS one <src.apps.statistics.wildfires.gwis.wildfire_statistics>` can be
put side by side and the difference read as a difference between the *datasets*
rather than between two ways of counting. The two derive fire events from the
same MODIS burnt-area product by different algorithms, so comparing them is a
real question — and it is only answerable if the method is held fixed.

How the area is measured
------------------------

The perimeters are stored in EPSG:4326, whose units are degrees, so an area has
to come from somewhere that yields metres. Two ways are offered and they agree:

``geodesic`` (the default)
    ``ST_Area(perimeter::geography)`` — the true area on the WGS84 ellipsoid, in
    square metres, divided by 10,000. No projection is chosen, so none has to be
    justified. This is what the GWIS report uses.

``equal-area``
    ``ST_Area(ST_Transform(perimeter, 6933))`` — projected into NSIDC EASE-Grid
    2.0 Global, a cylindrical **equal-area** projection defined worldwide, and
    measured there in metres.

Measured against each other on the same polygons they differ by at most **0.003%**
— at the equator, in Spain, in Sweden, at 70°N and in Tasmania alike. The choice
therefore does not move any number in this report, and the default is geodesic
only because it needs no CRS argued for.

.. warning::

   What *does* move the numbers is projecting into something that is not
   equal-area. The same polygons measured in Web Mercator (EPSG:3857) come out
   **76% too large in Spain, 82% too large in Tasmania and 759% too large at
   70°N**, because Mercator's area distortion grows as ``sec²(latitude)``. For a
   dataset that spans every latitude MODIS sees, "convert to projected
   coordinates and compute the surface" is only safe if the projection is chosen
   for area. That is why the option is named ``equal-area`` and not ``projected``.

The published size is not used
------------------------------

:attr:`~src.providers.gfa.wildfire.GfaWildfire.size_km2` is the Atlas's own
figure for the fire, and it is *not* what this reports. It is kept as published
and is worth having as an independent check — a systematic gap between it and
the measured area would say something about either the perimeter or the Atlas —
but a report that mixed the two would be comparing a measurement with a claim.

Which country a fire counts towards
-----------------------------------

Chosen with ``--country-source``:

``geometry`` (the default)
    Ask the database which country actually contains the fire, at report time, by
    testing an interior point of its perimeter against the real country polygons.
    A fire inside no country is excluded.

``reported``
    Use the ``admin_boundary_id`` the import stored.

The default is the cautious one, and it is the default because of datasets other
than this one. GFA resolves its country by containment at import, so the two
modes give the same answer here and ``reported`` is simply the faster path — an
index lookup on a foreign key instead of a point-in-polygon test per fire, about
ten times quicker, which on the whole dataset is minutes rather than tens of
minutes.

Elsewhere they do not agree. EGIF resolves its boundary from an INE municipal
code rather than from a coordinate, so a *parte* filed in Ourense whose published
northing is missing three digits keeps its Spanish boundary while its point sits
in the Gulf of Guinea. Under ``reported`` that fire is in Spain's total; under
``geometry`` it is in nobody's.

For GFA the two can also differ for a reason that is nobody's error: ``reported``
follows the **ignition point** — that is what the import resolves the country
from — while ``geometry`` follows an interior point of the **perimeter**. A fire
that ignites one side of a border and burns across it is attributed differently
by the two, which makes the pair a useful way to find such fires.

.. note::

   Both modes attribute a fire's **whole** area to one country. Splitting a
   border-crossing fire between the countries it actually burnt in is a different
   and larger question, and this report does not attempt it: ``Total (ha)`` is
   the area of fires attributed to a country, not the area burnt inside its
   borders.

Which fires are counted
-----------------------

Fires with no country are **excluded** — mid-ocean perimeters, and any fire that
matched no OCHA boundary. So is any fire with no perimeter. A run therefore does
not necessarily account for every hectare in the database, and the totals here
are totals *of attributable burnt area*.

Which year a fire counts towards
--------------------------------

The year of its **local** start date, ``start_date_time AT TIME ZONE
time_zone``, which is the year of the ``start_date`` GFA published and so the
year of the file the fire came from. Using the raw UTC instant instead would move
fires across the New Year boundary: a fire starting on 1 January in Sydney is
still 31 December in UTC.

One year at a time
------------------

The report is not one statement. The years are found first, then each is
measured by a statement of its own and the summary rows are computed from their
results.

It was one statement, with ``GROUPING SETS`` producing every level from a single
pass, which is the better shape and is still what each year's statement does
internally. It does not survive this dataset. Under ``geometry`` the Atlas's
**21.5 million** perimeters mean 21.5 million point-in-polygon tests against
country polygons of millions of vertices each, and the memory that goes into
them is only released when the statement ends. On a 30 GB machine the backend
reached 29.2 GB after four and a half hours and was killed by the OOM killer,
taking the whole cluster with it and leaving nothing to show for the run.

Per year it is a gigabyte or so at a time, released between statements, and
there is a year count to watch instead of a bar that turns for hours. It is not
free: the years cost a pass of their own, and each year's statement scans the
fires where one statement scanned them once — the local year cannot be indexed,
since ``AT TIME ZONE`` on a column is not immutable. Both are cheap beside the
point-in-polygon tests, whose number does not change, and a run that finishes
slightly slower is worth any number of runs that do not finish at all.

Nothing about the figures changes. ``count``, ``sum``, ``min`` and ``max`` all
decompose over a partition of the fires, so a country's ``Total`` row and the
whole World block are exactly the numbers ``GROUPING SETS`` returned, from the
same rows. Every statement runs in one transaction and so against one snapshot,
which is what keeps a report built from twenty-five queries as consistent as one
built from a single query.
"""

from __future__ import annotations

import argparse
import csv
import logging
import math
import os
import sys

from dataclasses import dataclass
from pathlib import Path

from geoalchemy2 import Geography
from sqlalchemy import ColumnElement
from sqlalchemy import Engine
from sqlalchemy import Integer
from sqlalchemy import Select
from sqlalchemy import cast
from sqlalchemy import create_engine
from sqlalchemy import func
from sqlalchemy import or_
from sqlalchemy import select
from sqlalchemy import true
from sqlalchemy.orm import Session

import src.settings  # noqa: F401  (imported for the side effect of loading .env)

from src.apps.imports import common
from src.data_model.geography.admin_boundary import AdminBoundary
from src.data_model.wildfire import Wildfire
from src.providers.gfa.wildfire import GfaWildfire
from src.providers.ocha.admin_boundary import OchaAdminBoundary

#: Label used in the ``Year`` column for a summary row.
TOTAL_LABEL = "Total"

#: Label used in the ``Country`` column for the rows that summarise every country.
#:
#: The block those rows form is what a reader wants first — the dataset does not
#: hold "the whole planet" in any tidy sense, and until you have seen the global
#: figures you have nothing to read a single country's against.
WORLD_LABEL = "World"

#: The report's columns, in order, shared by both output formats so that a change
#: to one cannot silently leave the other behind.
#:
#: ``Fires`` sits between the grouping and the areas because it is the count of
#: the things the three area figures summarise: the row reads "in 2021 Spain had
#: 1,204 fires, the smallest 25.34 ha, the largest 12,904.11 ha, 481,203.77 ha in
#: all". After ``Total`` it would follow a sum of hectares with a count of events.
COLUMNS = ("Country", "Year", "Fires", "Minimum (ha)", "Maximum (ha)", "Total (ha)")

#: Index of the first column that holds a number, and so is right-aligned in the
#: Word table. ``Year`` is left alone: it is a label, and one of its values is
#: :data:`TOTAL_LABEL`.
FIRST_NUMERIC_COLUMN = 2

#: The two ways of deciding which country a fire counts towards.
COUNTRY_SOURCE_GEOMETRY = "geometry"
COUNTRY_SOURCE_REPORTED = "reported"
COUNTRY_SOURCES = (COUNTRY_SOURCE_GEOMETRY, COUNTRY_SOURCE_REPORTED)

#: Administrative level of a country in ``admin_boundary``.
COUNTRY_LEVEL = 0

#: A point guaranteed to lie inside the burnt perimeter, for the containment test.
#:
#: ``ST_PointOnSurface`` and not ``ST_Centroid``: the centroid of a crescent or a
#: ring-shaped burn can fall outside the polygon entirely, and would then be
#: tested against a country the fire never reached — or against no country at all.
#: It is the same expression the import uses, so ``geometry`` and ``reported``
#: agree on every GWIS fire whose boundaries have not changed underneath it.
LOCATOR = func.ST_PointOnSurface(Wildfire.perimeter)

#: Square metres in a hectare.
SQUARE_METRES_PER_HECTARE = 10_000.0

#: NSIDC EASE-Grid 2.0 Global — a cylindrical equal-area projection in metres,
#: defined for the whole world. The CRS behind ``--area-method equal-area``.
#:
#: Equal-area is the whole requirement: the report sums and compares areas across
#: every latitude MODIS covers, and a conformal projection such as Web Mercator
#: would inflate a fire at 70°N by a factor of eight. EASE-Grid 2.0 is used rather
#: than an equal-area projection centred on each fire because one fixed CRS can be
#: applied in SQL to every row at once, and the difference between the two is
#: smaller than the difference between either and the geodesic answer.
EQUAL_AREA_SRID = 6933

#: The two ways of turning a perimeter in degrees into hectares.
AREA_METHOD_GEODESIC = "geodesic"
AREA_METHOD_EQUAL_AREA = "equal-area"
AREA_METHODS = (AREA_METHOD_GEODESIC, AREA_METHOD_EQUAL_AREA)


def burnt_area(method: str) -> ColumnElement:
    """Burnt area of one fire in hectares, by whichever method was asked for.

    Parameters
    ----------
    method : str
        One of :data:`AREA_METHODS`.

    Returns
    -------
    ColumnElement
        The SQL expression yielding hectares.

    Raises
    ------
    ValueError
        If ``method`` is not one of :data:`AREA_METHODS`.

    Notes
    -----
    The geodesic cast carries the geometry type and SRID rather than being a bare
    ``Geography()``: that renders as ``geography(GEOMETRY,-1)``, which PostGIS
    rejects, because -1 is not a SRID it knows.
    """
    if method == AREA_METHOD_GEODESIC:
        square_metres = func.ST_Area(
            cast(Wildfire.perimeter, Geography(geometry_type="MULTIPOLYGON", srid=4326))
        )
    elif method == AREA_METHOD_EQUAL_AREA:
        square_metres = func.ST_Area(func.ST_Transform(Wildfire.perimeter, EQUAL_AREA_SRID))
    else:
        raise ValueError(
            f"unknown area method {method!r}; expected one of {', '.join(AREA_METHODS)}"
        )
    return square_metres / SQUARE_METRES_PER_HECTARE


#: The year a fire counts towards: the year of its *local* start date.
#:
#: ``AT TIME ZONE`` has no construct of its own in SQLAlchemy, so it is applied
#: as an operator. ``EXTRACT`` yields a numeric, hence the cast — without it the
#: year comes back as ``Decimal`` and lands in the report as ``2021.0``.
LOCAL_YEAR = cast(
    func.extract(
        "year",
        Wildfire.start_date_time.op("AT TIME ZONE")(func.coalesce(Wildfire.time_zone, "UTC")),
    ),
    Integer,
)


def country_columns(source: str) -> tuple[ColumnElement, ColumnElement, list]:
    """Where a fire's country comes from, as ``(name, iso_3, joins)``.

    Parameters
    ----------
    source : str
        One of :data:`COUNTRY_SOURCES`.

    Returns
    -------
    tuple
        The name and ISO alpha-3 expressions, and the joins that have to be
        applied to a ``select`` over ``wildfire`` for them to resolve.

    Raises
    ------
    ValueError
        If ``source`` is not one of :data:`COUNTRY_SOURCES`.

    Notes
    -----
    **``geometry``** ignores ``admin_boundary_id`` and asks the database which
    country actually contains the fire, at report time, against the real polygons.
    It is the default because it cannot be wrong about a dataset whose own
    attribution is wrong — and datasets whose attribution *is* wrong are the
    normal case once you leave the two global Atlases. EGIF resolves its boundary
    from an INE municipal code rather than from a coordinate, so a *parte* filed
    in Ourense whose published northing is missing three digits keeps its Spanish
    boundary while its point sits in the Gulf of Guinea. In ``geometry`` mode that
    fire is simply not in Spain's total.

    **``reported``** trusts what the import stored. For GWIS and GFA that *is* a
    geometric answer — both resolve the country by containment at import time —
    so the two modes agree, and this one is far cheaper: it is an index lookup on
    a foreign key instead of a point-in-polygon test per fire. It is also the
    honest choice when the question is "what does this dataset claim", as opposed
    to "where did these fires actually burn".

    The two modes can disagree for a reason that is nobody's error: **a fire that
    crosses a border**. Both modes attribute the whole burnt area to one country,
    and for GFA they may pick different ones — ``reported`` follows the ignition
    point, ``geometry`` follows an interior point of the perimeter. Splitting the
    area between the countries it actually burnt in is a different and larger
    question, and this report does not attempt it: ``Total (ha)`` is the area of
    fires attributed to a country, not the area burnt inside its borders.
    """
    ocha_boundary = OchaAdminBoundary.__table__

    if source == COUNTRY_SOURCE_REPORTED:
        return (
            AdminBoundary.name,
            ocha_boundary.c.iso_3,
            [(AdminBoundary, AdminBoundary.id == Wildfire.admin_boundary_id, False),
             (ocha_boundary, ocha_boundary.c.id == AdminBoundary.id, True)],
        )

    if source != COUNTRY_SOURCE_GEOMETRY:
        raise ValueError(
            f"unknown country source {source!r}; expected one of {', '.join(COUNTRY_SOURCES)}"
        )

    # LATERAL ... LIMIT 1 rather than a plain join: a point on a shared border can
    # satisfy ST_Contains for two countries, and one fire must not become two rows.
    # The join is inner, which is what drops a fire that is inside no country at
    # all — the same rule ``reported`` gets from its inner join to admin_boundary.
    containing = (
        select(AdminBoundary.name.label("name"), ocha_boundary.c.iso_3.label("iso_3"))
        .select_from(AdminBoundary)
        .join(ocha_boundary, ocha_boundary.c.id == AdminBoundary.id)
        .where(AdminBoundary.level == COUNTRY_LEVEL)
        .where(func.ST_Contains(AdminBoundary.geometry, LOCATOR))
        .limit(1)
        .lateral("containing")
    )
    return containing.c.name, containing.c.iso_3, [(containing, true(), False)]


def years_query() -> Select:
    """The years the dataset holds fires in, newest first.

    Returns
    -------
    Select
        A query yielding one ``int`` per year.

    Notes
    -----
    Run before anything else, because each of those years is then measured by a
    statement of its own — see the module docstring for why the report is built
    that way.

    It carries neither the ``--country`` filter nor any join: resolving the
    country is the expensive half of this report and there is nothing to gain by
    paying for it twice. A year that turns out to hold no fire in the country
    asked for simply measures to nothing and never reaches the report.

    ``DISTINCT`` over the whole table rather than ``min(year)`` to ``max(year)``:
    a gap in the record is a gap in the report, and a range would fill it with
    statements that can only return nothing.
    """
    gfa_fire = GfaWildfire.__table__
    year = LOCAL_YEAR.label("year")
    return (
        select(year)
        .select_from(Wildfire)
        .join(gfa_fire, gfa_fire.c.id == Wildfire.id)
        .where(Wildfire.perimeter.is_not(None))
        .distinct()
        .order_by(year.desc())
    )


def statistics_query(country: str | None, year: int,
                     method: str = AREA_METHOD_GEODESIC,
                     country_source: str = COUNTRY_SOURCE_GEOMETRY) -> Select:
    """Build the statistics query for one year.

    Returns
    -------
    Select
        A query yielding ``country, minimum, maximum, total, fires``: one row per
        country that had a fire in ``year``, unordered. The summary rows and the
        report's order are :func:`summarise`'s work.

    Notes
    -----
    Built against the mapped classes rather than written as SQL text, so a column
    renamed on a model breaks this at import time rather than in front of a user.
    It also lets the country filter be a plain conditional: written as text it
    would have to become an "unset, or matching" disjunction so that one statement
    could serve every combination, leaving a branch in the SQL that is dead on
    every actual run.

    The inner query computes each area exactly once. Folded into the outer
    aggregate instead, the area expression would be evaluated three times per row
    — for the minimum, the maximum and the sum — and it is by far the most
    expensive thing here.

    Whichever way the country is resolved the join is inner, and that is what
    drops the fires belonging to no country. ``gfa_wildfire`` is joined — by
    table, to keep SQLAlchemy from adding the polymorphic join of its own —
    rather than filtering on ``wildfire.type``, so this stays a GFA report even if
    another provider ever adopts the same discriminator.

    A year with no fires in scope produces no rows at all, which is what keeps
    ``--year 1999`` on a dataset that starts in 2002 an empty report — and so the
    caller's "no wildfires matched" message — rather than one row of NULL
    aggregates over zero fires.

    See :func:`country_columns` for what ``country_source`` changes.
    """
    gfa_fire = GfaWildfire.__table__
    country_name, country_iso_3, joins = country_columns(country_source)

    fires = (
        select(
            country_name.label("country"),
            burnt_area(method).label("hectares"),
        )
        .select_from(Wildfire)
        .join(gfa_fire, gfa_fire.c.id == Wildfire.id)
        .where(Wildfire.perimeter.is_not(None))
        .where(LOCAL_YEAR == year)
    )
    for target, condition, is_outer in joins:
        fires = fires.outerjoin(target, condition) if is_outer \
            else fires.join(target, condition)

    if country is not None:
        fires = fires.where(or_(
            func.lower(country_name) == country.lower(),
            func.upper(country_iso_3) == country.upper(),
        ))

    fire = fires.subquery("fire")
    return (
        select(
            fire.c.country,
            func.min(fire.c.hectares).label("minimum"),
            func.max(fire.c.hectares).label("maximum"),
            func.sum(fire.c.hectares).label("total"),
            func.count().label("fires"),
        )
        .group_by(fire.c.country)
    )


@dataclass(frozen=True)
class Row:
    """One line of the report.

    Attributes
    ----------
    country : str
        Name of the country the fires burnt in, or :data:`WORLD_LABEL` for a row
        that summarises every country.
    year : int or None
        The year, or ``None`` for a summary row.
    minimum, maximum, total : float
        Smallest single fire, largest single fire and sum of every fire, in
        hectares.
    fires : int
        How many fires the three area figures were computed from — the count of
        wildfire events in this country and year, or in the country's whole
        period for a summary row.

        Worth reading beside the areas rather than instead of them: a year whose
        total is small because it had few fires is a different year from one
        whose total is small because its fires were.
    """

    country: str
    year: int | None
    minimum: float
    maximum: float
    total: float
    fires: int
    is_world: bool = False

    @property
    def is_total(self) -> bool:
        """Whether this is a summary row rather than one of the years."""
        return self.year is None

    @property
    def is_summary(self) -> bool:
        """Whether the row summarises rather than reports: emphasised in the Word table.

        True for a country's ``Total`` and for every row of the World block, which
        is a summary of the countries below it however it is grouped.
        """
        return self.is_total or self.is_world

    @property
    def year_label(self) -> str:
        return TOTAL_LABEL if self.is_total else str(self.year)


def combine(rows: list[Row], country: str, year: int | None,
            is_world: bool = False) -> Row:
    """One row summarising several: the four figures taken over all of them.

    Notes
    -----
    This is what makes measuring a year at a time cost nothing. All four figures
    decompose over a partition of the fires — a minimum of minima is a minimum, a
    sum of sums is a sum — so a country's ``Total`` and the World block are the
    same numbers ``GROUPING SETS`` produced from a single pass, and no fire is
    counted twice or left out.

    ``fsum`` rather than ``sum`` because the values being added are themselves
    sums of millions of hectares; an exact accumulation over a handful of them
    costs nothing and cannot drift from what one pass would have returned.
    """
    return Row(
        country=country,
        year=year,
        is_world=is_world,
        minimum=min(row.minimum for row in rows),
        maximum=max(row.maximum for row in rows),
        total=math.fsum(row.total for row in rows),
        fires=sum(row.fires for row in rows),
    )


def ordered_countries(session: Session, names: set[str]) -> list[str]:
    """The countries of the report, in the order the database would have sorted them.

    Notes
    -----
    Sorted by PostgreSQL and not by Python because the two do not agree. Python
    compares code points, which puts every accented name after every unaccented
    one — Côte d'Ivoire after Cyprus rather than between Costa Rica and Croatia —
    while the database sorts by its collation, which is what this report has
    always been ordered by and what a reader of a country list expects. It is one
    round trip for a couple of hundred names.

    A name the query does not return could only come of ``admin_boundary``
    changing underneath the run, but it is appended rather than dropped: losing a
    country from a report over an ordering detail would be a poor trade.
    """
    if not names:
        return []
    ordered = list(session.scalars(
        select(AdminBoundary.name)
        .distinct()
        .where(AdminBoundary.level == COUNTRY_LEVEL)
        .where(AdminBoundary.name.in_(names))
        .order_by(AdminBoundary.name)
    ))
    return ordered + sorted(names - set(ordered))


def summarise(measured: list[Row], countries: list[str], with_world: bool) -> list[Row]:
    """Build the report from the years measured: the summary rows, in order.

    Parameters
    ----------
    measured : list of Row
        One row per country and year, as the per-year statements returned them.
    countries : list of str
        The countries in scope, in the order they are to be reported.
    with_world : bool
        Whether to open with the World block. It is dropped when ``--country`` is
        given, where it could only repeat that country's rows word for word.

    Returns
    -------
    list of Row
        The World block first — its years newest first, then its total — and then
        each country, its years newest first and its summary row last.
    """
    report: list[Row] = []
    if with_world and measured:
        for year in sorted({row.year for row in measured}, reverse=True):
            report.append(combine([row for row in measured if row.year == year],
                                  WORLD_LABEL, year, is_world=True))
        report.append(combine(measured, WORLD_LABEL, None, is_world=True))

    for name in countries:
        rows = [row for row in measured if row.country == name]
        report += sorted(rows, key=lambda row: row.year, reverse=True)
        report.append(combine(rows, name, None))
    return report


def parse_arguments(argv: list[str] | None = None) -> argparse.Namespace:
    """Parse the command line."""
    parser = argparse.ArgumentParser(
        description="Burnt-area statistics for the Global Fire Atlas wildfires.",
        epilog="Areas are in hectares, geodesic on the WGS84 ellipsoid by default. Fires "
               "with no country are not counted. Database settings not given here are "
               "read from the environment (.env).",
    )
    selection = parser.add_argument_group("selection", "report on everything unless narrowed")
    selection.add_argument("-c", "--country",
                           help="restrict to one country, by name ('Spain') or ISO 3166-1 "
                                "alpha-3 code ('ESP'); case-insensitive")
    selection.add_argument("-y", "--year", type=int, help="restrict to one year, e.g. 2021")

    parser.add_argument("--area-method", default=AREA_METHOD_GEODESIC, choices=AREA_METHODS,
                        help="how to turn the EPSG:4326 perimeter into hectares: "
                             "'geodesic' measures on the WGS84 ellipsoid (default, and what "
                             "the GWIS report uses); 'equal-area' projects to EPSG:6933 and "
                             "measures there. They agree to within 0.003%%")

    parser.add_argument("--country-source", default=COUNTRY_SOURCE_GEOMETRY,
                        choices=COUNTRY_SOURCES,
                        help="how to decide which country a fire counts towards: "
                             "'geometry' (default) tests the perimeter against the real "
                             "country polygons at report time, so a dataset that "
                             "mis-attributes a fire cannot mislead the report; 'reported' "
                             "trusts the admin_boundary_id the import stored, which is "
                             "much faster and, for this dataset, the same answer")

    output = parser.add_argument_group("output", "at least one is required")
    output.add_argument("--csv", type=Path, help="write the report to this .csv")
    output.add_argument("--docx", type=Path, help="write the report to this .docx (MS Word)")

    common.add_database_arguments(parser)
    parser.add_argument("--log-level", default=os.getenv("GISFIRE_LOG_LEVEL", "INFO"),
                        choices=["CRITICAL", "ERROR", "WARNING", "INFO", "DEBUG"],
                        help="verbosity (env: GISFIRE_LOG_LEVEL, default INFO)")

    arguments = parser.parse_args(argv)
    if arguments.csv is None and arguments.docx is None:
        parser.error("nothing to write: pass --csv, --docx, or both")
    return arguments


def compute(session: Session, country: str | None, year: int | None,
            logger: logging.Logger,
            method: str = AREA_METHOD_GEODESIC,
            country_source: str = COUNTRY_SOURCE_GEOMETRY) -> list[Row]:
    """Measure the fires a year at a time, returning the report's rows in order.

    Notes
    -----
    One statement per year, for the reason given in the module docstring, each
    under a spinner of its own so that a run of hours says which year it is on
    rather than only that it is alive.

    Every one of them runs in ``session``'s transaction, and so against a single
    snapshot: a report assembled from twenty-five queries is then exactly as
    consistent as one assembled from a single query, and an import running
    alongside it cannot have a fire counted in one year's statement and not in
    another's.
    """
    if year is not None:
        years = [year]
    else:
        with common.Spinner("Finding the years the GFA fires cover", logger):
            years = list(session.scalars(years_query()))

    scope = country or "every country"
    measured: list[Row] = []
    for index, measuring in enumerate(years, start=1):
        with common.Spinner(f"Measuring the burnt area of the GFA fires "
                            f"({scope}, {measuring}: {index} of {len(years)})", logger):
            measured += [
                Row(country=record.country, year=measuring,
                    minimum=float(record.minimum),
                    maximum=float(record.maximum),
                    total=float(record.total),
                    fires=record.fires)
                for record in session.execute(
                    statistics_query(country, measuring, method, country_source))
            ]

    # Counted over the country rows alone: the World block is a summary of them,
    # not a country, and reporting "3 countries" for two would be a small lie in
    # the one line a user is most likely to read.
    countries = ordered_countries(session, {row.country for row in measured})
    rows = summarise(measured, countries, with_world=country is None)
    logger.info("Computed %d rows over %d countries (%s areas, country from %s)",
                len(rows), len(countries), method, country_source)
    return rows


def write_csv(rows: list[Row], path: Path, logger: logging.Logger) -> None:
    """Write the report as CSV.

    The numbers go out unformatted apart from being rounded to two decimals — no
    thousands separators — because a CSV is read by another program far more often
    than by a person, and a separator would make every figure a string.
    """
    path.parent.mkdir(parents=True, exist_ok=True)
    with path.open("w", newline="", encoding="utf-8") as handle:
        writer = csv.writer(handle)
        writer.writerow(COLUMNS)
        for row in rows:
            writer.writerow([row.country, row.year_label, row.fires,
                             f"{row.minimum:.2f}", f"{row.maximum:.2f}", f"{row.total:.2f}"])
    logger.info("Wrote %s", path)


def write_docx(rows: list[Row], path: Path, country: str | None, year: int | None,
               logger: logging.Logger,
               method: str = AREA_METHOD_GEODESIC) -> None:
    """Write the report as a Word document.

    One table, with each country's summary row in bold so the blocks read apart
    the way they would on paper. Numbers get thousands separators here — the
    opposite of the CSV, and for the opposite reason: this one is for reading.
    """
    # Imported here rather than at module scope so that --csv keeps working if
    # python-docx is not installed, which matters because it is the only
    # dependency this application adds.
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt

    document = Document()
    document.add_heading("Global Fire Atlas wildfire burnt area", level=1)

    scope = []
    if country is not None:
        scope.append(f"country: {country}")
    if year is not None:
        scope.append(f"year: {year}")
    subtitle = "; ".join(scope) if scope else "all countries, all years"
    measured = ("geodesically on the WGS84 ellipsoid" if method == AREA_METHOD_GEODESIC
                else f"in the equal-area projection EPSG:{EQUAL_AREA_SRID}")
    document.add_paragraph(
        f"Areas in hectares, computed {measured}. "
        f"Fires not attributable to a country are excluded. Scope: {subtitle}."
    )

    table = document.add_table(rows=1, cols=len(COLUMNS))
    table.style = "Table Grid"
    for cell, heading in zip(table.rows[0].cells, COLUMNS):
        cell.text = heading
        cell.paragraphs[0].runs[0].bold = True

    for row in rows:
        cells = table.add_row().cells
        values = [row.country, row.year_label, f"{row.fires:,}",
                  f"{row.minimum:,.2f}", f"{row.maximum:,.2f}", f"{row.total:,.2f}"]
        for index, (cell, value) in enumerate(zip(cells, values)):
            cell.text = value
            paragraph = cell.paragraphs[0]
            if index >= FIRST_NUMERIC_COLUMN:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            for run in paragraph.runs:
                run.bold = row.is_summary
                run.font.size = Pt(9)

    path.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(path))
    logger.info("Wrote %s", path)


def report(args: argparse.Namespace, engine: Engine, logger: logging.Logger) -> list[Row]:
    """Compute the statistics and write whichever outputs were asked for."""
    # No spinner here: compute runs one statement per year and turns one of its
    # own for each, which is the only honest place to say how far along it is.
    with Session(engine) as session:
        rows = compute(session, args.country, args.year, logger, args.area_method,
                       args.country_source)

    if not rows:
        # An empty report is almost always a mistyped country or a year with no
        # data, and writing an empty file would hide that.
        raise RuntimeError(
            "No wildfires matched. Check --country (a name or an ISO alpha-3 code) and "
            "--year, and that the GFA fires and the OCHA boundaries are both imported "
            "— fires with no country are not counted."
        )

    if args.csv is not None:
        write_csv(rows, args.csv, logger)
    if args.docx is not None:
        write_docx(rows, args.docx, args.country, args.year, logger, args.area_method)
    return rows


def main(argv: list[str] | None = None) -> int:
    args = parse_arguments(argv)
    logging.basicConfig(level=args.log_level, format="%(asctime)s %(levelname)s %(message)s")
    logger = logging.getLogger("gfa-statistics")

    try:
        settings = common.resolve_database_settings(args)
    except RuntimeError as error:
        logger.error("%s", error)
        return 1

    engine = create_engine(common.database_url(settings))
    try:
        report(args, engine, logger)
    except Exception as error:  # noqa: BLE001  (the CLI boundary: report, do not traceback)
        logger.error("%s", error)
        return 1
    finally:
        engine.dispose()
    return 0


if __name__ == "__main__":  # pragma nocover
    sys.exit(main())
