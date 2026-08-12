#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Wildfire counts by cause for the Mexican CONAFOR burnt area cartography.

Reports, per year, how many fires there were, how many carry a cause at all, how
many were classified ``Naturales``, and — for the years that publish a specific
cause — how many were **lightning**::

    Country     Year    Fires  Classified   Natural   Natural %   Lightning
    Mexico      2023     7513        7513       159        2.12
    Mexico      2019     4265        4265        90        2.11          78
    Mexico      2012      224         224         3        1.34           3
    Mexico      Total   44804       44644       894        2.00         356

Run it over everything, or narrow it to one year::

    python3 -m src.apps.statistics.wildfires.mexico_conafor.wildfire_causes --csv causes.csv
    python3 -m src.apps.statistics.wildfires.mexico_conafor.wildfire_causes \\
        --year 2023 --csv 2023.csv --docx 2023.docx
    python3 -m src.apps.statistics.wildfires.mexico_conafor.wildfire_causes \\
        --cause Intencional --csv arson.csv

At least one of ``--csv`` and ``--docx`` is required.

The companion of
:mod:`~src.apps.statistics.wildfires.mexico_conafor.wildfire_statistics`, over the
same fires and the same years — so the two can be read side by side. What that one
measures in hectares, this one counts by cause.

This dataset **does** publish lightning
----------------------------------------

The obvious question is how many fires lightning started, and unlike
:mod:`the ICNF <src.apps.statistics.wildfires.portugal_icnf.wildfire_causes>` —
whose classification names lightning nowhere, so that ``Natural`` is the finest
answer it supports — CONAFOR answers it directly. ``CAUSAESP`` has ``Rayos``, and
the 2011 layers spell the same thing ``Descargas electricas``. Both translate to
``Lightning`` in
:data:`~src.providers.mexico_conafor.fire_cause.SPECIFIC_CAUSE_TRANSLATIONS`, and
that is what the ``Lightning`` column counts.

The count is taken on
:attr:`~src.providers.mexico_conafor.fire_cause.ConaforFireCause.specific_cause_en`
rather than on a list of Spanish spellings, deliberately: the two published
wordings are a decade apart, both mean lightning, and a third would be caught the
moment it is added to that table rather than needing an edit here as well.

.. warning::

   **``CAUSAESP`` is published in 2010 and 2012-2019 and in no other year.** 2011,
   and every year from 2020, publish a cause and no specific cause — 27,624 fires,
   three in five.

   For those years the ``Lightning`` cell is left **empty**, not zero. A zero would
   say that lightning started no fires in Mexico in 2021, which is not what the
   absence of a column means. Never compare a ``Lightning`` count across that
   boundary, and never sum the column as though the blanks were zeros.

   The ``Natural`` column has no such gap: ``CAUSA`` is published by all fourteen
   layers.

``Natural`` and ``Lightning`` are not the same number, and neither contains the
other cleanly: a fire can be ``Naturales`` with no specific cause published, and
the specific causes include ``Erupciones volcanicas`` — one fire, in 2019 — which
is natural and is not lightning.

Counting on the reconciled cause, not the published one
--------------------------------------------------------

``--cause`` and the ``Natural`` column match on
:attr:`~src.providers.mexico_conafor.fire_cause.ConaforFireCause.cause_normalised`,
the canonical Spanish, and **not** on the published string.

That is the whole reason that column exists. CONAFOR publishes no cause code and
the cause is free text, typed sixty-four ways over fourteen years for about twenty
real causes — ``'Naturales'`` in the later layers is ``'Tormenta Electrica'``,
``'Tormenta Elcetrica'`` and ``'Descargas Electricas'`` in 2011. A report matching
on the published text would find none of those and would say that no natural fire
burnt in 2011.

The choices offered are therefore the canonical causes of
:data:`~src.providers.mexico_conafor.fire_cause.CAUSE_TRANSLATIONS`, and the
column heading uses the English from it.

.. warning::

   **A canonical cause can still be zero for a run of years because CONAFOR
   renamed it.** Reconciling the spellings does not reconcile the *categories*,
   and the archive has at least one category that changed name outright:

   .. code-block:: text

      --cause Intencional          2013-2019   551 to 1,106 fires a year
                                   2020-2022   zero, in all three years
                                   2023        2,726

      --cause "Actividades ilícitas"
                                   2020-2022   1,564 / 2,363 / 2,030
                                   every other year   zero

   *Intencional* and *Actividades ilícitas* are the same act filed under two
   administrative names three years apart, and the three zeros in the middle of an
   ``Intencional`` series are the rename, not a collapse in arson.

   They are kept as two causes rather than merged, because *actividades ilícitas*
   is the broader phrase — it can cover fires set to clear illicit crops, which the
   archive also files separately as ``Cultivos ilícitos`` — and merging two
   published categories on a guess would be a worse error than reporting them
   apart. Run the report twice and add the columns if a continuous series is what
   is wanted.

   The same caution applies less dramatically elsewhere: ``Actividades
   agropecuarias`` is split into ``Actividades agrícolas`` and ``Actividades
   pecuarias`` from 2018.

.. note::

   A fire whose published cause reached no canonical form has
   ``cause_normalised`` ``NULL``: it is **classified** and can never **match**.
   Three fires of the published archive are in that state, all in 2011, all of
   them a bare ``'12'`` typed into the cause field. The run says so at ``WARNING``
   when any are in scope. See
   :mod:`src.providers.mexico_conafor.fire_cause`.

Why ``Classified`` is a column and the percentage is of it
----------------------------------------------------------

160 fires of the published archive carry no cause at all — 2010 writes ``'0'``
into ``CAUSA`` seven times and 2011 writes ``'No'`` 153 times, and both are null
tokens rather than causes.

That is a much smaller hole than the ICNF's, where two fires in three are
unclassified, but the denominator is the classified fires all the same and for the
same reason: a percentage of *all* fires would be a statement partly about how
complete the classification is. Where nothing is classified there is no percentage
to give, and the cell is left empty rather than filled with a zero that would be a
claim.

.. warning::

   Take the classified share as a floor. ``Desconocidas`` is the second largest
   cause in the archive — 6,247 fires, one in seven — and a fire whose cause was
   never determined may well have been natural.

Which fires are counted
-----------------------

**Every fire**, with no filter at all: no country test (see the companion report),
and no requirement to have a perimeter.

This differs by a handful from the companion, which counts every fire the chosen
``--area-method`` can measure — nine fewer under a measured method, being the 2012
features that publish attributes and an empty shape, and one fewer under
``reported``. The difference is stated rather than engineered away: a fire with no
polygon still has a cause, and a causes report that dropped it would be answering a
question about polygons.

Which year a fire counts towards
--------------------------------

:attr:`~src.providers.mexico_conafor.wildfire.ConaforWildfire.year`, the year of
the archive the fire was published in, exactly as in the companion report and for
the same reason: nine fires carry a start date in a different calendar year from
the archive that filed them.

Shared with the companion report
--------------------------------

:data:`~...wildfire_statistics.COUNTRY_NAME`,
:data:`~...wildfire_statistics.PUBLISHED_YEAR`,
:data:`~...wildfire_statistics.TOTAL_LABEL` and
:func:`~...wildfire_statistics.years_query` are **imported from**
:mod:`~src.apps.statistics.wildfires.mexico_conafor.wildfire_statistics` rather
than copied. Two reports over one dataset that disagreed about what a year is would
be worse than one report, and a copy is a thing that drifts.
"""

from __future__ import annotations

import argparse
import csv
import logging
import os
import sys

from dataclasses import dataclass
from pathlib import Path

from sqlalchemy import Engine
from sqlalchemy import Select
from sqlalchemy import create_engine
from sqlalchemy import func
from sqlalchemy import literal
from sqlalchemy import select
from sqlalchemy.orm import Session

import src.settings  # noqa: F401  (imported for the side effect of loading .env)

from src.apps.imports import common
from src.apps.statistics.wildfires.mexico_conafor.wildfire_statistics import COUNTRY_NAME
from src.apps.statistics.wildfires.mexico_conafor.wildfire_statistics import PUBLISHED_YEAR
from src.apps.statistics.wildfires.mexico_conafor.wildfire_statistics import TOTAL_LABEL
from src.apps.statistics.wildfires.mexico_conafor.wildfire_statistics import years_query
from src.data_model.wildfire import Wildfire
from src.providers.mexico_conafor.fire_cause import CAUSE_TRANSLATIONS
from src.providers.mexico_conafor.fire_cause import ConaforFireCause
from src.providers.mexico_conafor.fire_cause import SPECIFIC_CAUSE_TRANSLATIONS
from src.providers.mexico_conafor.wildfire import ConaforWildfire

#: The canonical cause counted unless ``--cause`` says otherwise.
#:
#: Unlike its ICNF counterpart this is *not* the nearest thing the dataset has to
#: lightning — CONAFOR names lightning outright, and the ``Lightning`` column
#: counts it. This is the broader natural category, which is reported beside it
#: because the two answer different questions and because ``Natural`` is published
#: for every year while lightning is not.
DEFAULT_CAUSE = "Naturales"

#: The canonical causes, taken from the translation table on the model so that a
#: cause added there becomes selectable here without a second edit.
CAUSES = tuple(CAUSE_TRANSLATIONS)

#: The English that marks a specific cause as lightning.
#:
#: Matched on the translation rather than on the published Spanish on purpose: two
#: wordings a decade apart mean lightning — ``Rayos`` and ``Descargas
#: electricas`` — and a third would be picked up the moment it is added to
#: :data:`~src.providers.mexico_conafor.fire_cause.SPECIFIC_CAUSE_TRANSLATIONS`,
#: rather than needing an edit here as well.
LIGHTNING_LABEL = "Lightning"

#: The published specific causes that mean lightning, for the documentation and
#: for the tests. Derived, never used to build the query: the query matches the
#: label above.
LIGHTNING_SPECIFIC_CAUSES = tuple(
    published for published, english in SPECIFIC_CAUSE_TRANSLATIONS.items()
    if english == LIGHTNING_LABEL
)

#: Index of the first column that holds a number, and so is right-aligned in the
#: Word table.
FIRST_NUMERIC_COLUMN = 2


def columns(cause: str = DEFAULT_CAUSE) -> tuple[str, ...]:
    """The report's columns, in order, for a given canonical cause.

    Notes
    -----
    A function rather than a constant, unlike the companion report's ``COLUMNS``,
    because two of the headings name the cause being counted: a file of
    ``Intencional`` counts under a heading saying ``Natural`` would be a trap.

    Both output formats read them from here, so a change to one cannot silently
    leave the other behind. The first three are the companion report's first three,
    unchanged, which is what lets the two be read side by side.
    """
    label = CAUSE_TRANSLATIONS.get(cause, cause)
    return ("Country", "Year", "Fires", "Classified", label,
            f"{label} (% of classified)", LIGHTNING_LABEL)


def counts_query(year: int, cause: str = DEFAULT_CAUSE) -> Select:
    """Build the counting query for one year.

    Parameters
    ----------
    year : int
        The published year to count.
    cause : str
        The canonical cause to count, in Spanish — a key of
        :data:`~src.providers.mexico_conafor.fire_cause.CAUSE_TRANSLATIONS`.

    Returns
    -------
    Select
        A query yielding ``country, fires, classified, matching, detailed,
        lightning``: at most one row. The summary rows and the report's order are
        :func:`summarise`'s work.

    Notes
    -----
    Five counts in one pass, which is what makes this report cheap.

    The join to ``conafor_fire_cause`` is **outer**, and has to be: ``cause_id`` is
    ``NULL`` for the 160 fires whose published cause was a null token, and an inner
    join would silently turn this into a report of the classified fires alone — the
    ``Fires`` column would stop matching the companion report's and nothing would
    say why.

    ``count(cause_id)`` and not ``count(*)`` for the classified column: counting a
    nullable column counts the rows where it is filled in, which is the definition
    of classified here.

    ``detailed`` counts the fires that publish a **specific** cause. It is not a
    column of the report; it is what decides whether the ``Lightning`` cell is a
    number or a blank. A year in which no fire publishes one is a year whose layer
    does not have the field — ``CAUSAESP`` is all-or-nothing per layer — and its
    lightning count is unknown rather than zero. See the module docstring.

    The cause is compared against ``cause_normalised``, the canonical Spanish, and
    not against the published ``cause``: matching the published text would miss the
    2011 spellings entirely. Lightning is compared against ``specific_cause_en``,
    for the reason given at :data:`LIGHTNING_LABEL`.
    """
    conafor = ConaforWildfire.__table__
    fire_cause = ConaforFireCause.__table__
    country = literal(COUNTRY_NAME)

    return (
        select(
            country.label("country"),
            func.count().label("fires"),
            func.count(conafor.c.cause_id).label("classified"),
            func.count().filter(fire_cause.c.cause_normalised == cause).label("matching"),
            func.count(fire_cause.c.specific_cause).label("detailed"),
            func.count().filter(
                fire_cause.c.specific_cause_en == LIGHTNING_LABEL).label("lightning"),
        )
        .select_from(Wildfire)
        .join(conafor, conafor.c.id == Wildfire.id)
        .outerjoin(fire_cause, fire_cause.c.id == conafor.c.cause_id)
        .where(PUBLISHED_YEAR == year)
        .group_by(country)
    )


def unreconciled_query(year: int) -> Select:
    """How many of a year's fires are classified but reached no canonical cause.

    Returns
    -------
    Select
        A query yielding one ``int``.

    Notes
    -----
    Such a fire can never match ``--cause``, whatever is asked for, so it is
    invisible in the report's columns while still being counted in ``Classified``.
    Three fires of the published archive are in that state and all three are a bare
    ``'12'`` typed into the 2011 cause field; the run reports them rather than
    letting them be a silent discrepancy between two columns.
    """
    conafor = ConaforWildfire.__table__
    fire_cause = ConaforFireCause.__table__
    return (
        select(func.count())
        .select_from(Wildfire)
        .join(conafor, conafor.c.id == Wildfire.id)
        .join(fire_cause, fire_cause.c.id == conafor.c.cause_id)
        .where(PUBLISHED_YEAR == year)
        .where(fire_cause.c.cause_normalised.is_(None))
    )


@dataclass(frozen=True)
class Row:
    """One line of the report.

    Attributes
    ----------
    country : str
        Always :data:`~...wildfire_statistics.COUNTRY_NAME`. A label; see the
        companion report.
    year : int or None
        The published year, or ``None`` for the summary row.
    fires : int
        Every fire of the year.
    classified : int
        How many of them carry a cause at all.
    matching : int
        How many carry the canonical cause asked for.
    detailed : int
        How many publish a specific cause. Not a column of the report: it is what
        makes :attr:`lightning_label` a number or a blank.
    lightning : int
        How many were started by lightning, of the ``detailed`` ones. Meaningless
        unless :attr:`detailed` is non-zero.

    Notes
    -----
    ``matching <= classified <= fires`` always, and ``lightning <= detailed <=
    classified``. ``matching`` and ``lightning`` can each be zero for entirely
    different reasons: nothing of that cause, or nothing classified in that way.
    """

    country: str
    year: int | None
    fires: int
    classified: int
    matching: int
    detailed: int
    lightning: int

    @property
    def is_total(self) -> bool:
        """Whether this is the summary row rather than one of the years."""
        return self.year is None

    @property
    def year_label(self) -> str:
        return TOTAL_LABEL if self.is_total else str(self.year)

    @property
    def share(self) -> float | None:
        """``matching`` as a percentage of ``classified``, or ``None`` if none is.

        ``None`` and not zero: a year in which nothing was classified has no share
        to report, and zero would be a claim that none of its fires was natural.
        The writers turn it into an empty cell.
        """
        if not self.classified:
            return None
        return 100.0 * self.matching / self.classified

    @property
    def share_label(self) -> str:
        """The percentage as it is written out, empty where there is none."""
        return "" if self.share is None else f"{self.share:.2f}"

    @property
    def lightning_label(self) -> str:
        """The lightning count, **empty** where the year publishes no specific cause.

        This is the single most important cell in the report to get right. 2011 and
        every year from 2020 publish ``CAUSA`` and no ``CAUSAESP``, so nothing in
        them can be identified as lightning — and writing ``0`` there would say that
        lightning started no fires in Mexico in 2021, which is not what a missing
        column means. See the module docstring.
        """
        return "" if not self.detailed else f"{self.lightning:,}"


def combine(rows: list[Row], country: str, year: int | None) -> Row:
    """One row summarising several: the five counts added up.

    Notes
    -----
    Counts decompose over a partition of the fires, so the ``Total`` row is what a
    single pass over the same fires would have returned.

    The share is deliberately **not** averaged over the years: it is recomputed by
    :attr:`Row.share` from the summed counts, which is the ratio of the totals
    rather than the mean of the ratios. A year with eleven classified fires and a
    year with eleven thousand must not weigh the same in the answer.

    The summed ``lightning`` is a total over the years that publish a specific
    cause, and its ``detailed`` is summed with it — so the summary row's cell is a
    number whenever any year in scope had one. It is a total of *those* years and
    not of the period, which is why the report says so in its preamble rather than
    leaving a reader to add the blanks up as zeros.
    """
    return Row(
        country=country,
        year=year,
        fires=sum(row.fires for row in rows),
        classified=sum(row.classified for row in rows),
        matching=sum(row.matching for row in rows),
        detailed=sum(row.detailed for row in rows),
        lightning=sum(row.lightning for row in rows),
    )


def summarise(measured: list[Row]) -> list[Row]:
    """Build the report from the years counted: the summary row, in order.

    Returns
    -------
    list of Row
        The years newest first, and the summary row last.
    """
    if not measured:
        return []
    report = sorted(measured, key=lambda row: row.year, reverse=True)
    return report + [combine(measured, COUNTRY_NAME, None)]


def parse_arguments(argv: list[str] | None = None) -> argparse.Namespace:
    """Parse the command line."""
    parser = argparse.ArgumentParser(
        description="Wildfire counts by cause for the Mexican CONAFOR burnt area "
                    "cartography.",
        epilog="CONAFOR publishes lightning outright, in CAUSAESP — but only in 2010 and "
               "2012-2019, so the Lightning column is blank for the other years and must "
               "not be read as a zero. Causes are matched on the reconciled canonical "
               "form, not on the published text. Database settings not given here are "
               "read from the environment (.env).",
    )
    parser.add_argument("-y", "--year", type=int,
                        help="restrict to one year, e.g. 2023; this is the year of the "
                             "archive the fire was published in, not the year of its "
                             "start date")
    parser.add_argument("--cause", default=DEFAULT_CAUSE, choices=CAUSES,
                        help="which canonical cause to count, in Spanish: 'Naturales' "
                             "(default). Matched on cause_normalised, the reconciled "
                             "form — the published text is typed 64 ways over the "
                             "fourteen layers, so matching it directly would miss most "
                             "of them")

    # Accepted only so that they can be refused clearly, exactly as in the
    # companion report: anyone reaching for either here has copied a command line
    # from the GWIS, GFA or ICNF report, which is a reasonable thing to have done.
    parser.add_argument("--country", help=argparse.SUPPRESS)
    parser.add_argument("--country-source", help=argparse.SUPPRESS)
    # And this one because the ICNF report calls the same option --cause-type.
    parser.add_argument("--cause-type", help=argparse.SUPPRESS)

    output = parser.add_argument_group("output", "at least one is required")
    output.add_argument("--csv", type=Path, help="write the report to this .csv")
    output.add_argument("--docx", type=Path, help="write the report to this .docx (MS Word)")

    common.add_database_arguments(parser)
    parser.add_argument("--log-level", default=os.getenv("GISFIRE_LOG_LEVEL", "INFO"),
                        choices=["CRITICAL", "ERROR", "WARNING", "INFO", "DEBUG"],
                        help="verbosity (env: GISFIRE_LOG_LEVEL, default INFO)")

    arguments = parser.parse_args(argv)
    if arguments.country is not None:
        parser.error(
            "there is no --country here: CONAFOR publishes one country's fires, so "
            "there is nothing to select between. Every fire is counted under Mexico."
        )
    if arguments.country_source is not None:
        parser.error(
            "there is no --country-source here: every CONAFOR perimeter is inside "
            "Mexico, so this report runs no containment test and needs no boundaries "
            "imported. The Country column is a label, not a resolved answer."
        )
    if arguments.cause_type is not None:
        parser.error(
            "the option is --cause here, not --cause-type: CONAFOR publishes no cause "
            "type, only a cause and a specific cause. Pass --cause with one of: "
            + ", ".join(CAUSES)
        )
    if arguments.csv is None and arguments.docx is None:
        parser.error("nothing to write: pass --csv, --docx, or both")
    return arguments


def compute(session: Session, year: int | None, logger: logging.Logger,
            cause: str = DEFAULT_CAUSE) -> list[Row]:
    """Count the fires a year at a time, returning the report's rows in order.

    Notes
    -----
    One statement per year, under a spinner of its own — the same shape as the
    companion report.

    Every one of them runs in ``session``'s transaction, and so against a single
    snapshot.

    Two things are reported at ``WARNING`` rather than left to be noticed in the
    table: a scope in which nothing is classified, and a scope containing fires
    whose published cause reached no canonical form. The second is the one that
    would otherwise be invisible — such a fire is counted in ``Classified`` and can
    never appear in the cause column beside it.
    """
    if year is not None:
        years = [year]
    else:
        with common.Spinner("Finding the years the CONAFOR fires cover", logger):
            years = list(session.scalars(years_query()))

    measured: list[Row] = []
    unreconciled = 0
    for index, counting in enumerate(years, start=1):
        with common.Spinner(f"Counting the CONAFOR fires by cause "
                            f"({counting}: {index} of {len(years)})", logger):
            measured += [
                Row(country=record.country, year=counting,
                    fires=record.fires,
                    classified=record.classified,
                    matching=record.matching,
                    detailed=record.detailed,
                    lightning=record.lightning)
                for record in session.execute(counts_query(counting, cause))
            ]
            unreconciled += session.scalar(unreconciled_query(counting))

    rows = summarise(measured)

    classified = sum(row.classified for row in measured)
    if measured and not classified:
        logger.warning(
            "No fire in scope carries a cause at all, so every %s count is zero and no "
            "percentage can be given", cause)
    if unreconciled:
        logger.warning(
            "%d fire(s) in scope carry a published cause that reached no canonical form. "
            "They are counted under Classified and can never appear under %s, whatever "
            "--cause asks for. Add their spelling to "
            "src.providers.mexico_conafor.fire_cause.CAUSE_NORMALISATIONS.",
            unreconciled, CAUSE_TRANSLATIONS.get(cause, cause))

    blank = sorted(row.year for row in measured if not row.detailed)
    if blank:
        logger.info("No specific cause is published for %d year(s) (%s), so their "
                    "Lightning cell is blank rather than zero: CONAFOR publishes CAUSAESP "
                    "in 2010 and 2012-2019 only",
                    len(blank), ", ".join(str(one) for one in blank))

    logger.info("Counted %d rows over %d year(s) (%s fires)",
                len(rows), len({row.year for row in rows if not row.is_total}), cause)
    return rows


def write_csv(rows: list[Row], path: Path, logger: logging.Logger,
              cause: str = DEFAULT_CAUSE) -> None:
    """Write the report as CSV.

    The percentage and the lightning count are written bare and are **empty** where
    there is no answer — an empty field reads as no answer to whatever parses this,
    which is what it is, while a zero would read as an answer.
    """
    path.parent.mkdir(parents=True, exist_ok=True)
    with path.open("w", newline="", encoding="utf-8") as handle:
        writer = csv.writer(handle)
        writer.writerow(columns(cause))
        for row in rows:
            writer.writerow([row.country, row.year_label, row.fires, row.classified,
                             row.matching, row.share_label,
                             "" if not row.detailed else row.lightning])
    logger.info("Wrote %s", path)


def write_docx(rows: list[Row], path: Path, year: int | None,
               logger: logging.Logger, cause: str = DEFAULT_CAUSE) -> None:
    """Write the report as a Word document.

    One table, with the summary row in bold. Counts get thousands separators here —
    the opposite of the CSV, and for the opposite reason: this one is for reading.

    The opening paragraphs say which years publish a specific cause and what the
    blank ``Lightning`` cells mean. Both belong in the document and not only in the
    manual: a table of lightning counts read without them would be misread as a
    collapse to zero after 2019.
    """
    # Imported here rather than at module scope so that --csv keeps working if
    # python-docx is not installed, which matters because it is the only dependency
    # this application adds.
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt

    heading = columns(cause)[4]
    document = Document()
    document.add_heading(f"CONAFOR wildfire counts — {heading} causes (Mexico)", level=1)

    scope = f"year: {year}" if year is not None else "all published years"
    document.add_paragraph(
        f"Counts of CONAFOR wildfires whose reconciled cause is {cause} ({heading}), "
        f"with the lightning fires beside them. Every fire is counted; no country test "
        f"is applied and no perimeter is required. Years are the year of the archive the "
        f"fire was published in. Scope: {scope}."
    )
    document.add_paragraph(
        "CONAFOR publishes lightning outright, as the specific cause Rayos (Descargas "
        "electricas in the earlier layers) — but only in 2010 and 2012-2019. For 2011 "
        "and for every year from 2020 the field does not exist, and the Lightning cell "
        "is left blank. A blank is not a zero: it means the question cannot be asked of "
        "that year. The Total is a total of the years that answer it."
    )
    document.add_paragraph(
        "Causes are matched on the reconciled canonical form and not on the published "
        "text, which is typed sixty-four ways across the fourteen layers. The percentage "
        "is of the classified fires and not of all of them. Treat it as a floor: "
        "Desconocidas is the second largest cause in the archive."
    )

    table = document.add_table(rows=1, cols=len(columns(cause)))
    table.style = "Table Grid"
    for cell, title in zip(table.rows[0].cells, columns(cause)):
        cell.text = title
        cell.paragraphs[0].runs[0].bold = True

    for row in rows:
        cells = table.add_row().cells
        values = [row.country, row.year_label, f"{row.fires:,}", f"{row.classified:,}",
                  f"{row.matching:,}", row.share_label, row.lightning_label]
        for index, (cell, value) in enumerate(zip(cells, values)):
            cell.text = value
            paragraph = cell.paragraphs[0]
            if index >= FIRST_NUMERIC_COLUMN:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            for run in paragraph.runs:
                run.bold = row.is_total
                run.font.size = Pt(9)

    path.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(path))
    logger.info("Wrote %s", path)


def report(args: argparse.Namespace, engine: Engine, logger: logging.Logger) -> list[Row]:
    """Count the fires and write whichever outputs were asked for."""
    with Session(engine) as session:
        rows = compute(session, args.year, logger, args.cause)

    if not rows:
        # An empty report is almost always a year with no data, and writing an empty
        # file would hide that. Note that a year with no fires of the cause asked for
        # is not empty — it is a row of zeros, which is a different thing and is
        # reported as one.
        raise RuntimeError(
            "No wildfires matched. Check --year — a year whose archive was never "
            "imported has nothing to find — and that the CONAFOR fires are imported."
        )

    if args.csv is not None:
        write_csv(rows, args.csv, logger, args.cause)
    if args.docx is not None:
        write_docx(rows, args.docx, args.year, logger, args.cause)
    return rows


def main(argv: list[str] | None = None) -> int:
    args = parse_arguments(argv)
    logging.basicConfig(level=args.log_level, format="%(asctime)s %(levelname)s %(message)s")
    logger = logging.getLogger("conafor-causes")

    try:
        settings = common.resolve_database_settings(args)
    except RuntimeError as error:
        logger.error("%s", error)
        return 1

    engine = create_engine(common.database_url(settings))
    try:
        report(args, engine, logger)
    except Exception as error:  # noqa: BLE001  (the CLI boundary: report, do not traceback)
        logger.error("%s", error)
        return 1
    finally:
        engine.dispose()
    return 0


if __name__ == "__main__":  # pragma nocover
    sys.exit(main())
