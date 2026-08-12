#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Burnt-area statistics for the Mexican CONAFOR burnt area cartography.

Reports, per year, how many fires there were and the smallest, largest and total
area burnt, in hectares::

    Country          Year    Fires     Minimum      Maximum        Total
    Mexico           2023      7513        0.01     13182.83     363448.71
    Mexico           2022      6718        0.01     15000.00     282183.06
    Mexico           Total    44804        0.01     23809.00    3187726.94

Run it over everything, or narrow it to one year, or to the fires above a size::

    python3 -m src.apps.statistics.wildfires.mexico_conafor.wildfire_statistics --csv burnt.csv
    python3 -m src.apps.statistics.wildfires.mexico_conafor.wildfire_statistics \\
        --year 2023 --csv 2023.csv --docx 2023.docx
    python3 -m src.apps.statistics.wildfires.mexico_conafor.wildfire_statistics \\
        --area-method reported --csv as-published.csv

At least one of ``--csv`` and ``--docx`` is required: an application that computed
a report and then printed nothing would be a strange thing.

The application only reads. Database settings come from the environment
(``.env``, see :mod:`src.settings`); every one of them can be overridden with a
command-line argument.

One country, no ``--country``, and no containment test
-------------------------------------------------------

CONAFOR publishes one country's fires, so there is nothing to select between and
the option is not offered — the same as the
:mod:`ICNF <src.apps.statistics.wildfires.portugal_icnf.wildfire_statistics>`
report.

Unlike that one, there is **no ``--country-source`` either, and no country join at
all**. The ICNF report offers a geometry mode because its perimeters can and do
fall outside Portugal — into the sea, or across the Spanish border — and a fire
that is not in the country should not be in the country's total. CONAFOR's do not:
the published extent of all fourteen archives is inside Mexico, so a containment
test would cost a point-in-polygon per fire to confirm what is already known.

:data:`COUNTRY_NAME` is therefore a **label, not a computed answer**. The column
stays so that this report's CSV has the same shape as the GWIS, GFA and ICNF ones
and the four can be concatenated and compared; it does not stay because anything
was resolved.

The consequence worth stating: this report needs **no OCHA boundaries imported**,
and a fire whose ``admin_boundary_id`` is ``NULL`` is counted like any other. The
nine 2012 features that carry attributes and no geometry have no country for that
reason, and they are in the ``reported`` totals below.

Which year a fire counts towards
--------------------------------

:attr:`~src.providers.mexico_conafor.wildfire.ConaforWildfire.year` — the year of
the archive the fire was published in — and **not** the year of
``start_date_time``.

Unlike the ICNF report, that is not because the dates are placeholders: every
CONAFOR fire is dated to the day, and grouping on the start date would very nearly
agree. It is because the two do not agree *exactly*: **nine fires carry a start
date in a different calendar year from the archive that published them** — five
2016 dates in the 2017 layer, and four others. The published year is what CONAFOR
filed the fire under, is ``NOT NULL``, needs no time zone applied to it, and is
the unit an import replaces.

.. warning::

   **The counts are not comparable across 2016, and 2010's areas are not
   comparable with anything.**

   The feature count steps by an order of magnitude at 2016 — 628 polygons in 2014
   against 3,244 in 2016 — because before 2016 CONAFOR published only the fires it
   had drawn and from 2016 it publishes the season. That is a change in what was
   mapped, not in what burnt.

   And the published areas of the 2010 layer do not describe the 2010 polygons at
   all: the median ratio between them is 3.0 and the 90th percentile is 65. Under
   ``--area-method reported`` that year's figures are a different measurement from
   every other year's; under the two measured methods they are the polygons', which
   are consistent but which are 311 sketches of five to twelve vertices.

   :data:`~src.providers.mexico_conafor.FIRST_YEAR_WITH_MEASURED_AREA` is the year
   from which the published area and the polygon agree.

How the area is measured
------------------------

Three ways, chosen with ``--area-method``, and this is the only report in the
project that offers the third:

``geodesic``
    On the WGS84 ellipsoid, from the stored perimeter. The default, and what the
    GWIS, GFA and ICNF reports mean by an area.
``equal-area``
    Projected to EPSG:6933 and measured there. Agrees with ``geodesic`` to within
    a few thousandths of a percent.
``reported``
    :attr:`~src.providers.mexico_conafor.wildfire.ConaforWildfire.area_ha`, the
    ``AREA_HA`` CONAFOR published, measured by nobody here.

The third exists because this dataset is the one that makes it interesting.
CONAFOR publishes **both** a perimeter and a burnt area, and from 2016 the second
*is* the first's own area: the median ratio between the published figure and the
polygon's geodesic area is 1.000, and four rows in five agree to within 1%. So the
two are a check on each other, and running the report twice is how the 2010 warning
above becomes a number rather than a claim.

Why not a Mexican national grid
--------------------------------

Mexico's own projected CRS, ITRF2008 / LCC (EPSG:6362), is a **Lambert conformal
conic** — conformal, not equal-area — so its ``ST_Area`` is not a burnt area, for
the same reason the ICNF report declines EPSG:3763. It is not offered. Nothing is
lost by that: unlike Portugal, CONAFOR publishes in EPSG:4326 and nothing is stored
in a national grid to reproduce.

Which fires are counted
-----------------------

Every fire of the dataset **that the chosen method can measure**, and the two cases
differ:

* ``geodesic`` and ``equal-area`` need a perimeter, so the nine 2012 features that
  publish attributes and an empty shape are not counted.
* ``reported`` needs a published area, so the one fire that leaves ``AREA_HA``
  empty — ``21-24-0078``, which publishes everything else — is not counted, and
  the nine shapeless ones are.

So the ``Fires`` column can differ by a handful between two runs of this report,
and that is the truth about the dataset rather than an inconsistency. No fire is
excluded for having no country: see above.

``--min-area`` narrows it further, to the fires of at least that many hectares. By
default there is no threshold and every measurable fire counts.

One year at a time
------------------

The report is not one statement. The years are found first, then each is measured
by a statement of its own and the summary rows are computed from their results.

This dataset does not need it — forty-five thousand perimeters are seconds' work
however they are grouped — and it is built this way because its GWIS and GFA
counterparts have to be, at twenty million perimeters. Three reports meant to be
read side by side are worth keeping as one program over several datasets.

Nothing about the figures changes. ``count``, ``sum``, ``min`` and ``max`` all
decompose over a partition of the fires, so the ``Total`` row is exactly the number
a single pass would have returned, from the same rows. Every statement runs in one
transaction and so against one snapshot.
"""

from __future__ import annotations

import argparse
import csv
import logging
import math
import os
import sys

from dataclasses import dataclass
from pathlib import Path

from geoalchemy2 import Geography
from sqlalchemy import ColumnElement
from sqlalchemy import Engine
from sqlalchemy import Select
from sqlalchemy import cast
from sqlalchemy import create_engine
from sqlalchemy import func
from sqlalchemy import literal
from sqlalchemy import select
from sqlalchemy.orm import Session

import src.settings  # noqa: F401  (imported for the side effect of loading .env)

from src.apps.imports import common
from src.data_model.wildfire import Wildfire
from src.providers.mexico_conafor.wildfire import ConaforWildfire

#: Label used in the ``Year`` column for the summary row.
TOTAL_LABEL = "Total"

#: The country every CONAFOR fire is reported under.
#:
#: A **label, not a computed answer**. See the module docstring: the published
#: extent of all fourteen archives is inside Mexico, so this report runs no
#: containment test and needs no boundaries imported. The column exists so that
#: this CSV has the same shape as the GWIS, GFA and ICNF ones.
COUNTRY_NAME = "Mexico"

#: The report's columns, in order, shared by both output formats so that a change
#: to one cannot silently leave the other behind. Deliberately identical to the
#: GWIS, GFA and ICNF reports' columns, so the four CSVs can be concatenated.
COLUMNS = ("Country", "Year", "Fires", "Minimum (ha)", "Maximum (ha)", "Total (ha)")

#: Index of the first column that holds a number, and so is right-aligned in the
#: Word table.
FIRST_NUMERIC_COLUMN = 2

#: Square metres in a hectare.
SQUARE_METRES_PER_HECTARE = 10_000.0

#: NSIDC EASE-Grid 2.0 Global — a cylindrical equal-area projection in metres,
#: defined for the whole world. The CRS behind ``--area-method equal-area``.
EQUAL_AREA_SRID = 6933

#: The three ways of getting a burnt area in hectares. See the module docstring;
#: ``reported`` is offered here and in no other report in the project.
AREA_METHOD_GEODESIC = "geodesic"
AREA_METHOD_EQUAL_AREA = "equal-area"
AREA_METHOD_REPORTED = "reported"
AREA_METHODS = (AREA_METHOD_GEODESIC, AREA_METHOD_EQUAL_AREA, AREA_METHOD_REPORTED)

#: The methods that measure the stored polygon, and so need one.
MEASURED_METHODS = (AREA_METHOD_GEODESIC, AREA_METHOD_EQUAL_AREA)

#: The year a fire counts towards: the year of the archive it was published in.
#:
#: See the module docstring. Not derived from ``start_date_time``, which for nine
#: fires falls in a different calendar year from the archive that published them.
PUBLISHED_YEAR = ConaforWildfire.__table__.c.year

#: The burnt area CONAFOR published, behind ``--area-method reported``.
REPORTED_AREA = ConaforWildfire.__table__.c.area_ha


def burnt_area(method: str) -> ColumnElement:
    """Burnt area of one fire in hectares, by whichever method was asked for.

    Parameters
    ----------
    method : str
        One of :data:`AREA_METHODS`.

    Returns
    -------
    ColumnElement
        The SQL expression yielding hectares.

    Raises
    ------
    ValueError
        If ``method`` is not one of :data:`AREA_METHODS`.

    Notes
    -----
    The geodesic cast carries the geometry type and SRID rather than being a bare
    ``Geography()``: that renders as ``geography(GEOMETRY,-1)``, which PostGIS
    rejects, because -1 is not a SRID it knows.

    ``reported`` returns the published column untouched. It is not converted,
    scaled or measured — the whole point of the option is that it is CONAFOR's own
    number — and it is the one method whose value can be present for a fire with
    no polygon, and absent for a fire that has one. See :func:`measurable`.
    """
    if method == AREA_METHOD_GEODESIC:
        square_metres = func.ST_Area(
            cast(Wildfire.perimeter, Geography(geometry_type="MULTIPOLYGON", srid=4326))
        )
    elif method == AREA_METHOD_EQUAL_AREA:
        square_metres = func.ST_Area(func.ST_Transform(Wildfire.perimeter, EQUAL_AREA_SRID))
    elif method == AREA_METHOD_REPORTED:
        return REPORTED_AREA
    else:
        raise ValueError(
            f"unknown area method {method!r}; expected one of {', '.join(AREA_METHODS)}"
        )
    return square_metres / SQUARE_METRES_PER_HECTARE


def measurable(method: str) -> ColumnElement:
    """What a fire needs for this method to yield a number.

    Parameters
    ----------
    method : str
        One of :data:`AREA_METHODS`.

    Returns
    -------
    ColumnElement
        The SQL predicate a fire has to satisfy to be counted.

    Notes
    -----
    Not the same fires for every method, and deliberately not forced to be. A
    measured method needs the polygon, which nine 2012 features do not have; the
    published area needs ``AREA_HA``, which one fire of 45,914 does not have.
    Intersecting the two so that every method reported the same ``Fires`` would
    make three of the columns wrong for nine fires in order to make one column
    consistent — see the module docstring.
    """
    if method in MEASURED_METHODS:
        return Wildfire.perimeter.is_not(None)
    if method == AREA_METHOD_REPORTED:
        return REPORTED_AREA.is_not(None)
    raise ValueError(
        f"unknown area method {method!r}; expected one of {', '.join(AREA_METHODS)}"
    )


def years_query() -> Select:
    """The years the dataset holds fires in, newest first.

    Returns
    -------
    Select
        A query yielding one ``int`` per year.

    Notes
    -----
    Run before anything else, because each of those years is then measured by a
    statement of its own — see the module docstring for why the report is built
    that way.

    ``DISTINCT`` over the whole table rather than ``min(year)`` to ``max(year)``:
    a year the database happens not to hold — because its archive has not been
    imported, or because CONAFOR never published it — is then a gap in the report
    rather than a row reading zero, which would say that nothing burnt that year.

    The series is complete as published today, 2010 to 2023. It has not always
    looked complete: the 2015 archive is distributed separately from the others
    and is easy to miss, and a run against a database without it should say so by
    leaving the year out, not by inventing it.

    No filter on the perimeter or the published area, unlike the ICNF report's:
    which years the dataset covers is a fact about the dataset and not about the
    ``--area-method`` a particular run chose. A year none of whose fires the method
    can measure simply returns no rows and is absent from the report.
    """
    conafor = ConaforWildfire.__table__
    year = PUBLISHED_YEAR.label("year")
    return (
        select(year)
        .select_from(Wildfire)
        .join(conafor, conafor.c.id == Wildfire.id)
        .distinct()
        .order_by(year.desc())
    )


def statistics_query(year: int,
                     method: str = AREA_METHOD_GEODESIC,
                     min_area: float | None = None) -> Select:
    """Build the statistics query for one year.

    Parameters
    ----------
    year : int
        The published year to measure.
    method : str
        One of :data:`AREA_METHODS`.
    min_area : float, optional
        Count only fires of at least this many hectares. ``None``, the default,
        counts every fire the method can measure.

    Returns
    -------
    Select
        A query yielding ``country, minimum, maximum, total, fires``: at most one
        row. The summary rows and the report's order are :func:`summarise`'s work.

    Notes
    -----
    Built against the mapped classes rather than written as SQL text, so a column
    renamed on a model breaks this at import time rather than in front of a user.

    The inner query computes each area exactly once. Folded into the outer
    aggregate instead, the area expression would be evaluated three times per row —
    for the minimum, the maximum and the sum — and under a measured method it is by
    far the most expensive thing here.

    The country is a literal and there is no join for it, unlike every other report
    in this family: see the module docstring. It is still grouped on, so that the
    shape of the result is the one :func:`summarise` and the writers expect.

    ``conafor_wildfire`` is joined — by table, to keep SQLAlchemy from adding the
    polymorphic join of its own — rather than filtering on ``wildfire.type``, and it
    has to be joined in any case: :data:`PUBLISHED_YEAR` lives on it.

    ``min_area`` filters the subquery's column rather than repeating the area
    expression in a ``WHERE`` of its own, for the same reason the subquery exists at
    all, and it is applied before the aggregates rather than as a ``HAVING``: the
    threshold selects the fires the figures are computed from, it does not discard a
    year whose total came out small.
    """
    conafor = ConaforWildfire.__table__
    country = literal(COUNTRY_NAME)

    fires = (
        select(
            country.label("country"),
            burnt_area(method).label("hectares"),
        )
        .select_from(Wildfire)
        .join(conafor, conafor.c.id == Wildfire.id)
        .where(measurable(method))
        .where(PUBLISHED_YEAR == year)
    )

    fire = fires.subquery("fire")
    statistics = (
        select(
            fire.c.country,
            func.min(fire.c.hectares).label("minimum"),
            func.max(fire.c.hectares).label("maximum"),
            func.sum(fire.c.hectares).label("total"),
            func.count().label("fires"),
        )
        .group_by(fire.c.country)
    )
    if min_area is not None:
        statistics = statistics.where(fire.c.hectares >= min_area)
    return statistics


@dataclass(frozen=True)
class Row:
    """One line of the report.

    Attributes
    ----------
    country : str
        Always :data:`COUNTRY_NAME`. A label; see the module docstring.
    year : int or None
        The published year, or ``None`` for the summary row.
    minimum, maximum, total : float
        Smallest single fire, largest single fire and sum of every fire, in
        hectares.
    fires : int
        How many fires the three area figures were computed from — the count of
        wildfire events in this year, or in the whole period for a summary row.
    """

    country: str
    year: int | None
    minimum: float
    maximum: float
    total: float
    fires: int

    @property
    def is_total(self) -> bool:
        """Whether this is the summary row rather than one of the years."""
        return self.year is None

    @property
    def year_label(self) -> str:
        return TOTAL_LABEL if self.is_total else str(self.year)


def combine(rows: list[Row], country: str, year: int | None) -> Row:
    """One row summarising several: the four figures taken over all of them.

    Notes
    -----
    This is what makes measuring a year at a time cost nothing. All four figures
    decompose over a partition of the fires — a minimum of minima is a minimum, a
    sum of sums is a sum — so the ``Total`` row is the same number a single pass
    would have produced, and no fire is counted twice or left out.

    ``fsum`` rather than ``sum``: an exact accumulation over a handful of partial
    totals costs nothing and cannot drift from what one pass would have returned.
    """
    return Row(
        country=country,
        year=year,
        minimum=min(row.minimum for row in rows),
        maximum=max(row.maximum for row in rows),
        total=math.fsum(row.total for row in rows),
        fires=sum(row.fires for row in rows),
    )


def summarise(measured: list[Row]) -> list[Row]:
    """Build the report from the years measured: the summary row, in order.

    Parameters
    ----------
    measured : list of Row
        One row per year, as the per-year statements returned them.

    Returns
    -------
    list of Row
        The years newest first, and the summary row last.

    Notes
    -----
    Takes no list of countries, unlike its ICNF counterpart. There is one country
    and it is a literal, so there is nothing to order and nothing to ask the
    database about.
    """
    if not measured:
        return []
    report = sorted(measured, key=lambda row: row.year, reverse=True)
    return report + [combine(measured, COUNTRY_NAME, None)]


def hectares(text: str) -> float:
    """Argparse type for ``--min-area``: a finite, non-negative number of hectares.

    Raises
    ------
    argparse.ArgumentTypeError
        If the text is not a number, or is negative, or is a non-finite float.

    Notes
    -----
    A bare ``type=float`` would accept ``-5``, ``nan`` and ``inf``. The first two
    are almost certainly a typo and would silently produce the unfiltered report —
    ``nan`` compares false against every area, so it would produce an empty one —
    and none of the three is a size a fire can have.
    """
    try:
        area = float(text)
    except ValueError:
        raise argparse.ArgumentTypeError(f"{text!r} is not a number of hectares")
    if not math.isfinite(area):
        raise argparse.ArgumentTypeError(f"{text!r} is not a finite number of hectares")
    if area < 0:
        raise argparse.ArgumentTypeError(
            f"a burnt area cannot be negative, and {area:g} ha is: pass 0 or more, or "
            f"leave --min-area out to count every fire"
        )
    return area


def parse_arguments(argv: list[str] | None = None) -> argparse.Namespace:
    """Parse the command line."""
    parser = argparse.ArgumentParser(
        description="Burnt-area statistics for the Mexican CONAFOR burnt area cartography.",
        epilog="Areas are in hectares, geodesic on the WGS84 ellipsoid by default; "
               "--area-method reported gives CONAFOR's own published figure instead. "
               "CONAFOR publishes one country's fires and every perimeter is inside it, "
               "so there is no --country option and no containment test. Database "
               "settings not given here are read from the environment (.env).",
    )
    parser.add_argument("-y", "--year", type=int,
                        help="restrict to one year, e.g. 2023; this is the year of the "
                             "archive the fire was published in, not the year of its "
                             "start date")
    parser.add_argument("--min-area", type=hectares, default=None, metavar="HECTARES",
                        help="count only fires that burnt at least this many hectares, "
                             "e.g. 5; by default every fire counts. The threshold is "
                             "applied to the area this report uses (see --area-method)")

    parser.add_argument("--area-method", default=AREA_METHOD_GEODESIC, choices=AREA_METHODS,
                        help="where the hectares come from: 'geodesic' measures the "
                             "perimeter on the WGS84 ellipsoid (default); 'equal-area' "
                             "projects it to EPSG:6933 and measures there; 'reported' "
                             "uses CONAFOR's published AREA_HA and measures nothing. "
                             "From 2016 the published figure is the polygon's own area "
                             "to three decimals, so the three agree — in 2010 they do "
                             "not, by a factor of three at the median")

    # Accepted only so that they can be refused clearly. Without them argparse's
    # prefix matching would resolve "--country" to "--country-source"-like
    # neighbours or report an unrecognised argument, neither of which names the
    # real problem. Anyone reaching for either here has copied a command line from
    # the GWIS, GFA or ICNF report, which is a reasonable thing to have done.
    parser.add_argument("--country", help=argparse.SUPPRESS)
    parser.add_argument("--country-source", help=argparse.SUPPRESS)

    output = parser.add_argument_group("output", "at least one is required")
    output.add_argument("--csv", type=Path, help="write the report to this .csv")
    output.add_argument("--docx", type=Path, help="write the report to this .docx (MS Word)")

    common.add_database_arguments(parser)
    parser.add_argument("--log-level", default=os.getenv("GISFIRE_LOG_LEVEL", "INFO"),
                        choices=["CRITICAL", "ERROR", "WARNING", "INFO", "DEBUG"],
                        help="verbosity (env: GISFIRE_LOG_LEVEL, default INFO)")

    arguments = parser.parse_args(argv)
    if arguments.country is not None:
        parser.error(
            "there is no --country here: CONAFOR publishes one country's fires, so "
            "there is nothing to select between. Every fire is reported under Mexico."
        )
    if arguments.country_source is not None:
        parser.error(
            "there is no --country-source here: every CONAFOR perimeter is inside "
            "Mexico, so this report runs no containment test and needs no boundaries "
            "imported. The Country column is a label, not a resolved answer."
        )
    if arguments.csv is None and arguments.docx is None:
        parser.error("nothing to write: pass --csv, --docx, or both")
    return arguments


def compute(session: Session, year: int | None, logger: logging.Logger,
            method: str = AREA_METHOD_GEODESIC,
            min_area: float | None = None) -> list[Row]:
    """Measure the fires a year at a time, returning the report's rows in order.

    Notes
    -----
    One statement per year, for the reason given in the module docstring, each
    under a spinner of its own.

    Every one of them runs in ``session``'s transaction, and so against a single
    snapshot: a report assembled from many queries is then exactly as consistent as
    one assembled from a single query.

    ``min_area`` does not narrow the list of years: which years the dataset holds
    fires in is a fact about the dataset, and a year whose fires are all below the
    threshold simply returns no rows and is absent from the report.
    """
    if year is not None:
        years = [year]
    else:
        with common.Spinner("Finding the years the CONAFOR fires cover", logger):
            years = list(session.scalars(years_query()))

    measured: list[Row] = []
    for index, measuring in enumerate(years, start=1):
        with common.Spinner(f"Measuring the burnt area of the CONAFOR fires "
                            f"({measuring}: {index} of {len(years)})", logger):
            measured += [
                Row(country=record.country, year=measuring,
                    minimum=float(record.minimum),
                    maximum=float(record.maximum),
                    total=float(record.total),
                    fires=record.fires)
                for record in session.execute(statistics_query(measuring, method, min_area))
            ]

    rows = summarise(measured)
    logger.info("Computed %d rows over %d year(s) (%s areas, %s)",
                len(rows), len({row.year for row in rows if not row.is_total}), method,
                "every fire" if min_area is None else f"fires of {min_area:g} ha or more")
    return rows


def write_csv(rows: list[Row], path: Path, logger: logging.Logger) -> None:
    """Write the report as CSV.

    The numbers go out unformatted apart from being rounded to two decimals — no
    thousands separators — because a CSV is read by another program far more often
    than by a person, and a separator would make every figure a string.
    """
    path.parent.mkdir(parents=True, exist_ok=True)
    with path.open("w", newline="", encoding="utf-8") as handle:
        writer = csv.writer(handle)
        writer.writerow(COLUMNS)
        for row in rows:
            writer.writerow([row.country, row.year_label, row.fires,
                             f"{row.minimum:.2f}", f"{row.maximum:.2f}", f"{row.total:.2f}"])
    logger.info("Wrote %s", path)


def write_docx(rows: list[Row], path: Path, year: int | None,
               logger: logging.Logger,
               method: str = AREA_METHOD_GEODESIC,
               min_area: float | None = None) -> None:
    """Write the report as a Word document.

    One table, with the summary row in bold. Numbers get thousands separators here
    — the opposite of the CSV, and for the opposite reason: this one is for reading.

    The two caveats go into the document and not only into the manual: a table of
    these figures read without them would be misread twice over, once across 2016
    and once at 2010.
    """
    # Imported here rather than at module scope so that --csv keeps working if
    # python-docx is not installed, which matters because it is the only dependency
    # this application adds.
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt

    document = Document()
    document.add_heading("CONAFOR wildfire burnt area (Mexico)", level=1)

    if method == AREA_METHOD_GEODESIC:
        measured = "geodesically on the WGS84 ellipsoid"
    elif method == AREA_METHOD_EQUAL_AREA:
        measured = f"in the equal-area projection EPSG:{EQUAL_AREA_SRID}"
    else:
        measured = "as published by CONAFOR (AREA_HA), measured by nobody here"

    # The threshold belongs in the document and not only in the command line that
    # produced it: a table of the fires over 5 ha and a table of every fire look
    # exactly alike, and the difference is the whole meaning of the figures.
    scope = [f"year: {year}" if year is not None else "all published years"]
    if min_area is not None:
        scope.append(f"only fires of {min_area:g} ha or more")
    document.add_paragraph(
        f"Areas in hectares, {measured}. Years are the year of the archive the fire was "
        f"published in. Scope: {'; '.join(scope)}."
    )
    document.add_paragraph(
        "The series runs 2010 to 2023; a year missing from the table is a year not "
        "imported. The counts are not comparable across 2016: before then CONAFOR "
        "published only the fires it had drawn, and from 2016 it publishes the season."
    )
    document.add_paragraph(
        "The 2010 figures stand apart. The areas that layer publishes do not describe "
        "its polygons — the median ratio between them is 3.0 and the 90th percentile 65 "
        "— so its row means a different thing under 'reported' than under either "
        "measured method. From 2016 the two agree to three decimals."
    )

    table = document.add_table(rows=1, cols=len(COLUMNS))
    table.style = "Table Grid"
    for cell, heading in zip(table.rows[0].cells, COLUMNS):
        cell.text = heading
        cell.paragraphs[0].runs[0].bold = True

    for row in rows:
        cells = table.add_row().cells
        values = [row.country, row.year_label, f"{row.fires:,}",
                  f"{row.minimum:,.2f}", f"{row.maximum:,.2f}", f"{row.total:,.2f}"]
        for index, (cell, value) in enumerate(zip(cells, values)):
            cell.text = value
            paragraph = cell.paragraphs[0]
            if index >= FIRST_NUMERIC_COLUMN:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            for run in paragraph.runs:
                run.bold = row.is_total
                run.font.size = Pt(9)

    path.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(path))
    logger.info("Wrote %s", path)


def report(args: argparse.Namespace, engine: Engine, logger: logging.Logger) -> list[Row]:
    """Compute the statistics and write whichever outputs were asked for."""
    # No spinner here: compute runs one statement per year and turns one of its own
    # for each, which is the only honest place to say how far along it is.
    with Session(engine) as session:
        rows = compute(session, args.year, logger, args.area_method, args.min_area)

    if not rows:
        # An empty report is almost always a year with no data, and writing an empty
        # file would hide that. A threshold is named when there is one, because then
        # it is at least as likely to be the reason as the year is.
        threshold = "" if args.min_area is None else \
            f" No fire reached the --min-area of {args.min_area:g} ha."
        raise RuntimeError(
            "No wildfires matched. Check --year — a year whose archive was never "
            "imported has nothing to find — and that the CONAFOR fires are imported."
            + threshold
        )

    if args.csv is not None:
        write_csv(rows, args.csv, logger)
    if args.docx is not None:
        write_docx(rows, args.docx, args.year, logger, args.area_method, args.min_area)
    return rows


def main(argv: list[str] | None = None) -> int:
    args = parse_arguments(argv)
    logging.basicConfig(level=args.log_level, format="%(asctime)s %(levelname)s %(message)s")
    logger = logging.getLogger("conafor-statistics")

    try:
        settings = common.resolve_database_settings(args)
    except RuntimeError as error:
        logger.error("%s", error)
        return 1

    engine = create_engine(common.database_url(settings))
    try:
        report(args, engine, logger)
    except Exception as error:  # noqa: BLE001  (the CLI boundary: report, do not traceback)
        logger.error("%s", error)
        return 1
    finally:
        engine.dispose()
    return 0


if __name__ == "__main__":  # pragma nocover
    sys.exit(main())
