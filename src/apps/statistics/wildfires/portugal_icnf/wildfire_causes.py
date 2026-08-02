#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Wildfire counts by cause for the Portuguese ICNF burnt area cartography.

Reports, per year, how many fires there were, how many of them carry a cause at
all, and how many of those were classified ``Natural``::

    Country          Year    Fires  Classified   Natural   Natural %
    Portugal         2024      6820        6820       152        2.23
    Portugal         2023      7541        7541       198        2.63
    Portugal         2013      5106           0         0
    Portugal         Total    68435       18955       841        4.44

Run it over everything, or narrow it to one year::

    python3 -m src.apps.statistics.wildfires.portugal_icnf.wildfire_causes --csv causes.csv
    python3 -m src.apps.statistics.wildfires.portugal_icnf.wildfire_causes \\
        --year 2024 --csv 2024.csv --docx 2024.docx
    python3 -m src.apps.statistics.wildfires.portugal_icnf.wildfire_causes \\
        --cause-type Intencional --csv arson.csv

At least one of ``--csv`` and ``--docx`` is required.

The companion of
:mod:`~src.apps.statistics.wildfires.portugal_icnf.wildfire_statistics`, over the
same fires, the same years and the same rule for which fires are counted — so the
``Country``, ``Year`` and ``Fires`` columns of the two reports agree row for row
and the pair can be read side by side. What that one measures in hectares, this
one counts by cause.

There is no lightning in this dataset
-------------------------------------

The obvious question is how many fires lightning started, and **the ICNF does not
publish that**. Its ``Causa_Tipo`` has five values — ``Negligente``,
``Intencional``, ``Desconhecida``, ``Reacendimento`` and ``Natural`` — and its
twenty-four ``Causa_Desc`` values include exactly one natural description,
``Naturais``, which is not broken down any further. Lightning is not named
anywhere in the classification.

``Natural`` is therefore the finest answer this dataset supports, and it is what
this report counts by default. On the Iberian peninsula a natural fire is
overwhelmingly a lightning fire — Spain's EGIF, which *does* separate them, has
``Rayo`` as its whole ``100`` family — so ``Natural`` is a good proxy. It is a
proxy all the same, and the column is named ``Natural`` rather than ``Lightning``
so that nothing downstream mistakes one for the other.

``--cause-type`` counts a different one of the five instead. The type is matched
on the published Portuguese, which is what the column holds; the report's column
heading uses the English from
:data:`~src.providers.portugal_icnf.fire_cause.TYPE_TRANSLATIONS`.

Why ``Classified`` is a column and the percentage is of it
----------------------------------------------------------

**Only fires from 2014 on carry a cause: 18,955 of 68,435.** Every fire before
that has ``cause_id`` ``NULL``, and so does an unclassified fire in a year that
otherwise has them.

A percentage of *all* fires would therefore be a statement about how much of the
archive has been classified rather than about what caused the fires — it would
read 0.00 for every year from 1975 to 2013, and those years are not years without
natural fires. So the denominator is the classified fires, and the count of them
is a column of its own so that the denominator is never out of sight.

Where nothing is classified there is no percentage to give, and the cell is left
**empty** rather than filled with a zero that would be a claim.

.. warning::

   For the same reason, do not compare a ``Natural %`` from before 2014 with one
   from after: there is no such thing as the first. And take even the classified
   share as a floor — ``Desconhecida`` and ``Indeterminadas`` are real categories
   here, and a fire whose cause was never determined may well have been natural.

Which fires are counted
-----------------------

Exactly the rule its companion report uses, so the two agree: a fire needs a
perimeter and a country. Fires with no country are excluded — a perimeter in the
sea, and any fire that matched no OCHA boundary — and so is any fire with no
perimeter.

Counting a fire the other report cannot measure would be defensible on its own,
but it would mean the ``Fires`` column of the two reports quietly disagreeing, and
the whole value of a companion report is that it does not.

Which year a fire counts towards
--------------------------------

:attr:`~src.providers.portugal_icnf.wildfire.IcnfWildfire.year`, the published
``Ano``, exactly as in the companion report and for the same reason: 71% of these
fires publish no date and carry a 1 January placeholder built from this column in
the first place.

Shared with the companion report
--------------------------------

``--country-source``, the years query and the country ordering are **imported
from** :mod:`~src.apps.statistics.wildfires.portugal_icnf.wildfire_statistics`
rather than copied. Two reports over one dataset that disagreed about which
country a fire is in would be worse than one report, and a copy is a thing that
drifts.
"""

from __future__ import annotations

import argparse
import csv
import logging
import os
import sys

from dataclasses import dataclass
from pathlib import Path

from sqlalchemy import Engine
from sqlalchemy import Select
from sqlalchemy import create_engine
from sqlalchemy import func
from sqlalchemy import select
from sqlalchemy.orm import Session

import src.settings  # noqa: F401  (imported for the side effect of loading .env)

from src.apps.imports import common
from src.apps.statistics.wildfires.portugal_icnf.wildfire_statistics import COUNTRY_SOURCES
from src.apps.statistics.wildfires.portugal_icnf.wildfire_statistics import (
    COUNTRY_SOURCE_GEOMETRY,
)
from src.apps.statistics.wildfires.portugal_icnf.wildfire_statistics import PUBLISHED_YEAR
from src.apps.statistics.wildfires.portugal_icnf.wildfire_statistics import TOTAL_LABEL
from src.apps.statistics.wildfires.portugal_icnf.wildfire_statistics import country_columns
from src.apps.statistics.wildfires.portugal_icnf.wildfire_statistics import ordered_countries
from src.apps.statistics.wildfires.portugal_icnf.wildfire_statistics import years_query
from src.data_model.wildfire import Wildfire
from src.providers.portugal_icnf.fire_cause import TYPE_TRANSLATIONS
from src.providers.portugal_icnf.fire_cause import IcnfFireCause
from src.providers.portugal_icnf.wildfire import IcnfWildfire

#: The ``Causa_Tipo`` counted unless ``--cause-type`` says otherwise, and the
#: nearest thing the ICNF publishes to a lightning fire. See the module docstring:
#: there is no lightning category in this dataset, and this is not one.
DEFAULT_CAUSE_TYPE = "Natural"

#: The five types the ICNF publishes, taken from the translation table on the model
#: so that a type added there becomes selectable here without a second edit.
CAUSE_TYPES = tuple(TYPE_TRANSLATIONS)

#: Index of the first column that holds a number, and so is right-aligned in the
#: Word table.
FIRST_NUMERIC_COLUMN = 2


def columns(cause_type: str = DEFAULT_CAUSE_TYPE) -> tuple[str, ...]:
    """The report's columns, in order, for a given cause type.

    Notes
    -----
    A function rather than a constant, unlike the companion report's ``COLUMNS``,
    because two of the headings name the type being counted: a file of
    ``Intencional`` counts under a heading saying ``Natural`` would be a trap.

    Both output formats read them from here, so a change to one cannot silently
    leave the other behind. The first three are the companion report's first three,
    unchanged, which is what lets the two be read side by side.
    """
    label = TYPE_TRANSLATIONS.get(cause_type, cause_type)
    return ("Country", "Year", "Fires", "Classified", label, f"{label} (% of classified)")


def counts_query(year: int,
                 cause_type: str = DEFAULT_CAUSE_TYPE,
                 country_source: str = COUNTRY_SOURCE_GEOMETRY) -> Select:
    """Build the counting query for one year.

    Parameters
    ----------
    year : int
        The published ``Ano`` to count.
    cause_type : str
        The ``Causa_Tipo`` to count, in the published Portuguese.
    country_source : str
        One of :data:`~...wildfire_statistics.COUNTRY_SOURCES`.

    Returns
    -------
    Select
        A query yielding ``country, fires, classified, matching``: one row per
        country the year's fires turn out to be in, unordered. The summary rows and
        the report's order are :func:`summarise`'s work.

    Notes
    -----
    Three counts in one pass, which is what makes this report cheap: whatever the
    country source costs, it is paid once per fire and not once per count.

    The join to ``icnf_fire_cause`` is **outer**, and has to be: ``cause_id`` is
    ``NULL`` on every fire before 2014, and an inner join would silently turn this
    into a report of the classified fires alone — the ``Fires`` column would stop
    matching the companion report's and nothing would say why.

    ``count(cause_id)`` and not ``count(*)`` for the classified column: counting a
    nullable column counts the rows where it is filled in, which is the definition
    of classified here.

    The cause type is compared against the **published Portuguese**
    (:attr:`~src.providers.portugal_icnf.fire_cause.IcnfFireCause.type`) and not
    against the English translation beside it, because the Portuguese is what the
    provider published and is ``NOT NULL``; ``type_en`` is ``None`` for any category
    a later release invents.
    """
    icnf = IcnfWildfire.__table__
    cause = IcnfFireCause.__table__
    country_name, joins = country_columns(country_source)

    counts = (
        select(
            country_name.label("country"),
            func.count().label("fires"),
            func.count(icnf.c.cause_id).label("classified"),
            func.count().filter(cause.c.type == cause_type).label("matching"),
        )
        .select_from(Wildfire)
        .join(icnf, icnf.c.id == Wildfire.id)
        .outerjoin(cause, cause.c.id == icnf.c.cause_id)
        .where(Wildfire.perimeter.is_not(None))
        .where(PUBLISHED_YEAR == year)
    )
    for target, condition, is_outer in joins:
        counts = counts.outerjoin(target, condition) if is_outer \
            else counts.join(target, condition)

    return counts.group_by(country_name)


@dataclass(frozen=True)
class Row:
    """One line of the report.

    Attributes
    ----------
    country : str
        Name of the country the fires burnt in.
    year : int or None
        The published year, or ``None`` for the summary row.
    fires : int
        Every fire counted — with a perimeter, and in this country.
    classified : int
        How many of them carry a cause at all. Zero for every year before 2014.
    matching : int
        How many carry the cause type asked for.

    Notes
    -----
    ``matching <= classified <= fires`` always, and the first two can be zero for
    entirely different reasons: nothing of that cause, or nothing classified.
    """

    country: str
    year: int | None
    fires: int
    classified: int
    matching: int

    @property
    def is_total(self) -> bool:
        """Whether this is the summary row rather than one of the years."""
        return self.year is None

    @property
    def year_label(self) -> str:
        return TOTAL_LABEL if self.is_total else str(self.year)

    @property
    def share(self) -> float | None:
        """``matching`` as a percentage of ``classified``, or ``None`` if none is.

        ``None`` and not zero: a year in which nothing was classified has no share
        to report, and zero would be a claim that none of its fires was natural.
        The writers turn it into an empty cell.
        """
        if not self.classified:
            return None
        return 100.0 * self.matching / self.classified

    @property
    def share_label(self) -> str:
        """The percentage as it is written out, empty where there is none."""
        return "" if self.share is None else f"{self.share:.2f}"


def combine(rows: list[Row], country: str, year: int | None) -> Row:
    """One row summarising several: the three counts added up.

    Notes
    -----
    Counts decompose over a partition of the fires, so the ``Total`` row is what a
    single pass over the same fires would have returned.

    The share is deliberately **not** averaged over the years: it is recomputed by
    :attr:`Row.share` from the summed counts, which is the ratio of the totals
    rather than the mean of the ratios. A year with eleven classified fires and a
    year with eleven thousand must not weigh the same in the answer.
    """
    return Row(
        country=country,
        year=year,
        fires=sum(row.fires for row in rows),
        classified=sum(row.classified for row in rows),
        matching=sum(row.matching for row in rows),
    )


def summarise(measured: list[Row], countries: list[str]) -> list[Row]:
    """Build the report from the years counted: the summary rows, in order.

    Returns
    -------
    list of Row
        Each country, its years newest first and its summary row last.
    """
    report: list[Row] = []
    for name in countries:
        rows = [row for row in measured if row.country == name]
        if not rows:
            continue
        report += sorted(rows, key=lambda row: row.year, reverse=True)
        report.append(combine(rows, name, None))
    return report


def parse_arguments(argv: list[str] | None = None) -> argparse.Namespace:
    """Parse the command line."""
    parser = argparse.ArgumentParser(
        description="Wildfire counts by cause for the Portuguese ICNF burnt area "
                    "cartography.",
        epilog="The ICNF publishes no lightning category; 'Natural' is the closest it "
               "has and is what is counted by default. Only fires from 2014 on carry a "
               "cause at all. The ICNF publishes one country's fires, so there is no "
               "--country option. Database settings not given here are read from the "
               "environment (.env).",
    )
    parser.add_argument("-y", "--year", type=int,
                        help="restrict to one year, e.g. 2024; this is the published Ano, "
                             "not the year of the start date")
    parser.add_argument("--cause-type", default=DEFAULT_CAUSE_TYPE, choices=CAUSE_TYPES,
                        help="which published Causa_Tipo to count, in Portuguese: "
                             "'Natural' (default) is the nearest the ICNF comes to a "
                             "lightning fire — it publishes no lightning category, and "
                             "'Naturais' is its only natural description")
    parser.add_argument("--country-source", default=COUNTRY_SOURCE_GEOMETRY,
                        choices=COUNTRY_SOURCES,
                        help="how to decide which country a fire counts towards: "
                             "'geometry' (default) tests the perimeter against the real "
                             "country polygons at report time, so a perimeter digitised "
                             "into the sea drops out; 'reported' trusts the "
                             "admin_boundary_id the import stored, which is much faster "
                             "and, for this dataset, the same answer")

    # Accepted only so that it can be refused clearly, exactly as in the companion
    # report: without it argparse's prefix matching resolves "--country" to
    # "--country-source" and complains about an invalid choice.
    parser.add_argument("--country", help=argparse.SUPPRESS)

    output = parser.add_argument_group("output", "at least one is required")
    output.add_argument("--csv", type=Path, help="write the report to this .csv")
    output.add_argument("--docx", type=Path, help="write the report to this .docx (MS Word)")

    common.add_database_arguments(parser)
    parser.add_argument("--log-level", default=os.getenv("GISFIRE_LOG_LEVEL", "INFO"),
                        choices=["CRITICAL", "ERROR", "WARNING", "INFO", "DEBUG"],
                        help="verbosity (env: GISFIRE_LOG_LEVEL, default INFO)")

    arguments = parser.parse_args(argv)
    if arguments.country is not None:
        parser.error(
            "there is no --country here: the ICNF publishes one country's fires, so "
            "there is nothing to select between. Every fire is counted, and the "
            "Country column says which country each one turned out to be in — see "
            "--country-source."
        )
    if arguments.csv is None and arguments.docx is None:
        parser.error("nothing to write: pass --csv, --docx, or both")
    return arguments


def compute(session: Session, year: int | None, logger: logging.Logger,
            cause_type: str = DEFAULT_CAUSE_TYPE,
            country_source: str = COUNTRY_SOURCE_GEOMETRY) -> list[Row]:
    """Count the fires a year at a time, returning the report's rows in order.

    Notes
    -----
    One statement per year, under a spinner of its own — the same shape as the
    companion report, and for the same reason: whatever the point-in-polygon test
    of ``--country-source geometry`` costs, it is released when its statement ends.

    Every one of them runs in ``session``'s transaction, and so against a single
    snapshot.
    """
    if year is not None:
        years = [year]
    else:
        with common.Spinner("Finding the years the ICNF fires cover", logger):
            years = list(session.scalars(years_query()))

    measured: list[Row] = []
    for index, counting in enumerate(years, start=1):
        with common.Spinner(f"Counting the ICNF fires by cause "
                            f"({counting}: {index} of {len(years)})", logger):
            measured += [
                Row(country=record.country, year=counting,
                    fires=record.fires,
                    classified=record.classified,
                    matching=record.matching)
                for record in session.execute(
                    counts_query(counting, cause_type, country_source))
            ]

    rows = summarise(measured, ordered_countries(session, {row.country for row in measured}))

    classified = sum(row.classified for row in measured)
    if measured and not classified:
        # Almost always a run over years before 2014, and a table of zeros with no
        # explanation would look like an answer rather than an absence.
        logger.warning(
            "No fire in scope carries a cause at all, so every %s count is zero and no "
            "percentage can be given: the ICNF publishes a cause only from 2014 on",
            cause_type)

    logger.info("Counted %d rows over %d year(s) (%s fires, country from %s)",
                len(rows), len({row.year for row in rows if not row.is_total}),
                cause_type, country_source)
    return rows


def write_csv(rows: list[Row], path: Path, logger: logging.Logger,
              cause_type: str = DEFAULT_CAUSE_TYPE) -> None:
    """Write the report as CSV.

    The percentage is written bare and rounded to two decimals, and is **empty**
    where nothing was classified — an empty field reads as no answer to whatever
    parses this, which is what it is, while a zero would read as an answer.
    """
    path.parent.mkdir(parents=True, exist_ok=True)
    with path.open("w", newline="", encoding="utf-8") as handle:
        writer = csv.writer(handle)
        writer.writerow(columns(cause_type))
        for row in rows:
            writer.writerow([row.country, row.year_label, row.fires,
                             row.classified, row.matching, row.share_label])
    logger.info("Wrote %s", path)


def write_docx(rows: list[Row], path: Path, year: int | None,
               logger: logging.Logger,
               cause_type: str = DEFAULT_CAUSE_TYPE) -> None:
    """Write the report as a Word document.

    One table, with the summary row in bold. Counts get thousands separators here —
    the opposite of the CSV, and for the opposite reason: this one is for reading.

    The opening paragraphs say that the ICNF publishes no lightning category and
    that only fires from 2014 on carry a cause. Both belong in the document and not
    only in the manual: a table of natural fires read without either would be
    misread twice over.
    """
    # Imported here rather than at module scope so that --csv keeps working if
    # python-docx is not installed, which matters because it is the only
    # dependency this application adds.
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt

    heading = columns(cause_type)[4]
    document = Document()
    document.add_heading(f"ICNF wildfire counts — {heading} causes (Portugal)", level=1)

    scope = f"year: {year}" if year is not None else "all years"
    document.add_paragraph(
        f"Counts of ICNF wildfires classified as {cause_type} ({heading}). Fires not "
        f"attributable to a country are excluded, as is any fire with no perimeter. "
        f"Years are the published Ano. Scope: {scope}."
    )
    document.add_paragraph(
        "The ICNF publishes no lightning category: Natural is the closest its "
        "classification comes, and it is not broken down further. These counts are "
        "therefore a proxy for lightning fires, not a count of them."
    )
    document.add_paragraph(
        "Only fires from 2014 on carry a cause — 18,955 of 68,435 in the published "
        "archive. The percentage is of the classified fires and not of all of them, and "
        "is left blank for a year in which none was classified."
    )

    table = document.add_table(rows=1, cols=len(columns(cause_type)))
    table.style = "Table Grid"
    for cell, title in zip(table.rows[0].cells, columns(cause_type)):
        cell.text = title
        cell.paragraphs[0].runs[0].bold = True

    for row in rows:
        cells = table.add_row().cells
        values = [row.country, row.year_label, f"{row.fires:,}",
                  f"{row.classified:,}", f"{row.matching:,}", row.share_label]
        for index, (cell, value) in enumerate(zip(cells, values)):
            cell.text = value
            paragraph = cell.paragraphs[0]
            if index >= FIRST_NUMERIC_COLUMN:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            for run in paragraph.runs:
                run.bold = row.is_total
                run.font.size = Pt(9)

    path.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(path))
    logger.info("Wrote %s", path)


def report(args: argparse.Namespace, engine: Engine, logger: logging.Logger) -> list[Row]:
    """Count the fires and write whichever outputs were asked for."""
    with Session(engine) as session:
        rows = compute(session, args.year, logger, args.cause_type, args.country_source)

    if not rows:
        # An empty report is almost always a year with no data, and writing an
        # empty file would hide that. Note that a year with no *classified* fires
        # is not empty — it is a row of zeros, which is a different thing and is
        # reported as one.
        raise RuntimeError(
            "No wildfires matched. Check --year, and that the ICNF fires and the OCHA "
            "boundaries are both imported — fires with no country are not counted."
        )

    if args.csv is not None:
        write_csv(rows, args.csv, logger, args.cause_type)
    if args.docx is not None:
        write_docx(rows, args.docx, args.year, logger, args.cause_type)
    return rows


def main(argv: list[str] | None = None) -> int:
    args = parse_arguments(argv)
    logging.basicConfig(level=args.log_level, format="%(asctime)s %(levelname)s %(message)s")
    logger = logging.getLogger("icnf-causes")

    try:
        settings = common.resolve_database_settings(args)
    except RuntimeError as error:
        logger.error("%s", error)
        return 1

    engine = create_engine(common.database_url(settings))
    try:
        report(args, engine, logger)
    except Exception as error:  # noqa: BLE001  (the CLI boundary: report, do not traceback)
        logger.error("%s", error)
        return 1
    finally:
        engine.dispose()
    return 0


if __name__ == "__main__":  # pragma nocover
    sys.exit(main())
