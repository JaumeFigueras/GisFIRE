#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Burnt-area statistics for the Portuguese ICNF burnt area cartography.

Reports, per year, how many fires there were and the smallest, largest and total
area burnt, in hectares::

    Country          Year    Fires     Minimum      Maximum        Total
    Portugal         2024      6820        0.51     11485.32     135872.44
    Portugal         2023      7541        0.50      8802.11      88214.06
    Portugal         Total   68435        0.50     97462.10    4218774.51

Run it over everything, or narrow it to one year, or to the fires above a size::

    python3 -m src.apps.statistics.wildfires.portugal_icnf.wildfire_statistics --csv burnt.csv
    python3 -m src.apps.statistics.wildfires.portugal_icnf.wildfire_statistics \\
        --year 2024 --csv 2024.csv --docx 2024.docx
    python3 -m src.apps.statistics.wildfires.portugal_icnf.wildfire_statistics \\
        --min-area 5 --csv over-5-ha.csv

At least one of ``--csv`` and ``--docx`` is required: an application that
computed a report and then printed nothing would be a strange thing.

The application only reads. Database settings come from the environment
(``.env``, see :mod:`src.settings`); every one of them can be overridden with a
command-line argument.

One country, and no ``--country``
---------------------------------

The ICNF publishes one country's fires, so there is nothing to select between and
the option is not offered. The **column** stays, for two reasons: it keeps this
report's CSV the same shape as the
:mod:`GWIS <src.apps.statistics.wildfires.gwis.wildfire_statistics>` and
:mod:`GFA <src.apps.statistics.wildfires.gfa.wildfire_statistics>` ones, so the
three can be concatenated and compared; and under ``--country-source geometry``
it is not always ``Portugal``. A perimeter that reaches across the Spanish border,
or one that falls in the sea, is exactly what that mode is for, and the column is
where the answer shows up.

Which year a fire counts towards
--------------------------------

:attr:`~src.providers.portugal_icnf.wildfire.IcnfWildfire.year` — the ``Ano`` the ICNF
published, which is also the layer the fire came from — and **not** the year of
``start_date_time`` as in the other two reports.

That is the same rule reached more directly, and here it matters:
**48,860 of the 68,435 fires publish no date at all.** Those rows carry
``date_time_precision = 'year'`` and a ``start_date_time`` of the 1st of January
of their year, which is a placeholder satisfying a ``NOT NULL`` column rather than
a claim about the fire. Grouping on it would work — the year in the placeholder is
the right year — but it would mean routing 71% of the dataset through a value that
exists only because the column could not be null. ``year`` is the published
answer, is ``NOT NULL``, and needs no timezone applied to it.

.. warning::

   The consequence is worth stating plainly: **this report is sound, but a report
   grouped by month or day over the same data would not be.** 71% of these fires
   would all fall on the 1st of January. Anything finer than a year has to filter
   on ``date_time_precision`` first.

One year at a time
------------------

The report is not one statement. The years are found first, then each is
measured by a statement of its own and the summary rows are computed from their
results.

This dataset does not need it. Sixty-eight thousand perimeters are a minute's
work however they are grouped, and ``GROUPING SETS`` over a single pass — which
is what each year's statement still does internally — would do here. It is built
this way because its
:mod:`GWIS <src.apps.statistics.wildfires.gwis.wildfire_statistics>` and
:mod:`GFA <src.apps.statistics.wildfires.gfa.wildfire_statistics>` counterparts
have to be: at twenty million perimeters the single statement took a 30 GB
machine to 29.2 GB in four and a half hours and was killed by the OOM killer,
because the memory a point-in-polygon test against a country polygon needs is
only released when the statement ends. Three reports meant to be read side by
side are worth keeping as one program over three datasets, and the difference
here costs one extra pass over an indexed column.

Nothing about the figures changes. ``count``, ``sum``, ``min`` and ``max`` all
decompose over a partition of the fires, so the ``Total`` row is exactly the
number ``GROUPING SETS`` returned, from the same rows. Every statement runs in
one transaction and so against one snapshot.

How the area is measured
------------------------

The same two ways as the GFA report, chosen with ``--area-method``: ``geodesic``
on the WGS84 ellipsoid (the default) or ``equal-area`` in EPSG:6933. They agree to
within 0.003%.

Why not the CRS the ICNF publishes in
--------------------------------------

Tempting, and wrong. The perimeters are also stored as published, in EPSG:3763
(ETRS89 / Portugal TM06), on
:attr:`~src.providers.portugal_icnf.wildfire.IcnfWildfire.perimeter_etrs89_tm06` — a
projected national grid in metres, whose ``ST_Area`` reproduces the ICNF's own
``AreaHaSIG``. It is not offered as an area method because **PT-TM06 is a
transverse Mercator, not an equal-area projection**, and its distortion away from
the central meridian is not negligible for this dataset:

.. code-block:: text

   the same polygon, geodesic vs EPSG:3763
     Lisbon      (-9.1, 38.7)     +0.017%
     Bragança    (-6.8, 41.8)     +0.030%
     Faro        (-7.9, 37.0)     +0.001%
     Madeira    (-17.0, 32.7)     +1.717%
     Azores     (-28.0, 38.5)     +7.631%

On the mainland it is fine. On the islands it is not, and a report cannot know in
advance that no island fire will ever appear in it. See
:attr:`~src.providers.portugal_icnf.wildfire.IcnfWildfire.area_ha_gis` below for the right
way to reproduce the published figures.

The published areas are not used
--------------------------------

:attr:`~src.providers.portugal_icnf.wildfire.IcnfWildfire.area_ha_gis` (``AreaHaSIG``,
measured from the polygon) and
:attr:`~src.providers.portugal_icnf.wildfire.IcnfWildfire.area_ha_sgif` (``AreaHaSGIF``,
what the fire database recorded) are kept as published and are **not** what this
reports. They are two independent measurements of the same fire and are worth
having as a check on this one, but a report that mixed them with a measured area
would be comparing three different things in one column.

Which fires are counted
-----------------------

Fires with no country are **excluded** — a perimeter in the sea, and any fire that
matched no OCHA boundary. So is any fire with no perimeter. A run therefore does
not necessarily account for every hectare in the database, and the totals here are
totals *of attributable burnt area*.

``--min-area`` narrows it further, to the fires of at least that many hectares.
By default there is no threshold and every attributable fire counts.

.. note::

   The 1975-1999 layers only map fires of 5 ha or more; from 2000 on the small
   ones are mapped too. A count of fires per year is therefore **not comparable
   across 1999**, and neither is a minimum. The totals are much less affected,
   small fires being small.

   This is what ``--min-area 5`` is for: applying the old layers' mapping rule to
   the new ones puts every year on the same footing and makes the counts
   comparable again. It is a threshold on the area *this report measures* — the
   one ``--area-method`` selects — and not on the published ``AreaHaSIG``, so that
   the fires counted are exactly the fires the figures beside them are computed
   from.
"""

from __future__ import annotations

import argparse
import csv
import logging
import math
import os
import sys

from dataclasses import dataclass
from pathlib import Path

from geoalchemy2 import Geography
from sqlalchemy import ColumnElement
from sqlalchemy import Engine
from sqlalchemy import Select
from sqlalchemy import cast
from sqlalchemy import create_engine
from sqlalchemy import func
from sqlalchemy import select
from sqlalchemy import true
from sqlalchemy.orm import Session

import src.settings  # noqa: F401  (imported for the side effect of loading .env)

from src.apps.imports import common
from src.data_model.geography.admin_boundary import AdminBoundary
from src.data_model.wildfire import Wildfire
from src.providers.ocha.admin_boundary import OchaAdminBoundary
from src.providers.portugal_icnf.wildfire import IcnfWildfire

#: Label used in the ``Year`` column for the summary row.
TOTAL_LABEL = "Total"

#: The report's columns, in order, shared by both output formats so that a change
#: to one cannot silently leave the other behind. Deliberately identical to the
#: GWIS and GFA reports' columns, so the three CSVs can be concatenated.
COLUMNS = ("Country", "Year", "Fires", "Minimum (ha)", "Maximum (ha)", "Total (ha)")

#: Index of the first column that holds a number, and so is right-aligned in the
#: Word table.
FIRST_NUMERIC_COLUMN = 2

#: The two ways of deciding which country a fire counts towards.
COUNTRY_SOURCE_GEOMETRY = "geometry"
COUNTRY_SOURCE_REPORTED = "reported"
COUNTRY_SOURCES = (COUNTRY_SOURCE_GEOMETRY, COUNTRY_SOURCE_REPORTED)

#: Administrative level of a country in ``admin_boundary``.
COUNTRY_LEVEL = 0

#: A point guaranteed to lie inside the burnt perimeter, for the containment test.
#:
#: ``ST_PointOnSurface`` and not ``ST_Centroid``: the centroid of a crescent or a
#: ring-shaped burn can fall outside the polygon entirely, and would then be
#: tested against a country the fire never reached — or against no country at all.
LOCATOR = func.ST_PointOnSurface(Wildfire.perimeter)

#: Square metres in a hectare.
SQUARE_METRES_PER_HECTARE = 10_000.0

#: NSIDC EASE-Grid 2.0 Global — a cylindrical equal-area projection in metres,
#: defined for the whole world. The CRS behind ``--area-method equal-area``.
EQUAL_AREA_SRID = 6933

#: The two ways of turning a perimeter in degrees into hectares.
AREA_METHOD_GEODESIC = "geodesic"
AREA_METHOD_EQUAL_AREA = "equal-area"
AREA_METHODS = (AREA_METHOD_GEODESIC, AREA_METHOD_EQUAL_AREA)

#: The year a fire counts towards: the one the ICNF published.
#:
#: See the module docstring. Unlike the GWIS and GFA reports this does not derive
#: the year from ``start_date_time``, because for 71% of these fires that instant
#: is a placeholder built *from* this column in the first place.
PUBLISHED_YEAR = IcnfWildfire.__table__.c.year


def burnt_area(method: str) -> ColumnElement:
    """Burnt area of one fire in hectares, by whichever method was asked for.

    Parameters
    ----------
    method : str
        One of :data:`AREA_METHODS`.

    Returns
    -------
    ColumnElement
        The SQL expression yielding hectares.

    Raises
    ------
    ValueError
        If ``method`` is not one of :data:`AREA_METHODS`.

    Notes
    -----
    Always measured from the EPSG:4326 perimeter on the parent ``wildfire`` row,
    never from the published EPSG:3763 copy — see the module docstring for why the
    national grid is not an option here.

    The geodesic cast carries the geometry type and SRID rather than being a bare
    ``Geography()``: that renders as ``geography(GEOMETRY,-1)``, which PostGIS
    rejects, because -1 is not a SRID it knows.
    """
    if method == AREA_METHOD_GEODESIC:
        square_metres = func.ST_Area(
            cast(Wildfire.perimeter, Geography(geometry_type="MULTIPOLYGON", srid=4326))
        )
    elif method == AREA_METHOD_EQUAL_AREA:
        square_metres = func.ST_Area(func.ST_Transform(Wildfire.perimeter, EQUAL_AREA_SRID))
    else:
        raise ValueError(
            f"unknown area method {method!r}; expected one of {', '.join(AREA_METHODS)}"
        )
    return square_metres / SQUARE_METRES_PER_HECTARE


def country_columns(source: str) -> tuple[ColumnElement, list]:
    """Where a fire's country comes from, as ``(name, joins)``.

    Parameters
    ----------
    source : str
        One of :data:`COUNTRY_SOURCES`.

    Returns
    -------
    tuple
        The country-name expression, and the joins that have to be applied to a
        ``select`` over ``wildfire`` for it to resolve.

    Raises
    ------
    ValueError
        If ``source`` is not one of :data:`COUNTRY_SOURCES`.

    Notes
    -----
    Returns no ISO code, unlike its GWIS and GFA counterparts: those need one to
    resolve ``--country``, and this report has no such option.

    **``geometry``** ignores ``admin_boundary_id`` and asks the database which
    country actually contains the fire, at report time, against the real polygons.
    It is the default. For a single-country dataset its job is not to choose
    between countries but to catch the fires that are not in the country at all —
    a perimeter digitised into the Atlantic keeps its Portuguese
    ``admin_boundary_id`` and is silently in the total under ``reported``.

    **``reported``** trusts what the import stored. For ICNF that *is* a geometric
    answer — the import resolves the country by containment from a point on the
    perimeter — so the two modes agree, and this one is far cheaper: an index
    lookup on a foreign key instead of a point-in-polygon test per fire.
    """
    ocha_boundary = OchaAdminBoundary.__table__

    if source == COUNTRY_SOURCE_REPORTED:
        return AdminBoundary.name, [
            (AdminBoundary, AdminBoundary.id == Wildfire.admin_boundary_id, False),
        ]

    if source != COUNTRY_SOURCE_GEOMETRY:
        raise ValueError(
            f"unknown country source {source!r}; expected one of {', '.join(COUNTRY_SOURCES)}"
        )

    # LATERAL ... LIMIT 1 rather than a plain join: a point on a shared border can
    # satisfy ST_Contains for two countries, and one fire must not become two rows.
    # The join is inner, which is what drops a fire that is inside no country at
    # all — the same rule ``reported`` gets from its inner join to admin_boundary.
    containing = (
        select(AdminBoundary.name.label("name"))
        .select_from(AdminBoundary)
        .join(ocha_boundary, ocha_boundary.c.id == AdminBoundary.id)
        .where(AdminBoundary.level == COUNTRY_LEVEL)
        .where(func.ST_Contains(AdminBoundary.geometry, LOCATOR))
        .limit(1)
        .lateral("containing")
    )
    return containing.c.name, [(containing, true(), False)]


def years_query() -> Select:
    """The years the dataset holds fires in, newest first.

    Returns
    -------
    Select
        A query yielding one ``int`` per year.

    Notes
    -----
    Run before anything else, because each of those years is then measured by a
    statement of its own — see the module docstring for why the report is built
    that way.

    ``DISTINCT`` over the whole table rather than ``min(year)`` to ``max(year)``:
    a gap in the record is a gap in the report, and a range would fill it with
    statements that can only return nothing.
    """
    icnf = IcnfWildfire.__table__
    year = PUBLISHED_YEAR.label("year")
    return (
        select(year)
        .select_from(Wildfire)
        .join(icnf, icnf.c.id == Wildfire.id)
        .where(Wildfire.perimeter.is_not(None))
        .distinct()
        .order_by(year.desc())
    )


def statistics_query(year: int,
                     method: str = AREA_METHOD_GEODESIC,
                     country_source: str = COUNTRY_SOURCE_GEOMETRY,
                     min_area: float | None = None) -> Select:
    """Build the statistics query for one year.

    Parameters
    ----------
    year : int
        The published ``Ano`` to measure.
    method : str
        One of :data:`AREA_METHODS`.
    country_source : str
        One of :data:`COUNTRY_SOURCES`.
    min_area : float, optional
        Count only fires of at least this many hectares. ``None``, the default,
        counts every fire.

    Returns
    -------
    Select
        A query yielding ``country, minimum, maximum, total, fires``: one row per
        country the year's fires turn out to be in, unordered. The summary rows
        and the report's order are :func:`summarise`'s work.

    Notes
    -----
    Built against the mapped classes rather than written as SQL text, so a column
    renamed on a model breaks this at import time rather than in front of a user.

    The inner query computes each area exactly once. Folded into the outer
    aggregate instead, the area expression would be evaluated three times per row
    — for the minimum, the maximum and the sum — and it is by far the most
    expensive thing here.

    It groups on the country as well, even though there is normally only one:
    under ``geometry`` there may not be, and a fire that turns out to be in Spain
    should not be silently folded into Portugal's total.

    ``icnf_wildfire`` is joined — by table, to keep SQLAlchemy from adding the
    polymorphic join of its own — rather than filtering on ``wildfire.type``, and
    it has to be joined in any case: :data:`PUBLISHED_YEAR` lives on it.

    ``min_area`` filters the subquery's column rather than repeating the area
    expression in a ``WHERE`` of its own, for the same reason the subquery exists
    at all, and it is applied before the aggregates rather than as a ``HAVING``:
    the threshold selects the fires the figures are computed from, it does not
    discard countries whose total came out small.
    """
    icnf = IcnfWildfire.__table__
    country_name, joins = country_columns(country_source)

    fires = (
        select(
            country_name.label("country"),
            burnt_area(method).label("hectares"),
        )
        .select_from(Wildfire)
        .join(icnf, icnf.c.id == Wildfire.id)
        .where(Wildfire.perimeter.is_not(None))
        .where(PUBLISHED_YEAR == year)
    )
    for target, condition, is_outer in joins:
        fires = fires.outerjoin(target, condition) if is_outer \
            else fires.join(target, condition)

    fire = fires.subquery("fire")
    statistics = (
        select(
            fire.c.country,
            func.min(fire.c.hectares).label("minimum"),
            func.max(fire.c.hectares).label("maximum"),
            func.sum(fire.c.hectares).label("total"),
            func.count().label("fires"),
        )
        .group_by(fire.c.country)
    )
    if min_area is not None:
        statistics = statistics.where(fire.c.hectares >= min_area)
    return statistics


@dataclass(frozen=True)
class Row:
    """One line of the report.

    Attributes
    ----------
    country : str
        Name of the country the fires burnt in. Normally ``"Portugal"``; see the
        module docstring for when it is not.
    year : int or None
        The published year, or ``None`` for the summary row.
    minimum, maximum, total : float
        Smallest single fire, largest single fire and sum of every fire, in
        hectares.
    fires : int
        How many fires the three area figures were computed from — the count of
        wildfire events in this year, or in the whole period for a summary row.
    """

    country: str
    year: int | None
    minimum: float
    maximum: float
    total: float
    fires: int

    @property
    def is_total(self) -> bool:
        """Whether this is the summary row rather than one of the years."""
        return self.year is None

    @property
    def year_label(self) -> str:
        return TOTAL_LABEL if self.is_total else str(self.year)


def combine(rows: list[Row], country: str, year: int | None) -> Row:
    """One row summarising several: the four figures taken over all of them.

    Notes
    -----
    This is what makes measuring a year at a time cost nothing. All four figures
    decompose over a partition of the fires — a minimum of minima is a minimum, a
    sum of sums is a sum — so the ``Total`` row is the same number ``GROUPING
    SETS`` produced from a single pass, and no fire is counted twice or left out.

    ``fsum`` rather than ``sum``: an exact accumulation over a handful of partial
    totals costs nothing and cannot drift from what one pass would have returned.
    """
    return Row(
        country=country,
        year=year,
        minimum=min(row.minimum for row in rows),
        maximum=max(row.maximum for row in rows),
        total=math.fsum(row.total for row in rows),
        fires=sum(row.fires for row in rows),
    )


def ordered_countries(session: Session, names: set[str]) -> list[str]:
    """The countries of the report, in the order the database would have sorted them.

    Notes
    -----
    Normally one name, ``Portugal``; under ``--country-source geometry`` there can
    be a second, which is the whole point of that mode.

    Sorted by PostgreSQL and not by Python because the two do not agree: Python
    compares code points, which puts every accented name after every unaccented
    one, while the database sorts by its collation, which is what this report has
    always been ordered by.

    A name the query does not return could only come of ``admin_boundary``
    changing underneath the run, but it is appended rather than dropped: losing a
    country from a report over an ordering detail would be a poor trade.
    """
    if not names:
        return []
    ordered = list(session.scalars(
        select(AdminBoundary.name)
        .distinct()
        .where(AdminBoundary.level == COUNTRY_LEVEL)
        .where(AdminBoundary.name.in_(names))
        .order_by(AdminBoundary.name)
    ))
    return ordered + sorted(names - set(ordered))


def summarise(measured: list[Row], countries: list[str]) -> list[Row]:
    """Build the report from the years measured: the summary rows, in order.

    Parameters
    ----------
    measured : list of Row
        One row per country and year, as the per-year statements returned them.
    countries : list of str
        The countries in scope, in the order they are to be reported.

    Returns
    -------
    list of Row
        Each country, its years newest first and its summary row last.
    """
    report: list[Row] = []
    for name in countries:
        rows = [row for row in measured if row.country == name]
        report += sorted(rows, key=lambda row: row.year, reverse=True)
        report.append(combine(rows, name, None))
    return report



def hectares(text: str) -> float:
    """Argparse type for ``--min-area``: a finite, non-negative number of hectares.

    Raises
    ------
    argparse.ArgumentTypeError
        If the text is not a number, or is negative, or is a non-finite float.

    Notes
    -----
    A bare ``type=float`` would accept ``-5``, ``nan`` and ``inf``. The first two
    are almost certainly a typo and would silently produce the unfiltered report —
    ``nan`` compares false against every area, so it would produce an empty one —
    and none of the three is a size a fire can have.
    """
    try:
        area = float(text)
    except ValueError:
        raise argparse.ArgumentTypeError(f"{text!r} is not a number of hectares")
    if not math.isfinite(area):
        raise argparse.ArgumentTypeError(f"{text!r} is not a finite number of hectares")
    if area < 0:
        raise argparse.ArgumentTypeError(
            f"a burnt area cannot be negative, and {area:g} ha is: pass 0 or more, or "
            f"leave --min-area out to count every fire"
        )
    return area


def parse_arguments(argv: list[str] | None = None) -> argparse.Namespace:
    """Parse the command line."""
    parser = argparse.ArgumentParser(
        description="Burnt-area statistics for the Portuguese ICNF burnt area cartography.",
        epilog="Areas are in hectares, geodesic on the WGS84 ellipsoid by default. The "
               "ICNF publishes one country's fires, so there is no --country option. "
               "Database settings not given here are read from the environment (.env).",
    )
    parser.add_argument("-y", "--year", type=int,
                        help="restrict to one year, e.g. 2024; this is the published Ano, "
                             "not the year of the start date")
    parser.add_argument("--min-area", type=hectares, default=None, metavar="HECTARES",
                        help="count only fires that burnt at least this many hectares, "
                             "e.g. 5; by default every fire counts. The threshold is "
                             "applied to the area this report measures (see "
                             "--area-method), not to the published AreaHaSIG. Useful "
                             "against the 1975-1999 layers, which only mapped fires of "
                             "5 ha or more")

    parser.add_argument("--area-method", default=AREA_METHOD_GEODESIC, choices=AREA_METHODS,
                        help="how to turn the EPSG:4326 perimeter into hectares: "
                             "'geodesic' measures on the WGS84 ellipsoid (default); "
                             "'equal-area' projects to EPSG:6933 and measures there. They "
                             "agree to within 0.003%%. The published EPSG:3763 grid is not "
                             "offered: it is conformal, not equal-area, and is 7.6%% out "
                             "in the Azores")
    parser.add_argument("--country-source", default=COUNTRY_SOURCE_GEOMETRY,
                        choices=COUNTRY_SOURCES,
                        help="how to decide which country a fire counts towards: "
                             "'geometry' (default) tests the perimeter against the real "
                             "country polygons at report time, so a perimeter digitised "
                             "into the sea drops out; 'reported' trusts the "
                             "admin_boundary_id the import stored, which is much faster "
                             "and, for this dataset, the same answer")

    # Accepted only so that it can be refused clearly. Without it argparse's prefix
    # matching resolves "--country" to "--country-source" and reports
    # "invalid choice: 'Portugal'", which names neither the real problem nor the
    # option the user actually wanted. Anyone reaching for --country here has
    # copied a command line from the GWIS or GFA report, which is a reasonable
    # thing to have done.
    parser.add_argument("--country", help=argparse.SUPPRESS)

    output = parser.add_argument_group("output", "at least one is required")
    output.add_argument("--csv", type=Path, help="write the report to this .csv")
    output.add_argument("--docx", type=Path, help="write the report to this .docx (MS Word)")

    common.add_database_arguments(parser)
    parser.add_argument("--log-level", default=os.getenv("GISFIRE_LOG_LEVEL", "INFO"),
                        choices=["CRITICAL", "ERROR", "WARNING", "INFO", "DEBUG"],
                        help="verbosity (env: GISFIRE_LOG_LEVEL, default INFO)")

    arguments = parser.parse_args(argv)
    if arguments.country is not None:
        parser.error(
            "there is no --country here: the ICNF publishes one country's fires, so "
            "there is nothing to select between. Every fire is reported, and the "
            "Country column says which country each one turned out to be in — see "
            "--country-source."
        )
    if arguments.csv is None and arguments.docx is None:
        parser.error("nothing to write: pass --csv, --docx, or both")
    return arguments


def compute(session: Session, year: int | None, logger: logging.Logger,
            method: str = AREA_METHOD_GEODESIC,
            country_source: str = COUNTRY_SOURCE_GEOMETRY,
            min_area: float | None = None) -> list[Row]:
    """Measure the fires a year at a time, returning the report's rows in order.

    Notes
    -----
    One statement per year, for the reason given in the module docstring, each
    under a spinner of its own.

    Every one of them runs in ``session``'s transaction, and so against a single
    snapshot: a report assembled from many queries is then exactly as consistent
    as one assembled from a single query.

    ``min_area`` does not narrow the list of years: which years the dataset holds
    fires in is a fact about the dataset, and a year whose fires are all below the
    threshold simply returns no rows and is absent from the report. Measuring it is
    one statement that finds nothing, which is cheaper than a second pass computing
    every area to decide in advance.
    """
    if year is not None:
        years = [year]
    else:
        with common.Spinner("Finding the years the ICNF fires cover", logger):
            years = list(session.scalars(years_query()))

    measured: list[Row] = []
    for index, measuring in enumerate(years, start=1):
        with common.Spinner(f"Measuring the burnt area of the ICNF fires "
                            f"({measuring}: {index} of {len(years)})", logger):
            measured += [
                Row(country=record.country, year=measuring,
                    minimum=float(record.minimum),
                    maximum=float(record.maximum),
                    total=float(record.total),
                    fires=record.fires)
                for record in session.execute(
                    statistics_query(measuring, method, country_source, min_area))
            ]

    rows = summarise(measured, ordered_countries(session, {row.country for row in measured}))
    logger.info("Computed %d rows over %d year(s) (%s areas, country from %s, %s)",
                len(rows), len({row.year for row in rows if not row.is_total}),
                method, country_source,
                "every fire" if min_area is None else f"fires of {min_area:g} ha or more")
    return rows


def write_csv(rows: list[Row], path: Path, logger: logging.Logger) -> None:
    """Write the report as CSV.

    The numbers go out unformatted apart from being rounded to two decimals — no
    thousands separators — because a CSV is read by another program far more often
    than by a person, and a separator would make every figure a string.
    """
    path.parent.mkdir(parents=True, exist_ok=True)
    with path.open("w", newline="", encoding="utf-8") as handle:
        writer = csv.writer(handle)
        writer.writerow(COLUMNS)
        for row in rows:
            writer.writerow([row.country, row.year_label, row.fires,
                             f"{row.minimum:.2f}", f"{row.maximum:.2f}", f"{row.total:.2f}"])
    logger.info("Wrote %s", path)


def write_docx(rows: list[Row], path: Path, year: int | None,
               logger: logging.Logger,
               method: str = AREA_METHOD_GEODESIC,
               min_area: float | None = None) -> None:
    """Write the report as a Word document.

    One table, with the summary row in bold. Numbers get thousands separators here
    — the opposite of the CSV, and for the opposite reason: this one is for
    reading.
    """
    # Imported here rather than at module scope so that --csv keeps working if
    # python-docx is not installed, which matters because it is the only
    # dependency this application adds.
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt

    document = Document()
    document.add_heading("ICNF wildfire burnt area (Portugal)", level=1)

    measured = ("geodesically on the WGS84 ellipsoid" if method == AREA_METHOD_GEODESIC
                else f"in the equal-area projection EPSG:{EQUAL_AREA_SRID}")
    # The threshold belongs in the document and not only in the command line that
    # produced it: a table of the fires over 5 ha and a table of every fire look
    # exactly alike, and the difference is the whole meaning of the figures.
    scope = [f"year: {year}" if year is not None else "all years"]
    if min_area is not None:
        scope.append(f"only fires of {min_area:g} ha or more")
    document.add_paragraph(
        f"Areas in hectares, computed {measured}. Fires not attributable to a country "
        f"are excluded. Years are the published Ano. Scope: {'; '.join(scope)}."
    )

    table = document.add_table(rows=1, cols=len(COLUMNS))
    table.style = "Table Grid"
    for cell, heading in zip(table.rows[0].cells, COLUMNS):
        cell.text = heading
        cell.paragraphs[0].runs[0].bold = True

    for row in rows:
        cells = table.add_row().cells
        values = [row.country, row.year_label, f"{row.fires:,}",
                  f"{row.minimum:,.2f}", f"{row.maximum:,.2f}", f"{row.total:,.2f}"]
        for index, (cell, value) in enumerate(zip(cells, values)):
            cell.text = value
            paragraph = cell.paragraphs[0]
            if index >= FIRST_NUMERIC_COLUMN:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            for run in paragraph.runs:
                run.bold = row.is_total
                run.font.size = Pt(9)

    path.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(path))
    logger.info("Wrote %s", path)


def report(args: argparse.Namespace, engine: Engine, logger: logging.Logger) -> list[Row]:
    """Compute the statistics and write whichever outputs were asked for."""
    # No spinner here: compute runs one statement per year and turns one of its
    # own for each, which is the only honest place to say how far along it is.
    with Session(engine) as session:
        rows = compute(session, args.year, logger, args.area_method, args.country_source,
                       args.min_area)

    if not rows:
        # An empty report is almost always a year with no data, and writing an
        # empty file would hide that. A threshold is named when there is one,
        # because then it is at least as likely to be the reason as the year is.
        threshold = "" if args.min_area is None else \
            f" No fire reached the --min-area of {args.min_area:g} ha."
        raise RuntimeError(
            "No wildfires matched. Check --year, and that the ICNF fires and the OCHA "
            "boundaries are both imported — fires with no country are not counted."
            + threshold
        )

    if args.csv is not None:
        write_csv(rows, args.csv, logger)
    if args.docx is not None:
        write_docx(rows, args.docx, args.year, logger, args.area_method, args.min_area)
    return rows


def main(argv: list[str] | None = None) -> int:
    args = parse_arguments(argv)
    logging.basicConfig(level=args.log_level, format="%(asctime)s %(levelname)s %(message)s")
    logger = logging.getLogger("icnf-statistics")

    try:
        settings = common.resolve_database_settings(args)
    except RuntimeError as error:
        logger.error("%s", error)
        return 1

    engine = create_engine(common.database_url(settings))
    try:
        report(args, engine, logger)
    except Exception as error:  # noqa: BLE001  (the CLI boundary: report, do not traceback)
        logger.error("%s", error)
        return 1
    finally:
        engine.dispose()
    return 0


if __name__ == "__main__":  # pragma nocover
    sys.exit(main())
