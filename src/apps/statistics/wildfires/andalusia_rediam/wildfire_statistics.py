#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Burnt-area statistics for the Andalusian REDIAM burnt area cartography.

Reports, per year, how many fires there were, the smallest, largest and total area
burnt, in hectares, and **how many of those fires are bound to the EGIF *parte* for
the same fire**::

    Country   Year   Fires  Minimum   Maximum      Total  EGIF matched  EGIF matched (%)
    Spain     2025      97    10.24   1439.68   10015.74             0              0.00
    Spain     2024      36    10.56   2169.34    8328.29             0              0.00
    Spain     2023      40    10.45    425.49    2442.62            31             77.50
    Spain     2022      58     3.05   5198.20   18689.75            58            100.00
    Spain     Total    907     0.02  15249.70  165522.45           759             83.68

Run it over everything, or narrow it to one year, to the fires above a size, or to
the bindings you are willing to trust::

    python3 -m src.apps.statistics.wildfires.andalusia_rediam.wildfire_statistics --csv burnt.csv
    python3 -m src.apps.statistics.wildfires.andalusia_rediam.wildfire_statistics \\
        --year 2022 --csv 2022.csv --docx 2022.docx
    python3 -m src.apps.statistics.wildfires.andalusia_rediam.wildfire_statistics \\
        --min-area 5 --csv over-5-ha.csv
    python3 -m src.apps.statistics.wildfires.andalusia_rediam.wildfire_statistics \\
        --min-confidence 0.9 --csv identifier-matches.csv

…or over the hectares the service publishes rather than the ones this measures::

    python3 -m src.apps.statistics.wildfires.andalusia_rediam.wildfire_statistics \\
        --surface published --csv published.csv

At least one of ``--csv`` and ``--docx`` is required: an application that computed a
report and then printed nothing would be a strange thing.

The application only reads. Database settings come from the environment (``.env``,
see :mod:`src.settings`); every one of them can be overridden with a command-line
argument.

The sixth of the burnt-area reports, and the twin of
:mod:`Catalonia's <src.apps.statistics.wildfires.catalonia_darpa.wildfire_statistics>`
— same columns, same EGIF-match columns, same refusal to test anything against a
boundary. Its first six columns are the
:mod:`GWIS <src.apps.statistics.wildfires.gwis.wildfire_statistics>`,
:mod:`GFA <src.apps.statistics.wildfires.gfa.wildfire_statistics>`,
:mod:`ICNF <src.apps.statistics.wildfires.portugal_icnf.wildfire_statistics>` and
:mod:`EGIF <src.apps.statistics.wildfires.spain_egif.wildfire_statistics>` reports',
in their order, so the CSVs can still be concatenated on them.

Measured *or* published, which no other report can offer
---------------------------------------------------------

This is the one dataset in GisFIRE that publishes **both** a perimeter and a burnt
area, and ``--surface`` chooses which the report is of.

``measured`` (default)
    The area of the published polygon, computed by ``--area-method``. This is what
    the four perimeter reports do, and it is the default because it is what makes a
    row here comparable with a row of theirs.
``wooded``, ``scrub``, ``grassland``
    ``SUP_ARBOLA``, ``SUP_MATORR`` and ``SUP_PASTIZ`` as published, in hectares.
``published``
    The three added together, which is the nearest thing the service publishes to a
    total. There is no published total column; see
    :mod:`src.providers.andalusia_rediam.wildfire`.

.. warning::

   **The two are different quantities and neither is a correction of the other.**
   Over the 907 fires the published hectares sum to 152,696 and the measured
   perimeter to 165,522 geodesically (165,582 on the service's own grid) — a 7.8%
   difference, which is what one expects of three vegetation classes against an
   outline that also encloses everything that is none of them.

   So a ``measured`` run and a ``published`` run answer different questions, and
   adding them, or quoting one as a correction of the other, is a mistake the column
   headings cannot prevent. The ``.docx`` says which it is of on its front page.

``--area-method`` applies to ``measured`` and to nothing else. Passing it with a
published surface is refused rather than ignored: nothing is measured there, so a
choice of how to measure would be a claim about a number that was read off a form.

How the measured area is measured
----------------------------------

The same two ways as the GFA, ICNF and Catalan reports: ``geodesic`` on the WGS84
ellipsoid (the default) or ``equal-area`` in EPSG:6933. They agree to within 0.003%.

The published EPSG:25830 grid is **not** offered, for the reason the Catalan report
declines EPSG:25831 and the ICNF one EPSG:3763: UTM is a transverse Mercator, which
is conformal and not equal-area. Andalusia spans two UTM zones' worth of longitude
squeezed into one — the grid runs from 1.6°W to 7.5°W against a central meridian at
3°W — so the distortion is not uniform across it and does not cancel over a year
whose fires are not evenly spread. Measure there explicitly if a figure on the
service's own grid is what is wanted:

.. code-block:: sql

   SELECT year, sum(ST_Area(perimeter_etrs89_utm30n) / 10000.0) AS grid_ha
   FROM rediam_wildfire GROUP BY year ORDER BY year DESC;

How many matched the EGIF data
------------------------------

``EGIF matched`` counts the fires of the year that carry a link to an EGIF *parte*,
:attr:`~src.providers.andalusia_rediam.wildfire.RediamWildfire.egif_wildfire_id`, and
``EGIF matched (%)`` is that as a share of the ``Fires`` beside it.

Three things about it, exactly as in the Catalan report:

* **It is a column, not a filter.** An unbound fire is still a fire and still
  contributes its hectares to the row it is in.
* **It follows the scope.** Whatever ``--year``, ``--surface`` and ``--min-area``
  select, the matched count is counted over exactly those fires, so the percentage
  always has the ``Fires`` column as its denominator.
* **A zero is usually a fact about EGIF's coverage.** The binding can only reach the
  campaigns that have been imported, and the exports stop at 2023 — so 2024 and 2025
  match nothing at all, however good the rules are.

The links themselves are written by
:mod:`~src.apps.bindings.wildfires.andalusia_rediam.bind_egif_wildfires`. On the
published data 759 of the 907 perimeters are bound — 83.7% — and 133 of the 148 that
are not are 2024 and 2025.

.. note::

   If nothing has ever been bound the column is zero everywhere, which looks exactly
   like a dataset that failed to match. The log says which it is at ``WARNING``.

Not every binding is the same claim
^^^^^^^^^^^^^^^^^^^^^^^^^^^^^^^^^^^^

``--min-confidence`` counts only the bindings at or above a given
:attr:`~src.providers.andalusia_rediam.wildfire.RediamWildfire.match_confidence`.
Without it every binding counts.

Here that filter changes very little, and the fact that it changes little is worth
seeing: **749 of the 759 links rest on the published identifier**, because ``CODIGO``
*is* the EGIF ``report_number``. ``--min-confidence 0.9`` therefore removes ten fires,
against Catalonia's 177.

.. warning::

   The confidences are **an ordering, not probabilities**. Nothing has been
   calibrated against ground truth, so ``--min-confidence 0.75`` selects a class of
   matching rule and does not mean "matches that are 75% likely to be right".

Which year a fire counts towards
--------------------------------

:attr:`~src.providers.andalusia_rediam.wildfire.RediamWildfire.year` — the year of the
published ``FECHA_INC``, which is also the year the code names, checked at import on
every one of the 962 published features. It is ``NOT NULL`` and indexed, and needs no
timezone applied to it, which
:attr:`~src.data_model.wildfire.Wildfire.start_date_time` would: an Andalusian fire's
instant is local midnight, the dataset publishing no time of day anywhere.

Which fires are counted
-----------------------

Under ``measured``, every imported Andalusian fire with a perimeter, which is every
one of them. Under a published surface, the fires that **report** that surface: a
``NULL`` there is a form that does not say, not a fire that burnt none of it. No fire
in the 2008-2025 archive fails either test, which is exactly why the report counts
what it dropped rather than assuming it dropped nothing.

``--min-area`` narrows it to the fires of at least that many hectares **of the surface
being reported**, so the fires counted are always the fires the figures beside them
were computed from.

.. note::

   A fire is one published ``(code, date)`` and not one feature. 962 published
   features are 907 fires — 55 codes are published twice — and the import dissolves
   them, so nothing here counts or sums a fire twice.

No country test, and no ``--country``
-------------------------------------

Every fire in this dataset is Andalusian, and therefore Spanish, because the service
publishes the fires of Andalusia and nothing else. The ``Country`` column is the
constant :data:`COUNTRY_NAME` on every row and **nothing is tested against a
boundary**, so there is no ``--country`` and no ``--country-source``: both are refused
with a message saying why.

.. warning::

   The column says ``Spain`` and the report is **Andalusia's fires alone**. A total
   here is not a Spanish total: it is one autonomous community of seventeen. The
   region is in the title of the ``.docx`` and in this page, not in a column, so that
   the CSV keeps the shape the other five reports have.

One statement
-------------

The GWIS, GFA and ICNF reports measure one year per statement, because the memory a
point-in-polygon test against a country polygon needs is only released when the
statement ends. This report is one statement, like the EGIF and Catalan ones: it tests
nothing against a boundary and the whole archive is **907 fires**. The ``Total`` row is
arithmetic over the years, by :func:`combine`.
"""

from __future__ import annotations

import argparse
import csv
import logging
import math
import os
import sys

from dataclasses import dataclass
from pathlib import Path

from geoalchemy2 import Geography
from sqlalchemy import ColumnElement
from sqlalchemy import Engine
from sqlalchemy import Select
from sqlalchemy import cast
from sqlalchemy import create_engine
from sqlalchemy import func
from sqlalchemy import literal
from sqlalchemy import or_
from sqlalchemy import select
from sqlalchemy.orm import Session

import src.settings  # noqa: F401  (imported for the side effect of loading .env)

from src.apps.imports import common
from src.data_model.wildfire import Wildfire
from src.providers.andalusia_rediam.wildfire import RediamWildfire

#: Label used in the ``Year`` column for the summary row.
TOTAL_LABEL = "Total"

#: The six columns this report shares with the GWIS, GFA, ICNF, EGIF and Catalan
#: ones, in their order, so the CSVs can still be concatenated on them.
SHARED_COLUMNS = ("Country", "Year", "Fires", "Minimum (ha)", "Maximum (ha)", "Total (ha)")

#: The report's columns, in order, shared by both output formats so that a change to
#: one cannot silently leave the other behind. The last two are this dataset's and
#: Catalonia's own: how many of the year's fires are bound to an EGIF *parte*, and that
#: as a share of the fires counted.
COLUMNS = SHARED_COLUMNS + ("EGIF matched", "EGIF matched (%)")

#: Index of the first column that holds a number, and so is right-aligned in the Word
#: table.
FIRST_NUMERIC_COLUMN = 2

#: The country every fire in this dataset is in, and the whole of the ``Country``
#: column. Spelled as the OCHA boundaries spell it, so a row of this report sorts and
#: groups with the rows of the other five. Nothing is tested against a boundary to
#: arrive at it — see the module docstring.
COUNTRY_NAME = "Spain"

#: The autonomous community the dataset covers. Named in the ``.docx`` and in the log,
#: and deliberately **not** a column: the CSV keeps the other reports' shape, and every
#: row of this one would carry the same value.
REGION_NAME = "Andalusia"

#: Square metres in a hectare.
SQUARE_METRES_PER_HECTARE = 10_000.0

#: NSIDC EASE-Grid 2.0 Global — a cylindrical equal-area projection in metres, defined
#: for the whole world. The CRS behind ``--area-method equal-area``.
EQUAL_AREA_SRID = 6933

#: The two ways of turning a perimeter in degrees into hectares.
AREA_METHOD_GEODESIC = "geodesic"
AREA_METHOD_EQUAL_AREA = "equal-area"
AREA_METHODS = (AREA_METHOD_GEODESIC, AREA_METHOD_EQUAL_AREA)

#: What the report is of, as ``--surface`` accepts it.
#:
#: ``measured`` is the polygon and the other four are the hectares the service
#: publishes. This dataset is the only one in GisFIRE that has both — see the module
#: docstring on why they are different quantities rather than two estimates of one.
SURFACE_MEASURED = "measured"
SURFACE_WOODED = "wooded"
SURFACE_SCRUB = "scrub"
SURFACE_GRASSLAND = "grassland"
SURFACE_PUBLISHED = "published"
SURFACES = (SURFACE_MEASURED, SURFACE_WOODED, SURFACE_SCRUB, SURFACE_GRASSLAND,
            SURFACE_PUBLISHED)

#: How each surface reads in prose, for the Word document's opening paragraph. A report
#: of scrub hectares that said only "areas in hectares" would be indistinguishable from
#: a report of measured ones.
SURFACE_PROSE = {
    SURFACE_MEASURED: "the area of the published perimeter",
    SURFACE_WOODED: "burnt wooded area as published (SUP_ARBOLA)",
    SURFACE_SCRUB: "burnt scrub area as published (SUP_MATORR)",
    SURFACE_GRASSLAND: "burnt grassland area as published (SUP_PASTIZ)",
    SURFACE_PUBLISHED: "all the burnt area the service publishes: wooded plus scrub "
                       "plus grassland",
}

#: The year a fire counts towards: the year of the published date, which is also the
#: year the code names. See the module docstring.
PUBLISHED_YEAR = RediamWildfire.__table__.c.year

#: The confidence at or above which a binding rests on the published identifier rather
#: than on a date and a name — the value ``--min-confidence`` is reached for most often.
#:
#: Not a default: by default every binding counts. It is here so the help text and the
#: documentation cannot drift from
#: :data:`~src.providers.andalusia_rediam.wildfire.MATCH_METHOD_CONFIDENCE`.
IDENTIFIER_CONFIDENCE = 0.9


def burnt_area(surface: str = SURFACE_MEASURED,
               method: str = AREA_METHOD_GEODESIC) -> tuple[ColumnElement, ColumnElement]:
    """The burnt area of one fire in hectares, as ``(hectares, is_reported)``.

    Parameters
    ----------
    surface : str
        One of :data:`SURFACES`.
    method : str
        One of :data:`AREA_METHODS`. Used only when ``surface`` is
        :data:`SURFACE_MEASURED`.

    Returns
    -------
    tuple
        The SQL expression yielding hectares, and the condition a fire has to satisfy
        for that expression to be an answer rather than a silence.

    Raises
    ------
    ValueError
        If ``surface`` is not one of :data:`SURFACES`, or ``method`` not one of
        :data:`AREA_METHODS`.

    Notes
    -----
    The second element is what keeps the ``Fires`` column honest, and it means
    different things for the two kinds of surface. For ``measured`` it is a perimeter
    that exists — true of every fire here, and asserted rather than assumed. For a
    published one it is a figure that was filled in: the three columns are nullable,
    ``sum`` and ``min`` skip nulls and ``count`` does not, so a report that did not
    filter would count fires whose area it had not included.

    A published **zero is counted**, and is a real answer: a fire that burnt no wooded
    land has ``SUP_ARBOLA`` of 0.00. Under ``published`` a fire counts if any of the
    three is reported, the unreported ones contributing nothing.

    The measured area is always computed from the EPSG:4326 perimeter on the parent
    ``wildfire`` row, never from the published EPSG:25830 copy — see the module
    docstring for why the service's own grid is not an option here. The geodesic cast
    carries the geometry type and SRID rather than being a bare ``Geography()``: that
    renders as ``geography(GEOMETRY,-1)``, which PostGIS rejects.
    """
    fire = RediamWildfire.__table__.c

    if surface == SURFACE_MEASURED:
        if method == AREA_METHOD_GEODESIC:
            square_metres = func.ST_Area(
                cast(Wildfire.perimeter, Geography(geometry_type="MULTIPOLYGON", srid=4326))
            )
        elif method == AREA_METHOD_EQUAL_AREA:
            square_metres = func.ST_Area(
                func.ST_Transform(Wildfire.perimeter, EQUAL_AREA_SRID))
        else:
            raise ValueError(
                f"unknown area method {method!r}; expected one of {', '.join(AREA_METHODS)}"
            )
        return square_metres / SQUARE_METRES_PER_HECTARE, Wildfire.perimeter.is_not(None)

    simple = {
        SURFACE_WOODED: fire.area_ha_wooded,
        SURFACE_SCRUB: fire.area_ha_scrub,
        SURFACE_GRASSLAND: fire.area_ha_grassland,
    }
    if surface in simple:
        column = simple[surface]
        return column, column.is_not(None)

    if surface == SURFACE_PUBLISHED:
        components = (fire.area_ha_wooded, fire.area_ha_scrub, fire.area_ha_grassland)
        hectares = (func.coalesce(components[0], 0.0)
                    + func.coalesce(components[1], 0.0)
                    + func.coalesce(components[2], 0.0))
        return hectares, or_(*[component.is_not(None) for component in components])

    raise ValueError(
        f"unknown surface {surface!r}; expected one of {', '.join(SURFACES)}"
    )


def is_matched(min_confidence: float | None = None) -> ColumnElement:
    """Whether a fire counts as bound to an EGIF *parte*.

    Parameters
    ----------
    min_confidence : float, optional
        Count only bindings of at least this confidence. ``None``, the default, counts
        every binding whatever rule produced it.

    Returns
    -------
    ColumnElement
        A boolean expression, ``True`` for a fire the ``EGIF matched`` column counts.

    Notes
    -----
    The link and not the method is what is tested, because the two cannot disagree: a
    check constraint on the model makes ``egif_wildfire_id`` and ``match_method`` null
    together, so a row with one and not the other cannot exist to be counted wrongly.

    With a threshold the confidence is tested as well, and both conditions are stated
    rather than the second alone. ``match_confidence >= 0.5`` would in fact select the
    same rows today, every method having a confidence — but it would be relying on
    that, where this relies on the constraint.
    """
    rediam = RediamWildfire.__table__
    bound = rediam.c.egif_wildfire_id.is_not(None)
    if min_confidence is None:
        return bound
    return bound & (rediam.c.match_confidence >= min_confidence)


def statistics_query(surface: str = SURFACE_MEASURED,
                     method: str = AREA_METHOD_GEODESIC,
                     year: int | None = None,
                     min_area: float | None = None,
                     min_confidence: float | None = None) -> Select:
    """Build the statistics query: one row per year, newest first.

    Parameters
    ----------
    surface : str
        One of :data:`SURFACES`.
    method : str
        One of :data:`AREA_METHODS`, used only under :data:`SURFACE_MEASURED`.
    year : int, optional
        Restrict to one published year. ``None``, the default, reports every year.
    min_area : float, optional
        Count only fires of at least this many hectares of the chosen surface.
        ``None``, the default, counts every fire.
    min_confidence : float, optional
        Count as matched only the bindings of at least this confidence. ``None``, the
        default, counts every binding.

    Returns
    -------
    Select
        A query yielding ``country, year, minimum, maximum, total, fires, matched``.
        The summary row is :func:`summarise`'s work.

    Notes
    -----
    Built against the mapped classes rather than written as SQL text, so a column
    renamed on a model breaks this at import time rather than in front of a user.

    ``rediam_wildfire`` is joined by table, to keep SQLAlchemy from adding a
    polymorphic join of its own, and it has to be joined in any case: the year, the
    three published areas and the EGIF link all live on it. The parent ``wildfire`` row
    is only needed for the measured surface, and is joined either way so that the
    statement is one shape.

    The inner query computes each area exactly once. Folded into the outer aggregate
    instead, the expression would be evaluated three times per row — for the minimum,
    the maximum and the sum.

    The country is grouped on although it is a literal: by the time the outer aggregate
    sees it, it is a column of the subquery rather than a constant, and PostgreSQL
    requires every one of those in the ``GROUP BY``.

    ``min_area`` filters the subquery's column rather than repeating the expression,
    and is applied before the aggregates rather than as a ``HAVING``: the threshold
    selects the fires the figures are computed from, it does not discard years whose
    total came out small. The matched count is aggregated over the same filtered rows,
    which is what keeps ``EGIF matched`` a share of the ``Fires`` beside it.
    """
    rediam = RediamWildfire.__table__
    hectares, is_reported = burnt_area(surface, method)

    fires = (
        select(
            literal(COUNTRY_NAME).label("country"),
            PUBLISHED_YEAR.label("year"),
            hectares.label("hectares"),
            is_matched(min_confidence).label("matched"),
        )
        .select_from(Wildfire)
        .join(rediam, rediam.c.id == Wildfire.id)
        .where(is_reported)
    )
    if year is not None:
        fires = fires.where(PUBLISHED_YEAR == year)

    fire = fires.subquery("fire")
    statistics = (
        select(
            fire.c.country,
            fire.c.year,
            func.min(fire.c.hectares).label("minimum"),
            func.max(fire.c.hectares).label("maximum"),
            func.sum(fire.c.hectares).label("total"),
            func.count().label("fires"),
            func.count().filter(fire.c.matched).label("matched"),
        )
        .group_by(fire.c.country, fire.c.year)
        .order_by(fire.c.year.desc())
    )
    if min_area is not None:
        statistics = statistics.where(fire.c.hectares >= min_area)
    return statistics


def unreported_query(surface: str, year: int | None = None) -> Select:
    """How many fires in scope do **not** report the chosen surface.

    Returns
    -------
    Select
        A query yielding one integer.

    Notes
    -----
    Zero on the whole published archive, for every surface: every fire has a perimeter
    and all three areas. That is why it is asked rather than assumed — a later
    publication that leaves a column out would otherwise quietly shrink the ``Fires``
    column with nothing to say it had.
    """
    rediam = RediamWildfire.__table__
    _, is_reported = burnt_area(surface)

    fires = (
        select(literal(1))
        .select_from(Wildfire)
        .join(rediam, rediam.c.id == Wildfire.id)
        .where(~is_reported)
    )
    if year is not None:
        fires = fires.where(PUBLISHED_YEAR == year)
    return select(func.count()).select_from(fires.subquery("unreported"))


def surface_label(surface: str) -> str:
    """How a surface reads in a log line: ``burnt``, ``published``, ``published scrub``.

    A function rather than a dictionary entry because two of the five names would
    otherwise say themselves twice — ``published published hectares`` — and a log line
    that reads like that is a log line nobody trusts the rest of.
    """
    if surface == SURFACE_MEASURED:
        return "burnt"
    if surface == SURFACE_PUBLISHED:
        return "published"
    return f"published {surface}"


def share(part: int, whole: int) -> float | None:
    """``part`` as a percentage of ``whole``, or ``None`` where there is no whole.

    ``None`` and not zero: a percentage of nothing is not zero percent, it is no
    answer, and the writers turn it into an empty cell.
    """
    if not whole:
        return None
    return 100.0 * part / whole


def share_label(part: int, whole: int) -> str:
    """A percentage as it is written out, empty where there is none."""
    value = share(part, whole)
    return "" if value is None else f"{value:.2f}"


@dataclass(frozen=True)
class Row:
    """One line of the report.

    Attributes
    ----------
    country : str
        Always :data:`COUNTRY_NAME`. Nothing is tested against a boundary to arrive at
        it — the service publishes Andalusia's fires and nothing else.
    year : int or None
        The published year, or ``None`` for the summary row.
    minimum, maximum, total : float
        Smallest single fire, largest single fire and sum of every fire, in hectares of
        the surface asked for.
    fires : int
        How many fires the three area figures were computed from — a fire being one
        published ``(code, date)``, not one feature.
    matched : int
        How many of those ``fires`` are bound to an EGIF *parte*, at or above
        ``--min-confidence`` where one was given. Always at most ``fires``: it is
        counted over the same rows.
    """

    country: str
    year: int | None
    minimum: float
    maximum: float
    total: float
    fires: int
    matched: int

    @property
    def is_total(self) -> bool:
        """Whether this is the summary row rather than one of the years."""
        return self.year is None

    @property
    def year_label(self) -> str:
        return TOTAL_LABEL if self.is_total else str(self.year)

    @property
    def matched_share(self) -> float | None:
        """``matched`` as a percentage of ``fires``: how much of the year joins."""
        return share(self.matched, self.fires)

    @property
    def values(self) -> tuple[str, ...]:
        """The row as the CSV writes it, in :data:`COLUMNS` order."""
        return (self.country, self.year_label, str(self.fires),
                f"{self.minimum:.2f}", f"{self.maximum:.2f}", f"{self.total:.2f}",
                str(self.matched), share_label(self.matched, self.fires))

    @property
    def readable_values(self) -> tuple[str, ...]:
        """The row as the Word document writes it: the numbers with separators."""
        return (self.country, self.year_label, f"{self.fires:,}",
                f"{self.minimum:,.2f}", f"{self.maximum:,.2f}", f"{self.total:,.2f}",
                f"{self.matched:,}", share_label(self.matched, self.fires))


def combine(rows: list[Row], country: str = COUNTRY_NAME, year: int | None = None) -> Row:
    """One row summarising several: every figure taken over all of them.

    Notes
    -----
    All five decompose over a partition of the fires — a minimum of minima is a
    minimum, a sum of sums is a sum, a count of counts is a count — so the ``Total`` row
    is the number a second aggregate over the same rows would have returned, and no fire
    is counted twice or left out.

    The percentage is deliberately **not** averaged: :class:`Row` recomputes it from the
    summed counts, which is the ratio of the totals rather than the mean of the ratios.
    A year of seven fires and a year of ninety-seven must not weigh the same in the
    answer.

    ``fsum`` rather than ``sum``: an exact accumulation over a handful of partial totals
    costs nothing and cannot drift from what one pass would have returned.
    """
    return Row(
        country=country,
        year=year,
        minimum=min(row.minimum for row in rows),
        maximum=max(row.maximum for row in rows),
        total=math.fsum(row.total for row in rows),
        fires=sum(row.fires for row in rows),
        matched=sum(row.matched for row in rows),
    )


def summarise(measured: list[Row]) -> list[Row]:
    """Build the report from the years measured: the years newest first, then the total.

    Returns
    -------
    list of Row
        The years newest first with the summary row last. Empty if nothing was measured
        — a report of no fires has no total either.

    Notes
    -----
    No grouping by country and no lookup of one, unlike the four reports this is read
    beside: there is exactly one country here and it is a constant, so ordering the
    countries would be sorting a list of one against the database's collation.
    """
    if not measured:
        return []
    rows = sorted(measured, key=lambda row: row.year, reverse=True)
    return rows + [combine(rows)]


def hectares(text: str) -> float:
    """Argparse type for ``--min-area``: a finite, non-negative number of hectares.

    Raises
    ------
    argparse.ArgumentTypeError
        If the text is not a number, or is negative, or is a non-finite float.

    Notes
    -----
    A bare ``type=float`` would accept ``-5``, ``nan`` and ``inf``. The first two are
    almost certainly a typo and would silently produce the unfiltered report — ``nan``
    compares false against every area, so it would produce an empty one — and none of
    the three is a size a fire can have.
    """
    try:
        area = float(text)
    except ValueError:
        raise argparse.ArgumentTypeError(f"{text!r} is not a number of hectares")
    if not math.isfinite(area):
        raise argparse.ArgumentTypeError(f"{text!r} is not a finite number of hectares")
    if area < 0:
        raise argparse.ArgumentTypeError(
            f"a burnt area cannot be negative, and {area:g} ha is: pass 0 or more, or "
            f"leave --min-area out to count every fire"
        )
    return area


def confidence(text: str) -> float:
    """Argparse type for ``--min-confidence``: a number from 0 to 1 inclusive.

    Raises
    ------
    argparse.ArgumentTypeError
        If the text is not a number, or is not finite, or is outside ``[0, 1]``.

    Notes
    -----
    The range is closed at both ends on purpose: ``0`` counts every binding, which is
    the default said explicitly, and ``1`` counts only the exact-identifier ones.

    A value above 1 is refused rather than quietly counting nothing. Someone passing
    ``90`` means ninety percent, and a report that answered them with a column of zeros
    would be worse than one that stopped.
    """
    try:
        value = float(text)
    except ValueError:
        raise argparse.ArgumentTypeError(f"{text!r} is not a number")
    if not math.isfinite(value):
        raise argparse.ArgumentTypeError(f"{text!r} is not a finite confidence")
    if not 0.0 <= value <= 1.0:
        raise argparse.ArgumentTypeError(
            f"a match confidence is between 0 and 1, and {value:g} is not: "
            f"{IDENTIFIER_CONFIDENCE:g} is the boundary between the identifier matches "
            f"and the name matches"
        )
    return value


def parse_arguments(argv: list[str] | None = None) -> argparse.Namespace:
    """Parse the command line.

    ``--area-method`` defaults to ``None`` rather than to ``geodesic`` so that
    :func:`report` can tell "not given" from "given the default", which is what lets it
    be refused against a published surface without also refusing every plain run.
    """
    parser = argparse.ArgumentParser(
        description="Burnt-area statistics for the Andalusian REDIAM burnt area "
                    "cartography, with the number of fires bound to the Spanish EGIF "
                    "statistics.",
        epilog="Areas are in hectares measured from the published perimeter by default; "
               "--surface reports the hectares the service publishes instead, which are "
               "a different quantity and not a correction. Every fire is Andalusian, so "
               "there is no --country and nothing is tested against a boundary. Database "
               "settings not given here are read from the environment (.env).",
    )
    parser.add_argument("-y", "--year", type=int,
                        help="restrict to one year, e.g. 2022; this is the year of the "
                             "published FECHA_INC, which this dataset always agrees with "
                             "the year inside the code")
    parser.add_argument("--surface", default=SURFACE_MEASURED, choices=SURFACES,
                        help="what to report: 'measured' (default) is the area of the "
                             "published perimeter; 'wooded', 'scrub' and 'grassland' are "
                             "the hectares the service publishes (SUP_ARBOLA, "
                             "SUP_MATORR, SUP_PASTIZ); 'published' adds those three, "
                             "which is the nearest thing to a published total. Measured "
                             "and published are different quantities — over the archive "
                             "they differ by 7.8%%")
    parser.add_argument("--min-area", type=hectares, default=None, metavar="HECTARES",
                        help="count only fires that burnt at least this many hectares of "
                             "the chosen surface; by default every fire counts, "
                             "including a published zero")
    parser.add_argument("--min-confidence", type=confidence, default=None,
                        metavar="CONFIDENCE",
                        help=f"count a fire as matched only when its binding to an EGIF "
                             f"parte has at least this confidence, between 0 and 1; by "
                             f"default every binding counts. "
                             f"{IDENTIFIER_CONFIDENCE:g} is the boundary between the "
                             f"matches resting on the published identifier and those "
                             f"resting on a date and a municipality name — which here "
                             f"is 749 of the 759 bindings")
    parser.add_argument("--area-method", default=None, choices=AREA_METHODS,
                        help="how to turn the EPSG:4326 perimeter into hectares: "
                             "'geodesic' measures on the WGS84 ellipsoid (default); "
                             "'equal-area' projects to EPSG:6933 and measures there. "
                             "They agree to within 0.003%%. Applies to --surface "
                             "measured and to nothing else; the published EPSG:25830 "
                             "grid is not offered, being conformal rather than "
                             "equal-area")

    # Accepted only so that they can be refused clearly. Anyone reaching for either has
    # copied a command line from one of the other five reports, which is a reasonable
    # thing to have done, and argparse's own message would not say why this report is
    # different.
    parser.add_argument("--country", help=argparse.SUPPRESS)
    parser.add_argument("--country-source", help=argparse.SUPPRESS)

    output = parser.add_argument_group("output", "at least one is required")
    output.add_argument("--csv", type=Path, help="write the report to this .csv")
    output.add_argument("--docx", type=Path, help="write the report to this .docx (MS Word)")

    common.add_database_arguments(parser)
    parser.add_argument("--log-level", default=os.getenv("GISFIRE_LOG_LEVEL", "INFO"),
                        choices=["CRITICAL", "ERROR", "WARNING", "INFO", "DEBUG"],
                        help="verbosity (env: GISFIRE_LOG_LEVEL, default INFO)")

    arguments = parser.parse_args(argv)
    if arguments.country is not None:
        parser.error(
            f"there is no --country here: the service publishes the fires of "
            f"{REGION_NAME} and nothing else, so there is nothing to select between. "
            f"Every fire is counted and the Country column is {COUNTRY_NAME} on every "
            f"row."
        )
    if arguments.country_source is not None:
        parser.error(
            f"there is no --country-source here: these perimeters are {REGION_NAME}'s "
            f"own cartography of its own territory, so nothing is tested against a "
            f"boundary and there is nothing for a containment test to find. The Country "
            f"column is the constant {COUNTRY_NAME}."
        )
    if arguments.area_method is not None and arguments.surface != SURFACE_MEASURED:
        parser.error(
            f"--area-method applies to --surface {SURFACE_MEASURED} and to nothing else: "
            f"the {arguments.surface} hectares are read off the service's own figures, "
            f"not measured, so no CRS is involved in producing them."
        )
    if arguments.csv is None and arguments.docx is None:
        parser.error("nothing to write: pass --csv, --docx, or both")
    return arguments


def compute(session: Session, year: int | None, logger: logging.Logger,
            surface: str = SURFACE_MEASURED,
            method: str = AREA_METHOD_GEODESIC,
            min_area: float | None = None,
            min_confidence: float | None = None) -> list[Row]:
    """Run the statement and return the report's rows in order.

    Notes
    -----
    One statement, under one spinner — see the module docstring for why this report
    does not need the year-at-a-time machinery three of the other five are built on.
    The ``Total`` row is arithmetic over its result, not a second query.

    A second, cheap statement follows it wherever a fire could have been left out:
    :func:`unreported_query` counts the fires that do not report the chosen surface, so
    that a ``Fires`` column smaller than the archive says so rather than being noticed
    later. It returns 0 on every surface of the published data.

    The matched fires are logged as well as reported, and a report in which nothing at
    all is bound says so at ``WARNING``: a column of zeros is what an unrun binding
    application and a dataset that matched nothing look like alike, and they are not the
    same thing.
    """
    with common.Spinner(f"Measuring the {surface_label(surface)} area of the "
                        f"{REGION_NAME} fires and their EGIF matches", logger):
        measured = [
            Row(country=record.country, year=record.year,
                minimum=float(record.minimum),
                maximum=float(record.maximum),
                total=float(record.total),
                fires=record.fires,
                matched=record.matched)
            for record in session.execute(
                statistics_query(surface, method, year, min_area, min_confidence))
        ]
        unreported = session.scalar(unreported_query(surface, year))

    rows = summarise(measured)
    logger.info("Computed %d rows over %d year(s) (%s, %s)",
                len(rows), len(measured),
                f"{method} areas" if surface == SURFACE_MEASURED
                else f"{surface_label(surface)} hectares",
                "every fire" if min_area is None else f"fires of {min_area:g} ha or more")
    if unreported:
        logger.warning(
            "%d fire(s) in scope do not report a %s area and are in none of the figures "
            "above; the Fires column counts only the fires the areas were computed from",
            unreported, surface)
    if rows:
        total = rows[-1]
        logger.info("%d of %d fire(s) are bound to an EGIF parte (%s%%)%s",
                    total.matched, total.fires, share_label(total.matched, total.fires),
                    "" if min_confidence is None
                    else f", counting only bindings of confidence {min_confidence:g} or more")
        if not total.matched:
            logger.warning(
                "No fire in scope is bound to an EGIF parte, so every EGIF matched "
                "column is zero. Nothing here can tell an unrun binding from a real "
                "absence of matches — run "
                "src.apps.bindings.wildfires.andalusia_rediam.bind_egif_wildfires, and "
                "note that it can only reach the EGIF campaigns that are imported")
    return rows


def write_csv(rows: list[Row], path: Path, logger: logging.Logger) -> None:
    """Write the report as CSV.

    The numbers go out unformatted apart from being rounded to two decimals — no
    thousands separators — because a CSV is read by another program far more often than
    by a person, and a separator would make every figure a string.
    """
    path.parent.mkdir(parents=True, exist_ok=True)
    with path.open("w", newline="", encoding="utf-8") as handle:
        writer = csv.writer(handle)
        writer.writerow(COLUMNS)
        for row in rows:
            writer.writerow(row.values)
    logger.info("Wrote %s", path)


def write_docx(rows: list[Row], path: Path, year: int | None,
               logger: logging.Logger,
               surface: str = SURFACE_MEASURED,
               method: str = AREA_METHOD_GEODESIC,
               min_area: float | None = None,
               min_confidence: float | None = None) -> None:
    """Write the report as a Word document.

    One table, with the summary row in bold. Numbers get thousands separators here —
    the opposite of the CSV, and for the opposite reason: this one is for reading.

    The opening paragraphs name the region and the surface, because a table headed
    ``Spain`` that is in fact one autonomous community, and a table of published
    hectares that looks exactly like a table of measured ones, are both things a reader
    should not have to remember.
    """
    # Imported here rather than at module scope so that --csv keeps working if
    # python-docx is not installed, which matters because it is the only dependency this
    # application adds.
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt

    document = Document()
    document.add_heading(f"REDIAM wildfire burnt area ({REGION_NAME})", level=1)

    if surface == SURFACE_MEASURED:
        how = ("computed geodesically on the WGS84 ellipsoid"
               if method == AREA_METHOD_GEODESIC
               else f"computed in the equal-area projection EPSG:{EQUAL_AREA_SRID}")
    else:
        how = "as published by the service, not measured from the perimeter"
    scope = [f"year: {year}" if year is not None else "all years"]
    if min_area is not None:
        scope.append(f"only fires of {min_area:g} ha or more")
    if min_confidence is not None:
        scope.append(f"only EGIF bindings of confidence {min_confidence:g} or more")

    document.add_paragraph(
        f"Areas in hectares — {SURFACE_PROSE[surface]} — {how}. Years are the published "
        f"FECHA_INC. A fire is one published (code, date), which for 55 of them is two "
        f"published features. Scope: {'; '.join(scope)}."
    )
    document.add_paragraph(
        f"The Country column is {COUNTRY_NAME} on every row and nothing is tested "
        f"against a boundary: the service publishes the fires of {REGION_NAME} and "
        f"nothing else. These totals are therefore one autonomous community's and are "
        f"not a Spanish total."
    )
    # Not a footnote: this dataset is the only one with both kinds of number, and a
    # reader who takes one for the other will conclude the archive is inconsistent.
    document.add_paragraph(
        "This dataset publishes a burnt area as well as a perimeter, and the two are "
        "different quantities rather than two estimates of one: over the whole archive "
        "the published hectares sum to 152,696 and the measured perimeter to 165,522. "
        "Neither is a correction of the other, and they must not be added."
    )
    document.add_paragraph(
        "EGIF matched counts the fires linked to the Spanish parte for the same fire. It "
        "is a column and not a filter — an unbound fire still contributes its hectares — "
        "and a year the EGIF exports do not cover matches nothing however good the rules "
        "are."
    )

    table = document.add_table(rows=1, cols=len(COLUMNS))
    table.style = "Table Grid"
    for cell, heading in zip(table.rows[0].cells, COLUMNS):
        cell.text = heading
        cell.paragraphs[0].runs[0].bold = True

    for row in rows:
        cells = table.add_row().cells
        for index, (cell, value) in enumerate(zip(cells, row.readable_values)):
            cell.text = value
            paragraph = cell.paragraphs[0]
            if index >= FIRST_NUMERIC_COLUMN:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            for run in paragraph.runs:
                run.bold = row.is_total
                run.font.size = Pt(9)

    path.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(path))
    logger.info("Wrote %s", path)


def report(args: argparse.Namespace, engine: Engine, logger: logging.Logger) -> list[Row]:
    """Compute the statistics and write whichever outputs were asked for."""
    method = args.area_method or AREA_METHOD_GEODESIC
    with Session(engine) as session:
        rows = compute(session, args.year, logger, args.surface, method, args.min_area,
                       args.min_confidence)

    if not rows:
        # An empty report is almost always a year with no data, and writing an empty
        # file would hide that. A threshold is named when there is one, because then it
        # is at least as likely to be the reason as the year is.
        threshold = "" if args.min_area is None else \
            f" No fire reached the --min-area of {args.min_area:g} ha."
        raise RuntimeError(
            f"No wildfires matched. Check --year, and that the {REGION_NAME} fires are "
            f"imported and report a {args.surface} area — the published layers run from "
            f"2008." + threshold
        )

    if args.csv is not None:
        write_csv(rows, args.csv, logger)
    if args.docx is not None:
        write_docx(rows, args.docx, args.year, logger, args.surface, method,
                   args.min_area, args.min_confidence)
    return rows


def main(argv: list[str] | None = None) -> int:
    args = parse_arguments(argv)
    logging.basicConfig(level=args.log_level, format="%(asctime)s %(levelname)s %(message)s")
    logger = logging.getLogger("rediam-statistics")

    try:
        settings = common.resolve_database_settings(args)
    except RuntimeError as error:
        logger.error("%s", error)
        return 1

    engine = create_engine(common.database_url(settings))
    try:
        report(args, engine, logger)
    except Exception as error:  # noqa: BLE001  (the CLI boundary: report, do not traceback)
        logger.error("%s", error)
        return 1
    finally:
        engine.dispose()
    return 0


if __name__ == "__main__":  # pragma nocover
    sys.exit(main())
