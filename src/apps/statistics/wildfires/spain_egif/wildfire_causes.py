#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Wildfire counts by cause for the Spanish EGIF fire statistics.

Reports, per campaign, how many fires there were, how many carry a cause at all,
how many of those were started by **lightning**, and then the same three counts
again over the fires whose published ignition point really falls inside Spain::

    Country  Year   Fires  Classified  Lightning  Lightning %  Fires inside  ...
    Spain    2023    6294        6294        271         4.31          6290  ...
    Spain    2022    7362        7362        198         2.69          7358  ...
    Spain    Total  13656       13656        469         3.43         13648  ...

Run it over everything, or narrow it to one campaign, or count a different
family of causes::

    python3 -m src.apps.statistics.wildfires.spain_egif.wildfire_causes --csv causes.csv
    python3 -m src.apps.statistics.wildfires.spain_egif.wildfire_causes \\
        --year 2023 --csv 2023.csv --docx 2023.docx
    python3 -m src.apps.statistics.wildfires.spain_egif.wildfire_causes \\
        --cause-family intentional --csv arson.csv

At least one of ``--csv`` and ``--docx`` is required.

The companion of
:mod:`~src.apps.statistics.wildfires.spain_egif.wildfire_statistics`, over the
same archive and the same campaigns. What that one measures in hectares, this one
counts by cause.

EGIF does publish lightning
---------------------------

Unlike the ICNF — whose
:mod:`counterpart <src.apps.statistics.wildfires.portugal_icnf.wildfire_causes>`
has to fall back on ``Natural`` and say at length that it is a proxy — **EGIF names
lightning outright**. ``idcausa`` is a three-digit code whose first digit is the
family, and the ``1`` family is :data:`~src.providers.spain_egif.CAUSE_LIGHTNING`,
*Rayo*. So the default column of this report is ``Lightning`` and it means
lightning; there is no proxy to apologise for.

Matching is on the **family digit** and not on the exact code. ``100`` is today the
whole of the family, but the catalogue is versioned — the *Instrucciones para
cumplimentar el parte de incendio forestal* are at v3.6, *9ª actualización*, and
every revision so far has added subcodes — so a future ``101`` *Rayo seco* is
counted the day it appears rather than the day someone notices this file.

``--cause-family`` counts one of the other five instead; see
:data:`CAUSE_FAMILIES`. The digits of the four families the provider module names
are taken from its constants rather than written out again here.

Which fires are counted, and why it is not the companion's rule
---------------------------------------------------------------

**Every EGIF fire of the campaign.** There is no ``--surface`` and no
``--min-area``: this report counts fires, and a fire whose report form leaves the
burnt area blank is still a fire.

That is a deliberate divergence from the ICNF pair, whose two reports are built to
agree row for row. They can here, because the companion report **excludes any fire
that does not report the surface asked for** — it has to, or its ``Fires`` column
would stop counting the fires its three area figures were computed from. So:

.. warning::

   Do not expect this report's ``Fires`` column to match
   :mod:`~src.apps.statistics.wildfires.spain_egif.wildfire_statistics`'s. That one
   counts the fires that reported a burnt area; this one counts the fires. The
   difference is exactly the fires with a blank surface on the form, and it is a
   real property of the archive rather than a disagreement.

The points outside the border
-----------------------------

EGIF has no perimeter, but it does have a **point** — the published ignition
coordinate on :class:`~src.providers.spain_egif.ignition.EgifIgnition` — and that
point can be somewhere a Spanish fire is not. The importer's only geometric guard
is a plausibility box on the published easting and northing
(:data:`~src.providers.spain_egif.PLAUSIBLE_UTM_EASTING`), which is a rectangle
containing a great deal of Atlantic and Mediterranean; and where the published
*huso* is not a zone Spain lies in, the zone is replaced with the modal one for the
province, which can walk a coastal point out to sea or over a border.

So the report gives the counts **twice**. The first three columns are every fire
EGIF filed. The ``... inside`` columns are the fires whose point the database finds
inside the real Spanish polygon, tested against the OCHA country outlines at report
time. The gap between the two blocks is the answer to "how much of this can be
placed on the ground", and ``Fires inside (%)`` is that gap as a number.

Unlike the companion report this is **not an option**, because nothing is dropped
by it: a fire that is not inside still counts in the first block. The companion has
a ``--country-source`` precisely because there the test *removes* fires, and
switching it on silently halves the archive; here both answers are in the same row
and neither hides the other.

A fire is *outside* for two quite different reasons. The report counts the one that
is free to count and does not go looking for the other:

*no point published*
    An ordinary property of the archive, and the big number: **294,052 of the
    586,157 fires publish no coordinate at all**, every fire before 1998 among them.
    It says nothing whatsoever about the fires that do have one, and it is a null
    test on a column already in hand, so the log always reports it.
*a point that is not in Spain*
    A coordinate in the sea, or over the French or Portuguese border — the data
    fault the plausibility box lets through. It is the difference between the two
    numbers above and needs no extra work to see.

**Which** other country such a point falls in is deliberately not asked. It is a
different question — "where is this coordinate" rather than "is this fire in
Spain" — and answering it means testing every point against every country in the
world instead of against one polygon. That is what this report used to do, and it
is why it took hours; see *One polygon* below.

The fire is Spanish either way: it is a Spanish *parte*, so the ``Country`` column
is the constant :data:`~...wildfire_statistics.COUNTRY_NAME` on every row and a fire
is never moved into France's total.

.. note::

   If the OCHA country boundaries are not imported there is no Spanish polygon to
   be inside, every ``... inside`` column is zero, and the report says so at
   ``WARNING`` rather than letting a table of zeros pass for an answer.

Which year a fire counts towards
--------------------------------

:attr:`~src.providers.spain_egif.wildfire.EgifWildfire.campaign`, the filed
``Campania``, exactly as in the companion report and for the same reasons: it is
what a published yearly total is a total of, it is ``NOT NULL`` and indexed, and it
needs no timezone applied to it.

One polygon, one campaign at a time
-----------------------------------

The report is eight numbers per campaign — the three counts, the three counts
inside, and the two percentages — and it is computed exactly that way: one
statement per campaign, then the ``Total`` row by addition.

Six aggregates over the same rows come out of that one statement together, because
splitting them would read the same fires six times. A campaign at a time, rather
than all at once, because it bounds what any statement holds and because a run over
586,157 fires has to be able to say which year it is on.

.. warning::

   **The geometry is one ``ST_Contains`` per fire against one prepared polygon**,
   and it has to stay that way.

   This report once asked, for every fire, *which* of the world's countries
   contained its point — a lateral over every level-0 boundary, ordered so that
   Spain would win a tie, taking the first. That is 318 polygon tests per fire
   instead of one, and the ``ORDER BY`` meant the ``LIMIT 1`` could not stop early.
   Measured on the real archive it cost **0.6 seconds per located fire**: 200 fires
   in 120 seconds, which over the 292,105 fires that publish a point is about
   **49 hours**. The whole report now takes **17 seconds**.

   The fix was not a faster query but a smaller question. ``Is this point in
   Spain?`` needs Spain's polygon and nothing else, and
   :data:`SPAIN_SQL` fetches exactly that one row. A test asserts that the statement
   still selects the country by name, because a change that dropped the filter would
   give the same answers and quietly restore the hours.

Shared with the companion report
--------------------------------

The country name, the country level, the campaign column, the ``Total`` label and
the country ordering are **imported from**
:mod:`~src.apps.statistics.wildfires.spain_egif.wildfire_statistics` rather than
copied. Two reports over one dataset that disagreed about which campaign a fire is
filed under would be worse than one report, and a copy is a thing that drifts.
"""

from __future__ import annotations

import argparse
import csv
import logging
import os
import sys

from dataclasses import dataclass
from pathlib import Path

from sqlalchemy import Engine
from sqlalchemy import create_engine
from sqlalchemy import text
from sqlalchemy.orm import Session

import src.settings  # noqa: F401  (imported for the side effect of loading .env)

from src.apps.imports import common
from src.apps.statistics.wildfires.spain_egif.wildfire_statistics import COUNTRY_LEVEL
from src.apps.statistics.wildfires.spain_egif.wildfire_statistics import COUNTRY_NAME
from src.apps.statistics.wildfires.spain_egif.wildfire_statistics import TOTAL_LABEL
from src.providers.spain_egif import CAUSE_INTENTIONAL
from src.providers.spain_egif import CAUSE_LIGHTNING
from src.providers.spain_egif import CAUSE_REKINDLE
from src.providers.spain_egif import CAUSE_UNKNOWN



@dataclass(frozen=True)
class Family:
    """One family of EGIF causes: the leading digit of ``idcausa``, and its names.

    Attributes
    ----------
    digit : str
        The first character of every ``idcausa`` in the family, which is what the
        report matches on. See the module docstring for why the family and not the
        exact code.
    label : str
        English name, used for the report's column headings.
    spanish : str
        The family as the EGIF classification names it, for the prose that has to
        be checkable against the source.
    """

    digit: str
    label: str
    spanish: str


#: The six families of ``idcausa``, keyed by what ``--cause-family`` accepts.
#:
#: The digits of the four families :mod:`src.providers.spain_egif` names are taken
#: from its constants rather than written out again, so a renumbering there cannot
#: leave this report counting the old digit. ``2`` and ``3`` have no constant —
#: they are families of subcodes with no bare parent — and are spelled out.
#:
#: The ``2``/``3`` boundary is close to "with or without an implicit use of fire"
#: but is not clean at the edges: ``292`` *Fuegos artificiales* is in the first and
#: ``300`` *Quema de cables para extraer cobre* in the second.
CAUSE_FAMILIES = {
    "lightning": Family(CAUSE_LIGHTNING[0], "Lightning", "Rayo"),
    "negligence": Family("2", "Negligence", "Negligencias (uso del fuego)"),
    "accident": Family("3", "Accident", "Causas accidentales (sin uso del fuego)"),
    "intentional": Family(CAUSE_INTENTIONAL[0], "Intentional", "Intencionado"),
    "unknown": Family(CAUSE_UNKNOWN[0], "Unknown", "Desconocida"),
    "rekindle": Family(CAUSE_REKINDLE[0], "Rekindle", "Reproducido"),
}

#: The family counted unless ``--cause-family`` says otherwise. EGIF names it
#: outright, so this column is lightning and not a proxy for it — see the module
#: docstring, and contrast the ICNF report, which has no such category to count.
DEFAULT_FAMILY = "lightning"

#: Index of the first column that holds a number, and so is right-aligned in the
#: Word table.
FIRST_NUMERIC_COLUMN = 2


def columns(family: str = DEFAULT_FAMILY) -> tuple[str, ...]:
    """The report's columns, in order, for a given cause family.

    Notes
    -----
    A function rather than a constant, for the ICNF companion's reason: three of the
    headings name the family being counted, and a file of intentional fires counted
    under a heading saying ``Lightning`` would be a trap.

    Both output formats read them from here, so a change to one cannot silently
    leave the other behind. The first three are the companion report's first three,
    unchanged.
    """
    label = cause_family(family).label
    return (
        "Country", "Year", "Fires", "Classified", label,
        f"{label} (% of classified)",
        "Fires inside", "Fires inside (%)", "Classified inside",
        f"{label} inside", f"{label} inside (% of classified inside)",
    )


def cause_family(family: str) -> Family:
    """Look a family up by the name ``--cause-family`` accepts.

    Raises
    ------
    ValueError
        If ``family`` is not one of :data:`CAUSE_FAMILIES`.
    """
    try:
        return CAUSE_FAMILIES[family]
    except KeyError:
        raise ValueError(
            f"unknown cause family {family!r}; expected one of "
            f"{', '.join(CAUSE_FAMILIES)}"
        ) from None


#: Spain's outline, as one row.
#:
#: The whole of the geometry this report needs. The question each fire asks is "is
#: this point inside Spain", which is **one** polygon test — not "which of the
#: world's countries contains this point", which is 250 of them.
#:
#: Fetched once per statement and cross-joined, so PostGIS prepares the polygon
#: once and reuses it for every point in the campaign rather than rebuilding it per
#: row.
SPAIN_SQL = """
    SELECT boundary.geometry AS geometry
    FROM admin_boundary AS boundary
    JOIN ocha_admin_boundary AS ocha ON ocha.id = boundary.id
    WHERE boundary.level = :country_level
      AND boundary.name = :country_name
    LIMIT 1
"""

#: The eight numbers, for one campaign, in one statement.
#:
#: Steps 1-3 are the three counts over every fire the campaign filed; steps 5-7 are
#: the same three restricted to the fires whose published point is inside Spain. The
#: percentages (steps 4 and 8) are arithmetic and are :class:`Row`'s work — SQL that
#: returned them would only have to be divided again for the ``Total`` row.
#:
#: One pass over one campaign, and one polygon test per fire that has a point:
#:
#: * ``egif_wildfire`` is filtered by ``campaign`` first, which is indexed;
#: * ``egif_fire_cause`` is joined for the code alone, and outer, because an
#:   unclassified fire is still a fire;
#: * ``ignition`` is joined outer for the same reason — half the archive publishes
#:   no coordinate, and those fires belong in steps 1-3;
#: * ``spain`` is one row, cross-joined, so ``ST_Contains`` is called once per fire
#:   against a polygon PostGIS has already prepared.
#:
#: ``count(cause_id)`` and not ``count(*)`` for the classified counts: counting a
#: nullable column counts the rows where it is filled in, which is the definition of
#: classified here.
#:
#: ``no_point`` costs nothing — it is a null test on a column already in hand — and
#: it is the first thing to look at when ``inside`` is far below ``fires``.
COUNTS_SQL = """
WITH spain AS MATERIALIZED (""" + SPAIN_SQL + """)
SELECT
    count(*) AS fires,
    count(fire.cause_id) AS classified,
    count(*) FILTER (WHERE fire.family = :family) AS matching,
    count(*) FILTER (WHERE fire.inside) AS inside_fires,
    count(fire.cause_id) FILTER (WHERE fire.inside) AS inside_classified,
    count(*) FILTER (WHERE fire.inside AND fire.family = :family) AS inside_matching,
    count(*) FILTER (WHERE fire.ignition_id IS NULL) AS no_point
FROM (
    SELECT egif.cause_id AS cause_id,
           egif.ignition_id AS ignition_id,
           substr(cause.code, 1, 1) AS family,
           (ignition.geometry IS NOT NULL
            AND ST_Contains(spain.geometry, ignition.geometry)) AS inside
    FROM egif_wildfire AS egif
    LEFT JOIN egif_fire_cause AS cause ON cause.id = egif.cause_id
    LEFT JOIN ignition ON ignition.id = egif.ignition_id
    LEFT JOIN spain ON TRUE
    WHERE egif.campaign = :year
) AS fire
"""

#: Every campaign the archive holds, oldest first.
YEARS_SQL = """
SELECT DISTINCT campaign FROM egif_wildfire ORDER BY campaign
"""

#: Whether Spain's outline is there to be tested against at all.
COUNTRY_GEOMETRY_SQL = """
SELECT count(*)
FROM admin_boundary AS boundary
JOIN ocha_admin_boundary AS ocha ON ocha.id = boundary.id
WHERE boundary.level = :country_level AND boundary.name = :country_name
"""


def counts_query(year: int, family: str = DEFAULT_FAMILY):
    """The statement for one campaign, with its parameters bound.

    Parameters
    ----------
    year : int
        The campaign to count.
    family : str
        One of :data:`CAUSE_FAMILIES`.

    Returns
    -------
    tuple
        The SQL text object and the parameter dictionary, ready for
        ``session.execute(*counts_query(...))``.
    """
    return text(COUNTS_SQL), {
        "year": year,
        "family": cause_family(family).digit,
        "country_level": COUNTRY_LEVEL,
        "country_name": COUNTRY_NAME,
    }


def years_query():
    """Every campaign in the archive, oldest first."""
    return text(YEARS_SQL)


def country_geometry_query():
    """How many level-0 OCHA boundaries are named :data:`COUNTRY_NAME`.

    Asked once per run. With no boundaries imported there is nothing to be inside
    of, every ``... inside`` column is zero, and a table of zeros looks exactly like
    an answer — so it is worth one cheap statement to tell that apart from an
    archive whose coordinates really are all outside the country.
    """
    return text(COUNTRY_GEOMETRY_SQL), {
        "country_level": COUNTRY_LEVEL,
        "country_name": COUNTRY_NAME,
    }


def share(part: int, whole: int) -> float | None:
    """``part`` as a percentage of ``whole``, or ``None`` where there is no whole.

    ``None`` and not zero: a percentage of nothing is not zero percent, it is no
    answer, and the writers turn it into an empty cell. A campaign in which no fire
    was classified has no lightning share to report, and a zero there would be a
    claim that none of its fires was a lightning fire.
    """
    if not whole:
        return None
    return 100.0 * part / whole


def share_label(part: int, whole: int) -> str:
    """A percentage as it is written out, empty where there is none."""
    value = share(part, whole)
    return "" if value is None else f"{value:.2f}"


@dataclass(frozen=True)
class Row:
    """One line of the report: the three counts, and the three counts inside.

    Attributes
    ----------
    country : str
        The country the fires are filed in. Always
        :data:`~...wildfire_statistics.COUNTRY_NAME` — see the module docstring on
        why a point over a border does not move the fire.
    year : int or None
        The filed campaign, or ``None`` for the summary row.
    fires : int
        Every fire EGIF filed under this campaign, whatever its form says and
        wherever its coordinate is.
    classified : int
        How many of them carry a cause at all.
    matching : int
        How many carry a cause of the family asked for.
    inside_fires : int
        How many of ``fires`` publish a coordinate that really is inside the
        country's polygon.
    inside_classified, inside_matching : int
        ``classified`` and ``matching`` over those fires alone.

    Notes
    -----
    ``matching <= classified <= fires``, and each ``inside_*`` is at most its
    counterpart. Both of the first two can be zero for entirely different reasons —
    nothing of that cause, or nothing classified — which is why they are separate
    columns and why the percentage names its denominator.
    """

    country: str
    year: int | None
    fires: int
    classified: int
    matching: int
    inside_fires: int
    inside_classified: int
    inside_matching: int

    @property
    def is_total(self) -> bool:
        """Whether this is the summary row rather than one of the campaigns."""
        return self.year is None

    @property
    def year_label(self) -> str:
        return TOTAL_LABEL if self.is_total else str(self.year)

    @property
    def share(self) -> float | None:
        """``matching`` as a percentage of ``classified``."""
        return share(self.matching, self.classified)

    @property
    def inside_share(self) -> float | None:
        """``inside_matching`` as a percentage of ``inside_classified``."""
        return share(self.inside_matching, self.inside_classified)

    @property
    def located_share(self) -> float | None:
        """``inside_fires`` as a percentage of ``fires``.

        How much of the campaign can be placed on the ground at all — the number
        that says how much weight the second block of counts will bear.
        """
        return share(self.inside_fires, self.fires)

    @property
    def values(self) -> tuple[str, ...]:
        """The row as the CSV writes it, in :func:`columns` order."""
        return (
            self.country, self.year_label, str(self.fires), str(self.classified),
            str(self.matching), share_label(self.matching, self.classified),
            str(self.inside_fires), share_label(self.inside_fires, self.fires),
            str(self.inside_classified), str(self.inside_matching),
            share_label(self.inside_matching, self.inside_classified),
        )

    @property
    def readable_values(self) -> tuple[str, ...]:
        """The row as the Word document writes it: the counts with separators."""
        return (
            self.country, self.year_label, f"{self.fires:,}", f"{self.classified:,}",
            f"{self.matching:,}", share_label(self.matching, self.classified),
            f"{self.inside_fires:,}", share_label(self.inside_fires, self.fires),
            f"{self.inside_classified:,}", f"{self.inside_matching:,}",
            share_label(self.inside_matching, self.inside_classified),
        )


def total(rows: list[Row], country: str) -> Row:
    """The summary row: every count added up over the campaigns.

    Counts decompose over a partition of the fires, so this is what a single pass
    over all of them would have returned, and no fire is counted twice or left out.

    The percentages are deliberately **not** averaged. :class:`Row` recomputes them
    from the summed counts, which is the ratio of the totals rather than the mean of
    the ratios — a campaign with eleven classified fires and one with eleven
    thousand must not weigh the same in the answer.
    """
    return Row(
        country=country,
        year=None,
        fires=sum(row.fires for row in rows),
        classified=sum(row.classified for row in rows),
        matching=sum(row.matching for row in rows),
        inside_fires=sum(row.inside_fires for row in rows),
        inside_classified=sum(row.inside_classified for row in rows),
        inside_matching=sum(row.inside_matching for row in rows),
    )


def parse_arguments(argv: list[str] | None = None) -> argparse.Namespace:
    """Parse the command line."""
    parser = argparse.ArgumentParser(
        description="Wildfire counts by cause for the Spanish EGIF fire statistics.",
        epilog="EGIF names lightning outright — the 100 family, Rayo — so the default "
               "column is a count of lightning fires and not a proxy for one. Every fire "
               "is counted, and the '... inside' columns repeat the counts over the fires "
               "whose published ignition point really falls inside Spain. Spain is the "
               "only country, so there is no --country. Database settings not given here "
               "are read from the environment (.env).",
    )
    parser.add_argument("-y", "--year", type=int,
                        help="restrict to one campaign, e.g. 2023; this is the filed "
                             "Campania, not the year of the detection date")
    parser.add_argument("--cause-family", default=DEFAULT_FAMILY, choices=list(CAUSE_FAMILIES),
                        help="which family of idcausa to count, by the leading digit of "
                             "the code: 'lightning' (default) is the 100 family, Rayo. The "
                             "family and not the exact code, so a subcode a later edition "
                             "of the classification adds is counted as soon as it appears")

    # Accepted only so that they can be refused clearly. Anyone reaching for one has
    # copied a command line from the EGIF or ICNF statistics report, which is a
    # reasonable thing to have done, and argparse's own "unrecognized arguments"
    # would not say why this report is different.
    parser.add_argument("--country", help=argparse.SUPPRESS)
    parser.add_argument("--country-source", help=argparse.SUPPRESS)
    parser.add_argument("--cause-type", help=argparse.SUPPRESS)
    parser.add_argument("--surface", help=argparse.SUPPRESS)

    output = parser.add_argument_group("output", "at least one is required")
    output.add_argument("--csv", type=Path, help="write the report to this .csv")
    output.add_argument("--docx", type=Path, help="write the report to this .docx (MS Word)")

    common.add_database_arguments(parser)
    parser.add_argument("--log-level", default=os.getenv("GISFIRE_LOG_LEVEL", "INFO"),
                        choices=["CRITICAL", "ERROR", "WARNING", "INFO", "DEBUG"],
                        help="verbosity (env: GISFIRE_LOG_LEVEL, default INFO)")

    arguments = parser.parse_args(argv)
    if arguments.country is not None:
        parser.error(
            "there is no --country here: EGIF is the Spanish national statistic, so "
            f"every fire in it is filed in {COUNTRY_NAME} and there is nothing to select "
            "between."
        )
    if arguments.country_source is not None:
        parser.error(
            "there is no --country-source here: this report gives both answers at once. "
            "The Fires, Classified and cause columns count every filed fire, and the "
            "'... inside' columns repeat them over the fires whose published point really "
            "falls inside the country. Nothing is dropped, so there is nothing to choose."
        )
    if arguments.cause_type is not None:
        parser.error(
            "there is no --cause-type here: that is the ICNF report's option, and EGIF "
            "publishes no Causa_Tipo. Its causes are a hierarchical three-digit idcausa — "
            "use --cause-family, whose choices are "
            f"{', '.join(CAUSE_FAMILIES)}."
        )
    if arguments.surface is not None:
        parser.error(
            "there is no --surface here: this report counts fires, not hectares, and a "
            "fire whose report form leaves the burnt area blank is still a fire. Use the "
            "companion wildfire_statistics report for the areas."
        )
    if arguments.csv is None and arguments.docx is None:
        parser.error("nothing to write: pass --csv, --docx, or both")
    return arguments


def compute(session: Session, year: int | None, logger: logging.Logger,
            family: str = DEFAULT_FAMILY) -> list[Row]:
    """Count the fires a campaign at a time, returning the report's rows in order.

    The eight steps of the report, done once per campaign and then summed:

    1. how many fires the campaign filed;
    2. how many of them carry a cause;
    3. how many of those are of the family asked for;
    4. the percentage, which is :class:`Row`'s;
    5. how many of the fires have a published point inside Spain;
    6. how many of *those* carry a cause;
    7. how many of those are of the family;
    8. the percentage again.

    One statement per campaign — steps 1-3 and 5-7 come out of it together, because
    they are six aggregates over the same rows and splitting them into six
    statements would read the same fires six times.

    Notes
    -----
    A campaign at a time, under a spinner naming it, for two reasons. It bounds what
    any one statement holds, which is the rule every other wildfire report here
    follows; and it means a long run says which year it is on rather than sitting
    silent, which is what a report over 586,157 fires has to do.

    The geometry is the only part that could be expensive, and it is one
    ``ST_Contains`` per fire against **one** prepared polygon — see
    :data:`SPAIN_SQL`. Asking instead which of the world's countries contains each
    point, which is what this report used to do, is the same question multiplied by
    250 and is why it once took hours.
    """
    label = cause_family(family).label

    if year is not None:
        years = [year]
    else:
        with common.Spinner("Finding the campaigns the EGIF fires cover", logger):
            years = list(session.scalars(years_query()))

    measured: list[Row] = []
    no_point = 0
    for index, campaign in enumerate(years, start=1):
        with common.Spinner(f"Counting the EGIF fires by cause "
                            f"({campaign}: {index} of {len(years)})", logger):
            counted = session.execute(*counts_query(campaign, family)).one()
        if not counted.fires:
            continue
        no_point += counted.no_point
        measured.append(Row(
            country=COUNTRY_NAME,
            year=campaign,
            fires=counted.fires,
            classified=counted.classified,
            matching=counted.matching,
            inside_fires=counted.inside_fires,
            inside_classified=counted.inside_classified,
            inside_matching=counted.inside_matching,
        ))

    if not measured:
        return []

    rows = sorted(measured, key=lambda row: row.year, reverse=True)
    rows.append(total(measured, COUNTRY_NAME))
    summary = rows[-1]

    if not summary.classified:
        # An XML import into a database whose cause catalogue was never seeded,
        # almost always. A table of zeros with no explanation would look like an
        # answer rather than an absence.
        logger.warning(
            "No fire in scope carries a cause at all, so every %s count is zero and no "
            "percentage can be given: an XML import cannot resolve idcausa unless the "
            "cause catalogue has been seeded first", label)

    logger.info("Of %d fire(s), %d have a point inside %s (%s); %d publish no point at all",
                summary.fires, summary.inside_fires, COUNTRY_NAME,
                "—" if summary.located_share is None
                else f"{summary.located_share:.2f}%", no_point)
    if summary.fires and not summary.inside_fires:
        if not session.scalar(*country_geometry_query()):
            logger.warning(
                "No OCHA level-0 boundary named %s is imported, so no fire can be found "
                "inside one and every '... inside' column is zero: import the OCHA "
                "country boundaries before reading them", COUNTRY_NAME)

    logger.info("Counted %d rows over %d campaign(s) (%s fires)",
                len(rows), len(measured), label)
    return rows


def write_csv(rows: list[Row], path: Path, logger: logging.Logger,
              family: str = DEFAULT_FAMILY) -> None:
    """Write the report as CSV.

    The counts go out bare and the percentages rounded to two decimals, with no
    thousands separators, because a CSV is read by another program far more often
    than by a person and a separator would make every figure a string.

    A percentage with no denominator is written **empty** — an empty field reads as
    no answer to whatever parses this, which is what it is, while a zero would read
    as an answer.
    """
    path.parent.mkdir(parents=True, exist_ok=True)
    with path.open("w", newline="", encoding="utf-8") as handle:
        writer = csv.writer(handle)
        writer.writerow(columns(family))
        for row in rows:
            writer.writerow(row.values)
    logger.info("Wrote %s", path)


def write_docx(rows: list[Row], path: Path, year: int | None,
               logger: logging.Logger, family: str = DEFAULT_FAMILY) -> None:
    """Write the report as a Word document.

    One table, with the summary row in bold, on a **landscape** page: eleven
    columns do not fit across a portrait one, and a table that wraps is a table
    nobody reads.

    Counts get thousands separators here — the opposite of the CSV, and for the
    opposite reason: this one is for reading.

    The opening paragraphs say what the two blocks of counts are and why they
    differ. Both belong in the document and not only in the manual: a reader who
    took the ``... inside`` columns for the whole archive would conclude that half
    of Spain's fires never happened.
    """
    # Imported here rather than at module scope so that --csv keeps working if
    # python-docx is not installed, which matters because it is the only
    # dependency this application adds.
    from docx import Document
    from docx.enum.section import WD_ORIENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt

    kind = cause_family(family)
    headings = columns(family)

    document = Document()
    section = document.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = section.page_height, section.page_width

    document.add_heading(
        f"EGIF wildfire counts — {kind.label} causes ({COUNTRY_NAME})", level=1)

    scope = f"campaign: {year}" if year is not None else "all campaigns"
    document.add_paragraph(
        f"Counts of EGIF wildfires whose idcausa is of the {kind.digit} family "
        f"({kind.spanish} — {kind.label}). Every filed fire is counted, whether or not "
        f"its report form gives a burnt area and wherever its coordinate is. Years are "
        f"the filed Campania. Scope: {scope}."
    )
    document.add_paragraph(
        "The percentage is of the classified fires and not of all of them, and is left "
        "blank where nothing was classified."
    )
    document.add_paragraph(
        f"The four '... inside' columns repeat the counts over the fires whose published "
        f"ignition point the database finds inside the real {COUNTRY_NAME} polygon. They "
        f"are not the archive: EGIF publishes no coordinate at all for about half of the "
        f"1982-2023 fires, including every fire before 1998, and a published coordinate "
        f"can land in the sea or over a border. 'Fires inside (%)' is how much of each "
        f"campaign can be placed on the ground."
    )

    table = document.add_table(rows=1, cols=len(headings))
    table.style = "Table Grid"
    for cell, title in zip(table.rows[0].cells, headings):
        cell.text = title
        cell.paragraphs[0].runs[0].bold = True
        cell.paragraphs[0].runs[0].font.size = Pt(8)

    for row in rows:
        cells = table.add_row().cells
        for index, (cell, value) in enumerate(zip(cells, row.readable_values)):
            cell.text = value
            paragraph = cell.paragraphs[0]
            if index >= FIRST_NUMERIC_COLUMN:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            for run in paragraph.runs:
                run.bold = row.is_total
                run.font.size = Pt(8)

    path.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(path))
    logger.info("Wrote %s", path)


def report(args: argparse.Namespace, engine: Engine, logger: logging.Logger) -> list[Row]:
    """Count the fires and write whichever outputs were asked for."""
    with Session(engine) as session:
        rows = compute(session, args.year, logger, args.cause_family)

    if not rows:
        # An empty report is almost always a campaign with no data, and writing an
        # empty file would hide that. Note that a campaign with no fire of the
        # family asked for is not empty — it is a row of zeros, which is a
        # different thing and is reported as one.
        raise RuntimeError(
            "No wildfires matched. Check --year, and that the EGIF fires are imported — "
            "every filed fire is counted here, so an empty report means there are none "
            "in scope at all."
        )

    if args.csv is not None:
        write_csv(rows, args.csv, logger, args.cause_family)
    if args.docx is not None:
        write_docx(rows, args.docx, args.year, logger, args.cause_family)
    return rows


def main(argv: list[str] | None = None) -> int:
    args = parse_arguments(argv)
    logging.basicConfig(level=args.log_level, format="%(asctime)s %(levelname)s %(message)s")
    logger = logging.getLogger("egif-causes")

    try:
        settings = common.resolve_database_settings(args)
    except RuntimeError as error:
        logger.error("%s", error)
        return 1

    engine = create_engine(common.database_url(settings))
    try:
        report(args, engine, logger)
    except Exception as error:  # noqa: BLE001  (the CLI boundary: report, do not traceback)
        logger.error("%s", error)
        return 1
    finally:
        engine.dispose()
    return 0


if __name__ == "__main__":  # pragma nocover
    sys.exit(main())
