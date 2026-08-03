#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Burnt-area statistics for the Spanish EGIF fire statistics.

Reports, per campaign, how many fires there were and the smallest, largest and
total area burnt, in hectares::

    Country          Year    Fires     Minimum      Maximum        Total
    Spain            2023      6294        0.00      9520.00      75236.41
    Spain            2022      7362        0.00     31500.00     243610.28
    Spain            Total    13656        0.00     31500.00     318846.69

Run it over everything, or narrow it to one campaign, or to the fires above a
size::

    python3 -m src.apps.statistics.wildfires.spain_egif.wildfire_statistics --csv burnt.csv
    python3 -m src.apps.statistics.wildfires.spain_egif.wildfire_statistics \\
        --year 2023 --csv 2023.csv --docx 2023.docx
    python3 -m src.apps.statistics.wildfires.spain_egif.wildfire_statistics \\
        --surface burnt --min-area 5 --csv over-5-ha.csv

…or to one autonomous community::

    python3 -m src.apps.statistics.wildfires.spain_egif.wildfire_statistics \\
        --region Catalonia --csv catalonia.csv
    python3 -m src.apps.statistics.wildfires.spain_egif.wildfire_statistics \\
        --region Andalucía --year 2023 --csv andalucia-2023.csv

At least one of ``--csv`` and ``--docx`` is required: an application that
computed a report and then printed nothing would be a strange thing.

The application only reads. Database settings come from the environment
(``.env``, see :mod:`src.settings`); every one of them can be overridden with a
command-line argument.

The area is reported, not measured
----------------------------------

This is the one thing that makes this report different from its
:mod:`GWIS <src.apps.statistics.wildfires.gwis.wildfire_statistics>`,
:mod:`GFA <src.apps.statistics.wildfires.gfa.wildfire_statistics>` and
:mod:`ICNF <src.apps.statistics.wildfires.portugal_icnf.wildfire_statistics>`
counterparts, and it is not a detail of implementation. Those three measure a
polygon. **EGIF publishes no polygon, in any of its exports, and never will** —
:attr:`~src.data_model.wildfire.Wildfire.perimeter` is ``NULL`` on every EGIF
fire by design, see :mod:`src.providers.spain_egif.wildfire`. What EGIF publishes
is a burnt *area*, in hectares, as filled in on the report form.

So there is no ``--area-method`` here: nothing is being projected or measured on
an ellipsoid, and the figures cannot be made to agree with a perimeter dataset by
choosing a better CRS. They are a different kind of number, and the difference
between them and a measured area is a real quantity about the archive rather than
an error in either.

Which surface
-------------

EGIF does not publish *one* burnt area. It publishes five, and ``--surface``
picks which one the report is of:

``forest`` (default)
    ``SuperficieTotalForestal`` — total burnt forest area, and the figure the
    national statistic is quoted in. Summing it over fires reproduces the
    published national total, which is why it is the default.
``wooded``, ``non-wooded``
    ``SuperficieArbolada`` and ``SuperficieNoArbolada``, the two parts ``forest``
    is the sum of — exactly, checked on every row of the 2022-2023 sample.
``agricultural``, ``other-non-forest``
    ``SuperficieAgricola`` and ``OtrasSuperficiesNoforestales``. Both are
    **outside** the forest total: EGIF counts forest and non-forest separately.
``burnt``
    Everything that burnt: ``forest`` plus ``agricultural`` plus
    ``other-non-forest``. Not a figure EGIF publishes, and deliberately not the
    default — quoting it as *the* burnt area of a Spanish fire year would not
    match the national statistic. It is the right one to ask for when the
    question is about land burnt rather than about forest lost.

Never add two runs of this report together. ``wooded`` and ``non-wooded`` sum to
``forest``, and ``forest``, ``agricultural`` and ``other-non-forest`` sum to
``burnt``, but only in the ``Total (ha)`` column: a minimum of minima over two
different columns is not the minimum of their sum, because the two ends need not
belong to the same fire. Ask for ``burnt`` and let the database add the rows up.

Which year a fire counts towards
--------------------------------

:attr:`~src.providers.spain_egif.wildfire.EgifWildfire.campaign` — the campaign
the *parte* is filed under, which is also the first four characters of its report
number — and **not** the year of ``start_date_time``, exactly as the ICNF report
uses the published ``Ano``.

The two are normally the same year and can disagree at a New Year's Eve fire.
Where they do, the report follows the filing, because a published yearly total is
a total of what the service filed that year. ``campaign`` is also ``NOT NULL``
and indexed, and needs no timezone applied to it — which
``start_date_time`` would, EGIF's instants being local wall-clock readings in two
different zones.

Which fires are counted
-----------------------

Those whose chosen surface **is reported**. A fire whose
``SuperficieTotalForestal`` is ``NULL`` is not a fire that burnt no forest; it is
a fire whose form does not say, and it is left out of the count as well as out of
the areas — so the ``Fires`` column always counts exactly the fires the three
figures beside it were computed from.

A reported **zero is counted**, and is a real answer: a fire that burnt only
agricultural land has a forest total of 0.00 ha. That is why a minimum of
``0.00`` is normal in this report and not a sign of a missing value. Under
``--surface burnt`` a fire counts if *any* of the three components is reported,
the unreported ones contributing nothing.

``--min-area`` narrows it further, to the fires of at least that many hectares.
By default there is no threshold.

One autonomous community
------------------------

``--region`` reports one *comunidad autónoma* instead of the whole country —
``--region Catalonia``, ``--region Andalucía``, ``--region 09``. It takes the
published IGN name, either half of a bilingual one (``Cataluña/Catalunya``), the
English name, or the two-digit INE code; case and accents do not matter, and a name
that picks out exactly one community by being part of it (``Madrid``) works too. It
needs the IGN administrative boundaries imported, and an unrecognised name is
answered with the list of the ones that are.

**The selection is on the filing, not on the map.** A fire is in the report when its
*parte* is filed to a province of that community —
:attr:`~src.providers.spain_egif.wildfire.EgifWildfire.province_ine_code`, which is
``NOT NULL`` on every fire EGIF has ever published. The provinces of the community
are read from the IGN hierarchy at the start of the run, so the statement itself
carries a handful of two-character codes and joins nothing.

Testing the published ignition point against the community's polygon would be a
different question, and a much worse one to answer with: half the archive publishes
no coordinate at all, so it would silently drop 294,052 of the 586,157 fires. It is
therefore not what this does, and ``--region`` composes with ``--country-source``
without either changing what the other means.

.. note::

   The ``Country`` column still says ``Spain`` on every row, and the CSV keeps the
   shape the other four reports have. What the table is a table *of* is in the
   ``.docx`` heading and in the log — so name the region in the filename of a CSV
   you intend to keep.

.. warning::

   **A campaign's total is a floor, not the year's burnt area.** Every fire the
   service exports is in state *Cerrado Revisión*, and a region appears only once
   it has closed its fires, so a freshly exported year is missing whole regions —
   the 2022 sample here totals 243,610 ha against the ~306,000 ha eventually
   published. A year has to be re-exported and re-imported later. Comparing two
   campaigns of different vintages compares two different degrees of completeness,
   not two fire seasons.

Which country a fire counts towards, and the points in the sea
--------------------------------------------------------------

EGIF has no perimeter, but it does have a **point** — the published ignition
coordinate on :class:`~src.providers.spain_egif.ignition.EgifIgnition` — and that
point is testable against a country polygon exactly as the other three reports
test an interior point of their perimeter. ``--country-source`` chooses whether to
use it.

**``filed``** (the default) takes EGIF's own word for it: the report is a Spanish
*parte*, so the fire is in Spain, and the ``Country`` column is the constant
:data:`COUNTRY_NAME` on every row. Every fire that reports the surface is counted.

**``geometry``** asks the database which country actually contains the ignition
point, at report time, against the real polygons. It is worth reaching for because
**an EGIF coordinate can be badly wrong and still survive import**. The importer's
only geometric guard is a plausibility box on the published easting and northing
(:data:`~src.providers.spain_egif.PLAUSIBLE_UTM_EASTING`), which is a rectangle
containing a great deal of Atlantic and Mediterranean; and where the published
*huso* is not a zone Spain lies in, the zone is replaced with the modal one for the
province, which can walk a coastal point out to sea. Under ``geometry`` those fires
are dropped, and a point that lands over the French or Portuguese border is
reported as that country's row rather than folded into Spain's.

Two things to know before switching it on:

* **It drops every fire with no published point, which is half the archive.**
  293,710 of the 586,157 fires of 1982-2023 publish no coordinate at all, every
  fire before 1998 among them. A ``geometry`` run of the historical series is
  therefore not comparable with a ``filed`` one and is nowhere near the published
  national figure. It answers "what do the located fires say", not "what burnt".
* Whenever it is in force the log says how many fires were dropped and why, split
  between *no point published* and *point in no country* — see
  :func:`location_audit`. The second number is the one to watch: it is the count
  of coordinates that are somewhere a Spanish fire is not.

So the default is the opposite of the other three reports', which default to
``geometry``. There it costs nothing, every perimeter being present; here it would
silently halve the report.

One statement
-------------

The other three reports measure one year per statement, because the memory a
point-in-polygon test against a country polygon needs is only released when the
statement ends, and a single pass over twenty million perimeters took a 30 GB
machine to the OOM killer.

This report is one statement. Under ``filed`` it does no geometry whatsoever — an
indexed aggregate over one column of one table, which does not even join the
parent ``wildfire`` row. Under ``geometry`` it does test a point per fire, but
against 586,157 points at worst, two orders of magnitude short of the case that
died, and a point is a far cheaper thing to contain than a multipolygon. The
``Total`` row is arithmetic over the campaigns, by :func:`combine`, so the output
is the same shape as the other three either way.

The ``Country`` column is kept in both modes so that this report's CSV has the
same shape as the other three and the four can be concatenated and compared, which
is the whole reason for reporting a national statistic in the same table as three
perimeter datasets.
"""

from __future__ import annotations

import argparse
import csv
import logging
import math
import os
import sys
import unicodedata

from dataclasses import dataclass
from pathlib import Path

from sqlalchemy import ColumnElement
from sqlalchemy import Engine
from sqlalchemy import Select
from sqlalchemy import create_engine
from sqlalchemy import func
from sqlalchemy import literal
from sqlalchemy import or_
from sqlalchemy import select
from sqlalchemy import true
from sqlalchemy.orm import Session
from sqlalchemy.orm import aliased

import src.settings  # noqa: F401  (imported for the side effect of loading .env)

from src.apps.imports import common
from src.data_model.geography.admin_boundary import AdminBoundary
from src.data_model.ignition import Ignition
from src.providers.ocha.admin_boundary import OchaAdminBoundary
from src.providers.spain_egif.wildfire import EgifWildfire
from src.providers.spain_ign.admin_boundary import IgnAdminBoundary

#: Label used in the ``Year`` column for the summary row.
TOTAL_LABEL = "Total"

#: The report's columns, in order, shared by both output formats so that a change
#: to one cannot silently leave the other behind. Deliberately identical to the
#: GWIS, GFA and ICNF reports' columns, so the four CSVs can be concatenated.
COLUMNS = ("Country", "Year", "Fires", "Minimum (ha)", "Maximum (ha)", "Total (ha)")

#: Index of the first column that holds a number, and so is right-aligned in the
#: Word table.
FIRST_NUMERIC_COLUMN = 2

#: The country EGIF reports on, and the whole of the ``Country`` column under
#: ``--country-source filed``. Spelled as the OCHA boundaries spell it, so a row of
#: this report sorts and groups with the rows of the other three.
COUNTRY_NAME = "Spain"

#: The two ways of deciding which country a fire counts towards. Unlike the other
#: three reports the default is **not** ``geometry``: see the module docstring —
#: half the archive publishes no coordinate, so testing one would halve the report.
COUNTRY_SOURCE_FILED = "filed"
COUNTRY_SOURCE_GEOMETRY = "geometry"
COUNTRY_SOURCES = (COUNTRY_SOURCE_FILED, COUNTRY_SOURCE_GEOMETRY)

#: Administrative level of a country in ``admin_boundary``.
COUNTRY_LEVEL = 0

#: Administrative level of a *comunidad autónoma* in ``admin_boundary``, which is
#: what ``--region`` selects, and of a *provincia*, which is what an EGIF fire is
#: filed to and therefore how the selection is made. See :func:`resolve_region`.
REGION_LEVEL = 1
PROVINCE_LEVEL = 2

#: The English name of each *comunidad autónoma*, keyed by its two-digit INE code,
#: which is also digits 3-4 of the IGN ``NATCODE``.
#:
#: Accepted by ``--region`` **in addition to** the published name, never instead of
#: it: the IGN names are the answer to "what is this region called", and these are a
#: convenience so that ``--region Catalonia`` from an English-language command line
#: does not fail against ``Cataluña/Catalunya``. Keyed by the code rather than by the
#: name so that a renaming in a later IGN edition cannot silently unpair them.
REGION_ENGLISH_NAMES = {
    "01": "Andalusia",
    "02": "Aragon",
    "03": "Asturias",
    "04": "Balearic Islands",
    "05": "Canary Islands",
    "06": "Cantabria",
    "07": "Castile and León",
    "08": "Castile-La Mancha",
    "09": "Catalonia",
    "10": "Valencian Community",
    "11": "Extremadura",
    "12": "Galicia",
    "13": "Community of Madrid",
    "14": "Region of Murcia",
    "15": "Navarre",
    "16": "Basque Country",
    "17": "La Rioja",
    "18": "Ceuta",
    "19": "Melilla",
}

#: The burnt areas EGIF publishes, as ``--surface`` accepts them.
SURFACE_FOREST = "forest"
SURFACE_WOODED = "wooded"
SURFACE_NON_WOODED = "non-wooded"
SURFACE_AGRICULTURAL = "agricultural"
SURFACE_OTHER_NON_FOREST = "other-non-forest"
SURFACE_BURNT = "burnt"
SURFACES = (SURFACE_FOREST, SURFACE_WOODED, SURFACE_NON_WOODED,
            SURFACE_AGRICULTURAL, SURFACE_OTHER_NON_FOREST, SURFACE_BURNT)

#: How each surface reads in prose, for the Word document's opening paragraph. A
#: report of agricultural hectares that said only "areas in hectares" would be
#: indistinguishable from a report of forest ones.
SURFACE_PROSE = {
    SURFACE_FOREST: "total burnt forest area (SuperficieTotalForestal)",
    SURFACE_WOODED: "burnt wooded forest area (SuperficieArbolada)",
    SURFACE_NON_WOODED: "burnt non-wooded forest area (SuperficieNoArbolada)",
    SURFACE_AGRICULTURAL: "burnt agricultural area (SuperficieAgricola)",
    SURFACE_OTHER_NON_FOREST: "other burnt non-forest area (OtrasSuperficiesNoforestales)",
    SURFACE_BURNT: "all burnt area: forest plus agricultural plus other non-forest",
}

#: The campaign a fire counts towards. See the module docstring: the filing, not
#: the clock.
CAMPAIGN = EgifWildfire.__table__.c.campaign

#: The province the *parte* is filed to, which is what ``--region`` selects on. INE
#: two-digit codes, and ``NOT NULL`` on every fire in the archive — which is why the
#: region filter reaches all of it and a geometric one would reach half.
PROVINCE = EgifWildfire.__table__.c.province_ine_code


def reported_surface(surface: str) -> tuple[ColumnElement, ColumnElement]:
    """The burnt area of one fire in hectares, as ``(hectares, is_reported)``.

    Parameters
    ----------
    surface : str
        One of :data:`SURFACES`.

    Returns
    -------
    tuple
        The SQL expression yielding hectares, and the condition a fire has to
        satisfy for that expression to be an answer rather than a silence.

    Raises
    ------
    ValueError
        If ``surface`` is not one of :data:`SURFACES`.

    Notes
    -----
    The second element is what keeps the ``Fires`` column honest. Every one of
    these columns is nullable, ``sum`` and ``min`` skip nulls and ``count`` does
    not, so a report that did not filter would count fires whose area it had not
    included. ``NULL`` here means the form does not say — not zero hectares.

    ``burnt`` is the one composite: a fire counts if any of the three components
    is reported, and the ones that are not contribute nothing. Requiring all three
    would discard most of the archive over an unreported column that is zero on
    the overwhelming majority of fires anyway.
    """
    fire = EgifWildfire.__table__.c

    simple = {
        SURFACE_FOREST: fire.area_ha_forest_total,
        SURFACE_WOODED: fire.area_ha_wooded,
        SURFACE_NON_WOODED: fire.area_ha_non_wooded,
        SURFACE_AGRICULTURAL: fire.area_ha_agricultural,
        SURFACE_OTHER_NON_FOREST: fire.area_ha_other_non_forest,
    }
    if surface in simple:
        column = simple[surface]
        return column, column.is_not(None)

    if surface == SURFACE_BURNT:
        components = (fire.area_ha_forest_total, fire.area_ha_agricultural,
                      fire.area_ha_other_non_forest)
        hectares = (func.coalesce(components[0], 0.0)
                    + func.coalesce(components[1], 0.0)
                    + func.coalesce(components[2], 0.0))
        return hectares, or_(*[component.is_not(None) for component in components])

    raise ValueError(
        f"unknown surface {surface!r}; expected one of {', '.join(SURFACES)}"
    )


def country_columns(source: str) -> tuple[ColumnElement, list]:
    """Where a fire's country comes from, as ``(name, joins)``.

    Parameters
    ----------
    source : str
        One of :data:`COUNTRY_SOURCES`.

    Returns
    -------
    tuple
        The country-name expression, and the joins that have to be applied to a
        ``select`` over ``egif_wildfire`` for it to resolve.

    Raises
    ------
    ValueError
        If ``source`` is not one of :data:`COUNTRY_SOURCES`.

    Notes
    -----
    **``filed``** is a literal and needs no join. It is still grouped on, like the
    real column: by the time the outer aggregate sees it, it is a column of the
    subquery rather than a constant, and PostgreSQL requires every one of those in
    the ``GROUP BY``. One extra grouping key with a single distinct value costs
    nothing and keeps the two modes the same shape.

    **``geometry``** tests the published ignition point. Both of its joins are
    inner, and each drops a different kind of fire: the join to ``ignition`` drops
    the fires that publish no coordinate — half the archive — and the lateral drops
    the ones whose coordinate is inside no country, which is what a point in the
    sea is. :func:`location_audit` counts the two separately, because they mean
    entirely different things about the data.

    ``LATERAL ... LIMIT 1`` rather than a plain join: a point on a shared border
    can satisfy ``ST_Contains`` for two countries, and one fire must not become two
    rows. The join to ``ocha_admin_boundary`` is what keeps the test against the
    OCHA country outlines rather than against every level-0 boundary any provider
    has ever loaded.
    """
    if source == COUNTRY_SOURCE_FILED:
        return literal(COUNTRY_NAME), []

    if source != COUNTRY_SOURCE_GEOMETRY:
        raise ValueError(
            f"unknown country source {source!r}; expected one of {', '.join(COUNTRY_SOURCES)}"
        )

    ignition = Ignition.__table__
    ocha_boundary = OchaAdminBoundary.__table__
    containing = (
        select(AdminBoundary.name.label("name"))
        .select_from(AdminBoundary)
        .join(ocha_boundary, ocha_boundary.c.id == AdminBoundary.id)
        .where(AdminBoundary.level == COUNTRY_LEVEL)
        .where(func.ST_Contains(AdminBoundary.geometry, ignition.c.geometry))
        .limit(1)
        .lateral("containing")
    )
    joins = [
        (ignition, ignition.c.id == EgifWildfire.__table__.c.ignition_id),
        (containing, true()),
    ]
    return containing.c.name, joins


@dataclass(frozen=True)
class Region:
    """One *comunidad autónoma*, and the provinces an EGIF fire can be filed to in it.

    Attributes
    ----------
    source_id : str
        The IGN ``NATCODE``, eleven digits — ``34090000000`` for Cataluña.
    name : str
        The name the IGN publishes, which for four of the nineteen is bilingual and
        carries a slash: ``Cataluña/Catalunya``, ``País Vasco/Euskadi``.
    provinces : tuple of str
        The INE province codes of its *provincias*, in order. This is what the
        reports filter on — see :func:`resolve_region`.
    """

    source_id: str
    name: str
    provinces: tuple[str, ...]

    @property
    def code(self) -> str:
        """The two-digit INE code of the community, digits 3-4 of the ``NATCODE``."""
        return self.source_id[2:4]

    @property
    def english_name(self) -> str | None:
        """The English name ``--region`` also accepts, where there is one."""
        return REGION_ENGLISH_NAMES.get(self.code)

    @property
    def label(self) -> str:
        """How the region is named in prose, in the log and in the ``.docx``."""
        english = self.english_name
        if english is None or normalise(english) in {normalise(part)
                                                     for part in self.name.split("/")}:
            return self.name
        return f"{self.name} ({english})"

    @property
    def aliases(self) -> set[str]:
        """Everything ``--region`` accepts for this region, normalised.

        The code, the published name, each half of a bilingual one, and the English
        name. Compared after :func:`normalise`, so case and accents do not matter and
        ``Cataluna`` reaches ``Cataluña/Catalunya``.
        """
        names = [self.name, *self.name.split("/")]
        if self.english_name is not None:
            names.append(self.english_name)
        return {self.code} | {normalise(name) for name in names}


def normalise(text: str) -> str:
    """A name as it is compared: case-folded, unaccented, stripped.

    Notes
    -----
    Decomposed with NFKD and the combining marks dropped, rather than a table of
    substitutions, so that every accent in every one of the nineteen names is handled
    by the same rule. It is deliberately **not** applied to what is stored or
    reported: it exists to compare what someone typed with what the IGN published.
    """
    decomposed = unicodedata.normalize("NFKD", text.strip().casefold())
    return "".join(character for character in decomposed
                   if not unicodedata.combining(character))


def regions_query() -> Select:
    """Every IGN *comunidad autónoma* with the provinces under it, one row per province.

    Returns
    -------
    Select
        A query yielding ``source_id, name, province`` — the last one ``None`` for a
        community with no province imported under it.

    Notes
    -----
    Joined to ``ign_admin_boundary`` and not left at ``level = 1``, for the same
    reason the country queries join ``ocha_admin_boundary``: level 1 is *some*
    provider's first division below the country, and the Portuguese CAOP *distritos*
    are at that level too. Only the IGN publishes Spanish communities.

    The province's INE code is read out of its ``NATCODE`` rather than from
    :attr:`~src.providers.spain_ign.admin_boundary.IgnAdminBoundary.ine_code`, which
    the IGN fills on *municipios* only — above that level those five digits are
    padding and the column is ``NULL``. Digits 5-6 of a *provincia*'s ``NATCODE`` are
    its INE code, which is exactly what EGIF files a fire under.

    Outer-joined to the provinces, so a community whose *provincias* were never
    imported comes back with none rather than vanishing: that is a different failure
    from naming a region that does not exist, and :func:`resolve_region` reports it
    as one.
    """
    region = aliased(AdminBoundary, name="region")
    province = aliased(AdminBoundary, name="province")
    ign = IgnAdminBoundary.__table__

    return (
        select(region.source_id.label("source_id"),
               region.name.label("name"),
               func.substr(province.source_id, 5, 2).label("province"))
        .select_from(region)
        .join(ign, ign.c.id == region.id)
        .outerjoin(province, (province.parent_id == region.id)
                   & (province.level == PROVINCE_LEVEL))
        .where(region.level == REGION_LEVEL)
        .order_by(region.source_id, province.source_id)
    )


def available_regions(session: Session) -> list[Region]:
    """The *comunidades autónomas* in the database, in ``NATCODE`` order."""
    regions: dict[str, tuple[str, list[str]]] = {}
    for record in session.execute(regions_query()):
        name, provinces = regions.setdefault(record.source_id, (record.name, []))
        if record.province is not None:
            provinces.append(record.province)
    return [Region(source_id=source_id, name=name, provinces=tuple(provinces))
            for source_id, (name, provinces) in regions.items()]


def resolve_region(session: Session, wanted: str) -> Region:
    """The *comunidad autónoma* ``--region`` names.

    Parameters
    ----------
    session : Session
        An open session. The regions are read from the database rather than written
        out here, so the nineteen names are the IGN edition's own.
    wanted : str
        What was passed to ``--region``: the published name, either half of a
        bilingual one, the English name, or the two-digit INE code. Case and accents
        do not matter, and a name that identifies exactly one region by being part of
        it — ``Madrid`` for ``Comunidad de Madrid`` — is accepted too.

    Returns
    -------
    Region
        The community, with the INE province codes the reports filter on.

    Raises
    ------
    RuntimeError
        If no IGN community is imported at all, if nothing matches, if several
        regions match, or if the one that matched has no *provincias* under it.

    Notes
    -----
    **Why the provinces and not the polygon.** A fire is selected by the province its
    *parte* is filed to, which is
    :attr:`~src.providers.spain_egif.wildfire.EgifWildfire.province_ine_code`,
    ``NOT NULL`` on every fire in the archive. Testing the published ignition point
    against the community's polygon would be a different question with a much worse
    answer: half the archive publishes no coordinate at all, so it would drop 294,052
    of the 586,157 fires — see the module docstring on ``--country-source geometry``,
    which is the one place this report does ask a geometric question.

    So ``--region`` is a filter on the **filing**, and it composes with
    ``--country-source`` without either changing what the other means.

    Every error names what to do about it. Getting a region wrong is much likelier
    than the other failures this report can have — there are nineteen names, four of
    them bilingual — and an unhelpful message would be answered by reading the source.
    """
    regions = available_regions(session)
    if not regions:
        raise RuntimeError(
            "No comunidad autónoma is imported, so --region has nothing to select "
            "from: import the IGN administrative boundaries first, with "
            "src.apps.imports.admin_boundaries.ign.import_admin_boundaries"
        )

    asked = normalise(wanted)
    if asked.isdigit():
        asked = asked.zfill(2)

    matched = [region for region in regions if asked in region.aliases]
    if not matched:
        # Only when nothing matched exactly, so 'Rioja' cannot be dragged away from a
        # region actually called that by a longer one that happens to contain it.
        matched = [region for region in regions
                   if any(asked in alias for alias in region.aliases if not alias.isdigit())]
    if not matched:
        raise RuntimeError(
            f"No comunidad autónoma matches {wanted!r}. The regions imported are: "
            f"{', '.join(region.label for region in regions)}"
        )
    if len(matched) > 1:
        raise RuntimeError(
            f"{wanted!r} matches several comunidades autónomas — "
            f"{', '.join(region.label for region in matched)} — so it is refused rather "
            f"than one of them being guessed at. Give the whole name, or the INE code"
        )

    region = matched[0]
    if not region.provinces:
        raise RuntimeError(
            f"{region.label} is imported but none of its provincias is, and an EGIF fire "
            f"is filed to a province rather than to a community: import the IGN level-2 "
            f"boundaries as well"
        )
    return region


def statistics_query(surface: str = SURFACE_FOREST,
                     year: int | None = None,
                     min_area: float | None = None,
                     country_source: str = COUNTRY_SOURCE_FILED,
                     region: Region | None = None) -> Select:
    """Build the statistics query: one row per campaign, newest first.

    Parameters
    ----------
    surface : str
        One of :data:`SURFACES`.
    year : int, optional
        Restrict to one campaign. ``None``, the default, reports every campaign.
    min_area : float, optional
        Count only fires of at least this many hectares of the chosen surface.
        ``None``, the default, counts every fire.
    country_source : str
        One of :data:`COUNTRY_SOURCES`.
    region : Region, optional
        Restrict to the fires filed in one *comunidad autónoma*. ``None``, the
        default, reports the whole country.

    Returns
    -------
    Select
        A query yielding ``country, year, minimum, maximum, total, fires``. The
        summary rows are :func:`summarise`'s work.

    Notes
    -----
    Built against the mapped classes rather than written as SQL text, so a column
    renamed on a model breaks this at import time rather than in front of a user.

    ``region`` is an ``IN`` over the province the *parte* is filed to, and joins
    nothing: the provinces were resolved once, before the report ran (see
    :func:`resolve_region`), so the statement carries a handful of two-character
    codes rather than a join to the boundary tables.

    Under ``filed`` it is selected from ``egif_wildfire`` alone: everything the
    report reads — the campaign and the five reported areas — is on that table, so
    neither the parent ``wildfire`` row nor the ignition is joined at all. By table
    and not through the mapped class, to keep SQLAlchemy from adding a polymorphic
    join of its own.

    The inner query computes each area exactly once. Folded into the outer
    aggregate instead, the ``burnt`` expression would be evaluated three times per
    row, for the minimum, the maximum and the sum.

    ``min_area`` filters that subquery's column rather than repeating the
    expression, and is applied before the aggregates rather than as a ``HAVING``:
    the threshold selects the fires the figures are computed from, it does not
    discard years whose total came out small.
    """
    hectares, is_reported = reported_surface(surface)
    country_name, joins = country_columns(country_source)

    fires = (
        select(country_name.label("country"),
               CAMPAIGN.label("year"),
               hectares.label("hectares"))
        .select_from(EgifWildfire.__table__)
        .where(is_reported)
    )
    for target, condition in joins:
        fires = fires.join(target, condition)
    if year is not None:
        fires = fires.where(CAMPAIGN == year)
    if region is not None:
        fires = fires.where(PROVINCE.in_(region.provinces))

    fire = fires.subquery("fire")
    statistics = (
        select(
            fire.c.country,
            fire.c.year,
            func.min(fire.c.hectares).label("minimum"),
            func.max(fire.c.hectares).label("maximum"),
            func.sum(fire.c.hectares).label("total"),
            func.count().label("fires"),
        )
        .group_by(fire.c.country, fire.c.year)
        .order_by(fire.c.year.desc())
    )
    if min_area is not None:
        statistics = statistics.where(fire.c.hectares >= min_area)
    return statistics


def location_audit(surface: str = SURFACE_FOREST,
                   year: int | None = None,
                   min_area: float | None = None,
                   region: Region | None = None) -> Select:
    """Count the fires ``--country-source geometry`` leaves out, and why.

    Returns
    -------
    Select
        A query yielding one row, ``no_point, outside``: how many fires in scope
        publish no coordinate at all, and how many publish one that is inside no
        country.

    Notes
    -----
    The two are worth separating and the report does not add them up. *No point* is
    an ordinary property of the archive — half of it, and every fire before 1998 —
    and says nothing about the fires that do have one. *Outside* is a coordinate
    that is somewhere a Spanish fire is not, which is a data fault: the import's
    only geometric guard is a plausibility box on the published UTM easting and
    northing, and that rectangle contains a great deal of sea.

    Run under the same scope as the report it accompanies — same surface, same
    year, same threshold, same region — so that its two numbers and the ``Fires``
    column add up to the fires that reported the surface.
    """
    hectares, is_reported = reported_surface(surface)
    egif = EgifWildfire.__table__
    ignition = Ignition.__table__
    ocha_boundary = OchaAdminBoundary.__table__

    located = (
        select(AdminBoundary.id)
        .select_from(AdminBoundary)
        .join(ocha_boundary, ocha_boundary.c.id == AdminBoundary.id)
        .where(AdminBoundary.level == COUNTRY_LEVEL)
        .where(func.ST_Contains(AdminBoundary.geometry, ignition.c.geometry))
        .exists()
    )

    fires = (
        select(egif.c.ignition_id.label("ignition_id"),
               hectares.label("hectares"))
        .select_from(egif)
        .outerjoin(ignition, ignition.c.id == egif.c.ignition_id)
        .add_columns(located.label("in_a_country"))
        .where(is_reported)
    )
    if year is not None:
        fires = fires.where(CAMPAIGN == year)
    if region is not None:
        fires = fires.where(PROVINCE.in_(region.provinces))

    fire = fires.subquery("fire")
    audit = select(
        func.count().filter(fire.c.ignition_id.is_(None)).label("no_point"),
        func.count().filter(
            fire.c.ignition_id.is_not(None) & ~fire.c.in_a_country
        ).label("outside"),
    ).select_from(fire)
    if min_area is not None:
        audit = audit.where(fire.c.hectares >= min_area)
    return audit


@dataclass(frozen=True)
class Row:
    """One line of the report.

    Attributes
    ----------
    country : str
        Name of the country the fires burnt in. :data:`COUNTRY_NAME` under
        ``--country-source filed``; under ``geometry`` it is whichever country
        contains the ignition point, which is normally but not always Spain.
    year : int or None
        The campaign, or ``None`` for the summary row.
    minimum, maximum, total : float
        Smallest single fire, largest single fire and sum of every fire, in
        hectares of the surface asked for.
    fires : int
        How many fires the three area figures were computed from — the fires of
        this campaign that report that surface, or of the whole period for the
        summary row.
    """

    country: str
    year: int | None
    minimum: float
    maximum: float
    total: float
    fires: int

    @property
    def is_total(self) -> bool:
        """Whether this is the summary row rather than one of the campaigns."""
        return self.year is None

    @property
    def year_label(self) -> str:
        return TOTAL_LABEL if self.is_total else str(self.year)


def combine(rows: list[Row], country: str = COUNTRY_NAME, year: int | None = None) -> Row:
    """One row summarising several: the four figures taken over all of them.

    Notes
    -----
    All four decompose over a partition of the fires — a minimum of minima is a
    minimum, a sum of sums is a sum — so the ``Total`` row is the number a second
    aggregate over the same rows would have returned, and no fire is counted twice
    or left out.

    ``fsum`` rather than ``sum``: an exact accumulation over a handful of partial
    totals costs nothing and cannot drift from what one pass would have returned.
    """
    return Row(
        country=country,
        year=year,
        minimum=min(row.minimum for row in rows),
        maximum=max(row.maximum for row in rows),
        total=math.fsum(row.total for row in rows),
        fires=sum(row.fires for row in rows),
    )


def ordered_countries(session: Session, names: set[str]) -> list[str]:
    """The countries of the report, in the order the database would have sorted them.

    Notes
    -----
    One name, ``Spain``, under ``--country-source filed``; under ``geometry`` there
    can be a second, which is a fire whose ignition point is over a border.

    Sorted by PostgreSQL and not by Python because the two do not agree: Python
    compares code points, which puts every accented name after every unaccented
    one, while the database sorts by its collation.

    A name the query does not return is appended rather than dropped — under
    ``filed`` the constant ``Spain`` is not a row of ``admin_boundary`` at all
    unless the OCHA boundaries happen to be imported, and a report that vanished
    for want of them would be a poor trade.
    """
    if not names:
        return []
    ordered = list(session.scalars(
        select(AdminBoundary.name)
        .distinct()
        .where(AdminBoundary.level == COUNTRY_LEVEL)
        .where(AdminBoundary.name.in_(names))
        .order_by(AdminBoundary.name)
    ))
    return ordered + sorted(names - set(ordered))


def summarise(measured: list[Row], countries: list[str]) -> list[Row]:
    """Build the report from the campaigns measured: the summary rows, in order.

    Parameters
    ----------
    measured : list of Row
        One row per country and campaign, as the statement returned them.
    countries : list of str
        The countries in scope, in the order they are to be reported.

    Returns
    -------
    list of Row
        Each country, its campaigns newest first and its summary row last. Empty
        if nothing was measured — a report of no fires has no total either.
    """
    report: list[Row] = []
    for name in countries:
        rows = [row for row in measured if row.country == name]
        if not rows:
            continue
        report += sorted(rows, key=lambda row: row.year, reverse=True)
        report.append(combine(rows, name, None))
    return report


def hectares(text: str) -> float:
    """Argparse type for ``--min-area``: a finite, non-negative number of hectares.

    Raises
    ------
    argparse.ArgumentTypeError
        If the text is not a number, or is negative, or is a non-finite float.

    Notes
    -----
    A bare ``type=float`` would accept ``-5``, ``nan`` and ``inf``. The first two
    are almost certainly a typo and would silently produce the unfiltered report —
    ``nan`` compares false against every area, so it would produce an empty one —
    and none of the three is a size a fire can have.
    """
    try:
        area = float(text)
    except ValueError:
        raise argparse.ArgumentTypeError(f"{text!r} is not a number of hectares")
    if not math.isfinite(area):
        raise argparse.ArgumentTypeError(f"{text!r} is not a finite number of hectares")
    if area < 0:
        raise argparse.ArgumentTypeError(
            f"a burnt area cannot be negative, and {area:g} ha is: pass 0 or more, or "
            f"leave --min-area out to count every fire"
        )
    return area


def parse_arguments(argv: list[str] | None = None) -> argparse.Namespace:
    """Parse the command line."""
    parser = argparse.ArgumentParser(
        description="Burnt-area statistics for the Spanish EGIF fire statistics.",
        epilog="Areas are the hectares EGIF reports, not a measured perimeter: EGIF "
               "publishes no polygon, so there is no --area-method. Spain is the only "
               "country, so there is no --country. Database settings not given here are "
               "read from the environment (.env).",
    )
    parser.add_argument("-y", "--year", type=int,
                        help="restrict to one campaign, e.g. 2023; this is the filed "
                             "Campania, not the year of the detection date")
    parser.add_argument("--surface", default=SURFACE_FOREST, choices=SURFACES,
                        help="which reported burnt area to report: 'forest' (default) is "
                             "SuperficieTotalForestal, the figure the national statistic "
                             "is quoted in; 'wooded' and 'non-wooded' are the two parts it "
                             "is the sum of; 'agricultural' and 'other-non-forest' are "
                             "outside it; 'burnt' adds all three together, which EGIF does "
                             "not publish")
    parser.add_argument("--min-area", type=hectares, default=None, metavar="HECTARES",
                        help="count only fires that burnt at least this many hectares of "
                             "the chosen surface; by default every fire that reports one "
                             "counts, including a reported zero")
    parser.add_argument("-r", "--region", metavar="COMUNIDAD",
                        help="restrict to one comunidad autónoma, e.g. 'Cataluña', "
                             "'Catalonia', 'Andalucía' or the INE code '09'. Accents and "
                             "case do not matter, either half of a bilingual name works, "
                             "and so does a name that identifies exactly one region — "
                             "'Madrid'. The fires selected are those whose parte is filed "
                             "to a province of that community, so the whole archive is in "
                             "scope; it needs the IGN administrative boundaries imported")
    parser.add_argument("--country-source", default=COUNTRY_SOURCE_FILED,
                        choices=COUNTRY_SOURCES,
                        help="how to decide which country a fire counts towards: 'filed' "
                             "(default) takes EGIF's word for it, every fire being a "
                             "Spanish parte; 'geometry' tests the published ignition point "
                             "against the real country polygons, so a coordinate that "
                             "landed in the sea drops out and one over a border is "
                             "reported as that country. Note that 'geometry' also drops "
                             "every fire with no published point, which is half the "
                             "1982-2023 archive — the log says how many of each")

    # Accepted only so that they can be refused clearly. Anyone reaching for either
    # has copied a command line from the GWIS, GFA or ICNF report, which is a
    # reasonable thing to have done, and argparse's own "unrecognized arguments"
    # would not say why this report is different.
    parser.add_argument("--country", help=argparse.SUPPRESS)
    parser.add_argument("--area-method", help=argparse.SUPPRESS)

    output = parser.add_argument_group("output", "at least one is required")
    output.add_argument("--csv", type=Path, help="write the report to this .csv")
    output.add_argument("--docx", type=Path, help="write the report to this .docx (MS Word)")

    common.add_database_arguments(parser)
    parser.add_argument("--log-level", default=os.getenv("GISFIRE_LOG_LEVEL", "INFO"),
                        choices=["CRITICAL", "ERROR", "WARNING", "INFO", "DEBUG"],
                        help="verbosity (env: GISFIRE_LOG_LEVEL, default INFO)")

    arguments = parser.parse_args(argv)
    if arguments.country is not None:
        parser.error(
            "there is no --country here: EGIF is the Spanish national statistic, so "
            f"every fire in it is filed in {COUNTRY_NAME} and there is nothing to select "
            "between. Every fire is reported, and the Country column says which country "
            "each one turned out to be in — see --country-source."
        )
    if arguments.area_method is not None:
        parser.error(
            "there is no --area-method here: EGIF publishes no perimeter, so nothing is "
            "measured and no CRS is involved. The hectares are the ones on the report "
            "form — use --surface to choose which of them."
        )
    if arguments.csv is None and arguments.docx is None:
        parser.error("nothing to write: pass --csv, --docx, or both")
    return arguments


def compute(session: Session, year: int | None, logger: logging.Logger,
            surface: str = SURFACE_FOREST,
            min_area: float | None = None,
            country_source: str = COUNTRY_SOURCE_FILED,
            region: Region | None = None) -> list[Row]:
    """Run the statement and return the report's rows in order.

    Notes
    -----
    One statement, under one spinner — see the module docstring for why this
    report does not need the year-at-a-time machinery the other three are built
    on. The ``Total`` row is arithmetic over its result, not a second query.

    Under ``geometry`` a second, cheap statement follows it: :func:`location_audit`
    counts what the first one dropped and why. A report that quietly left out half
    its fires would be worse than one that did not offer the option.

    ``region`` narrows both of them, so the audit keeps accounting for exactly the
    fires the report is of.
    """
    scope = "the EGIF fires" if region is None else f"the EGIF fires of {region.name}"
    with common.Spinner(f"Measuring the reported {surface} area of {scope}", logger):
        measured = [
            Row(country=record.country, year=record.year,
                minimum=float(record.minimum),
                maximum=float(record.maximum),
                total=float(record.total),
                fires=record.fires)
            for record in session.execute(
                statistics_query(surface, year, min_area, country_source, region))
        ]

    if country_source == COUNTRY_SOURCE_GEOMETRY:
        with common.Spinner("Counting the fires with no usable point", logger):
            audit = session.execute(location_audit(surface, year, min_area, region)).one()
        counted = sum(row.fires for row in measured)
        logger.info(
            "Excluded %d of %d fires reporting a %s area: %d publish no point, "
            "%d publish a point in no country",
            audit.no_point + audit.outside, counted + audit.no_point + audit.outside,
            surface, audit.no_point, audit.outside)
        if audit.outside:
            logger.warning(
                "%d fire(s) have a published coordinate that is inside no country — "
                "a point in the sea survives import, whose only geometric guard is a "
                "plausibility box on the UTM easting and northing", audit.outside)

    rows = summarise(measured, ordered_countries(session, {row.country for row in measured}))
    logger.info("Computed %d rows over %d campaign(s) (%s area, country from %s, %s, %s)",
                len(rows), len({row.year for row in rows if not row.is_total}),
                surface, country_source,
                "every fire" if min_area is None else f"fires of {min_area:g} ha or more",
                "the whole country" if region is None
                else f"the fires filed in {region.label}")
    return rows


def write_csv(rows: list[Row], path: Path, logger: logging.Logger) -> None:
    """Write the report as CSV.

    The numbers go out unformatted apart from being rounded to two decimals — no
    thousands separators — because a CSV is read by another program far more often
    than by a person, and a separator would make every figure a string.
    """
    path.parent.mkdir(parents=True, exist_ok=True)
    with path.open("w", newline="", encoding="utf-8") as handle:
        writer = csv.writer(handle)
        writer.writerow(COLUMNS)
        for row in rows:
            writer.writerow([row.country, row.year_label, row.fires,
                             f"{row.minimum:.2f}", f"{row.maximum:.2f}", f"{row.total:.2f}"])
    logger.info("Wrote %s", path)


def write_docx(rows: list[Row], path: Path, year: int | None,
               logger: logging.Logger,
               surface: str = SURFACE_FOREST,
               min_area: float | None = None,
               country_source: str = COUNTRY_SOURCE_FILED,
               region: Region | None = None) -> None:
    """Write the report as a Word document.

    One table, with the summary row in bold. Numbers get thousands separators here
    — the opposite of the CSV, and for the opposite reason: this one is for
    reading.

    The opening paragraph names the surface, the scope and the incompleteness of a
    freshly exported campaign, because a table of agricultural hectares and a
    table of forest ones look exactly alike, and a total that is a floor rather
    than a year's burnt area should not have to be remembered by the reader.

    A region goes in the **heading** and not only in the scope line: the ``Country``
    column still says ``Spain`` on every row, and a table of one community's fires
    that did not say so on its front page would read as a national one.
    """
    # Imported here rather than at module scope so that --csv keeps working if
    # python-docx is not installed, which matters because it is the only
    # dependency this application adds.
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt

    document = Document()
    where = COUNTRY_NAME if region is None else f"{region.label}, {COUNTRY_NAME}"
    document.add_heading(f"EGIF wildfire burnt area ({where})", level=1)

    scope = [f"campaign: {year}" if year is not None else "all campaigns"]
    if region is not None:
        scope.append(f"only the fires filed in {region.label}")
    if min_area is not None:
        scope.append(f"only fires of {min_area:g} ha or more")
    if country_source == COUNTRY_SOURCE_GEOMETRY:
        scope.append("only fires whose published ignition point falls inside a country")
    document.add_paragraph(
        f"Areas in hectares as reported by EGIF — {SURFACE_PROSE[surface]} — and not "
        f"measured from a perimeter, which EGIF does not publish. Fires that do not "
        f"report this area are excluded; a reported zero is counted. Years are the filed "
        f"Campania. Scope: {'; '.join(scope)}."
    )
    document.add_paragraph(
        "A campaign's total is a floor: the service exports only fires it has closed and "
        "reviewed, so a recently exported year is missing whole regions and will grow "
        "when it is re-imported."
    )
    if region is not None:
        # The Country column says Spain on every row of this table, so what it is a
        # table of has to be said in prose.
        document.add_paragraph(
            f"These are the fires filed in {region.label} — those whose parte names one "
            f"of its provinces ({', '.join(region.provinces)}) — and not a national "
            f"total. The Country column is still {COUNTRY_NAME}, as in every one of "
            f"these reports. The selection is on the filing and not on the published "
            f"coordinate, so the fires with no coordinate are in it too."
        )
    if country_source == COUNTRY_SOURCE_GEOMETRY:
        # Not a footnote: this run left out half the archive, and a reader who does
        # not know that will compare its totals with a filed run's.
        document.add_paragraph(
            "These figures cover only the fires with a usable published coordinate. EGIF "
            "publishes none at all for about half of the 1982-2023 archive, including "
            "every fire before 1998, and those fires are not in this table — so its "
            "totals are well below the national figure and are not comparable with a run "
            "made without this restriction."
        )

    table = document.add_table(rows=1, cols=len(COLUMNS))
    table.style = "Table Grid"
    for cell, heading in zip(table.rows[0].cells, COLUMNS):
        cell.text = heading
        cell.paragraphs[0].runs[0].bold = True

    for row in rows:
        cells = table.add_row().cells
        values = [row.country, row.year_label, f"{row.fires:,}",
                  f"{row.minimum:,.2f}", f"{row.maximum:,.2f}", f"{row.total:,.2f}"]
        for index, (cell, value) in enumerate(zip(cells, values)):
            cell.text = value
            paragraph = cell.paragraphs[0]
            if index >= FIRST_NUMERIC_COLUMN:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            for run in paragraph.runs:
                run.bold = row.is_total
                run.font.size = Pt(9)

    path.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(path))
    logger.info("Wrote %s", path)


def report(args: argparse.Namespace, engine: Engine, logger: logging.Logger) -> list[Row]:
    """Compute the statistics and write whichever outputs were asked for.

    The region is resolved first, inside the same session, so that a name nobody
    recognises fails before any fire is measured — and against the database, so the
    error can list the communities that really are imported.
    """
    with Session(engine) as session:
        region = None if args.region is None else resolve_region(session, args.region)
        rows = compute(session, args.year, logger, args.surface, args.min_area,
                       args.country_source, region)

    if not rows:
        # An empty report is almost always a campaign with no data, and writing an
        # empty file would hide that. The surface is named because asking for one
        # the export does not fill in is the other likely reason, and so are the
        # threshold, the region and the point test when any is in force.
        extra = "" if args.min_area is None else \
            f" No fire reached the --min-area of {args.min_area:g} ha."
        if region is not None:
            extra += (f" The report is of {region.label} alone, and the EGIF exports do "
                      f"not cover every community in every campaign.")
        if args.country_source == COUNTRY_SOURCE_GEOMETRY:
            extra += (" --country-source geometry also needs the OCHA country boundaries "
                      "imported, and counts only fires with a published ignition point.")
        raise RuntimeError(
            f"No wildfires matched. Check --year, and that the EGIF fires are imported "
            f"and report a {args.surface} area — fires that do not report the surface "
            f"asked for are not counted." + extra
        )

    if args.csv is not None:
        write_csv(rows, args.csv, logger)
    if args.docx is not None:
        write_docx(rows, args.docx, args.year, logger, args.surface, args.min_area,
                   args.country_source, region)
    return rows


def main(argv: list[str] | None = None) -> int:
    args = parse_arguments(argv)
    logging.basicConfig(level=args.log_level, format="%(asctime)s %(levelname)s %(message)s")
    logger = logging.getLogger("egif-statistics")

    try:
        settings = common.resolve_database_settings(args)
    except RuntimeError as error:
        logger.error("%s", error)
        return 1

    engine = create_engine(common.database_url(settings))
    try:
        report(args, engine, logger)
    except Exception as error:  # noqa: BLE001  (the CLI boundary: report, do not traceback)
        logger.error("%s", error)
        return 1
    finally:
        engine.dispose()
    return 0


if __name__ == "__main__":  # pragma nocover
    sys.exit(main())
