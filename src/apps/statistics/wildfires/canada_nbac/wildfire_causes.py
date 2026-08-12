#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Wildfire counts by cause for the Canadian National Burned Area Composite.

Reports, per year, how many fires there were, how many of them have a **determined**
cause, how many of those were natural, and — the column that matters most here — how
many hectares those natural fires burnt::

    Country  Year   Fires  Determined  Natural  Natural (%)  Natural (ha)  Natural (% of ha)
    Canada   2025    1919        1783      958        53.73    5500397.75              76.33
    Canada   2024    1948        1864     1350        72.42    4752551.22              99.09
    Canada   2023    2215        1999     1399        69.98   13889929.12              97.09
    Canada   1977    1126         175      111        63.43     754594.83              92.00
    Canada   Total  51418       38656    25971        67.18  104179051.12              90.54

1977 is the row to read twice: 1,126 fires and **175 determined causes**. Its 63.43%
is a percentage of a seventh of the year.

Run it over everything, or narrow it to one year, or count the human-caused fires
instead::

    python3 -m src.apps.statistics.wildfires.canada_nbac.wildfire_causes --csv causes.csv
    python3 -m src.apps.statistics.wildfires.canada_nbac.wildfire_causes \\
        --year 2023 --csv 2023.csv --docx 2023.docx
    python3 -m src.apps.statistics.wildfires.canada_nbac.wildfire_causes \\
        --cause human --csv human.csv

At least one of ``--csv`` and ``--docx`` is required.

The companion of
:mod:`~src.apps.statistics.wildfires.canada_nbac.wildfire_statistics`, over the same
fires, the same years and the same scope — so the ``Country``, ``Year`` and ``Fires``
columns of the two reports agree row for row and the pair can be read side by side.
What that one measures in hectares, this one counts by cause.

There is no lightning in this dataset
--------------------------------------

The obvious question is how many fires lightning started, and **NBAC does not publish
that**. ``FIRECAUS`` takes exactly three values, and the published metadata glosses
the one that matters *"Ignition source by natural cause. Most often lightning."*

``Natural`` is therefore the finest answer this dataset supports, and it is what this
report counts by default. In the Canadian boreal a natural fire is overwhelmingly a
lightning fire, so it is a good proxy. It is a proxy all the same, and the column is
headed ``Natural`` rather than ``Lightning`` so that nothing downstream mistakes one
for the other. See :data:`~src.providers.canada_nbac.CAUSE_NATURAL`.

Nine hectares in ten, and the two datasets agree on it
-------------------------------------------------------

This is the finding the report exists to make visible, and it is why there is an area
column at all — the Portuguese and Spanish counts-by-cause reports have none.

Over the whole archive, **67.18% of the fires with a determined cause are natural and
they account for 90.54% of the burnt area.** The companion
:mod:`points report <src.apps.statistics.wildfires.canada_nfdb.wildfire_causes>` puts
the same two figures at **46.02% and 90.70%**.

So the two Canadian datasets disagree by twenty-one points about how many fires are
natural — because they count different populations, NBAC mapping what burnt and NFDB
recording every call-out an agency filed — and agree to **within two tenths of a
point** about how much of the *area* is. Two independently built archives converging
on the same answer to the question that matters is worth considerably more than either
of them saying it alone.

A report that counted only fires would show the twenty-one-point gap and hide the
agreement, and would read as the two datasets contradicting each other.

Why the denominator is the **determined** fires
------------------------------------------------

``Undetermined`` is a published category here and not a missing value — 12,762 of the
51,418 fires — and **it is not evenly spread**: 3,777 of the 1970s' 5,386 fires carry
it, against 766 of the 2020s' 9,964.

A percentage of *all* fires would therefore be a statement about how much of the
archive was ever investigated rather than about what caused the fires. It reads 17.1%
for the 1970s and 60.5% for the 2010s, and the difference between those two numbers is
almost entirely the difference in how many causes were determined:

=========  =======  ============  =====================  ==================
Decade     Fires    Undetermined  Natural, % of all      Natural, % of determined
=========  =======  ============  =====================  ==================
1970s        5,386         3,777                  17.1                    57.2
1980s        5,233         1,355                  48.4                    65.3
1990s        6,773         2,845                  44.6                    76.9
2000s       10,713         3,109                  49.5                    69.8
2010s       13,349           910                  61.1                    65.6
2020s        9,964           766                  60.5                    65.5
=========  =======  ============  =====================  ==================

The right-hand column is a fire statistic; the one beside it is a reporting statistic
wearing the same units. So the denominator is the determined fires, and their count is
a column of its own so that the denominator is never out of sight — the same decision
:mod:`the ICNF report <src.apps.statistics.wildfires.portugal_icnf.wildfire_causes>`
makes about its ``Classified`` column, for the same reason.

Where nothing is determined there is no percentage to give, and the cell is left
**empty** rather than filled with a zero that would be a claim.

.. warning::

   Even the determined share is a floor for the natural one. ``Undetermined`` fires
   are not causeless; they are uninvestigated, and in the remote boreal an
   uninvestigated fire is more likely to be a lightning fire than an average one —
   nobody was there to see it start or to be blamed for it. The bias in the missing
   quarter runs the same way as the answer.

Why ``undetermined`` cannot be counted
---------------------------------------

``--cause`` takes ``natural`` or ``human`` and not the third value. ``Undetermined``
is the complement of this report's denominator, not one of the things being compared
against it: counting it as a share of the determined fires would be zero by
construction, every time.

It is not hidden. ``Fires`` minus ``Determined`` is exactly the undetermined count,
and the report logs it.

The hectares are ``POLY_HA``, and that costs nothing
-----------------------------------------------------

The companion report measures the perimeter geodesically by default; this one reads
:attr:`~src.providers.canada_nbac.wildfire.NbacWildfire.area_ha_polygon` instead. They
are **the same quantity** — the service computes ``POLY_HA`` on an equal-area
projection, and over the whole archive the two differ by seven tenths of a hectare in
132.7 million, 0.0000005%. Reading the column costs nothing where ``ST_Area`` over a
geography is the most expensive thing the companion report does, and a counts-by-cause
report has no business being the slower of the two.

Which fires are counted
-----------------------

Exactly the rule its companion report uses, so the two agree: every imported perimeter
of the year, with the prescribed burns excluded unless ``--include-prescribed`` says
otherwise. Nothing is tested against a boundary — Natural Resources Canada maps
Canada's fires and nothing else — so the ``Country`` column is the constant ``Canada``
and there is no ``--country`` or ``--country-source``.

One year at a time
------------------

One statement per year, as in every report here: the years are found first, then each
is counted by a statement of its own and the summary row is arithmetic over the
results. Counts and sums both decompose over a partition of the fires, so the
``Total`` row is exactly what a single pass would have returned.

Shared with the companion report
---------------------------------

The years query, the scope conditions, the country name and the percentage helpers are
**imported from**
:mod:`~src.apps.statistics.wildfires.canada_nbac.wildfire_statistics` rather than
copied. Two reports over one dataset that disagreed about which fires are in scope
would be worse than one report, and a copy is a thing that drifts.
"""

from __future__ import annotations

import argparse
import csv
import logging
import math
import os
import sys

from dataclasses import dataclass
from pathlib import Path

from sqlalchemy import Engine
from sqlalchemy import Select
from sqlalchemy import create_engine
from sqlalchemy import func
from sqlalchemy import literal
from sqlalchemy import select
from sqlalchemy.orm import Session

import src.settings  # noqa: F401  (imported for the side effect of loading .env)

from src.apps.imports import common
from src.apps.statistics.wildfires.canada_nbac.wildfire_statistics import COUNTRY_NAME
from src.apps.statistics.wildfires.canada_nbac.wildfire_statistics import (
    FIRST_NUMERIC_COLUMN,
)
from src.apps.statistics.wildfires.canada_nbac.wildfire_statistics import PUBLISHED_YEAR
from src.apps.statistics.wildfires.canada_nbac.wildfire_statistics import TOTAL_LABEL
from src.apps.statistics.wildfires.canada_nbac.wildfire_statistics import scope_conditions
from src.apps.statistics.wildfires.canada_nbac.wildfire_statistics import share
from src.apps.statistics.wildfires.canada_nbac.wildfire_statistics import share_label
from src.apps.statistics.wildfires.canada_nbac.wildfire_statistics import years_query
from src.data_model.wildfire import Wildfire
from src.providers import canada_nbac
from src.providers.canada_nbac.wildfire import NbacWildfire

#: The causes this report can count, keyed to the published ``FIRECAUS`` value.
#:
#: :data:`~src.providers.canada_nbac.CAUSE_UNDETERMINED` is deliberately **not** here:
#: it is the complement of this report's denominator rather than one of the things
#: being compared against it, and counting it as a share of the determined fires would
#: be zero by construction. See the module docstring.
COUNTABLE_CAUSES = {
    "natural": canada_nbac.CAUSE_NATURAL,
    "human": canada_nbac.CAUSE_HUMAN,
}

#: The cause counted unless ``--cause`` says otherwise, and the nearest thing NBAC
#: publishes to a lightning fire. It is **not** a lightning category — see the module
#: docstring and :data:`~src.providers.canada_nbac.CAUSE_NATURAL`.
DEFAULT_CAUSE = "natural"

#: The causes that count towards the denominator: the ones somebody determined.
DETERMINED_CAUSES = (canada_nbac.CAUSE_NATURAL, canada_nbac.CAUSE_HUMAN)


def columns(cause: str = DEFAULT_CAUSE) -> tuple[str, ...]:
    """The report's columns, in order, for a given cause.

    Notes
    -----
    A function rather than a constant, unlike the companion report's ``COLUMNS``,
    because four of the headings name the cause being counted: a file of ``Human``
    counts under a heading saying ``Natural`` would be a trap.

    Both output formats read them from here, so a change to one cannot silently leave
    the other behind. The first three are the companion report's first three,
    unchanged, which is what lets the two be read side by side.
    """
    label = COUNTABLE_CAUSES[cause]
    return ("Country", "Year", "Fires", "Determined", label, f"{label} (%)",
            f"{label} (ha)", f"{label} (% of ha)")


def counts_query(year: int, cause: str = DEFAULT_CAUSE,
                 include_prescribed: bool = False) -> Select:
    """Build the counting query for one year.

    Parameters
    ----------
    year : int
        The published ``YEAR`` to count.
    cause : str
        One of :data:`COUNTABLE_CAUSES`' keys.
    include_prescribed : bool
        Count the prescribed burns as wildfires, as the companion report's option does.

    Returns
    -------
    Select
        A query yielding ``country, fires, determined, matching, hectares,
        determined_hectares``: at most one row. The summary row and the report's order
        are :func:`summarise`'s work.

    Notes
    -----
    Five aggregates in one pass, which is what makes this report cheap: the scan is
    paid once and not once per count.

    ``nbac_wildfire`` is joined by table, to keep SQLAlchemy from adding a polymorphic
    join of its own, and the parent ``wildfire`` is the ``FROM`` because a fire *is*
    the parent row — which is what makes the count a count of wildfires rather than of
    provider rows, and what keeps it equal to the companion report's ``Fires``.

    Both hectare figures are filtered sums over
    :attr:`~src.providers.canada_nbac.wildfire.NbacWildfire.area_ha_polygon`, and the
    denominator one is restricted to the determined fires for the same reason the count
    is: the undetermined hectares belong to fires nobody classified, and dividing by
    them would measure investigation rather than fire.
    """
    if cause not in COUNTABLE_CAUSES:
        raise ValueError(
            f"unknown cause {cause!r}; expected one of {', '.join(COUNTABLE_CAUSES)}"
        )
    nbac = NbacWildfire.__table__
    wanted = COUNTABLE_CAUSES[cause]
    determined = nbac.c.fire_cause.in_(DETERMINED_CAUSES)
    matching = nbac.c.fire_cause == wanted

    counts = (
        select(
            literal(COUNTRY_NAME).label("country"),
            func.count().label("fires"),
            func.count().filter(determined).label("determined"),
            func.count().filter(matching).label("matching"),
            func.coalesce(
                func.sum(nbac.c.area_ha_polygon).filter(matching), 0.0
            ).label("hectares"),
            func.coalesce(
                func.sum(nbac.c.area_ha_polygon).filter(determined), 0.0
            ).label("determined_hectares"),
        )
        .select_from(Wildfire)
        .join(nbac, nbac.c.id == Wildfire.id)
        .where(PUBLISHED_YEAR == year)
    )
    # The cause is what this report counts, so it is never a filter here: only the
    # prescribed-burn condition of the companion's scope is applied.
    for condition in scope_conditions(include_prescribed, None):
        counts = counts.where(condition)
    return counts.group_by(literal(COUNTRY_NAME))


@dataclass(frozen=True)
class Row:
    """One line of the report.

    Attributes
    ----------
    country : str
        Always :data:`COUNTRY_NAME`.
    year : int or None
        The published year, or ``None`` for the summary row.
    fires : int
        Every fire counted, determined or not.
    determined : int
        How many of them have a cause somebody determined — ``Natural`` or ``Human``.
        The denominator, and the difference from :attr:`fires` is the undetermined
        count.
    matching : int
        How many carry the cause asked for.
    hectares : float
        What those ``matching`` fires burnt, as ``POLY_HA``.
    determined_hectares : float
        What the determined fires burnt, which is what :attr:`area_share` divides by.

    Notes
    -----
    ``matching <= determined <= fires`` always, and the first two can be small for
    entirely different reasons: few fires of that cause, or few causes determined.
    """

    country: str
    year: int | None
    fires: int
    determined: int
    matching: int
    hectares: float
    determined_hectares: float

    @property
    def is_total(self) -> bool:
        """Whether this is the summary row rather than one of the years."""
        return self.year is None

    @property
    def year_label(self) -> str:
        return TOTAL_LABEL if self.is_total else str(self.year)

    @property
    def undetermined(self) -> int:
        """How many fires nobody classified. Not a column; logged, and derivable."""
        return self.fires - self.determined

    @property
    def share(self) -> float | None:
        """``matching`` as a percentage of ``determined``, or ``None`` if none is.

        ``None`` and not zero: a year in which nothing was determined has no share to
        report, and zero would be a claim that none of its fires was natural.
        """
        return share(self.matching, self.determined)

    @property
    def area_share(self) -> float | None:
        """The matching hectares as a percentage of the determined ones.

        The figure this report exists for: over the archive it is 90.54%, against
        67.18% of the fires, and the points report agrees with it to two tenths of a
        point while disagreeing about the fires by twenty-one.
        """
        if not self.determined_hectares:
            return None
        return 100.0 * self.hectares / self.determined_hectares

    @property
    def area_share_label(self) -> str:
        return "" if self.area_share is None else f"{self.area_share:.2f}"

    @property
    def values(self) -> tuple[str, ...]:
        """The row as the CSV writes it, in :func:`columns` order."""
        return (self.country, self.year_label, str(self.fires), str(self.determined),
                str(self.matching), share_label(self.matching, self.determined),
                f"{self.hectares:.2f}", self.area_share_label)

    @property
    def readable_values(self) -> tuple[str, ...]:
        """The row as the Word document writes it: the numbers with separators."""
        return (self.country, self.year_label, f"{self.fires:,}", f"{self.determined:,}",
                f"{self.matching:,}", share_label(self.matching, self.determined),
                f"{self.hectares:,.2f}", self.area_share_label)


def combine(rows: list[Row], country: str = COUNTRY_NAME,
            year: int | None = None) -> Row:
    """One row summarising several: the counts and the hectares added up.

    Notes
    -----
    Counts and sums both decompose over a partition of the fires, so the ``Total`` row
    is what a single pass over the same fires would have returned.

    Both shares are deliberately **not** averaged over the years: :class:`Row`
    recomputes them from the summed counts, which is the ratio of the totals rather
    than the mean of the ratios. A year with 235 determined fires and a year with 2,117
    must not weigh the same in the answer.

    ``fsum`` for the hectares, as in the companion report: an exact accumulation over a
    handful of partial totals costs nothing and cannot drift.
    """
    return Row(
        country=country,
        year=year,
        fires=sum(row.fires for row in rows),
        determined=sum(row.determined for row in rows),
        matching=sum(row.matching for row in rows),
        hectares=math.fsum(row.hectares for row in rows),
        determined_hectares=math.fsum(row.determined_hectares for row in rows),
    )


def summarise(measured: list[Row]) -> list[Row]:
    """Build the report from the years counted: the years newest first, then the total.

    Empty if nothing was counted — a report of no fires has no total either.
    """
    if not measured:
        return []
    rows = sorted(measured, key=lambda row: row.year, reverse=True)
    return rows + [combine(rows)]


def parse_arguments(argv: list[str] | None = None) -> argparse.Namespace:
    """Parse the command line."""
    parser = argparse.ArgumentParser(
        description="Wildfire counts by cause for the Canadian National Burned Area "
                    "Composite, with the area those fires burnt.",
        epilog="NBAC publishes no lightning category; 'Natural' is the closest it has "
               "and is what is counted by default. The percentages are of the fires "
               "whose cause somebody determined, because Undetermined is a published "
               "category here and is far commoner in the early years. Every perimeter "
               "is Canadian, so there is no --country. Database settings not given here "
               "are read from the environment (.env).",
    )
    parser.add_argument("-y", "--year", type=int,
                        help="restrict to one year, e.g. 2023; this is the published "
                             "YEAR, not the year of the resolved start date")
    parser.add_argument("--cause", default=DEFAULT_CAUSE, choices=sorted(COUNTABLE_CAUSES),
                        help="which published FIRECAUS to count: 'natural' (default) is "
                             "the nearest NBAC comes to a lightning fire — the metadata "
                             "glosses it 'most often lightning' — and 'human' is the "
                             "other determined cause. 'Undetermined' cannot be counted: "
                             "it is the complement of the denominator, so its share of "
                             "the determined fires is zero by construction")
    parser.add_argument("--include-prescribed", action="store_true",
                        help="count the fires flagged PRESCRIBED, as the companion "
                             "report's option of the same name does. Left out by default")

    # Accepted only so that they can be refused clearly, exactly as in the companion
    # report: anyone reaching for one has copied a command line from another dataset's.
    parser.add_argument("--country", help=argparse.SUPPRESS)
    parser.add_argument("--country-source", help=argparse.SUPPRESS)

    output = parser.add_argument_group("output", "at least one is required")
    output.add_argument("--csv", type=Path, help="write the report to this .csv")
    output.add_argument("--docx", type=Path, help="write the report to this .docx (MS Word)")

    common.add_database_arguments(parser)
    parser.add_argument("--log-level", default=os.getenv("GISFIRE_LOG_LEVEL", "INFO"),
                        choices=["CRITICAL", "ERROR", "WARNING", "INFO", "DEBUG"],
                        help="verbosity (env: GISFIRE_LOG_LEVEL, default INFO)")

    arguments = parser.parse_args(argv)
    if arguments.country is not None or arguments.country_source is not None:
        parser.error(
            f"there is no --country or --country-source here: Natural Resources Canada "
            f"maps the fires of {COUNTRY_NAME} and nothing else, so nothing is tested "
            f"against a boundary and the Country column is {COUNTRY_NAME} on every row."
        )
    if arguments.csv is None and arguments.docx is None:
        parser.error("nothing to write: pass --csv, --docx, or both")
    return arguments


def compute(session: Session, year: int | None, logger: logging.Logger,
            cause: str = DEFAULT_CAUSE,
            include_prescribed: bool = False) -> list[Row]:
    """Count the fires a year at a time, returning the report's rows in order.

    Notes
    -----
    One statement per year, under a spinner of its own — the same shape as the
    companion report, and every one of them in ``session``'s transaction and so
    against a single snapshot.

    The undetermined fires are logged as well as being derivable from the table,
    because a reader who does not know how many there are will read the percentage as
    a share of all the fires, which for the 1970s is wrong by forty points.
    """
    if year is not None:
        years = [year]
    else:
        with common.Spinner(f"Finding the years the {COUNTRY_NAME} perimeters cover",
                            logger):
            years = list(session.scalars(years_query(include_prescribed, None)))

    measured: list[Row] = []
    for index, counting in enumerate(years, start=1):
        with common.Spinner(f"Counting the {COUNTRY_NAME} perimeters by cause "
                            f"({cause}, {counting}: {index} of {len(years)})", logger):
            measured += [
                Row(country=record.country, year=counting,
                    fires=record.fires,
                    determined=record.determined,
                    matching=record.matching,
                    hectares=float(record.hectares),
                    determined_hectares=float(record.determined_hectares))
                for record in session.execute(
                    counts_query(counting, cause, include_prescribed))
            ]

    rows = summarise(measured)
    logger.info("Counted %d rows over %d year(s) (%s fires)",
                len(rows), len(measured), cause)
    if rows:
        total = rows[-1]
        logger.info("%d of %d fire(s) have a determined cause; %d of those are %s (%s%%)",
                    total.determined, total.fires, total.matching, cause,
                    share_label(total.matching, total.determined))
        logger.info("They burnt %.0f ha of the %.0f ha whose cause was determined (%s%%)",
                    total.hectares, total.determined_hectares, total.area_share_label)
        if total.undetermined:
            logger.info("%d fire(s) in scope are %s and are in neither figure; the "
                        "percentages are of the determined fires, never of all of them",
                        total.undetermined, canada_nbac.CAUSE_UNDETERMINED)
        if not total.determined:
            logger.warning(
                "No fire in scope has a determined cause, so no percentage can be "
                "given: every fire of this year is %s", canada_nbac.CAUSE_UNDETERMINED)
    return rows


def write_csv(rows: list[Row], path: Path, logger: logging.Logger,
              cause: str = DEFAULT_CAUSE) -> None:
    """Write the report as CSV.

    The percentages are written bare and rounded to two decimals, and are **empty**
    where nothing was determined — an empty field reads as no answer to whatever parses
    this, which is what it is, while a zero would read as an answer.
    """
    path.parent.mkdir(parents=True, exist_ok=True)
    with path.open("w", newline="", encoding="utf-8") as handle:
        writer = csv.writer(handle)
        writer.writerow(columns(cause))
        for row in rows:
            writer.writerow(row.values)
    logger.info("Wrote %s", path)


def write_docx(rows: list[Row], path: Path, year: int | None,
               logger: logging.Logger,
               cause: str = DEFAULT_CAUSE,
               include_prescribed: bool = False) -> None:
    """Write the report as a Word document.

    One table, with the summary row in bold. Counts get thousands separators here —
    the opposite of the CSV, and for the opposite reason: this one is for reading.

    The opening paragraphs say that NBAC publishes no lightning category, what the
    denominator is and why, and that the area share is the figure the points report
    corroborates. All three belong in the document and not only in the manual: a table
    of natural fires read without them would be misread three ways over.
    """
    # Imported here rather than at module scope so that --csv keeps working if
    # python-docx is not installed, which matters because it is the only dependency
    # this application adds.
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt

    label = COUNTABLE_CAUSES[cause]
    document = Document()
    document.add_heading(f"NBAC wildfires by cause ({COUNTRY_NAME})", level=1)

    scope = [f"year: {year}" if year is not None else "all years"]
    scope.append("prescribed burns counted" if include_prescribed
                 else "prescribed burns excluded")
    document.add_paragraph(
        f"Counts of the perimeters whose published FIRECAUS is {label}, and the area "
        f"they burnt as the service reports it (POLY_HA). Years are the published YEAR. "
        f"Scope: {'; '.join(scope)}."
    )
    document.add_paragraph(
        f"NBAC publishes no lightning category. {canada_nbac.CAUSE_NATURAL} is the "
        f"nearest it comes — the published metadata glosses it 'ignition source by "
        f"natural cause, most often lightning' — and in the Canadian boreal it is "
        f"dominated by lightning. It is a proxy all the same, which is why the column "
        f"is headed {canada_nbac.CAUSE_NATURAL} and not Lightning."
    )
    document.add_paragraph(
        f"The percentages are of the fires whose cause somebody determined, not of all "
        f"of them. {canada_nbac.CAUSE_UNDETERMINED} is a published category here rather "
        f"than a missing value, and it is far commoner in the early years — 3,777 of the "
        f"1970s' 5,386 fires against 766 of the 2020s' 9,964 — so a share of all fires "
        f"would measure how much of the archive was investigated rather than what caused "
        f"the fires. The Determined column is the denominator, and Fires minus "
        f"Determined is the undetermined count."
    )
    if cause == DEFAULT_CAUSE:
        document.add_paragraph(
            "Note the two percentages against each other. Over the archive 67.18% of "
            "the determined fires are natural and they account for 90.54% of the burnt "
            "area: natural fires are fewer than the count suggests and very much larger. "
            "The companion NFDB points report, built from the agencies' own records "
            "rather than from imagery, puts the same two figures at 46.02% and 90.70% — "
            "so the two archives disagree by twenty-one points about the fires and agree "
            "to two tenths of a point about the area."
        )

    table = document.add_table(rows=1, cols=len(columns(cause)))
    table.style = "Table Grid"
    for cell, heading in zip(table.rows[0].cells, columns(cause)):
        cell.text = heading
        cell.paragraphs[0].runs[0].bold = True

    for row in rows:
        cells = table.add_row().cells
        for index, (cell, value) in enumerate(zip(cells, row.readable_values)):
            cell.text = value
            paragraph = cell.paragraphs[0]
            if index >= FIRST_NUMERIC_COLUMN:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            for run in paragraph.runs:
                run.bold = row.is_total
                run.font.size = Pt(9)

    path.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(path))
    logger.info("Wrote %s", path)


def report(args: argparse.Namespace, engine: Engine, logger: logging.Logger) -> list[Row]:
    """Count the causes and write whichever outputs were asked for."""
    with Session(engine) as session:
        rows = compute(session, args.year, logger, args.cause, args.include_prescribed)

    if not rows:
        raise RuntimeError(
            f"No wildfires matched. Check --year, and that the {COUNTRY_NAME} "
            f"perimeters are imported — the published archive runs from "
            f"{canada_nbac.FIRST_YEAR}."
        )

    if args.csv is not None:
        write_csv(rows, args.csv, logger, args.cause)
    if args.docx is not None:
        write_docx(rows, args.docx, args.year, logger, args.cause,
                   args.include_prescribed)
    return rows


def main(argv: list[str] | None = None) -> int:
    args = parse_arguments(argv)
    logging.basicConfig(level=args.log_level, format="%(asctime)s %(levelname)s %(message)s")
    logger = logging.getLogger("nbac-causes")

    try:
        settings = common.resolve_database_settings(args)
    except RuntimeError as error:
        logger.error("%s", error)
        return 1

    engine = create_engine(common.database_url(settings))
    try:
        report(args, engine, logger)
    except Exception as error:  # noqa: BLE001  (the CLI boundary: report, do not traceback)
        logger.error("%s", error)
        return 1
    finally:
        engine.dispose()
    return 0


if __name__ == "__main__":  # pragma nocover
    sys.exit(main())
