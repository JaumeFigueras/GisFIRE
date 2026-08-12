#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Burnt-area statistics for the Canadian National Burned Area Composite.

Reports the archive per year: how many fires there were, the smallest single fire,
the largest single fire and the total area burnt in hectares — and, this dataset's
own pair, how many of the year's fires carry a real published date::

    Country   Year   Fires   Minimum (ha)   Maximum (ha)    Total (ha)  Dated  Dated (%)
    Canada    2025    1919           0.00      632231.48   7304324.18   1906      99.32
    Canada    2024    1948           0.00      451987.21   4916944.65   1911      98.10
    Canada    2023    2215           0.00     1146936.73  14796456.54   2176      98.24
    ...
    Canada    1977    1126           0.00       80309.35   1211977.72    176      15.63
    Canada    1976    1540           0.00      196940.57   2331894.63    320      20.78
    Canada    1973     512           4.78      139772.45   1845111.03    268      52.34
    Canada    Total  51418           0.00     1146936.73 132557203.23  41491      80.69

Run it over everything, or narrow it to one year, to the fires above a size, or to
one cause::

    python3 -m src.apps.statistics.wildfires.canada_nbac.wildfire_statistics --csv burnt.csv
    python3 -m src.apps.statistics.wildfires.canada_nbac.wildfire_statistics \\
        --year 2023 --csv 2023.csv --docx 2023.docx
    python3 -m src.apps.statistics.wildfires.canada_nbac.wildfire_statistics \\
        --min-area 200 --csv large-fires.csv
    python3 -m src.apps.statistics.wildfires.canada_nbac.wildfire_statistics \\
        --cause natural --csv natural.csv

…or over the hectares the service publishes rather than the ones this measures::

    python3 -m src.apps.statistics.wildfires.canada_nbac.wildfire_statistics \\
        --surface published --csv published.csv

At least one of ``--csv`` and ``--docx`` is required: an application that computed a
report and then printed nothing would be a strange thing.

The application only reads. Database settings come from the environment (``.env``,
see :mod:`src.settings`); every one of them can be overridden with a command-line
argument.

The eighth of the burnt-area reports, and GisFIRE's first outside Europe. Its first
six columns are the
:mod:`GWIS <src.apps.statistics.wildfires.gwis.wildfire_statistics>`,
:mod:`GFA <src.apps.statistics.wildfires.gfa.wildfire_statistics>`,
:mod:`ICNF <src.apps.statistics.wildfires.portugal_icnf.wildfire_statistics>`,
:mod:`EGIF <src.apps.statistics.wildfires.spain_egif.wildfire_statistics>`,
:mod:`DARPA <src.apps.statistics.wildfires.catalonia_darpa.wildfire_statistics>`,
:mod:`REDIAM <src.apps.statistics.wildfires.andalusia_rediam.wildfire_statistics>` and
:mod:`Greek <src.apps.statistics.wildfires.greece_ffa.wildfire_statistics>` reports',
in their order, so the CSVs can still be concatenated on them; the two after them are
this dataset's own.

Measured *or* published, and here the two agree
------------------------------------------------

Like the Andalusian report, this dataset publishes **both** a perimeter and a burnt
area, and ``--surface`` chooses which the report is of:

``measured`` (default)
    The area of the dissolved perimeter, computed by ``--area-method``. The default,
    because it is what makes a row here comparable with a row of the four other
    perimeter reports.
``published``
    ``POLY_HA``, the mapped area as the service computes it, **summed over the
    dissolved parts** — see :mod:`src.providers.canada_nbac`.
``adjusted``
    ``ADJ_HA``, the adjusted area burned from the models of
    `doi:10.1088/1748-9326/abfb2c <https://doi.org/10.1088/1748-9326/abfb2c>`_.

Unlike Andalusia, ``measured`` and ``published`` here are the **same number**. Over
the whole archive they are 132,738,031.02 ha and 132,738,030.32 ha — a difference of
seven tenths of a hectare in a hundred and thirty-two million, 0.0000005%. That is
not luck: the published metadata says ``POLY_HA`` is computed on the Canada Albers
Equal Area Conic projection, and an equal-area projection and a geodesic measurement
are two ways of answering the same question.

So there is no warning here of the kind the Andalusian report has to give. The two
can be quoted interchangeably, and ``published`` is worth asking for mainly as a
check on the import: a year where the two diverge is a year whose polygons did not
survive the dissolve or the reprojection intact.

``adjusted`` is a different quantity and is not interchangeable with either. It
equals ``POLY_HA`` exactly on the 49,306 fires no model was applied to, and on the
2,512 it was applied to it comes out **lower** — 5,341,008.62 ha against
6,099,383.05 — so the archive total is 131,979,655.87 rather than 132,738,030.32.
Read
:attr:`~src.providers.canada_nbac.wildfire.NbacWildfire.area_adjusted` before
quoting it: it is the flag saying whether a row's adjusted figure is a model output
at all.

How the measured area is measured
----------------------------------

The same two ways as the GFA, ICNF, Catalan and Andalusian reports: ``geodesic`` on
the WGS84 ellipsoid (the default) or ``equal-area`` in EPSG:6933. Over this archive
they agree to within 0.000004% — five hectares in a hundred and thirty-two million.

Why not the CRS the service publishes in
^^^^^^^^^^^^^^^^^^^^^^^^^^^^^^^^^^^^^^^^^

The perimeters are also stored as published, in
:data:`~src.providers.canada_nbac.SOURCE_SRID` (NAD83 / Canada Atlas Lambert), on
:attr:`~src.providers.canada_nbac.wildfire.NbacWildfire.perimeter_lambert`. It is
**not** offered as an area method, for the reason the ICNF report declines EPSG:3763
and the Andalusian one EPSG:25830, and here the reason is not a rounding detail:

.. code-block:: text

   the archive, measured four ways
     geodesic (WGS84 ellipsoid)      132,738,031 ha
     equal-area (EPSG:6933)          132,738,036 ha
     as published (POLY_HA)          132,738,030 ha
     on the published grid (3978)    127,114,627 ha     −4.2%

EPSG:3978 is a Lambert **conformal** conic — it preserves angles, not areas — and
over a country spanning 41°N to 83°N its area distortion does not cancel. Measuring
there would understate every Canadian fire year by about four percent, silently and
consistently, and would make this report incomparable with every other one in
GisFIRE. Measure there explicitly if a figure on the service's own grid is what is
wanted:

.. code-block:: sql

   SELECT year, sum(ST_Area(perimeter_lambert) / 10000.0) AS grid_ha
   FROM nbac_wildfire GROUP BY year ORDER BY year DESC;

Which year a fire counts towards
---------------------------------

The **published** ``YEAR``, which is
:attr:`~src.providers.canada_nbac.wildfire.NbacWildfire.year`. Not the year of
:attr:`~src.data_model.wildfire.Wildfire.start_date_time`.

That is the rule every report in this project follows, and here it is not a fine
distinction. 9,941 of the 51,818 fires publish no date at all: they carry
:data:`~src.providers.canada_nbac.PRECISION_YEAR` and a start of 1 January of the
published year, which is a placeholder satisfying a ``NOT NULL`` column rather than
a claim about the fire. Grouping on the clock would route a fifth of the archive
through a value that exists only because the column could not be null — and would
move 35 fires into a year the service did not file them under.

.. warning::

   The consequence is worth stating plainly, exactly as in the Portuguese report:
   **this report is sound, but a report grouped by month or day over the same data
   would not be.** A fifth of these fires would all fall on the 1st of January, and
   in 1976 it would be four fifths of them. Anything finer than a year has to filter
   on :attr:`~src.providers.canada_nbac.wildfire.NbacWildfire.date_time_precision`
   first — which is what the ``Dated`` column is here to make visible.

How many carry a real date
---------------------------

The two extra columns. ``Dated`` counts the fires of the year whose start came from
a published date — the agency's ``AG_SDATE`` or the first satellite hotspot
``HS_SDATE`` — rather than from the year alone, and ``Dated (%)`` is that as a share
of the ``Fires`` beside it.

It is not evenly spread, which is the whole reason it is a column and not a
footnote: 2010 onwards is over 90% dated every year and usually over 98%, while 1977
is 15.63% and 1978 21.93%. Satellite hotspots only start in 1989, so before that a
fire is dated if and only if an agency wrote a date down — and the years the
composite reconstructed from imagery long after the fact are the years where nobody
had.

Three things about it, as in the Catalan and Andalusian reports:

* **It is a column, not a filter.** An undated fire is still a fire and still
  contributes its hectares to the row it is in.
* **It follows the scope.** Whatever ``--year``, ``--surface``, ``--min-area``,
  ``--cause`` and ``--include-prescribed`` select, the dated count is counted over
  exactly those fires, so the percentage always has the ``Fires`` column as its
  denominator.
* **The percentage in the ``Total`` row is the ratio of the totals**, not the mean
  of the years' ratios — see :func:`combine`.

Prescribed burns are excluded by default
-----------------------------------------

400 of the 51,818 fires carry ``PRESCRIBED``, and a prescribed burn is a deliberate
fire rather than a wildfire —
:attr:`~src.providers.canada_nbac.wildfire.NbacWildfire.prescribed` says so
outright. They are left out and the log says how many were left out;
``--include-prescribed`` counts them.

They are not evenly spread either: none before 1980, 143 in the 2000s, and the 2000s
ones burnt 71,368 ha between them. On the archive total the exclusion is worth about
0.13%, but on a single year of a single province it can be most of the row.

One cause, if you want one
---------------------------

``--cause`` narrows the report to ``natural``, ``human`` or ``undetermined``
fires — the three values of ``FIRECAUS``, which
:mod:`src.providers.canada_nbac` defines.

.. warning::

   ``--cause natural`` is **not** a lightning filter. The published metadata glosses
   ``Natural`` as *"Ignition source by natural cause. Most often lightning"*, which
   is a proxy and not a category: see :data:`~src.providers.canada_nbac.CAUSE_NATURAL`.
   Anything counting lightning fires from this report is counting natural-cause fires
   and should say so.

   And the cause is not evenly reported across the archive: 1976 and 1977 are 80%
   ``Undetermined``, 2017 and 2018 barely 2%. A trend in ``--cause natural`` over
   fifty years is partly a trend in how willing an agency was to write ``Undetermined``.

No country test, and no ``--country``
--------------------------------------

Every perimeter in this dataset is Canadian, because Natural Resources Canada maps
Canada's fires and nothing else. The ``Country`` column is the constant
:data:`COUNTRY_NAME` on every row and **nothing is tested against a boundary**, so
there is no ``--country`` and no ``--country-source``: both are refused with a
message saying why.

That is a change from the GWIS, GFA and ICNF reports, and it is deliberate. Testing
these perimeters would cost the most expensive thing those reports do and would
answer a question nobody has: the import already resolved a boundary for 51,816 of
the 51,818 fires and every one of them is Canada. The two it did not are a 0.76 ha
burn off Nova Scotia and a 1.68 ha one beside it, which fall outside the OCHA
coastline rather than outside Canada — and under a containment test they would
simply vanish from the report.

The companion :mod:`points report <src.apps.statistics.wildfires.canada_nfdb.
wildfire_statistics>` **does** test, and needs to. Its coordinates are agency
reports rather than mapped polygons, and a good many of them land outside the
country.

One year at a time
-------------------

The report is not one statement. The years are found first, then each is measured by
a statement of its own and the summary rows are computed from their results — the
:mod:`GWIS <src.apps.statistics.wildfires.gwis.wildfire_statistics>` shape, for the
reason given there.

Fifty-one thousand perimeters do not need it the way twenty million do. They are
built this way because ``ST_Area`` over a geography is the expensive thing in this
report and its memory is only released when the statement ends, because these
polygons are large — a 1.1 million hectare burn is a lot of vertices — and because
five reports meant to be read side by side are worth keeping as one program over
five datasets.

Nothing about the figures changes. ``count``, ``sum``, ``min`` and ``max`` all
decompose over a partition of the fires, so the ``Total`` row is exactly the number
one pass would have returned, from the same rows. Every statement runs in one
transaction and so against one snapshot.
"""

from __future__ import annotations

import argparse
import csv
import logging
import math
import os
import sys

from dataclasses import dataclass
from pathlib import Path

from geoalchemy2 import Geography
from sqlalchemy import ColumnElement
from sqlalchemy import Engine
from sqlalchemy import Select
from sqlalchemy import cast
from sqlalchemy import create_engine
from sqlalchemy import func
from sqlalchemy import literal
from sqlalchemy import select
from sqlalchemy.orm import Session

import src.settings  # noqa: F401  (imported for the side effect of loading .env)

from src.apps.imports import common
from src.data_model.wildfire import Wildfire
from src.providers import canada_nbac
from src.providers.canada_nbac.wildfire import NbacWildfire

#: Label used in the ``Year`` column for the summary row.
TOTAL_LABEL = "Total"

#: The six columns this report shares with the GWIS, GFA, ICNF, EGIF, DARPA, REDIAM
#: and Greek ones, in their order, so the CSVs can still be concatenated on them.
SHARED_COLUMNS = ("Country", "Year", "Fires", "Minimum (ha)", "Maximum (ha)", "Total (ha)")

#: The report's columns, in order, shared by both output formats so that a change to
#: one cannot silently leave the other behind. The last two are this dataset's own:
#: how many of the year's fires carry a real published date, and that as a share.
COLUMNS = SHARED_COLUMNS + ("Dated", "Dated (%)")

#: Index of the first column that holds a number, and so is right-aligned in the Word
#: table.
FIRST_NUMERIC_COLUMN = 2

#: The country every fire in this dataset is in, and the whole of the ``Country``
#: column. Spelled as the OCHA boundaries spell it, so a row of this report sorts and
#: groups with the rows of the other seven. Nothing is tested against a boundary to
#: arrive at it — see the module docstring.
COUNTRY_NAME = "Canada"

#: Square metres in a hectare.
SQUARE_METRES_PER_HECTARE = 10_000.0

#: NSIDC EASE-Grid 2.0 Global — a cylindrical equal-area projection in metres, defined
#: for the whole world. The CRS behind ``--area-method equal-area``.
EQUAL_AREA_SRID = 6933

#: The two ways of turning a perimeter in degrees into hectares. The published
#: EPSG:3978 grid is deliberately not among them — it is conformal, and measuring
#: there understates this archive by 4.1%. See the module docstring.
AREA_METHOD_GEODESIC = "geodesic"
AREA_METHOD_EQUAL_AREA = "equal-area"
AREA_METHODS = (AREA_METHOD_GEODESIC, AREA_METHOD_EQUAL_AREA)

#: What the report is of, as ``--surface`` accepts it.
#:
#: ``measured`` is the polygon; the other two are the hectares the service publishes.
#: Unlike Andalusia's, the first two are the same quantity here — ``POLY_HA`` is
#: computed on an equal-area projection — and only ``adjusted`` is a different one.
SURFACE_MEASURED = "measured"
SURFACE_PUBLISHED = "published"
SURFACE_ADJUSTED = "adjusted"
SURFACES = (SURFACE_MEASURED, SURFACE_PUBLISHED, SURFACE_ADJUSTED)

#: How each surface reads in prose, for the Word document's opening paragraph. A
#: report of adjusted hectares that said only "areas in hectares" would be
#: indistinguishable from a report of mapped ones.
SURFACE_PROSE = {
    SURFACE_MEASURED: "the area of the dissolved published perimeter",
    SURFACE_PUBLISHED: "the mapped area as the service computes it (POLY_HA), summed "
                       "over the dissolved parts",
    SURFACE_ADJUSTED: "the adjusted area burned (ADJ_HA), which is a model output on "
                      "the 2,512 fires flagged ADJ_FLAG and a copy of POLY_HA on the rest",
}

#: The three causes, as ``--cause`` accepts them, keyed to the values ``FIRECAUS``
#: publishes. Lower-cased on the command line because a shell argument that has to be
#: capitalised exactly is a trap, and mapped rather than title-cased so that the
#: vocabulary stays :mod:`src.providers.canada_nbac`'s and not this module's.
CAUSES = {
    "natural": canada_nbac.CAUSE_NATURAL,
    "human": canada_nbac.CAUSE_HUMAN,
    "undetermined": canada_nbac.CAUSE_UNDETERMINED,
}

#: The year a fire counts towards: the published ``YEAR``. See the module docstring —
#: the filing, not the clock, and here a fifth of the clocks are placeholders.
PUBLISHED_YEAR = NbacWildfire.__table__.c.year


def burnt_area(surface: str = SURFACE_MEASURED,
               method: str = AREA_METHOD_GEODESIC) -> tuple[ColumnElement, ColumnElement]:
    """The burnt area of one fire in hectares, as ``(hectares, is_reported)``.

    Parameters
    ----------
    surface : str
        One of :data:`SURFACES`.
    method : str
        One of :data:`AREA_METHODS`. Used only when ``surface`` is
        :data:`SURFACE_MEASURED`.

    Returns
    -------
    tuple
        The SQL expression yielding hectares, and the condition a fire has to satisfy
        for that expression to be an answer rather than a silence.

    Raises
    ------
    ValueError
        If ``surface`` is not one of :data:`SURFACES`, or ``method`` not one of
        :data:`AREA_METHODS`.

    Notes
    -----
    The second element is what keeps the ``Fires`` column honest, and it means
    different things for the two kinds of surface. For ``measured`` it is a perimeter
    that exists — true of every fire here, and asserted rather than assumed. For a
    published one it is a figure that was filled in: both columns are nullable, ``sum``
    and ``min`` skip nulls and ``count`` does not, so a report that did not filter
    would count fires whose area it had not included.

    In practice neither filter bites on the published archive: every one of the 51,818
    fires has a perimeter, a ``POLY_HA`` and an ``ADJ_HA``. They are here because the
    next publication is not obliged to keep that up, and because a ``Fires`` count that
    silently stopped matching its own hectares would be very hard to notice.

    A published **zero is counted** and is a real answer — 427 fires are under a
    hundredth of a hectare and the smallest mapped polygon in the archive is 1e-7 ha,
    which is why a ``Minimum (ha)`` of ``0.00`` here is a rounded sliver rather than a
    missing value.

    The measured area is always computed from the EPSG:4326 perimeter on the parent
    ``wildfire`` row, never from the published EPSG:3978 copy — see the module
    docstring for why the service's own grid is not an option. The geodesic cast
    carries the geometry type and SRID rather than being a bare ``Geography()``: that
    renders as ``geography(GEOMETRY,-1)``, which PostGIS rejects, because -1 is not a
    SRID it knows.
    """
    fire = NbacWildfire.__table__.c

    if surface == SURFACE_MEASURED:
        if method == AREA_METHOD_GEODESIC:
            square_metres = func.ST_Area(
                cast(Wildfire.perimeter, Geography(geometry_type="MULTIPOLYGON", srid=4326))
            )
        elif method == AREA_METHOD_EQUAL_AREA:
            square_metres = func.ST_Area(
                func.ST_Transform(Wildfire.perimeter, EQUAL_AREA_SRID))
        else:
            raise ValueError(
                f"unknown area method {method!r}; expected one of {', '.join(AREA_METHODS)}"
            )
        return square_metres / SQUARE_METRES_PER_HECTARE, Wildfire.perimeter.is_not(None)

    published = {
        SURFACE_PUBLISHED: fire.area_ha_polygon,
        SURFACE_ADJUSTED: fire.area_ha_adjusted,
    }
    if surface in published:
        column = published[surface]
        return column, column.is_not(None)

    raise ValueError(
        f"unknown surface {surface!r}; expected one of {', '.join(SURFACES)}"
    )


def is_dated() -> ColumnElement:
    """Whether a fire's start came from a published date rather than from its year.

    Returns
    -------
    ColumnElement
        A boolean expression, ``True`` for a fire the ``Dated`` column counts.

    Notes
    -----
    The precision and not the source is what is tested. The two answer different
    questions — :attr:`~src.providers.canada_nbac.wildfire.NbacWildfire.date_source`
    says *who observed it*, an agency or a satellite, and
    :attr:`~src.providers.canada_nbac.wildfire.NbacWildfire.date_time_precision` says
    *how much of it is real* — and it is the second that decides whether the stored
    instant is a date or a placeholder. They cannot disagree today, ``year`` being the
    only source that yields a year-precision start; testing the precision is testing
    the property the column is about.

    Every date this dataset publishes is a bare date, so
    :data:`~src.providers.canada_nbac.PRECISION_DAY` is the best it ever gets and this
    is a two-valued question rather than the three-valued one the Portuguese archive
    poses.
    """
    return NbacWildfire.__table__.c.date_time_precision == canada_nbac.PRECISION_DAY


def is_a_wildfire(include_prescribed: bool = False) -> ColumnElement | None:
    """The condition excluding the prescribed burns, or ``None`` to keep them.

    Parameters
    ----------
    include_prescribed : bool
        ``True`` to count prescribed burns as wildfires. ``False``, the default,
        leaves them out.

    Returns
    -------
    ColumnElement or None
        ``None`` when nothing is to be filtered, so the caller can leave the ``WHERE``
        off altogether rather than adding a clause that is always true.

    Notes
    -----
    A plain ``IS NOT true`` is not needed and is not used: the column is ``NOT NULL``
    with a default of false, so ``~prescribed`` cannot silently drop a row whose flag
    was never set — which is exactly the trap
    :func:`src.apps.statistics.wildfires.greece_ffa.wildfire_statistics.is_a_fire` has
    to use ``IS DISTINCT FROM`` to avoid on a nullable column.
    """
    if include_prescribed:
        return None
    return ~NbacWildfire.__table__.c.prescribed


def cause_condition(cause: str | None) -> ColumnElement | None:
    """The condition selecting one ``FIRECAUS`` value, or ``None`` for every cause.

    Parameters
    ----------
    cause : str or None
        One of :data:`CAUSES`' keys, or ``None`` to count every fire.

    Raises
    ------
    ValueError
        If ``cause`` is neither ``None`` nor one of the three.

    Notes
    -----
    The column is nullable and is compared with ``=``, which drops the fires that
    publish no cause at all. That is the wanted behaviour and not an oversight: a fire
    whose ``FIRECAUS`` is unpublished is not a fire of the cause asked for. There are
    none in the published archive, every row carrying one of the three.
    """
    if cause is None:
        return None
    if cause not in CAUSES:
        raise ValueError(
            f"unknown cause {cause!r}; expected one of {', '.join(CAUSES)}"
        )
    return NbacWildfire.__table__.c.fire_cause == CAUSES[cause]


def scope_conditions(include_prescribed: bool = False,
                     cause: str | None = None) -> list[ColumnElement]:
    """Every condition narrowing which fires the report is of, as a list.

    Built once and applied to each year's statement, so that the two filters cannot
    drift apart between the query that finds the years and the queries that measure
    them.
    """
    conditions = [condition for condition in (is_a_wildfire(include_prescribed),
                                              cause_condition(cause))
                  if condition is not None]
    return conditions


def years_query(include_prescribed: bool = False, cause: str | None = None) -> Select:
    """The years the dataset holds fires in, newest first.

    Returns
    -------
    Select
        A query yielding one ``int`` per year.

    Notes
    -----
    Run before anything else, because each of those years is then measured by a
    statement of its own — see the module docstring.

    It carries the scope filters but no geometry: the years are read off
    ``nbac_wildfire`` alone, on an indexed ``NOT NULL`` column, so a run narrowed to
    ``--cause natural`` does not open with a statement measuring anything.

    ``DISTINCT`` over the table rather than ``min(year)`` to ``max(year)``: a gap in
    the record is a gap in the report, and a range would fill it with statements that
    can only return nothing. The published series has no gap today; 1972 is missing
    from the front of it rather than from the middle, which a range would not have
    told anyone either.
    """
    query = (
        select(PUBLISHED_YEAR.label("year"))
        .select_from(NbacWildfire.__table__)
        .distinct()
        .order_by(PUBLISHED_YEAR.desc())
    )
    for condition in scope_conditions(include_prescribed, cause):
        query = query.where(condition)
    return query


def statistics_query(year: int,
                     surface: str = SURFACE_MEASURED,
                     method: str = AREA_METHOD_GEODESIC,
                     min_area: float | None = None,
                     include_prescribed: bool = False,
                     cause: str | None = None) -> Select:
    """Build the statistics query for one year.

    Parameters
    ----------
    year : int
        The published year to measure.
    surface : str
        One of :data:`SURFACES`.
    method : str
        One of :data:`AREA_METHODS`. Used only under :data:`SURFACE_MEASURED`.
    min_area : float, optional
        Count only fires of at least this many hectares of the chosen surface.
    include_prescribed : bool
        Count the prescribed burns as wildfires.
    cause : str, optional
        Restrict to one ``FIRECAUS`` value.

    Returns
    -------
    Select
        A query yielding at most one row, ``country, minimum, maximum, total, fires,
        dated``. The summary row and the report's order are :func:`summarise`'s work.

    Notes
    -----
    Built against the mapped classes rather than written as SQL text, so a column
    renamed on a model breaks this at import time rather than in front of a user.

    ``nbac_wildfire`` is joined by table, to keep SQLAlchemy from adding a polymorphic
    join of its own, and it has to be joined in any case: the year, the published
    areas, the precision, the cause and the prescribed flag all live on it. The parent
    ``wildfire`` is the ``FROM`` because the perimeter is there and because a fire *is*
    the parent row — which is what makes the count a count of wildfires rather than of
    provider rows.

    The inner query computes the area once. Folded into the outer aggregate instead,
    ``ST_Area`` over a geography would be evaluated three times per row — for the
    minimum, the maximum and the sum — and it is by far the most expensive thing here.

    ``min_area`` filters the subquery's column rather than repeating the expression,
    and is applied before the aggregates rather than as a ``HAVING``: the threshold
    selects the fires the figures are computed from, it does not discard years whose
    total came out small. The dated count is aggregated over the same filtered rows,
    which is what keeps ``Dated`` a share of the ``Fires`` beside it.

    A year with no fires in scope produces no rows at all, rather than one row of
    ``NULL`` aggregates over zero fires — which is what keeps ``--year 1972`` on an
    archive starting in 1973 an empty report and so the caller's "no wildfires
    matched" message.
    """
    hectares_column, reported = burnt_area(surface, method)
    nbac = NbacWildfire.__table__

    fires = (
        select(
            literal(COUNTRY_NAME).label("country"),
            hectares_column.label("hectares"),
            is_dated().label("dated"),
        )
        .select_from(Wildfire)
        .join(nbac, nbac.c.id == Wildfire.id)
        .where(reported)
        .where(PUBLISHED_YEAR == year)
    )
    for condition in scope_conditions(include_prescribed, cause):
        fires = fires.where(condition)

    fire = fires.subquery("fire")
    statistics = (
        select(
            fire.c.country,
            func.min(fire.c.hectares).label("minimum"),
            func.max(fire.c.hectares).label("maximum"),
            func.sum(fire.c.hectares).label("total"),
            func.count().label("fires"),
            func.count().filter(fire.c.dated).label("dated"),
        )
        .group_by(fire.c.country)
    )
    if min_area is not None:
        statistics = statistics.where(fire.c.hectares >= min_area)
    return statistics


def prescribed_count(session: Session, year: int | None = None,
                     cause: str | None = None) -> int:
    """How many prescribed burns are in scope, for the line reporting the exclusion.

    Counted over the same cause and year the report is of, so that the number a run
    says it left out is the number that run really left out. Not narrowed by
    ``--min-area``, which is a threshold on an area this does not compute; the log
    line says "in scope" rather than "excluded from the table" for that reason.
    """
    nbac = NbacWildfire.__table__
    query = (
        select(func.count())
        .select_from(nbac)
        .where(nbac.c.prescribed)
    )
    if year is not None:
        query = query.where(PUBLISHED_YEAR == year)
    condition = cause_condition(cause)
    if condition is not None:
        query = query.where(condition)
    return session.scalar(query) or 0


def share(part: int, whole: int) -> float | None:
    """``part`` as a percentage of ``whole``, or ``None`` where there is no whole.

    ``None`` and not zero: a percentage of nothing is not zero percent, it is no
    answer, and the writers turn it into an empty cell.
    """
    if not whole:
        return None
    return 100.0 * part / whole


def share_label(part: int, whole: int) -> str:
    """A percentage as it is written out, empty where there is none."""
    value = share(part, whole)
    return "" if value is None else f"{value:.2f}"


@dataclass(frozen=True)
class Row:
    """One line of the report.

    Attributes
    ----------
    country : str
        Always :data:`COUNTRY_NAME`. Nothing is tested against a boundary to arrive at
        it — Natural Resources Canada maps Canada's fires and nothing else.
    year : int or None
        The published year, or ``None`` for the summary row.
    minimum, maximum, total : float
        Smallest single fire, largest single fire and sum of every fire, in hectares
        of the chosen surface.
    fires : int
        How many fires the three area figures were computed from.
    dated : int
        How many of those ``fires`` carry a real published date rather than being
        dated to 1 January of their year. Always at most ``fires``: it is counted over
        the same rows.
    """

    country: str
    year: int | None
    minimum: float
    maximum: float
    total: float
    fires: int
    dated: int

    @property
    def is_total(self) -> bool:
        """Whether this is the summary row rather than one of the years."""
        return self.year is None

    @property
    def year_label(self) -> str:
        return TOTAL_LABEL if self.is_total else str(self.year)

    @property
    def dated_share(self) -> float | None:
        """``dated`` as a percentage of ``fires``: how much of the year has a real date."""
        return share(self.dated, self.fires)

    @property
    def values(self) -> tuple[str, ...]:
        """The row as the CSV writes it, in :data:`COLUMNS` order."""
        return (self.country, self.year_label, str(self.fires),
                f"{self.minimum:.2f}", f"{self.maximum:.2f}", f"{self.total:.2f}",
                str(self.dated), share_label(self.dated, self.fires))

    @property
    def readable_values(self) -> tuple[str, ...]:
        """The row as the Word document writes it: the numbers with separators."""
        return (self.country, self.year_label, f"{self.fires:,}",
                f"{self.minimum:,.2f}", f"{self.maximum:,.2f}", f"{self.total:,.2f}",
                f"{self.dated:,}", share_label(self.dated, self.fires))


def combine(rows: list[Row], country: str = COUNTRY_NAME, year: int | None = None) -> Row:
    """One row summarising several: every figure taken over all of them.

    Notes
    -----
    This is what makes measuring a year at a time cost nothing. All five decompose
    over a partition of the fires — a minimum of minima is a minimum, a sum of sums is
    a sum, a count of counts is a count — so the ``Total`` row is the number one pass
    over the same rows would have returned, and no fire is counted twice or left out.

    The percentage is deliberately **not** averaged: :class:`Row` recomputes it from
    the summed counts, which is the ratio of the totals rather than the mean of the
    ratios. Fifty years of 20% and twenty years of 99% must not average to something
    neither of them is.

    ``fsum`` rather than ``sum``: an exact accumulation over a handful of partial
    totals costs nothing and cannot drift from what one pass would have returned.
    """
    return Row(
        country=country,
        year=year,
        minimum=min(row.minimum for row in rows),
        maximum=max(row.maximum for row in rows),
        total=math.fsum(row.total for row in rows),
        fires=sum(row.fires for row in rows),
        dated=sum(row.dated for row in rows),
    )


def summarise(measured: list[Row]) -> list[Row]:
    """Build the report from the years measured: the years newest first, then the total.

    Parameters
    ----------
    measured : list of Row
        One row per year, as the per-year statements returned them.

    Returns
    -------
    list of Row
        The years newest first with the summary row last. Empty if nothing was
        measured — a report of no fires has no total either.

    Notes
    -----
    No grouping by country and no lookup of one, unlike the four worldwide reports:
    there is exactly one country here and it is a constant, so ordering the countries
    would be sorting a list of one against the database's collation.
    """
    if not measured:
        return []
    rows = sorted(measured, key=lambda row: row.year, reverse=True)
    return rows + [combine(rows)]


def hectares(text: str) -> float:
    """Argparse type for ``--min-area``: a finite, non-negative number of hectares.

    Raises
    ------
    argparse.ArgumentTypeError
        If the text is not a number, or is negative, or is a non-finite float.

    Notes
    -----
    A bare ``type=float`` would accept ``-5``, ``nan`` and ``inf``. The first two are
    almost certainly a typo and would silently produce the unfiltered report — ``nan``
    compares false against every area, so it would produce an empty one — and none of
    the three is a size a fire can have.
    """
    try:
        area = float(text)
    except ValueError:
        raise argparse.ArgumentTypeError(f"{text!r} is not a number of hectares")
    if not math.isfinite(area):
        raise argparse.ArgumentTypeError(f"{text!r} is not a finite number of hectares")
    if area < 0:
        raise argparse.ArgumentTypeError(
            f"a burnt area cannot be negative, and {area:g} ha is: pass 0 or more, or "
            f"leave --min-area out to count every fire"
        )
    return area


def parse_arguments(argv: list[str] | None = None) -> argparse.Namespace:
    """Parse the command line."""
    parser = argparse.ArgumentParser(
        description="Burnt-area statistics for the Canadian National Burned Area "
                    "Composite, with the number of fires that carry a real published "
                    "date.",
        epilog="Areas are geodesic on the WGS84 ellipsoid by default, in hectares. Every "
               "perimeter is Canadian, so there is no --country and nothing is tested "
               "against a boundary. Database settings not given here are read from the "
               "environment (.env).",
    )
    parser.add_argument("-y", "--year", type=int,
                        help="restrict to one year, e.g. 2023; this is the published "
                             "YEAR, not the year of the resolved start date")
    parser.add_argument("--surface", default=SURFACE_MEASURED, choices=SURFACES,
                        help="what to report: 'measured' (default) is the area of the "
                             "dissolved perimeter; 'published' is POLY_HA as the service "
                             "computes it, which is the same quantity to within "
                             "0.0000005%%; 'adjusted' is ADJ_HA, a model output on the "
                             "2,512 fires flagged ADJ_FLAG and a copy of POLY_HA on the "
                             "rest, and is not interchangeable with either")
    parser.add_argument("--area-method", default=None, choices=AREA_METHODS,
                        help="how to turn the EPSG:4326 perimeter into hectares: "
                             "'geodesic' measures on the WGS84 ellipsoid (default); "
                             "'equal-area' projects to EPSG:6933 and measures there. They "
                             "agree to within 0.000004%%. Applies to --surface measured "
                             "and to nothing else; the published EPSG:3978 grid is not "
                             "offered, being a conformal conic — measuring there "
                             "understates this archive by 4.2%%")
    parser.add_argument("--min-area", type=hectares, default=None, metavar="HECTARES",
                        help="count only fires that burnt at least this many hectares of "
                             "the chosen surface; by default every fire counts. 200 is "
                             "the service's own large-fire threshold")
    parser.add_argument("--cause", default=None, choices=sorted(CAUSES),
                        help="restrict to one FIRECAUS value. 'natural' is the nearest "
                             "thing this dataset has to a lightning category and is not "
                             "one — the metadata glosses it 'most often lightning'. The "
                             "cause is unevenly reported across the archive: 1976 and "
                             "1977 are 80%% undetermined, 2017 and 2018 barely 2%%")
    parser.add_argument("--include-prescribed", action="store_true",
                        help="count the 400 fires flagged PRESCRIBED — a deliberate burn "
                             "rather than a wildfire. None before 1980, 143 in the 2000s. "
                             "Left out by default")

    # Accepted only so that they can be refused clearly. Anyone reaching for one has
    # copied a command line from one of the other seven reports, which is a reasonable
    # thing to have done, and argparse's own message would not say why this report is
    # different.
    parser.add_argument("--country", help=argparse.SUPPRESS)
    parser.add_argument("--country-source", help=argparse.SUPPRESS)

    output = parser.add_argument_group("output", "at least one is required")
    output.add_argument("--csv", type=Path, help="write the report to this .csv")
    output.add_argument("--docx", type=Path, help="write the report to this .docx (MS Word)")

    common.add_database_arguments(parser)
    parser.add_argument("--log-level", default=os.getenv("GISFIRE_LOG_LEVEL", "INFO"),
                        choices=["CRITICAL", "ERROR", "WARNING", "INFO", "DEBUG"],
                        help="verbosity (env: GISFIRE_LOG_LEVEL, default INFO)")

    arguments = parser.parse_args(argv)
    if arguments.country is not None:
        parser.error(
            f"there is no --country here: Natural Resources Canada maps the fires of "
            f"{COUNTRY_NAME} and nothing else, so there is nothing to select between. "
            f"Every fire is counted and the Country column is {COUNTRY_NAME} on every "
            f"row."
        )
    if arguments.country_source is not None:
        parser.error(
            f"there is no --country-source here: these are {COUNTRY_NAME}'s own "
            f"perimeters of its own territory, and the import already resolved a "
            f"boundary for all but two of them — every one of them {COUNTRY_NAME}. The "
            f"companion points report does test its coordinates, and needs to."
        )
    if arguments.area_method is not None and arguments.surface != SURFACE_MEASURED:
        parser.error(
            f"--area-method applies to --surface {SURFACE_MEASURED} and to nothing else: "
            f"the {arguments.surface} hectares are read off the service's own figures, "
            f"not measured, so no CRS is involved in producing them."
        )
    if arguments.csv is None and arguments.docx is None:
        parser.error("nothing to write: pass --csv, --docx, or both")
    return arguments


def compute(session: Session, year: int | None, logger: logging.Logger,
            surface: str = SURFACE_MEASURED,
            method: str = AREA_METHOD_GEODESIC,
            min_area: float | None = None,
            include_prescribed: bool = False,
            cause: str | None = None) -> list[Row]:
    """Measure the fires a year at a time, returning the report's rows in order.

    Notes
    -----
    One statement per year, for the reason given in the module docstring, each under a
    spinner of its own so that a long run says which year it is on rather than only
    that it is alive.

    Every one of them runs in ``session``'s transaction, and so against a single
    snapshot: a report assembled from many queries is then exactly as consistent as
    one assembled from a single query, and an import running alongside it cannot have
    a fire counted in one year's statement and not in another's.

    The dated fires are logged as well as reported, because the shape of that number
    is one of the first things anyone should know about this dataset: a ``Dated``
    column at a quarter in the 1970s is not a broken import, it is an archive whose
    early years were compiled from imagery long after the fact.
    """
    if year is not None:
        years = [year]
    else:
        with common.Spinner(f"Finding the years the {COUNTRY_NAME} perimeters cover",
                            logger):
            years = list(session.scalars(years_query(include_prescribed, cause)))

    scope = "every cause" if cause is None else f"{cause} fires"
    measured: list[Row] = []
    for index, measuring in enumerate(years, start=1):
        with common.Spinner(f"Measuring the burnt area of the {COUNTRY_NAME} perimeters "
                            f"({surface}, {scope}, {measuring}: {index} of {len(years)})",
                            logger):
            measured += [
                Row(country=record.country, year=measuring,
                    minimum=float(record.minimum),
                    maximum=float(record.maximum),
                    total=float(record.total),
                    fires=record.fires,
                    dated=record.dated)
                for record in session.execute(
                    statistics_query(measuring, surface, method, min_area,
                                     include_prescribed, cause))
            ]

    excluded = 0 if include_prescribed else prescribed_count(session, year, cause)

    rows = summarise(measured)
    logger.info("Computed %d rows over %d year(s) (%s hectares%s, %s, %s)",
                len(rows), len(measured), surface,
                f", {method}" if surface == SURFACE_MEASURED else "",
                "every fire" if min_area is None else f"fires of {min_area:g} ha or more",
                scope)
    if excluded:
        logger.info("Excluded %d prescribed burn(s) in scope; pass --include-prescribed "
                    "to count them", excluded)
    if rows:
        total = rows[-1]
        logger.info("%d of %d fire(s) carry a published date (%s%%); the rest are dated "
                    "to 1 January of their year",
                    total.dated, total.fires, share_label(total.dated, total.fires))
    return rows


def write_csv(rows: list[Row], path: Path, logger: logging.Logger) -> None:
    """Write the report as CSV.

    The numbers go out unformatted apart from being rounded to two decimals — no
    thousands separators — because a CSV is read by another program far more often
    than by a person, and a separator would make every figure a string.
    """
    path.parent.mkdir(parents=True, exist_ok=True)
    with path.open("w", newline="", encoding="utf-8") as handle:
        writer = csv.writer(handle)
        writer.writerow(COLUMNS)
        for row in rows:
            writer.writerow(row.values)
    logger.info("Wrote %s", path)


def write_docx(rows: list[Row], path: Path, year: int | None,
               logger: logging.Logger,
               surface: str = SURFACE_MEASURED,
               method: str = AREA_METHOD_GEODESIC,
               min_area: float | None = None,
               include_prescribed: bool = False,
               cause: str | None = None) -> None:
    """Write the report as a Word document.

    One table, with the summary row in bold. Numbers get thousands separators here —
    the opposite of the CSV, and for the opposite reason: this one is for reading.

    The opening paragraphs name the surface, the scope, what ``Dated`` is and what
    happened to the prescribed burns, because a table of adjusted hectares that looks
    like a table of mapped ones, a column that is a property of the archive rather
    than of the import, and a silently dropped set of deliberate fires are all things
    a reader should not have to remember.
    """
    # Imported here rather than at module scope so that --csv keeps working if
    # python-docx is not installed, which matters because it is the only dependency
    # this application adds.
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt

    document = Document()
    document.add_heading(f"NBAC wildfire burnt area ({COUNTRY_NAME})", level=1)

    scope = [f"year: {year}" if year is not None else "all years"]
    if cause is not None:
        scope.append(f"only {cause} fires")
    if min_area is not None:
        scope.append(f"only fires of {min_area:g} ha or more")
    scope.append("prescribed burns counted" if include_prescribed
                 else "prescribed burns excluded")
    measurement = (f", measured {method}" if surface == SURFACE_MEASURED else "")
    document.add_paragraph(
        f"Areas in hectares. Surface: {SURFACE_PROSE[surface]}{measurement}. A fire is "
        f"one published GID and not one polygon — the features are cut at provincial, "
        f"territorial and park boundaries, and the import dissolved them. Years are the "
        f"published YEAR, not the resolved start date's. Scope: {'; '.join(scope)}."
    )
    document.add_paragraph(
        f"The Country column is {COUNTRY_NAME} on every row and nothing is tested "
        f"against a boundary: Natural Resources Canada maps the fires of "
        f"{COUNTRY_NAME} and nothing else."
    )
    document.add_paragraph(
        "Dated counts the fires whose start came from a published date — the agency's "
        "AG_SDATE or the first satellite hotspot HS_SDATE — rather than from the year "
        "alone. It is a column and not a filter: an undated fire still contributes its "
        "hectares. Nearly a fifth of the archive is undated and it is not evenly "
        "spread, so any analysis finer than a year has to filter on the precision first."
    )
    if cause is not None:
        # Not a footnote: 'natural' is the reason most readers will reach for this
        # option, and it is not the thing they will assume it is.
        document.add_paragraph(
            f"These are the fires whose published FIRECAUS is {CAUSES[cause]}. Natural "
            f"is the nearest thing this dataset has to a lightning category and is not "
            f"one — the published metadata glosses it 'ignition source by natural cause, "
            f"most often lightning'. The cause is also unevenly reported: 1976 and 1977 "
            f"are 80% Undetermined and 2017 and 2018 barely 2%, so a trend across the "
            f"archive is partly a trend in reporting."
        )

    table = document.add_table(rows=1, cols=len(COLUMNS))
    table.style = "Table Grid"
    for cell, heading in zip(table.rows[0].cells, COLUMNS):
        cell.text = heading
        cell.paragraphs[0].runs[0].bold = True

    for row in rows:
        cells = table.add_row().cells
        for index, (cell, value) in enumerate(zip(cells, row.readable_values)):
            cell.text = value
            paragraph = cell.paragraphs[0]
            if index >= FIRST_NUMERIC_COLUMN:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            for run in paragraph.runs:
                run.bold = row.is_total
                run.font.size = Pt(9)

    path.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(path))
    logger.info("Wrote %s", path)


def report(args: argparse.Namespace, engine: Engine, logger: logging.Logger) -> list[Row]:
    """Compute the statistics and write whichever outputs were asked for."""
    # No spinner here: compute runs one statement per year and turns one of its own
    # for each, which is the only honest place to say how far along it is.
    method = args.area_method or AREA_METHOD_GEODESIC
    with Session(engine) as session:
        rows = compute(session, args.year, logger, args.surface, method, args.min_area,
                       args.include_prescribed, args.cause)

    if not rows:
        # An empty report is almost always a year with no data, and writing an empty
        # file would hide that. The narrowing options are named when they are in force,
        # because then each is at least as likely to be the reason as the year is.
        extra = "" if args.min_area is None else \
            f" No fire reached the --min-area of {args.min_area:g} ha."
        if args.cause is not None:
            extra += f" The report is of {args.cause} fires alone."
        raise RuntimeError(
            f"No wildfires matched. Check --year, and that the {COUNTRY_NAME} "
            f"perimeters are imported — the published archive runs from "
            f"{canada_nbac.FIRST_YEAR}." + extra
        )

    if args.csv is not None:
        write_csv(rows, args.csv, logger)
    if args.docx is not None:
        write_docx(rows, args.docx, args.year, logger, args.surface, method,
                   args.min_area, args.include_prescribed, args.cause)
    return rows


def main(argv: list[str] | None = None) -> int:
    args = parse_arguments(argv)
    logging.basicConfig(level=args.log_level, format="%(asctime)s %(levelname)s %(message)s")
    logger = logging.getLogger("nbac-statistics")

    try:
        settings = common.resolve_database_settings(args)
    except RuntimeError as error:
        logger.error("%s", error)
        return 1

    engine = create_engine(common.database_url(settings))
    try:
        report(args, engine, logger)
    except Exception as error:  # noqa: BLE001  (the CLI boundary: report, do not traceback)
        logger.error("%s", error)
        return 1
    finally:
        engine.dispose()
    return 0


if __name__ == "__main__":  # pragma nocover
    sys.exit(main())
