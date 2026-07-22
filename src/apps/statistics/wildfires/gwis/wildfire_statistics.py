#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Burnt-area statistics for the GWIS GlobFire wildfires.

Reports, per country and year, the smallest and largest single fire and the
total area burnt, in hectares::

    Country          Year     Minimum      Maximum        Total
    Spain            2021       25.34     12904.11    481203.77
    Spain            2020       25.34      8110.02    377411.20
    Spain            Total      25.34     12904.11    858614.97

    France           2021       25.31      6002.45    120884.10
    ...

The ``Total`` row of a country is the smallest and largest fire it has *ever*
had, and the sum of every year — not a total of the column above it, which for
the minimum and maximum would mean nothing.

Run it over everything, or narrow it to one country, one year, or both::

    python3 -m src.apps.statistics.wildfires.gwis.wildfire_statistics --csv burnt.csv
    python3 -m src.apps.statistics.wildfires.gwis.wildfire_statistics \\
        --country Spain --year 2021 --csv spain_2021.csv --docx spain_2021.docx

At least one of ``--csv`` and ``--docx`` is required: an application that
computed a report and then printed nothing would be a strange thing.

The application only reads. Database settings come from the environment
(``.env``, see :mod:`src.settings`); every one of them can be overridden with a
command-line argument.

How the area is measured
------------------------

``ST_Area(perimeter::geography)``: the true area on the WGS84 ellipsoid, in
square metres, divided by 10,000 for hectares.

Deliberately *not* a projected area. Every map projection distorts something,
and for a dataset spanning every latitude from the Arctic to Tasmania there is
no projection whose distortion is negligible everywhere — a fire measured in Web
Mercator at 70°N would come out roughly nine times too large. An equal-area
projection would fix that, but the geodesic calculation is at least as accurate
and needs no CRS to be chosen and justified in the first place.

Which fires are counted
-----------------------

Fires with no country are **excluded** — mid-ocean perimeters, and any fire
whose interior point matched no OCHA boundary. So is any fire with no perimeter.
A run therefore does not necessarily account for every hectare in the database,
and the totals here are totals *of attributable burnt area*.

Which year a fire counts towards
--------------------------------

The year of its **local** start date, ``start_date_time AT TIME ZONE
time_zone`` — that is, the year of the ``IDate`` GWIS published, and so the year
of the file the fire came from. Using the raw UTC instant instead would move
fires across the New Year boundary: a fire starting on 1 January in Sydney is
still 31 December in UTC.
"""

from __future__ import annotations

import argparse
import csv
import logging
import os
import sys

from dataclasses import dataclass
from pathlib import Path

from geoalchemy2 import Geography
from sqlalchemy import Engine
from sqlalchemy import Integer
from sqlalchemy import Select
from sqlalchemy import cast
from sqlalchemy import create_engine
from sqlalchemy import func
from sqlalchemy import or_
from sqlalchemy import select
from sqlalchemy import tuple_
from sqlalchemy.orm import Session

import src.settings  # noqa: F401  (imported for the side effect of loading .env)

from src.apps.imports import common
from src.data_model.geography.admin_boundary import AdminBoundary
from src.data_model.wildfire import Wildfire
from src.providers.gwis.wildfire import GwisWildfire
from src.providers.ocha.admin_boundary import OchaAdminBoundary

#: Label used in the ``Year`` column for a country's summary row.
TOTAL_LABEL = "Total"

#: Square metres in a hectare.
SQUARE_METRES_PER_HECTARE = 10_000.0

#: Burnt area of one fire, in hectares.
#:
#: The cast carries the geometry type and SRID rather than being a bare
#: ``Geography()``: that renders as ``geography(GEOMETRY,-1)``, which PostGIS
#: rejects, because -1 is not a SRID it knows.
BURNT_AREA = (
    func.ST_Area(cast(Wildfire.perimeter, Geography(geometry_type="MULTIPOLYGON", srid=4326)))
    / SQUARE_METRES_PER_HECTARE
)

#: The year a fire counts towards: the year of its *local* start date.
#:
#: ``AT TIME ZONE`` has no construct of its own in SQLAlchemy, so it is applied
#: as an operator. ``EXTRACT`` yields a numeric, hence the cast — without it the
#: year comes back as ``Decimal`` and lands in the report as ``2021.0``.
LOCAL_YEAR = cast(
    func.extract(
        "year",
        Wildfire.start_date_time.op("AT TIME ZONE")(func.coalesce(Wildfire.time_zone, "UTC")),
    ),
    Integer,
)


def statistics_query(country: str | None, year: int | None) -> Select:
    """Build the statistics query.

    Returns
    -------
    Select
        A query yielding ``country, year, is_total, minimum, maximum, total,
        fires``, ordered country by country with each country's years newest
        first and its summary row last.

    Notes
    -----
    Built against the mapped classes rather than written as SQL text, so a
    column renamed on a model breaks this at import time rather than in front of
    a user. It also lets the two filters be plain conditionals: written as text
    they would each have to become an "unset, or matching" disjunction so that
    one statement could serve every combination, leaving branches in the SQL
    that are dead on every actual run.

    The inner query computes each area exactly once. Folded into the outer
    aggregate instead, ``ST_Area`` would be evaluated three times per row — for
    the minimum, the maximum and the sum — and it is by far the most expensive
    thing here.

    ``GROUPING SETS`` then produces the per-year rows and the per-country
    summary row from that one pass, instead of aggregating twice or totalling in
    Python. ``GROUPING(year)`` is 0 for a real year and 1 for a summary row,
    which is what both sorts the two apart and tells them apart on the way out.

    The join to ``admin_boundary`` is inner, and that is what drops the fires
    with no country. ``gwis_wildfire`` is joined — by table, to keep SQLAlchemy
    from adding the polymorphic join of its own — rather than filtering on
    ``wildfire.type``, so this stays a GWIS report even if another provider ever
    adopts the same discriminator.
    """
    gwis = GwisWildfire.__table__
    ocha_boundary = OchaAdminBoundary.__table__

    fires = (
        select(
            AdminBoundary.name.label("country"),
            LOCAL_YEAR.label("year"),
            BURNT_AREA.label("hectares"),
        )
        .select_from(Wildfire)
        .join(gwis, gwis.c.id == Wildfire.id)
        .join(AdminBoundary, AdminBoundary.id == Wildfire.admin_boundary_id)
        .outerjoin(ocha_boundary, ocha_boundary.c.id == AdminBoundary.id)
        .where(Wildfire.perimeter.is_not(None))
    )
    if country is not None:
        fires = fires.where(or_(
            func.lower(AdminBoundary.name) == country.lower(),
            func.upper(ocha_boundary.c.iso_3) == country.upper(),
        ))
    if year is not None:
        fires = fires.where(LOCAL_YEAR == year)

    fire = fires.subquery("fire")
    grouping = func.grouping(fire.c.year)
    return (
        select(
            fire.c.country,
            fire.c.year,
            grouping.label("is_total"),
            func.min(fire.c.hectares).label("minimum"),
            func.max(fire.c.hectares).label("maximum"),
            func.sum(fire.c.hectares).label("total"),
            func.count().label("fires"),
        )
        .group_by(func.grouping_sets(
            tuple_(fire.c.country, fire.c.year),
            tuple_(fire.c.country),
        ))
        .order_by(fire.c.country, grouping, fire.c.year.desc())
    )


@dataclass(frozen=True)
class Row:
    """One line of the report.

    Attributes
    ----------
    country : str
        Name of the country the fires burnt in.
    year : int or None
        The year, or ``None`` for a country's summary row.
    minimum, maximum, total : float
        Smallest single fire, largest single fire and sum of every fire, in
        hectares.
    fires : int
        How many fires the three figures were computed from. Not printed — the
        report keeps to the columns asked for — but carried so that a caller
        using this module as a library does not have to query again for it.
    """

    country: str
    year: int | None
    minimum: float
    maximum: float
    total: float
    fires: int

    @property
    def is_total(self) -> bool:
        """Whether this is a country's summary row rather than one of its years."""
        return self.year is None

    @property
    def year_label(self) -> str:
        return TOTAL_LABEL if self.is_total else str(self.year)


def parse_arguments(argv: list[str] | None = None) -> argparse.Namespace:
    """Parse the command line."""
    parser = argparse.ArgumentParser(
        description="Burnt-area statistics for the GWIS GlobFire wildfires.",
        epilog="Areas are geodesic, on the WGS84 ellipsoid, in hectares. Fires with no "
               "country are not counted. Database settings not given here are read from "
               "the environment (.env).",
    )
    selection = parser.add_argument_group("selection", "report on everything unless narrowed")
    selection.add_argument("-c", "--country",
                           help="restrict to one country, by name ('Spain') or ISO 3166-1 "
                                "alpha-3 code ('ESP'); case-insensitive")
    selection.add_argument("-y", "--year", type=int, help="restrict to one year, e.g. 2021")

    output = parser.add_argument_group("output", "at least one is required")
    output.add_argument("--csv", type=Path, help="write the report to this .csv")
    output.add_argument("--docx", type=Path, help="write the report to this .docx (MS Word)")

    common.add_database_arguments(parser)
    parser.add_argument("--log-level", default=os.getenv("GISFIRE_LOG_LEVEL", "INFO"),
                        choices=["CRITICAL", "ERROR", "WARNING", "INFO", "DEBUG"],
                        help="verbosity (env: GISFIRE_LOG_LEVEL, default INFO)")

    arguments = parser.parse_args(argv)
    if arguments.csv is None and arguments.docx is None:
        parser.error("nothing to write: pass --csv, --docx, or both")
    return arguments


def compute(session: Session, country: str | None, year: int | None,
            logger: logging.Logger) -> list[Row]:
    """Run the statistics query, returning the report's rows in order."""
    result = session.execute(statistics_query(country, year))
    rows = [
        Row(country=record.country,
            year=None if record.is_total else record.year,
            minimum=float(record.minimum),
            maximum=float(record.maximum),
            total=float(record.total),
            fires=record.fires)
        for record in result
    ]
    countries = len({row.country for row in rows})
    logger.info("Computed %d rows over %d countries", len(rows), countries)
    return rows


def write_csv(rows: list[Row], path: Path, logger: logging.Logger) -> None:
    """Write the report as CSV.

    The numbers go out unformatted apart from being rounded to two decimals —
    no thousands separators — because a CSV is read by another program far more
    often than by a person, and a separator would make every figure a string.
    """
    path.parent.mkdir(parents=True, exist_ok=True)
    with path.open("w", newline="", encoding="utf-8") as handle:
        writer = csv.writer(handle)
        writer.writerow(["Country", "Year", "Minimum (ha)", "Maximum (ha)", "Total (ha)"])
        for row in rows:
            writer.writerow([row.country, row.year_label,
                             f"{row.minimum:.2f}", f"{row.maximum:.2f}", f"{row.total:.2f}"])
    logger.info("Wrote %s", path)


def write_docx(rows: list[Row], path: Path, country: str | None, year: int | None,
               logger: logging.Logger) -> None:
    """Write the report as a Word document.

    One table, with each country's summary row in bold so the blocks read apart
    the way they would on paper. Numbers get thousands separators here — the
    opposite of the CSV, and for the opposite reason: this one is for reading.
    """
    # Imported here rather than at module scope so that --csv keeps working if
    # python-docx is not installed, which matters because it is the only
    # dependency this application adds.
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt

    document = Document()
    document.add_heading("GWIS wildfire burnt area", level=1)

    scope = []
    if country is not None:
        scope.append(f"country: {country}")
    if year is not None:
        scope.append(f"year: {year}")
    subtitle = "; ".join(scope) if scope else "all countries, all years"
    document.add_paragraph(
        f"Areas in hectares, computed geodesically on the WGS84 ellipsoid. "
        f"Fires not attributable to a country are excluded. Scope: {subtitle}."
    )

    table = document.add_table(rows=1, cols=5)
    table.style = "Table Grid"
    for cell, heading in zip(table.rows[0].cells,
                             ["Country", "Year", "Minimum (ha)", "Maximum (ha)", "Total (ha)"]):
        cell.text = heading
        cell.paragraphs[0].runs[0].bold = True

    for row in rows:
        cells = table.add_row().cells
        values = [row.country, row.year_label,
                  f"{row.minimum:,.2f}", f"{row.maximum:,.2f}", f"{row.total:,.2f}"]
        for index, (cell, value) in enumerate(zip(cells, values)):
            cell.text = value
            paragraph = cell.paragraphs[0]
            if index >= 2:  # the three numeric columns
                paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            for run in paragraph.runs:
                run.bold = row.is_total
                run.font.size = Pt(9)

    path.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(path))
    logger.info("Wrote %s", path)


def report(args: argparse.Namespace, engine: Engine, logger: logging.Logger) -> list[Row]:
    """Compute the statistics and write whichever outputs were asked for."""
    with Session(engine) as session:
        rows = compute(session, args.country, args.year, logger)

    if not rows:
        # An empty report is almost always a mistyped country or a year with no
        # data, and writing an empty file would hide that.
        raise RuntimeError(
            "No wildfires matched. Check --country (a name or an ISO alpha-3 code) and "
            "--year, and that the GWIS fires and the OCHA boundaries are both imported "
            "— fires with no country are not counted."
        )

    if args.csv is not None:
        write_csv(rows, args.csv, logger)
    if args.docx is not None:
        write_docx(rows, args.docx, args.country, args.year, logger)
    return rows


def main(argv: list[str] | None = None) -> int:
    args = parse_arguments(argv)
    logging.basicConfig(level=args.log_level, format="%(asctime)s %(levelname)s %(message)s")
    logger = logging.getLogger("gwis-statistics")

    try:
        settings = common.resolve_database_settings(args)
    except RuntimeError as error:
        logger.error("%s", error)
        return 1

    engine = create_engine(common.database_url(settings))
    try:
        report(args, engine, logger)
    except Exception as error:  # noqa: BLE001  (the CLI boundary: report, do not traceback)
        logger.error("%s", error)
        return 1
    finally:
        engine.dispose()
    return 0


if __name__ == "__main__":  # pragma nocover
    sys.exit(main())
