#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Wildfire statistics for the Guatemalan INAB fire reports.

Reports the archive per year: how many fires INAB was told about, how many of those
reports were false alarms, how many carry a coordinate and how many fell inside a
protected area::

    python3 -m src.apps.statistics.wildfires.guatemala_inab.wildfire_statistics \\
        --csv guatemala.csv --docx guatemala.docx

The eighth of the burnt-area reports, alongside :doc:`GWIS <gwis_wildfire_statistics>`,
GFA, ICNF, EGIF, DARPA, REDIAM and Greece — and the first with **no hectares in it at
all**.

There is no burnt area, and the columns are empty rather than absent
---------------------------------------------------------------------

This is the whole shape of this report, so it is the first thing in it.

:mod:`src.providers.spain_egif` and :mod:`src.providers.greece_ffa` publish no
perimeter but do publish burnt areas — five figures and eight respectively — so their
reports sum hectares. **INAB publishes neither.** There is no perimeter on any row,
no ``area_ha`` on the model to sum, no land-cover split, and nothing in the
thirty-three published attributes that says how big a fire was. There is therefore no
``--area-method`` and no ``--surface``: there is nothing to measure and nothing to
choose between.

``Minimum (ha)``, ``Maximum (ha)`` and ``Total (ha)`` are nevertheless in this
report, in the position and the order the other seven put them, and **every one of
their cells is empty**:

* an empty cell is a claim that nothing was published, which is true. A zero would be
  a claim that nothing burnt, which is not — the 4,611 fires here burnt an unknown
  amount;
* the CSV still concatenates with the other seven on
  :data:`SHARED_COLUMNS`, so a reader comparing the eight countries sees Guatemala's
  gap in the table instead of having to notice its absence from it.

What this dataset answers is *where and when*, and the four columns after the empty
three are what it answers with.

Which year a fire counts towards
---------------------------------

The **Guatemalan calendar year** of the fire's own instant, resolved through
:data:`~src.providers.guatemala_inab.DEFAULT_TIME_ZONE`.

That is a departure from the rule the other reports follow, and it is forced: the ICNF
report groups on the published ``Ano``, the EGIF one on the filed ``Campania`` and the
Greek one on the sheet the row came from, because those sources publish a year.
**INAB publishes none.** ``fecha_hora_incendio`` is the only thing in the record that
says when, so the year has to be derived from it, and the local year is the one that
means something — a fire reported at nine in the evening on 31 December is a fire of
that year and not of the next.

It is the same arithmetic the import uses to decide which year a record replaces, so
a year in this report is exactly a year in
:mod:`src.apps.imports.wildfires.guatemala_inab.import_wildfires`. Both are six hours
away from the UTC year, which is what the ArcGIS server's own ``EXTRACT`` may be
counting in.

False alarms are excluded, and counted
---------------------------------------

140 of the 4,615 published records are
:data:`~src.providers.guatemala_inab.STATUS_FALSE` — the report was false, there was
no fire. They are records of a *call*, not of a fire, so ``Fires`` leaves them out, as
the Greek report leaves out ``ΨΕΥΔΗΣ ΑΝΑΓΓΕΛΙΑ``.

Unlike that report, this one gives them a column of their own rather than a line in
the log. They are 3% of the archive and they are not evenly spread, and a reader who
cannot see them in the table cannot tell a quiet year from a well-checked one.
``False alarms`` counts them whether or not ``--include-false-alarms`` puts them into
``Fires``; the Word document says which.

.. warning::

   The test is ``IS DISTINCT FROM`` and not ``<>``. ``report_status`` is ``NULL`` on
   the records that carry no attributes at all, and ``<>`` evaluates to ``NULL``
   there, so the obvious filter would silently drop them from the count.

Every fire is located, which is worth a column for the opposite reason
------------------------------------------------------------------------

``Located`` is in this report because it is in the Greek one, and it reads the other
way round: **all 4,615 published records carry an EPSG:4326 point**, so the column is
100% throughout rather than zero for twenty years and then 94%.

That makes this the best-located administrative fire statistic in the project, and it
is worth stating in the table rather than in a footnote. A column that stops being
100% is then a fact about a future publication, visible immediately.

``In protected area`` is Guatemala's own column: the fires whose point fell inside one
— 1,455 of them, close to one in three, which is a real property of Guatemalan fire
rather than a gap in the data. It counts a fire when ``nombre_ap_1`` is filled, which
depends on the import having folded this source's ``""`` to ``NULL``; see
:func:`~src.providers.guatemala_inab.blank_to_none`.

There is no country to choose, and no cause to count
------------------------------------------------------

No ``--country`` and no ``--country-source``, as for the Greek, Catalan and Andalusian
reports: INAB publishes Guatemala's fires and nothing else, so the ``Country`` column
is the constant :data:`COUNTRY_NAME` and nothing is tested against a boundary to
arrive at it. Three of the published points are not in Guatemala and all three are
already flagged ``falso``; they leave the count with the other false alarms.

And there is no counts-by-cause companion, as there is for Portugal, Spain and Canada:
**nothing in the thirty-three published attributes says why a fire started.** What
this dataset does carry is four published vocabularies describing the fire and the
report, and :mod:`the classification report
<src.apps.statistics.wildfires.guatemala_inab.wildfire_classification>` is what counts
those.

One statement
-------------

Like the EGIF, Greek, Catalan and Andalusian reports and unlike the GWIS and GFA ones,
this is a single aggregate: 4,611 rows and no geometry work at all, which is four
orders of magnitude short of the twenty-million-perimeter case that made the others
read a year at a time. The ``Total`` row is arithmetic over the years, by
:func:`combine`, so the output is the same shape either way.
"""

from __future__ import annotations

import argparse
import csv
import logging
import os
import sys

from dataclasses import dataclass
from pathlib import Path

from sqlalchemy import ColumnElement
from sqlalchemy import Engine
from sqlalchemy import Integer
from sqlalchemy import Select
from sqlalchemy import cast
from sqlalchemy import create_engine
from sqlalchemy import func
from sqlalchemy import literal
from sqlalchemy import select
from sqlalchemy import true
from sqlalchemy.orm import Session

import src.settings  # noqa: F401  (imported for the side effect of loading .env)

from src.apps.imports import common
from src.data_model.wildfire import Wildfire
from src.providers import guatemala_inab
from src.providers.guatemala_inab.wildfire import InabWildfire

#: Label used in the ``Year`` column for the summary row.
TOTAL_LABEL = "Total"

#: The six columns this report shares with the GWIS, GFA, ICNF, EGIF, DARPA, REDIAM
#: and Greek ones, in their order, so the CSVs can still be concatenated on them.
#:
#: The three hectare columns are here and are empty on every row — see the module
#: docstring. They are not dropped, because a gap a reader can see beats an absence a
#: reader has to notice.
SHARED_COLUMNS = ("Country", "Year", "Fires", "Minimum (ha)", "Maximum (ha)",
                  "Total (ha)")

#: The report's columns, in order, shared by both output formats so that a change to
#: one cannot silently leave the other behind. The last four are this dataset's own.
COLUMNS = SHARED_COLUMNS + ("False alarms", "Located", "Located (%)",
                            "In protected area")

#: Index of the first column that holds a number, and so is right-aligned in the Word
#: table.
FIRST_NUMERIC_COLUMN = 2

#: What a hectare cell contains. Empty, on every row, for every year.
#:
#: A constant rather than a bare ``""`` at three call sites, so that the reason is
#: attached to the value: nothing was published, and a zero would say something else
#: entirely. See the module docstring.
NO_AREA = ""

#: The country every fire in this dataset is in, and the whole of the ``Country``
#: column. Spelled as the OCHA boundaries spell it, so a row of this report sorts and
#: groups with the rows of the other seven. Nothing is tested against a boundary to
#: arrive at it — see the module docstring.
COUNTRY_NAME = "Guatemala"

#: The year a fire counts towards: the Guatemalan calendar year of its own instant.
#:
#: Derived and not published — this source has no year field of any kind, unlike every
#: other dataset with a report in this project. See the module docstring.
#:
#: ``AT TIME ZONE`` and not ``EXTRACT(YEAR FROM start_date_time)``: the column is
#: ``timestamptz``, so the bare extract would give the year in whatever the session's
#: ``TimeZone`` happens to be, and a report whose year boundaries moved with the
#: reader's locale would be a very hard thing to notice. Naming the zone makes the
#: figure the same everywhere, and the same as the import's.
LOCAL_YEAR = cast(
    func.extract("year",
                 func.timezone(guatemala_inab.DEFAULT_TIME_ZONE,
                               Wildfire.__table__.c.start_date_time)),
    Integer,
)


def is_located() -> ColumnElement:
    """Whether a fire carries a published coordinate.

    Returns
    -------
    ColumnElement
        A boolean expression, ``True`` for a fire the ``Located`` column counts.

    Notes
    -----
    The link and not the geometry is what is tested, as in the Greek report:
    ``ignition_id`` is written only for a record that published a point, so a row with
    a link has one and a row without had none. Joining through to the geometry would
    cost a join and answer the same question.

    True on every record published today. See the module docstring for why that is a
    column rather than a footnote.
    """
    return InabWildfire.__table__.c.ignition_id.is_not(None)


def is_in_protected_area() -> ColumnElement:
    """Whether the fire's point fell inside a protected area.

    Returns
    -------
    ColumnElement
        A boolean expression, ``True`` for a fire the ``In protected area`` column
        counts.

    Notes
    -----
    ``nombre_ap_1`` filled, which is the provider's own containment test rather than
    anything computed here — no Guatemalan protected-area boundaries are imported that
    a point could be tested against.

    This column is only correct because the import folded this source's ``""`` to
    ``NULL``: ``nombre_ap_1`` is ``null`` on 80 records and ``""`` on 3,080, so a
    database loaded without
    :func:`~src.providers.guatemala_inab.blank_to_none` would report 4,535 fires
    inside a protected area instead of 1,455. The report cannot detect that, which is
    why it is said here.
    """
    return InabWildfire.__table__.c.protected_area_name.is_not(None)


def is_a_false_alarm() -> ColumnElement:
    """Whether the report said there was no fire.

    Returns
    -------
    ColumnElement
        A boolean expression, ``True`` for the 140 records the ``False alarms`` column
        counts.

    Notes
    -----
    ``no_verificado`` is deliberately not folded in. It says nobody went to look,
    which is a different claim from *there was no fire*, and folding the two together
    would turn 90 unknowns into 90 non-events. See
    :func:`~src.providers.guatemala_inab.is_false_alarm`; the classification report
    counts both, separately.
    """
    return InabWildfire.__table__.c.report_status == guatemala_inab.STATUS_FALSE


def is_a_fire(include_false_alarms: bool = False) -> ColumnElement:
    """Whether a record counts towards ``Fires``.

    Parameters
    ----------
    include_false_alarms : bool
        ``True`` to count the false alarms as fires. ``False``, the default, leaves
        them out.

    Returns
    -------
    ColumnElement
        A boolean expression. Always an expression and never ``None``, unlike the
        Greek report's: this one is used as an aggregate ``FILTER`` rather than as a
        ``WHERE``, because the false alarms have to stay in the result set in order to
        be counted in their own column.

    Notes
    -----
    ``is_distinct_from`` and not ``!=``. ``report_status`` is ``NULL`` on the records
    that carry no attributes at all, where ``!=`` evaluates to ``NULL`` rather than to
    true, so the obvious filter would drop them from ``Fires`` while leaving them in
    every other column of their row.
    """
    if include_false_alarms:
        return true()
    return InabWildfire.__table__.c.report_status.is_distinct_from(
        guatemala_inab.STATUS_FALSE
    )


def statistics_query(year: int | None = None,
                     include_false_alarms: bool = False) -> Select:
    """Build the statistics query: one row per year, newest first.

    Parameters
    ----------
    year : int, optional
        Restrict to one year. ``None``, the default, reports every year.
    include_false_alarms : bool
        Count the false alarms as fires.

    Returns
    -------
    Select
        A query yielding ``country, year, fires, false_alarms, located, protected``.
        The summary row is :func:`summarise`'s work.

    Notes
    -----
    Built against the mapped classes rather than written as SQL text, so a column
    renamed on a model breaks this at import time rather than in front of a user.

    ``inab_wildfire`` is joined by table, to keep SQLAlchemy from adding a polymorphic
    join of its own, and it has to be joined in any case: the status, the protected
    area and the ignition link all live on it. The parent ``wildfire`` is the ``FROM``
    because a fire *is* the parent row and because the instant the year comes from is
    there.

    The false alarms are **not** filtered out in a ``WHERE``. Every count is a filtered
    aggregate over the same rows, which is what lets one pass produce both a ``Fires``
    that excludes them and a ``False alarms`` that counts them. It also means a year
    made up entirely of false alarms appears with ``Fires`` of zero rather than
    vanishing, which is the honest answer.

    ``Located`` and ``In protected area`` are filtered by the *same* condition as
    ``Fires``, which is what keeps ``Located`` a share of the ``Fires`` beside it
    rather than of some larger number.
    """
    inab = InabWildfire.__table__
    counted = is_a_fire(include_false_alarms)

    fires = (
        select(
            literal(COUNTRY_NAME).label("country"),
            LOCAL_YEAR.label("year"),
            counted.label("counted"),
            is_located().label("located"),
            is_in_protected_area().label("protected"),
            is_a_false_alarm().label("false_alarm"),
        )
        .select_from(Wildfire)
        .join(inab, inab.c.id == Wildfire.id)
    )
    if year is not None:
        fires = fires.where(LOCAL_YEAR == year)

    fire = fires.subquery("fire")
    return (
        select(
            fire.c.country,
            fire.c.year,
            func.count().filter(fire.c.counted).label("fires"),
            func.count().filter(fire.c.false_alarm).label("false_alarms"),
            func.count().filter(fire.c.counted & fire.c.located).label("located"),
            func.count().filter(fire.c.counted & fire.c.protected).label("protected"),
        )
        .group_by(fire.c.country, fire.c.year)
        .order_by(fire.c.year.desc())
    )


def share(part: int, whole: int) -> float | None:
    """``part`` as a percentage of ``whole``, or ``None`` where there is no whole.

    ``None`` and not zero: a percentage of nothing is not zero percent, it is no
    answer, and the writers turn it into an empty cell.
    """
    if not whole:
        return None
    return 100.0 * part / whole


def share_label(part: int, whole: int) -> str:
    """A percentage as it is written out, empty where there is none."""
    value = share(part, whole)
    return "" if value is None else f"{value:.2f}"


@dataclass(frozen=True)
class Row:
    """One line of the report.

    Attributes
    ----------
    country : str
        Always :data:`COUNTRY_NAME`. Nothing is tested against a boundary to arrive at
        it — INAB publishes Guatemala's fires and nothing else.
    year : int or None
        The Guatemalan calendar year, or ``None`` for the summary row.
    fires : int
        How many fires the year holds. Excludes the false alarms unless
        ``--include-false-alarms`` was passed.
    false_alarms : int
        How many of the year's records say there was no fire. Counted whether or not
        they are in :attr:`fires`, so ``fires + false_alarms`` is the published record
        count under the default scope and ``fires`` alone is under the other.
    located : int
        How many of those ``fires`` carry a published coordinate. Always at most
        ``fires``: it is counted over the same rows. Equal to it on every record
        published today.
    protected : int
        How many of those ``fires`` fell inside a protected area.

    Notes
    -----
    There are no ``minimum``, ``maximum`` or ``total`` attributes. The three hectare
    columns exist in the output and are empty in it, and there is no number behind
    them to carry — see the module docstring.
    """

    country: str
    year: int | None
    fires: int
    false_alarms: int
    located: int
    protected: int

    @property
    def is_total(self) -> bool:
        """Whether this is the summary row rather than one of the years."""
        return self.year is None

    @property
    def year_label(self) -> str:
        return TOTAL_LABEL if self.is_total else str(self.year)

    @property
    def located_share(self) -> float | None:
        """``located`` as a percentage of ``fires``: how much of the year is mappable."""
        return share(self.located, self.fires)

    @property
    def values(self) -> tuple[str, ...]:
        """The row as the CSV writes it, in :data:`COLUMNS` order."""
        return (self.country, self.year_label, str(self.fires),
                NO_AREA, NO_AREA, NO_AREA,
                str(self.false_alarms), str(self.located),
                share_label(self.located, self.fires), str(self.protected))

    @property
    def readable_values(self) -> tuple[str, ...]:
        """The row as the Word document writes it: the numbers with separators."""
        return (self.country, self.year_label, f"{self.fires:,}",
                NO_AREA, NO_AREA, NO_AREA,
                f"{self.false_alarms:,}", f"{self.located:,}",
                share_label(self.located, self.fires), f"{self.protected:,}")


def combine(rows: list[Row], country: str = COUNTRY_NAME,
            year: int | None = None) -> Row:
    """One row summarising several: every figure taken over all of them.

    Notes
    -----
    All four counts decompose over a partition of the fires — a count of counts is a
    count — so the ``Total`` row is the number a second aggregate over the same rows
    would have returned, and no fire is counted twice or left out.

    The percentage is deliberately **not** averaged: :class:`Row` recomputes it from
    the summed counts, which is the ratio of the totals rather than the mean of the
    ratios.
    """
    return Row(
        country=country,
        year=year,
        fires=sum(row.fires for row in rows),
        false_alarms=sum(row.false_alarms for row in rows),
        located=sum(row.located for row in rows),
        protected=sum(row.protected for row in rows),
    )


def summarise(measured: list[Row]) -> list[Row]:
    """Build the report from the years measured: the years newest first, then the total.

    Parameters
    ----------
    measured : list of Row
        One row per year, as the statement returned them.

    Returns
    -------
    list of Row
        The years newest first with the summary row last. Empty if nothing was
        measured — a report of no fires has no total either.

    Notes
    -----
    No grouping by country and no lookup of one, unlike the four worldwide reports:
    there is exactly one country here and it is a constant.
    """
    if not measured:
        return []
    rows = sorted(measured, key=lambda row: row.year, reverse=True)
    return rows + [combine(rows)]


def add_refused_arguments(parser: argparse.ArgumentParser) -> None:
    """Accept the other reports' options, so that they can be refused clearly.

    Anyone reaching for one has copied a command line from one of the seven reports
    that do have hectares, which is a reasonable thing to have done, and argparse's own
    *unrecognized arguments* would not say why this report is different. Shared with
    the classification report so both refuse them in the same words.
    """
    for option in ("--country", "--country-source", "--area-method", "--surface",
                   "--min-area"):
        parser.add_argument(option, help=argparse.SUPPRESS)


def check_refused_arguments(parser: argparse.ArgumentParser,
                            arguments: argparse.Namespace) -> None:
    """Fail with an explanation if one of the refused options was given."""
    if arguments.country is not None:
        parser.error(
            f"there is no --country here: INAB publishes the fires of {COUNTRY_NAME} "
            f"and nothing else, so there is nothing to select between. Every fire is "
            f"counted and the Country column is {COUNTRY_NAME} on every row."
        )
    if arguments.country_source is not None:
        parser.error(
            f"there is no --country-source here: these are {COUNTRY_NAME}'s own "
            f"records of its own fires, so nothing is tested against a boundary and "
            f"there is nothing for a containment test to find. The three published "
            f"points that fall outside the country are all already flagged "
            f"{guatemala_inab.STATUS_FALSE!r}."
        )
    for option, given in (("--area-method", arguments.area_method),
                          ("--surface", arguments.surface),
                          ("--min-area", arguments.min_area)):
        if given is not None:
            parser.error(
                f"there is no {option} here, and there is nothing it could select: "
                f"INAB publishes no perimeter, no burnt area and no land-cover split — "
                f"not one of its thirty-three attributes is a size. The hectare columns "
                f"of this report are empty on every row, which is what that looks like "
                f"written down."
            )


def parse_arguments(argv: list[str] | None = None) -> argparse.Namespace:
    """Parse the command line."""
    parser = argparse.ArgumentParser(
        description="Wildfire statistics for the Guatemalan INAB fire reports: how many "
                    "fires per year, how many false alarms, how many located and how "
                    "many inside a protected area.",
        epilog="There are no hectares in this report and the three area columns are "
               "empty on every row: INAB publishes no perimeter and no burnt area, so "
               "there is nothing to measure and no --area-method or --surface to choose "
               "with. Every fire is Guatemalan, so there is no --country. Database "
               "settings not given here are read from the environment (.env).",
    )
    parser.add_argument("-y", "--year", type=int,
                        help="restrict to one year, e.g. 2025; this is the Guatemalan "
                             "calendar year of the fire's own instant, this source "
                             "publishing no year field of its own")
    parser.add_argument("--include-false-alarms", action="store_true",
                        help=f"count the records whose report status is "
                             f"{guatemala_inab.STATUS_FALSE} — the report was false, "
                             f"there was no fire. 140 of the 4,615. Left out of Fires by "
                             f"default; the False alarms column reports them either way")

    add_refused_arguments(parser)

    output = parser.add_argument_group("output", "at least one is required")
    output.add_argument("--csv", type=Path, help="write the report to this .csv")
    output.add_argument("--docx", type=Path,
                        help="write the report to this .docx (MS Word)")

    common.add_database_arguments(parser)
    parser.add_argument("--log-level", default=os.getenv("GISFIRE_LOG_LEVEL", "INFO"),
                        choices=["CRITICAL", "ERROR", "WARNING", "INFO", "DEBUG"],
                        help="verbosity (env: GISFIRE_LOG_LEVEL, default INFO)")

    arguments = parser.parse_args(argv)
    check_refused_arguments(parser, arguments)
    if arguments.csv is None and arguments.docx is None:
        parser.error("nothing to write: pass --csv, --docx, or both")
    return arguments


def compute(session: Session, year: int | None, logger: logging.Logger,
            include_false_alarms: bool = False) -> list[Row]:
    """Run the statement and return the report's rows in order.

    Notes
    -----
    One statement, under one spinner. The ``Total`` row is arithmetic over its result,
    not a second query.

    The located share is logged as well as reported, because a value below 100% is
    news about this dataset — every record published today has a point — in the same
    way that the Greek report's zeros are news about that one.
    """
    with common.Spinner(f"Counting the {COUNTRY_NAME} fire reports", logger):
        measured = [
            Row(country=record.country, year=record.year, fires=record.fires,
                false_alarms=record.false_alarms, located=record.located,
                protected=record.protected)
            for record in session.execute(statistics_query(year, include_false_alarms))
        ]

    rows = summarise(measured)
    logger.info("Computed %d rows over %d year(s); no burnt area is reported, this "
                "source publishing none", len(rows), len(measured))
    if rows:
        total = rows[-1]
        if total.false_alarms:
            logger.info(
                "%d false alarm(s) (%s) %s Fires; %d fire(s) counted",
                total.false_alarms, guatemala_inab.STATUS_FALSE,
                "counted in" if include_false_alarms else
                "excluded from — pass --include-false-alarms to count them in",
                total.fires,
            )
        logger.info("%d of %d fire(s) publish a coordinate (%s%%), %d inside a "
                    "protected area",
                    total.located, total.fires,
                    share_label(total.located, total.fires), total.protected)
        if total.located < total.fires:
            logger.warning(
                "%d fire(s) in scope publish no coordinate. Every record published to "
                "date carries one, so this is a property of a newer publication rather "
                "than of the import",
                total.fires - total.located)
    return rows


def write_csv(rows: list[Row], path: Path, logger: logging.Logger) -> None:
    """Write the report as CSV.

    The numbers go out unformatted — no thousands separators — because a CSV is read
    by another program far more often than by a person, and a separator would make
    every figure a string. The three hectare fields go out **empty**, which is what
    every reader of a CSV takes as *not published*.
    """
    path.parent.mkdir(parents=True, exist_ok=True)
    with path.open("w", newline="", encoding="utf-8") as handle:
        writer = csv.writer(handle)
        writer.writerow(COLUMNS)
        for row in rows:
            writer.writerow(row.values)
    logger.info("Wrote %s", path)


def write_docx(rows: list[Row], path: Path, year: int | None, logger: logging.Logger,
               include_false_alarms: bool = False) -> None:
    """Write the report as a Word document.

    One table, with the summary row in bold. Numbers get thousands separators here —
    the opposite of the CSV, and for the opposite reason: this one is for reading.

    The opening paragraphs say why three columns are empty, which year is being
    grouped on, what happened to the false alarms and what ``Located`` means, because
    a table with three empty columns is the first thing a reader will ask about and
    the answer is not something they should have to guess.
    """
    # Imported here rather than at module scope so that --csv keeps working if
    # python-docx is not installed, which matters because it is the only dependency
    # this application adds.
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt

    document = Document()
    document.add_heading(f"INAB wildfire statistics ({COUNTRY_NAME})", level=1)

    scope = [f"year: {year}" if year is not None else "all years"]
    scope.append("false alarms counted in Fires" if include_false_alarms
                 else "false alarms excluded from Fires")
    document.add_paragraph(
        f"Counts of the fire reports INAB received. Scope: {'; '.join(scope)}. Years "
        f"are the Guatemalan calendar year of each fire's own instant "
        f"({guatemala_inab.DEFAULT_TIME_ZONE}): this source publishes no year field, "
        f"unlike every other dataset with a report in this project, so the year is "
        f"derived from fecha_hora_incendio."
    )
    document.add_paragraph(
        "The Minimum (ha), Maximum (ha) and Total (ha) columns are empty on every row, "
        "and that is the dataset rather than a fault. INAB publishes no perimeter, no "
        "burnt area and no land-cover split — not one of its thirty-three attributes "
        "is a size — so there is nothing to measure and nothing to sum. The columns are "
        "kept in the position the other seven burnt-area reports put them so that this "
        "table can be read beside theirs; an empty cell says nothing was published, "
        "where a zero would say nothing burnt."
    )
    document.add_paragraph(
        f"False alarms counts the records whose status is "
        f"{guatemala_inab.STATUS_FALSE} — the report was false, there was no fire. "
        f"They are counted here whether or not they are in Fires. Records whose status "
        f"is {guatemala_inab.STATUS_UNVERIFIED}, meaning nobody went to look, are a "
        f"different claim and are left in Fires; the classification report counts them "
        f"separately."
    )
    document.add_paragraph(
        f"Located counts the fires that publish a coordinate, and In protected area "
        f"those whose point fell inside one, as INAB reports it. Every record published "
        f"to date carries a point, so Located is normally 100% — the opposite of the "
        f"Greek report, where it is zero for twenty years. This is the best-located "
        f"administrative fire statistic in GisFIRE."
    )
    document.add_paragraph(
        "INAB publishes no cause for any fire, so there is no counts-by-cause companion "
        "to this report as there is for Portugal, Spain and Canada. The classification "
        "report counts the four published vocabularies instead."
    )

    table = document.add_table(rows=1, cols=len(COLUMNS))
    table.style = "Table Grid"
    for cell, heading in zip(table.rows[0].cells, COLUMNS):
        cell.text = heading
        cell.paragraphs[0].runs[0].bold = True

    for row in rows:
        cells = table.add_row().cells
        for index, (cell, value) in enumerate(zip(cells, row.readable_values)):
            cell.text = value
            paragraph = cell.paragraphs[0]
            if index >= FIRST_NUMERIC_COLUMN:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            for run in paragraph.runs:
                run.bold = row.is_total
                run.font.size = Pt(9)

    path.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(path))
    logger.info("Wrote %s", path)


def report(args: argparse.Namespace, engine: Engine,
           logger: logging.Logger) -> list[Row]:
    """Compute the statistics and write whichever outputs were asked for."""
    with Session(engine) as session:
        rows = compute(session, args.year, logger, args.include_false_alarms)

    if not rows:
        # An empty report is almost always a year with no data, and writing an empty
        # file would hide that.
        raise RuntimeError(
            f"No wildfires matched. Check --year, and that the {COUNTRY_NAME} fire "
            f"reports are imported — the published archive starts in 2023 — with "
            f"src.apps.imports.wildfires.guatemala_inab.import_wildfires."
        )

    if args.csv is not None:
        write_csv(rows, args.csv, logger)
    if args.docx is not None:
        write_docx(rows, args.docx, args.year, logger, args.include_false_alarms)
    return rows


def main(argv: list[str] | None = None) -> int:
    args = parse_arguments(argv)
    logging.basicConfig(level=args.log_level,
                        format="%(asctime)s %(levelname)s %(message)s")
    logger = logging.getLogger("inab-statistics")

    try:
        settings = common.resolve_database_settings(args)
    except RuntimeError as error:
        logger.error("%s", error)
        return 1

    engine = create_engine(common.database_url(settings))
    try:
        report(args, engine, logger)
    except Exception as error:  # noqa: BLE001  (the CLI boundary: report, do not traceback)
        logger.error("%s", error)
        return 1
    finally:
        engine.dispose()
    return 0


if __name__ == "__main__":
    sys.exit(main())
