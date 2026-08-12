#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Wildfire counts by published classification for the Guatemalan INAB fire reports.

The companion of
:mod:`~src.apps.statistics.wildfires.guatemala_inab.wildfire_statistics`, over the
same fires, the same years and the same scope — so the ``Country``, ``Year`` and
``Fires`` columns of the two reports agree row for row. What that one counts, this one
breaks down::

    Country    Year   Fires  Classified  Classified (%)  In forest  In forest (%)  Outside forest  Outside forest (%)
    Guatemala  2025    1735         188           10.83        147          78.19              41               21.81
    Guatemala  2024     704          78           11.08         61          78.21              17               21.79
    Guatemala  Total   4471         489           10.94        383          78.32             106               21.68

Four classifications, one of which is about the fire and three of which are about the
report::

    python3 -m src.apps.statistics.wildfires.guatemala_inab.wildfire_classification \\
        --classification location --csv location.csv --docx location.docx

    python3 -m src.apps.statistics.wildfires.guatemala_inab.wildfire_classification \\
        --classification status --csv status.csv

    python3 -m src.apps.statistics.wildfires.guatemala_inab.wildfire_classification \\
        --classification institution --year 2025 --csv institutions-2025.csv

At least one of ``--csv`` and ``--docx`` is required.

.. important::

   **This is not a causes report, and it must not be read as one.**

   Portugal, Spain and Canada have a counts-by-cause report because those sources
   publish a cause: EGIF has ``idcausa``, the NFDB has ``CAUSE``, the ICNF has its
   ``causa`` catalogue. **Nothing in INAB's thirty-three published attributes says why
   a fire started** — there is no cause column, no lightning category and no arson
   category, and no amount of aggregation over what is published will produce one.

   What this report counts is the four closed vocabularies the source *does* publish.
   One of them, ``tipo_incendio``, classifies the fire; the other three classify the
   report. None of them is a cause, which is why this application is named
   ``wildfire_classification`` and not ``wildfire_causes``, and why no column of it is
   ever headed *Cause*.

The four classifications
-------------------------

``location`` — ``tipo_incendio``, the default
    Whether the fire was inside or outside forest. **The only classification of the
    fire itself that this dataset carries**, and the nearest thing in it to the kind
    of column a causes report would count.

    It is filled on **489 of the 4,615 records, 10.6%**, which is the single most
    important thing to know about this report: a table of it describes one record in
    ten, and the ``Classified (%)`` column is there so that no one can read the
    breakdown without seeing the coverage that produced it.

``status`` — ``estado_aviso``
    What became of the report: closed, false, unverified, confirmed, still burning.
    Filled on effectively everything, and the one classification here with real
    coverage. It is also where the 140 false alarms and the 90 unverified reports
    become visible side by side, which
    :mod:`~src.apps.statistics.wildfires.guatemala_inab.wildfire_statistics` shows only
    for the first of the two.

``institution`` — ``institucion``
    Which organisation called the fire in. Fourteen values, led by ``conred`` and
    ``conap``.

``channel`` — ``forma_comunicacion``
    How the report reached INAB: by telephone, in person, through the app, through
    social media, by radio. A measure of how a country hears about its fires, and the
    only such measure in the project.

Where the columns come from, and why one of them is different
---------------------------------------------------------------

Three of the four vocabularies are **published constants**:
:data:`~src.providers.guatemala_inab.FIRE_LOCATIONS`,
:data:`~src.providers.guatemala_inab.REPORT_STATUSES` and
:data:`~src.providers.guatemala_inab.REPORT_CHANNELS`. Their columns come from those
tuples, in the order the provider module lists them, so the report has the same
columns whatever happens to be in the database — which is what lets two runs over
different scopes be compared, and what makes an empty column an answer rather than an
absence.

``institution`` has no such constant, because the provider module has none to give:
the fourteen values were observed once and never published as a list. Its columns are
therefore **built from the data in scope**, most frequent first. That is a real
difference and it is stated in the output: two runs of ``--classification institution``
over different years can have different columns, and their CSVs cannot be concatenated
on anything but the first five.

A value in the data that is not in the published tuple is **reported and then counted
anyway**, in a column of its own after the published ones. The provider module gives
these vocabularies no ``CHECK`` constraint precisely because they are one publication
observed once, so the first value INAB adds must not vanish from a report — see
:mod:`src.providers.guatemala_inab`.

Why the denominator is the classified fires
---------------------------------------------

``Classified`` counts the fires that carry any value at all for the chosen
classification, and every percentage after it is a share of that, not of ``Fires``.
This is the choice the Canadian causes reports make about their ``U`` category, and it
matters far more here than it does there: ``U`` is under 2% of the NFDB, while
``tipo_incendio`` is **absent from 89% of this archive**.

So ``In forest (%)`` says *of the fires somebody classified, this share were in
forest*. Reading it as a share of all Guatemalan fires would be wrong by a factor of
nine. Where nothing in a year is classified there is no percentage to give, and the
cell is left **empty** rather than filled with a zero that would be a claim.

.. warning::

   **The coverage is not evenly spread, and the report cannot correct for that.**
   ``tipo_incendio`` is filled by whoever handled the report, so the classified tenth
   is a sample of INAB's reporting practice rather than a random sample of Guatemalan
   fire. ``--year`` and the ``Classified (%)`` column are what make that visible; there
   is no weighting that would make the other nine tenths speak.

Shared with the companion report
---------------------------------

The country, the year expression, the false-alarm scope, the percentage helpers and
the refused options are **imported from**
:mod:`~src.apps.statistics.wildfires.guatemala_inab.wildfire_statistics` rather than
copied. Two reports over one dataset that disagreed about which fires are in scope, or
about which year one is in, would be worse than one report.
"""

from __future__ import annotations

import argparse
import csv
import logging
import os
import sys

from dataclasses import dataclass
from pathlib import Path

from sqlalchemy import ColumnElement
from sqlalchemy import Engine
from sqlalchemy import Select
from sqlalchemy import create_engine
from sqlalchemy import func
from sqlalchemy import literal
from sqlalchemy import select
from sqlalchemy.orm import Session

import src.settings  # noqa: F401  (imported for the side effect of loading .env)

from src.apps.imports import common
from src.apps.statistics.wildfires.guatemala_inab.wildfire_statistics import COUNTRY_NAME
from src.apps.statistics.wildfires.guatemala_inab.wildfire_statistics import (
    FIRST_NUMERIC_COLUMN,
)
from src.apps.statistics.wildfires.guatemala_inab.wildfire_statistics import LOCAL_YEAR
from src.apps.statistics.wildfires.guatemala_inab.wildfire_statistics import TOTAL_LABEL
from src.apps.statistics.wildfires.guatemala_inab.wildfire_statistics import (
    add_refused_arguments,
)
from src.apps.statistics.wildfires.guatemala_inab.wildfire_statistics import (
    check_refused_arguments,
)
from src.apps.statistics.wildfires.guatemala_inab.wildfire_statistics import is_a_fire
from src.apps.statistics.wildfires.guatemala_inab.wildfire_statistics import share
from src.apps.statistics.wildfires.guatemala_inab.wildfire_statistics import share_label
from src.data_model.wildfire import Wildfire
from src.providers import guatemala_inab
from src.providers.guatemala_inab.wildfire import InabWildfire


@dataclass(frozen=True)
class Classification:
    """One published vocabulary this report can count.

    Attributes
    ----------
    key : str
        What ``--classification`` takes.
    column : str
        The :class:`~src.providers.guatemala_inab.wildfire.InabWildfire` attribute
        holding the value.
    published_name : str
        The attribute as **INAB** publishes it. Kept beside the model's own name
        because the two differ on all four of these, and a reader who came from the
        provider documentation or from the GeoJSON is thinking in this one.
    published : tuple of str or None
        The vocabulary as the provider module publishes it, in its order, or ``None``
        for a classification with no published list — see the module docstring on
        ``institution``.
    labels : dict
        How each published value is written in a column heading. A published value
        with no entry is titled from its own slug, which is what a *new* value gets.
    prose : str
        One line naming the published attribute and what it classifies, for the Word
        document and the ``--list`` output.
    """

    key: str
    column: str
    published_name: str
    published: tuple[str, ...] | None
    labels: dict[str, str]
    prose: str

    @property
    def attribute(self) -> ColumnElement:
        """The mapped column, resolved once so a rename breaks this at import time."""
        return InabWildfire.__table__.c[self.column]

    def label(self, value: str) -> str:
        """The column heading for one published value."""
        return self.labels.get(value, value.replace("_", " ").capitalize())


#: Whether the fire was inside or outside forest — the only classification of the
#: **fire** in this dataset, and filled on one record in ten.
CLASSIFICATION_LOCATION = "location"

#: What became of the report. The only one of the four with real coverage.
CLASSIFICATION_STATUS = "status"

#: Which organisation called it in. The one vocabulary with no published list.
CLASSIFICATION_INSTITUTION = "institution"

#: How the report reached INAB.
CLASSIFICATION_CHANNEL = "channel"

#: Every classification this report can count, keyed by what ``--classification``
#: takes.
#:
#: The published tuples are the provider module's, not copies of them, so a value
#: added there appears here without an edit. ``institution`` has ``None`` because
#: there is no such tuple to point at — the fourteen values were observed once and
#: never published as a list.
CLASSIFICATIONS: dict[str, Classification] = {
    CLASSIFICATION_LOCATION: Classification(
        key=CLASSIFICATION_LOCATION,
        column="fire_location",
        published_name="tipo_incendio",
        published=guatemala_inab.FIRE_LOCATIONS,
        labels={
            guatemala_inab.LOCATION_IN_FOREST: "In forest",
            guatemala_inab.LOCATION_OUT_OF_FOREST: "Outside forest",
        },
        prose="tipo_incendio, whether the fire was inside or outside forest — the only "
              "classification of the fire itself this dataset carries, and filled on "
              "about one record in ten",
    ),
    CLASSIFICATION_STATUS: Classification(
        key=CLASSIFICATION_STATUS,
        column="report_status",
        published_name="estado_aviso",
        published=guatemala_inab.REPORT_STATUSES,
        labels={
            guatemala_inab.STATUS_CLOSED: "Closed",
            guatemala_inab.STATUS_FALSE: "False alarm",
            guatemala_inab.STATUS_UNVERIFIED: "Unverified",
            guatemala_inab.STATUS_TRUE: "Confirmed",
            guatemala_inab.STATUS_ACTIVE: "Active",
        },
        prose="estado_aviso, what became of the report — closed, false, unverified, "
              "confirmed or still burning",
    ),
    CLASSIFICATION_INSTITUTION: Classification(
        key=CLASSIFICATION_INSTITUTION,
        column="institution",
        published_name="institucion",
        published=None,
        labels={},
        prose="institucion, which organisation reported the fire — fourteen values, "
              "for which the provider publishes no list, so the columns come from the "
              "data in scope",
    ),
    CLASSIFICATION_CHANNEL: Classification(
        key=CLASSIFICATION_CHANNEL,
        column="report_channel",
        published_name="forma_comunicacion",
        published=guatemala_inab.REPORT_CHANNELS,
        labels={
            "telefono": "Telephone",
            "personal": "In person",
            "app": "App",
            "redes_sociales": "Social media",
            "radio": "Radio",
        },
        prose="forma_comunicacion, how the report reached INAB — by telephone, in "
              "person, through the app, through social media or by radio",
    ),
}

#: The classification counted unless ``--classification`` says otherwise.
#:
#: The one that classifies the fire rather than the report, and so the one a reader
#: who came here looking for causes is nearest to wanting. Its coverage is the reason
#: ``Classified (%)`` is the third column and not a footnote.
DEFAULT_CLASSIFICATION = CLASSIFICATION_LOCATION

#: The columns every run has, whatever is being counted. The first three are the
#: companion report's first three, unchanged.
LEADING_COLUMNS = ("Country", "Year", "Fires", "Classified", "Classified (%)")


def columns(values: tuple[str, ...], classification: Classification) -> tuple[str, ...]:
    """The report's columns, in order, for one classification and its values.

    A function rather than a constant, for the reason the Canadian causes reports'
    is: most of the headings name what is being counted, and a file of ``conap``
    counts under a heading saying ``In forest`` would be a trap.
    """
    headings: list[str] = list(LEADING_COLUMNS)
    for value in values:
        label = classification.label(value)
        headings += [label, f"{label} (%)"]
    return tuple(headings)


def observed_values(session: Session, classification: Classification,
                    year: int | None = None,
                    include_false_alarms: bool = False) -> list[str]:
    """The values actually present in scope, most frequent first.

    Returns
    -------
    list of str
        Every non-``NULL`` value of the classification's column, ordered by how many
        fires carry it and then by the value itself so the order is total.

    Notes
    -----
    Used two ways. For ``institution`` it *is* the column list, there being no
    published vocabulary to use instead. For the other three it is how an
    **unpublished** value is discovered — a value INAB has added since the vocabulary
    was observed — so that it can be reported and then given a column of its own
    rather than silently dropped from a report whose columns came from a fixed tuple.

    The tie-break on the value matters: without it two runs over the same data could
    order two equally common institutions differently, and their CSVs would not line
    up.
    """
    inab = InabWildfire.__table__
    counts = (
        select(classification.attribute.label("value"), func.count().label("fires"))
        .select_from(Wildfire)
        .join(inab, inab.c.id == Wildfire.id)
        .where(classification.attribute.is_not(None))
        .where(is_a_fire(include_false_alarms))
        .group_by(classification.attribute)
        .order_by(func.count().desc(), classification.attribute)
    )
    if year is not None:
        counts = counts.where(LOCAL_YEAR == year)
    return [record.value for record in session.execute(counts)]


def report_values(classification: Classification, observed: list[str],
                  logger: logging.Logger) -> tuple[str, ...]:
    """The values this run gives a column to, in the order it gives them.

    Parameters
    ----------
    classification : Classification
        What is being counted.
    observed : list of str
        What :func:`observed_values` found, most frequent first.
    logger : logging.Logger
        Where an unpublished value is reported.

    Returns
    -------
    tuple of str
        The published vocabulary in its published order, followed by any value found
        in the data that is not in it — or, for a classification with no published
        vocabulary, the observed values alone.

    Notes
    -----
    A published value that no fire in scope carries still gets a column, holding zero.
    That is the point of using the published tuple: a year in which nobody reported a
    fire as still burning should show that as a zero, not as a missing column, and two
    years' CSVs should have the same header.

    An **unpublished** value is a small event worth a log line. These vocabularies
    carry no ``CHECK`` constraint precisely so that INAB can add a value without the
    import failing, which means this report is where such a value first becomes
    visible.
    """
    if classification.published is None:
        return tuple(observed)

    extra = [value for value in observed if value not in classification.published]
    if extra:
        logger.warning(
            "%d value(s) of %s are not in the published vocabulary and are counted in "
            "columns of their own: %s. The provider gives this column no CHECK "
            "constraint for exactly this reason; consider adding them to "
            "src.providers.guatemala_inab",
            len(extra), classification.column, ", ".join(repr(value) for value in extra),
        )
    return tuple(classification.published) + tuple(extra)


def counts_query(classification: Classification, values: tuple[str, ...],
                 year: int | None = None,
                 include_false_alarms: bool = False) -> Select:
    """Build the counting query: one row per year, newest first.

    Returns
    -------
    Select
        A query yielding ``country, year, fires, classified`` and one count per value
        in ``values``, labelled ``value_0``, ``value_1`` and so on.

    Notes
    -----
    One pass, whatever the vocabulary's size: every count is a filtered aggregate over
    the same scan, so counting fourteen institutions costs what counting two locations
    does.

    The counts are labelled by **position** rather than by the published value.
    ``forma_comunicacion`` values are safe identifiers, but an institution's is
    whatever INAB typed, and a label taken from data is one SQL injection or one
    quoting bug away from being a problem. The position is enough: ``values`` is the
    order the columns are written in.

    The classification is what this report counts, so it is never a filter — only the
    false-alarm half of the companion's scope is applied, which is what keeps the
    ``Fires`` column equal to that report's.
    """
    inab = InabWildfire.__table__
    counted = is_a_fire(include_false_alarms)
    classified = classification.attribute.is_not(None)

    counts = (
        select(
            literal(COUNTRY_NAME).label("country"),
            LOCAL_YEAR.label("year"),
            func.count().filter(counted).label("fires"),
            func.count().filter(counted & classified).label("classified"),
            *[
                func.count()
                .filter(counted & (classification.attribute == value))
                .label(f"value_{index}")
                for index, value in enumerate(values)
            ],
        )
        .select_from(Wildfire)
        .join(inab, inab.c.id == Wildfire.id)
        .group_by(LOCAL_YEAR)
        .order_by(LOCAL_YEAR.desc())
    )
    if year is not None:
        counts = counts.where(LOCAL_YEAR == year)
    return counts


@dataclass(frozen=True)
class Row:
    """One line of the report.

    Attributes
    ----------
    country : str
        Always :data:`COUNTRY_NAME`.
    year : int or None
        The Guatemalan calendar year, or ``None`` for the summary row.
    fires : int
        Every fire counted, classified or not. Equal to the companion report's
        ``Fires`` for the same year and scope.
    classified : int
        How many carry a value for the chosen classification. **The denominator**;
        the difference from :attr:`fires` is how much of the year the breakdown does
        not describe.
    counts : tuple of int
        How many carry each value, in the order of the report's columns.
    """

    country: str
    year: int | None
    fires: int
    classified: int
    counts: tuple[int, ...]

    @property
    def is_total(self) -> bool:
        return self.year is None

    @property
    def year_label(self) -> str:
        return TOTAL_LABEL if self.is_total else str(self.year)

    @property
    def classified_share(self) -> float | None:
        """``classified`` as a percentage of ``fires``: the coverage of the breakdown."""
        return share(self.classified, self.fires)

    def value_share(self, index: int) -> float | None:
        """One value's share **of the classified fires**, not of every fire.

        See the module docstring: with ``tipo_incendio`` absent from 89% of the
        archive, a share of ``fires`` would be a different and much smaller number,
        and reading one for the other is the mistake this report is shaped to prevent.
        """
        return share(self.counts[index], self.classified)

    def _cells(self, thousands: bool) -> tuple[str, ...]:
        """The row, with or without thousands separators."""
        def number(value: int) -> str:
            return f"{value:,}" if thousands else str(value)

        cells = [self.country, self.year_label, number(self.fires),
                 number(self.classified), share_label(self.classified, self.fires)]
        for count in self.counts:
            cells += [number(count), share_label(count, self.classified)]
        return tuple(cells)

    @property
    def values(self) -> tuple[str, ...]:
        """The row as the CSV writes it."""
        return self._cells(thousands=False)

    @property
    def readable_values(self) -> tuple[str, ...]:
        """The row as the Word document writes it."""
        return self._cells(thousands=True)


def combine(rows: list[Row], country: str = COUNTRY_NAME,
            year: int | None = None) -> Row:
    """One row summarising several: every count taken over all of them.

    The percentages are deliberately **not** averaged: :class:`Row` recomputes them
    from the summed counts, which is the ratio of the totals rather than the mean of
    the ratios. A year of 11% coverage and a year of 3% must not average to 7%.
    """
    return Row(
        country=country,
        year=year,
        fires=sum(row.fires for row in rows),
        classified=sum(row.classified for row in rows),
        counts=tuple(sum(row.counts[index] for row in rows)
                     for index in range(len(rows[0].counts))),
    )


def summarise(measured: list[Row]) -> list[Row]:
    """The years newest first, then the total. Empty in, empty out."""
    if not measured:
        return []
    rows = sorted(measured, key=lambda row: row.year, reverse=True)
    return rows + [combine(rows)]


def parse_arguments(argv: list[str] | None = None) -> argparse.Namespace:
    """Parse the command line."""
    parser = argparse.ArgumentParser(
        description="Wildfire counts by published classification for the Guatemalan "
                    "INAB fire reports. NOT a causes report: this source publishes no "
                    "cause for any fire.",
        epilog="--classification location is the only one of the four that classifies "
               "the fire rather than the report, and it is filled on about one record "
               "in ten — read Classified (%) before reading anything after it. Database "
               "settings not given here are read from the environment (.env).",
    )
    parser.add_argument("-c", "--classification", default=DEFAULT_CLASSIFICATION,
                        choices=sorted(CLASSIFICATIONS),
                        help=f"which published vocabulary to count (default: "
                             f"{DEFAULT_CLASSIFICATION})")
    parser.add_argument("-y", "--year", type=int,
                        help="restrict to one year, e.g. 2025; this is the Guatemalan "
                             "calendar year of the fire's own instant")
    parser.add_argument("--include-false-alarms", action="store_true",
                        help=f"count the records whose report status is "
                             f"{guatemala_inab.STATUS_FALSE}. Left out by default, as "
                             f"in the companion report — note that leaving them out "
                             f"empties the False alarm column of --classification "
                             f"status, which is where they are meant to be looked at")
    parser.add_argument("--list-classifications", action="store_true",
                        help="print what each classification counts, and exit")

    # Accepted only so that they can be refused in the same words as the companion
    # report refuses them.
    add_refused_arguments(parser)
    # And this one, which is the whole point of the application's name.
    parser.add_argument("--cause", help=argparse.SUPPRESS)

    output = parser.add_argument_group("output", "at least one is required")
    output.add_argument("--csv", type=Path, help="write the report to this .csv")
    output.add_argument("--docx", type=Path,
                        help="write the report to this .docx (MS Word)")

    common.add_database_arguments(parser)
    parser.add_argument("--log-level", default=os.getenv("GISFIRE_LOG_LEVEL", "INFO"),
                        choices=["CRITICAL", "ERROR", "WARNING", "INFO", "DEBUG"],
                        help="verbosity (env: GISFIRE_LOG_LEVEL, default INFO)")

    arguments = parser.parse_args(argv)
    if arguments.list_classifications:
        return arguments
    if arguments.cause is not None:
        parser.error(
            "there is no --cause here, and there is nothing it could select: INAB "
            "publishes no cause for any fire. None of its thirty-three attributes says "
            "why a fire started — there is no equivalent of EGIF's idcausa or the "
            "NFDB's CAUSE, no lightning category and no arson category. What this "
            "report counts is --classification: what the fire was (location) or what "
            "became of the report (status, institution, channel)."
        )
    check_refused_arguments(parser, arguments)
    if arguments.csv is None and arguments.docx is None:
        parser.error("nothing to write: pass --csv, --docx, or both")
    return arguments


def compute(session: Session, classification: Classification, year: int | None,
            logger: logging.Logger,
            include_false_alarms: bool = False) -> tuple[list[Row], tuple[str, ...]]:
    """Run the statements and return the report's rows and the values they count.

    Returns
    -------
    tuple
        The rows in order, and the values each row's counts correspond to — which the
        writers need in order to build the header.

    Notes
    -----
    Two statements: one to find which values are present, one to count them. The first
    is what lets the second label its counts by position, and it is what discovers a
    value INAB has added since its vocabulary was observed.

    The coverage is logged as well as reported, because with ``location`` it is the
    number that decides whether the rest of the table means anything.
    """
    with common.Spinner(f"Counting the {COUNTRY_NAME} fire reports by "
                        f"{classification.column}", logger):
        observed = observed_values(session, classification, year, include_false_alarms)
        values = report_values(classification, observed, logger)
        measured = [
            Row(country=record.country, year=record.year, fires=record.fires,
                classified=record.classified,
                counts=tuple(getattr(record, f"value_{index}")
                             for index in range(len(values))))
            for record in session.execute(
                counts_query(classification, values, year, include_false_alarms))
        ]

    rows = summarise(measured)
    logger.info("Computed %d rows over %d year(s), counting %s across %d value(s)",
                len(rows), len(measured), classification.column, len(values))
    if rows:
        total = rows[-1]
        logger.info("%d of %d fire(s) carry a %s (%s%%)",
                    total.classified, total.fires, classification.column,
                    share_label(total.classified, total.fires))
        if not total.classified:
            logger.warning(
                "No fire in scope carries a %s, so every column after Classified is "
                "zero and every percentage is empty. That is the expected answer for a "
                "scope of fires nobody classified, and is not a sign of a failed import",
                classification.column)
        elif total.classified < total.fires / 2:
            logger.warning(
                "This breakdown describes %s%% of the fires in scope. It is a sample of "
                "who filled the form in, not of Guatemalan fire, and nothing here "
                "weights it back",
                share_label(total.classified, total.fires))
    return rows, values


def write_csv(rows: list[Row], values: tuple[str, ...],
              classification: Classification, path: Path,
              logger: logging.Logger) -> None:
    """Write the report as CSV, with no thousands separators."""
    path.parent.mkdir(parents=True, exist_ok=True)
    with path.open("w", newline="", encoding="utf-8") as handle:
        writer = csv.writer(handle)
        writer.writerow(columns(values, classification))
        for row in rows:
            writer.writerow(row.values)
    logger.info("Wrote %s", path)


def write_docx(rows: list[Row], values: tuple[str, ...],
               classification: Classification, path: Path, year: int | None,
               logger: logging.Logger, include_false_alarms: bool = False) -> None:
    """Write the report as a Word document.

    The opening paragraphs say what is being counted, that it is not a cause, what the
    denominator is and where the columns came from — all four of which a reader would
    otherwise have to assume, and would probably assume wrongly.
    """
    # Imported here rather than at module scope so that --csv keeps working if
    # python-docx is not installed.
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt

    headings = columns(values, classification)
    document = Document()
    document.add_heading(
        f"INAB wildfire counts by {classification.key} ({COUNTRY_NAME})", level=1)

    scope = [f"year: {year}" if year is not None else "all years"]
    scope.append("false alarms counted" if include_false_alarms
                 else "false alarms excluded")
    document.add_paragraph(
        f"Counting {classification.prose}. Scope: {'; '.join(scope)}. Years are the "
        f"Guatemalan calendar year of each fire's own instant "
        f"({guatemala_inab.DEFAULT_TIME_ZONE}), this source publishing no year field."
    )
    document.add_paragraph(
        "This is not a report of causes. INAB publishes no cause for any fire — none "
        "of its thirty-three attributes says why a fire started — so there is no "
        "counts-by-cause report for Guatemala as there is for Portugal, Spain and "
        "Canada. What is counted here is a published vocabulary describing the fire or "
        "the report."
    )
    document.add_paragraph(
        "Classified counts the fires that carry a value at all, and every percentage "
        "after it is a share of that rather than of Fires. Read Classified (%) first: "
        "where it is low, the breakdown beside it describes only that fraction of the "
        "fires, and it is a sample of who filled the form in rather than of Guatemalan "
        "fire. Where nothing is classified there is no percentage to give and the cell "
        "is left empty."
    )
    if classification.published is None:
        document.add_paragraph(
            "The columns of this classification come from the data in scope, most "
            "frequent first, because the provider publishes no list of its values. Two "
            "runs over different years can therefore have different columns."
        )
    else:
        document.add_paragraph(
            "The columns are the published vocabulary, in its published order, so a "
            "value no fire in scope carries still gets a column holding zero and two "
            "runs have the same header. A value found in the data but not in that "
            "vocabulary is counted in a column of its own at the end."
        )

    table = document.add_table(rows=1, cols=len(headings))
    table.style = "Table Grid"
    for cell, heading in zip(table.rows[0].cells, headings):
        cell.text = heading
        cell.paragraphs[0].runs[0].bold = True

    for row in rows:
        cells = table.add_row().cells
        for index, (cell, value) in enumerate(zip(cells, row.readable_values)):
            cell.text = value
            paragraph = cell.paragraphs[0]
            if index >= FIRST_NUMERIC_COLUMN:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            for run in paragraph.runs:
                run.bold = row.is_total
                run.font.size = Pt(9)

    path.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(path))
    logger.info("Wrote %s", path)


def list_classifications() -> None:
    """Print what each classification counts."""
    print()
    for classification in CLASSIFICATIONS.values():
        published = ("no published vocabulary; columns come from the data"
                     if classification.published is None
                     else f"{len(classification.published)} published value(s)")
        print(f"  {classification.key:<14} {classification.published_name} "
              f"(stored as {classification.column})")
        print(f"  {'':<14} {classification.prose}")
        print(f"  {'':<14} {published}")
        print()
    print("  None of these is a cause. INAB publishes none.\n")


def report(args: argparse.Namespace, engine: Engine,
           logger: logging.Logger) -> list[Row]:
    """Compute the counts and write whichever outputs were asked for."""
    classification = CLASSIFICATIONS[args.classification]
    with Session(engine) as session:
        rows, values = compute(session, classification, args.year, logger,
                               args.include_false_alarms)

    if not rows:
        raise RuntimeError(
            f"No wildfires matched. Check --year, and that the {COUNTRY_NAME} fire "
            f"reports are imported — the published archive starts in 2023 — with "
            f"src.apps.imports.wildfires.guatemala_inab.import_wildfires."
        )
    if not values:
        # Only reachable for a classification whose columns come from the data:
        # a published vocabulary always has columns, and reports zeros in them,
        # which is an answer. An empty header is not.
        raise RuntimeError(
            f"No fire in scope carries a {classification.column}, and this "
            f"classification takes its columns from the data, so there is nothing to "
            f"break down and the report would be five columns and no breakdown. Try "
            f"--classification status, which effectively every record publishes."
        )

    if args.csv is not None:
        write_csv(rows, values, classification, args.csv, logger)
    if args.docx is not None:
        write_docx(rows, values, classification, args.docx, args.year, logger,
                   args.include_false_alarms)
    return rows


def main(argv: list[str] | None = None) -> int:
    args = parse_arguments(argv)
    logging.basicConfig(level=args.log_level,
                        format="%(asctime)s %(levelname)s %(message)s")
    logger = logging.getLogger("inab-classification")

    if args.list_classifications:
        list_classifications()
        return 0

    try:
        settings = common.resolve_database_settings(args)
    except RuntimeError as error:
        logger.error("%s", error)
        return 1

    engine = create_engine(common.database_url(settings))
    try:
        report(args, engine, logger)
    except Exception as error:  # noqa: BLE001  (the CLI boundary: report, do not traceback)
        logger.error("%s", error)
        return 1
    finally:
        engine.dispose()
    return 0


if __name__ == "__main__":
    sys.exit(main())
