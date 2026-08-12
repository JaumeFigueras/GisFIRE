#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Burnt-area statistics for the Greek Fire Service wildfire records.

Reports the archive per year: how many fires there were, the smallest single fire,
the largest single fire and the total burnt area in hectares — and, this dataset's
own pair, how many of the year's fires carry a published coordinate::

    python3 -m src.apps.statistics.wildfires.greece_ffa.wildfire_statistics \\
        --csv burnt.csv --docx burnt.docx

The seventh of the burnt-area reports, alongside
:doc:`GWIS <gwis_wildfire_statistics>`, GFA, ICNF, EGIF, DARPA and REDIAM. Its first
six columns are theirs, in their order, so the CSVs can still be concatenated on
them; the two after them are this dataset's own.

The hectares are reported, not measured
----------------------------------------

Like the EGIF report and unlike the four perimeter ones. The Greek Fire Service
publishes no polygon in any year — it is an administrative statistic — so there is no
``--area-method`` here and nothing is computed from a geometry. What it publishes is
**eight** burnt areas, one per land cover, and ``--surface`` picks which of them the
report is of.

There is no published total. ``--surface burnt``, the default, is the sum of all
eight, which is the nearest thing to one; see :data:`SURFACES` for the rest and for
the one grouping this report invents rather than reads.

The figures are hectares because the import converted them. The service publishes
στρέμματα, a tenth of a hectare each — see
:mod:`src.apps.imports.wildfires.greece_ffa.import_wildfires`.

Which year a fire counts towards
---------------------------------

The **published** year: the sheet the record was read from, which is
:attr:`~src.providers.greece_ffa.wildfire.GreeceFfaWildfire.year`. Not the year of
:attr:`~src.data_model.wildfire.Wildfire.start_date_time`.

That is the rule every report in this project follows — the ICNF one groups on the
published ``Ano``, the EGIF one on the filed ``Campania``, the Catalan one on the
published layer — for the same reason: a published yearly total is a total of what
the service filed under that year, and a report that regrouped the fires by their
clocks would not reproduce any figure the service ever published.

Here the two agree on all but a handful of rows anyway, every sheet being a year.

How many were located
---------------------

The two extra columns, and the reason they are here rather than in a footnote:
**no year before 2020 publishes a coordinate at all**. 54,491 of the 260,194 fires
have a point; the other 205,703 are locatable only by the prefecture, municipality,
forest district and locality named on them.

So ``Located`` and ``Located (%)`` are 0 and 0.00 for twenty years and then jump to
about 94%, and any spatial analysis over this dataset is an analysis of its last six
years. A reader who cannot see that in the table will assume otherwise. It is a
column and not a filter: an unlocated fire still contributes its hectares.

False alarms are excluded by default
-------------------------------------

The 2025 file publishes ``Κατηγορία Συμβάντος``, and 1,255 of its 9,043 rows are
:data:`~src.providers.greece_ffa.CATEGORY_FALSE_ALARM` — a call-out that found no
fire. They are records of a *dispatch*, so this report leaves them out and says how
many it left out. ``--include-false-alarms`` counts them.

The test is ``IS DISTINCT FROM`` and not ``<>``, which matters more than it looks:
the column is ``NULL`` for every year before 2025, and ``<>`` evaluates to ``NULL``
there, so the obvious filter would silently drop twenty-five years of the archive.

There is no country to choose, and no cause to count
------------------------------------------------------

No ``--country`` and no ``--country-source``, as for the Catalan and Andalusian
reports: the Fire Service publishes Greece's fires and nothing else, so the
``Country`` column is the constant :data:`COUNTRY_NAME` and nothing is tested against
a boundary to arrive at it.

And there is no companion counts-by-cause report, unlike
:doc:`ICNF <icnf_wildfire_causes>` and :doc:`EGIF <egif_wildfire_causes>`. **Nothing
in any of the twenty-six published sheets says why a fire started** — there is no
equivalent of EGIF's ``idcausa``, no lightning category, and no cause column of any
kind. See :mod:`src.providers.greece_ffa`.

One statement
-------------

Like the EGIF, Catalan and Andalusian reports and unlike the GWIS and GFA ones, this
is a single aggregate: 260,194 rows with no geometry work in them at all, which is
four orders of magnitude short of the twenty-million-perimeter case that made the
others read a year at a time. The ``Total`` row is arithmetic over the years, by
:func:`combine`, so the output is the same shape either way.
"""

from __future__ import annotations

import argparse
import csv
import logging
import math
import operator
import os
import sys

from dataclasses import dataclass
from functools import reduce
from pathlib import Path

from sqlalchemy import ColumnElement
from sqlalchemy import Engine
from sqlalchemy import Select
from sqlalchemy import create_engine
from sqlalchemy import func
from sqlalchemy import literal
from sqlalchemy import or_
from sqlalchemy import select
from sqlalchemy.orm import Session

import src.settings  # noqa: F401  (imported for the side effect of loading .env)

from src.apps.imports import common
from src.data_model.wildfire import Wildfire
from src.providers import greece_ffa
from src.providers.greece_ffa.wildfire import GreeceFfaWildfire

#: Label used in the ``Year`` column for the summary row.
TOTAL_LABEL = "Total"

#: The six columns this report shares with the GWIS, GFA, ICNF, EGIF, DARPA and
#: REDIAM ones, in their order, so the CSVs can still be concatenated on them.
SHARED_COLUMNS = ("Country", "Year", "Fires", "Minimum (ha)", "Maximum (ha)", "Total (ha)")

#: The report's columns, in order, shared by both output formats so that a change to
#: one cannot silently leave the other behind. The last two are this dataset's own:
#: how many of the year's fires publish a coordinate, and that as a share.
COLUMNS = SHARED_COLUMNS + ("Located", "Located (%)")

#: Index of the first column that holds a number, and so is right-aligned in the
#: Word table.
FIRST_NUMERIC_COLUMN = 2

#: The country every fire in this dataset is in, and the whole of the ``Country``
#: column. Spelled as the OCHA boundaries spell it, so a row of this report sorts and
#: groups with the rows of the other six. Nothing is tested against a boundary to
#: arrive at it — see the module docstring.
COUNTRY_NAME = "Greece"

#: The eight land covers the service publishes a burnt area for, and the two
#: groupings this report offers over them.
#:
#: The eight are read straight off the row. The two composites are not published:
#:
#: ``burnt``
#:     all eight summed — the nearest thing to the total the service does not print,
#:     and the default, because "how much burnt" is the question this report is
#:     normally asked.
#: ``forest-total``
#:     ``Δάση`` plus ``Δασική Έκταση`` plus ``Άλση``, the three woodland classes.
#:     **This report's grouping and not the service's**: there is no published figure
#:     it corresponds to, and it is offered because it is the nearest analogue to
#:     EGIF's ``SuperficieTotalForestal``, which is what a comparison between the two
#:     countries would need. Treat it as this report's arithmetic, not as a Greek
#:     statistic.
SURFACE_BURNT = "burnt"
SURFACE_FOREST_TOTAL = "forest-total"
SURFACE_FOREST = "forest"
SURFACE_FOREST_LAND = "forest-land"
SURFACE_GROVE = "grove"
SURFACE_GRASSLAND = "grassland"
SURFACE_REEDS_MARSH = "reeds-marsh"
SURFACE_AGRICULTURAL = "agricultural"
SURFACE_CROP_RESIDUE = "crop-residue"
SURFACE_LANDFILL = "landfill"
SURFACES = (SURFACE_BURNT, SURFACE_FOREST_TOTAL, SURFACE_FOREST, SURFACE_FOREST_LAND,
            SURFACE_GROVE, SURFACE_GRASSLAND, SURFACE_REEDS_MARSH, SURFACE_AGRICULTURAL,
            SURFACE_CROP_RESIDUE, SURFACE_LANDFILL)

#: How each surface reads in prose, for the Word document's opening paragraph. A
#: report of grassland hectares that said only "areas in hectares" would be
#: indistinguishable from a report of forest ones.
SURFACE_PROSE = {
    SURFACE_BURNT: "all burnt area: the eight published land covers summed",
    SURFACE_FOREST_TOTAL: ("burnt woodland: Δάση plus Δασική Έκταση plus Άλση, this "
                           "report's grouping and not a published figure"),
    SURFACE_FOREST: "burnt forest (Δάση)",
    SURFACE_FOREST_LAND: "burnt forest land (Δασική Έκταση)",
    SURFACE_GROVE: "burnt groves (Άλση)",
    SURFACE_GRASSLAND: "burnt grassland (Χορτ/κές Εκτάσεις)",
    SURFACE_REEDS_MARSH: "burnt reed beds and marsh (Καλάμια - Βάλτοι)",
    SURFACE_AGRICULTURAL: "burnt agricultural land (Γεωργικές Εκτάσεις)",
    SURFACE_CROP_RESIDUE: "burnt crop residue (Υπολλείματα Καλλιεργειών)",
    SURFACE_LANDFILL: "burnt landfill (Σκουπιδότοποι)",
}

#: The year a fire counts towards: the sheet it was published in. See the module
#: docstring — the filing, not the clock.
PUBLISHED_YEAR = GreeceFfaWildfire.__table__.c.year


def reported_surface(surface: str) -> tuple[ColumnElement, ColumnElement]:
    """The burnt area of one fire in hectares, as ``(hectares, is_reported)``.

    Parameters
    ----------
    surface : str
        One of :data:`SURFACES`.

    Returns
    -------
    tuple
        The SQL expression yielding hectares, and the condition a fire has to satisfy
        for that expression to be an answer rather than a silence.

    Raises
    ------
    ValueError
        If ``surface`` is not one of :data:`SURFACES`.

    Notes
    -----
    The second element is what keeps the ``Fires`` column honest, exactly as in the
    EGIF report: every one of these columns is nullable, ``sum`` and ``min`` skip
    nulls and ``count`` does not, so a report that did not filter would count fires
    whose area it had not included. ``NULL`` here means *this year does not publish
    this column* — never zero hectares.

    In practice the filter almost never bites: all eight columns are published in all
    twenty-six years, and exactly one cell in the 260,194 rows is empty. It is here
    because the next publication is not obliged to keep that up, and because a
    ``Fires`` count that silently stopped matching its own hectares would be very
    hard to notice.

    The two composites count a fire if **any** component is reported, and the ones
    that are not contribute nothing. Requiring all of them would discard a year the
    moment it dropped a single column.
    """
    fire = GreeceFfaWildfire.__table__.c

    simple = {
        SURFACE_FOREST: fire.area_ha_forest,
        SURFACE_FOREST_LAND: fire.area_ha_forest_land,
        SURFACE_GROVE: fire.area_ha_grove,
        SURFACE_GRASSLAND: fire.area_ha_grassland,
        SURFACE_REEDS_MARSH: fire.area_ha_reeds_marsh,
        SURFACE_AGRICULTURAL: fire.area_ha_agricultural,
        SURFACE_CROP_RESIDUE: fire.area_ha_crop_residue,
        SURFACE_LANDFILL: fire.area_ha_landfill,
    }
    if surface in simple:
        column = simple[surface]
        return column, column.is_not(None)

    composite = {
        SURFACE_BURNT: (fire.area_ha_forest, fire.area_ha_forest_land, fire.area_ha_grove,
                        fire.area_ha_grassland, fire.area_ha_reeds_marsh,
                        fire.area_ha_agricultural, fire.area_ha_crop_residue,
                        fire.area_ha_landfill),
        SURFACE_FOREST_TOTAL: (fire.area_ha_forest, fire.area_ha_forest_land,
                               fire.area_ha_grove),
    }
    if surface in composite:
        components = composite[surface]
        hectares = reduce(operator.add,
                          (func.coalesce(component, 0.0) for component in components))
        return hectares, or_(*[component.is_not(None) for component in components])

    raise ValueError(
        f"unknown surface {surface!r}; expected one of {', '.join(SURFACES)}"
    )


def is_located() -> ColumnElement:
    """Whether a fire carries a published coordinate.

    Returns
    -------
    ColumnElement
        A boolean expression, ``True`` for a fire the ``Located`` column counts.

    Notes
    -----
    The link and not the geometry is what is tested. ``ignition_id`` is written only
    for a pair :func:`~src.providers.greece_ffa.is_located` accepted at import, so a
    row with a link has a real point and a row without one had none published or had
    the ``0``/``0`` the service writes for *not located*. Joining through to the
    geometry would cost a join and answer the same question.
    """
    return GreeceFfaWildfire.__table__.c.ignition_id.is_not(None)


def is_a_fire(include_false_alarms: bool = False) -> ColumnElement | None:
    """The condition excluding the 2025 false alarms, or ``None`` to keep them.

    Parameters
    ----------
    include_false_alarms : bool
        ``True`` to count the false alarms as fires. ``False``, the default, leaves
        them out.

    Returns
    -------
    ColumnElement or None
        ``None`` when nothing is to be filtered, so the caller can leave the
        ``WHERE`` off altogether rather than adding a clause that is always true.

    Notes
    -----
    ``is_distinct_from`` and not ``!=``. ``incident_category`` is ``NULL`` for every
    year before 2025, where ``!=`` evaluates to ``NULL`` rather than to true, so the
    obvious filter would drop twenty-five of the twenty-six years and report the 2025
    fires alone. That failure is silent — the report still builds, still totals, and
    is simply of the wrong archive.
    """
    if include_false_alarms:
        return None
    return GreeceFfaWildfire.__table__.c.incident_category.is_distinct_from(
        greece_ffa.CATEGORY_FALSE_ALARM
    )


def statistics_query(surface: str = SURFACE_BURNT,
                     year: int | None = None,
                     min_area: float | None = None,
                     include_false_alarms: bool = False) -> Select:
    """Build the statistics query: one row per year, newest first.

    Parameters
    ----------
    surface : str
        One of :data:`SURFACES`.
    year : int, optional
        Restrict to one published year. ``None``, the default, reports every year.
    min_area : float, optional
        Count only fires of at least this many hectares of the chosen surface.
        ``None``, the default, counts every fire.
    include_false_alarms : bool
        Count the 2025 false alarms as fires.

    Returns
    -------
    Select
        A query yielding ``country, year, minimum, maximum, total, fires, located``.
        The summary row is :func:`summarise`'s work.

    Notes
    -----
    Built against the mapped classes rather than written as SQL text, so a column
    renamed on a model breaks this at import time rather than in front of a user.

    ``greece_ffa_wildfire`` is joined by table, to keep SQLAlchemy from adding a
    polymorphic join of its own, and it has to be joined in any case: the year, the
    areas, the category and the ignition link all live on it. The parent ``wildfire``
    is joined for nothing this report reads — it is the ``FROM`` because a fire *is*
    the parent row, and the join is what makes the count a count of wildfires rather
    than of provider rows.

    The inner query computes the area once. Folded into the outer aggregate instead,
    an eight-column ``coalesce`` sum would be evaluated three times per row — for the
    minimum, the maximum and the sum.

    ``min_area`` filters the subquery's column rather than repeating the expression,
    and is applied before the aggregates rather than as a ``HAVING``: the threshold
    selects the fires the figures are computed from, it does not discard years whose
    total came out small. The located count is aggregated over the same filtered
    rows, which is what keeps ``Located`` a share of the ``Fires`` beside it.
    """
    hectares_column, reported = reported_surface(surface)
    greek = GreeceFfaWildfire.__table__

    fires = (
        select(
            literal(COUNTRY_NAME).label("country"),
            PUBLISHED_YEAR.label("year"),
            hectares_column.label("hectares"),
            is_located().label("located"),
        )
        .select_from(Wildfire)
        .join(greek, greek.c.id == Wildfire.id)
        .where(reported)
    )
    real = is_a_fire(include_false_alarms)
    if real is not None:
        fires = fires.where(real)
    if year is not None:
        fires = fires.where(PUBLISHED_YEAR == year)

    fire = fires.subquery("fire")
    statistics = (
        select(
            fire.c.country,
            fire.c.year,
            func.min(fire.c.hectares).label("minimum"),
            func.max(fire.c.hectares).label("maximum"),
            func.sum(fire.c.hectares).label("total"),
            func.count().label("fires"),
            func.count().filter(fire.c.located).label("located"),
        )
        .group_by(fire.c.country, fire.c.year)
        .order_by(fire.c.year.desc())
    )
    if min_area is not None:
        statistics = statistics.where(fire.c.hectares >= min_area)
    return statistics


def false_alarm_count(session: Session, year: int | None = None) -> int:
    """How many false alarms are in scope, for the line that reports the exclusion.

    A count of zero is the normal answer for any scope that does not reach 2025,
    which is the only year publishing the category at all.
    """
    greek = GreeceFfaWildfire.__table__
    query = (
        select(func.count())
        .select_from(greek)
        .where(greek.c.incident_category == greece_ffa.CATEGORY_FALSE_ALARM)
    )
    if year is not None:
        query = query.where(PUBLISHED_YEAR == year)
    return session.scalar(query) or 0


def share(part: int, whole: int) -> float | None:
    """``part`` as a percentage of ``whole``, or ``None`` where there is no whole.

    ``None`` and not zero: a percentage of nothing is not zero percent, it is no
    answer, and the writers turn it into an empty cell.
    """
    if not whole:
        return None
    return 100.0 * part / whole


def share_label(part: int, whole: int) -> str:
    """A percentage as it is written out, empty where there is none."""
    value = share(part, whole)
    return "" if value is None else f"{value:.2f}"


@dataclass(frozen=True)
class Row:
    """One line of the report.

    Attributes
    ----------
    country : str
        Always :data:`COUNTRY_NAME`. Nothing is tested against a boundary to arrive
        at it — the Fire Service publishes Greece's fires and nothing else.
    year : int or None
        The published year, or ``None`` for the summary row.
    minimum, maximum, total : float
        Smallest single fire, largest single fire and sum of every fire, in hectares
        of the chosen surface, as published.
    fires : int
        How many fires the three area figures were computed from.
    located : int
        How many of those ``fires`` carry a published coordinate. Always at most
        ``fires``: it is counted over the same rows. Zero for every year before 2020,
        which publishes none.
    """

    country: str
    year: int | None
    minimum: float
    maximum: float
    total: float
    fires: int
    located: int

    @property
    def is_total(self) -> bool:
        """Whether this is the summary row rather than one of the years."""
        return self.year is None

    @property
    def year_label(self) -> str:
        return TOTAL_LABEL if self.is_total else str(self.year)

    @property
    def located_share(self) -> float | None:
        """``located`` as a percentage of ``fires``: how much of the year is mappable."""
        return share(self.located, self.fires)

    @property
    def values(self) -> tuple[str, ...]:
        """The row as the CSV writes it, in :data:`COLUMNS` order."""
        return (self.country, self.year_label, str(self.fires),
                f"{self.minimum:.2f}", f"{self.maximum:.2f}", f"{self.total:.2f}",
                str(self.located), share_label(self.located, self.fires))

    @property
    def readable_values(self) -> tuple[str, ...]:
        """The row as the Word document writes it: the numbers with separators."""
        return (self.country, self.year_label, f"{self.fires:,}",
                f"{self.minimum:,.2f}", f"{self.maximum:,.2f}", f"{self.total:,.2f}",
                f"{self.located:,}", share_label(self.located, self.fires))


def combine(rows: list[Row], country: str = COUNTRY_NAME, year: int | None = None) -> Row:
    """One row summarising several: every figure taken over all of them.

    Notes
    -----
    All five decompose over a partition of the fires — a minimum of minima is a
    minimum, a sum of sums is a sum, a count of counts is a count — so the ``Total``
    row is the number a second aggregate over the same rows would have returned, and
    no fire is counted twice or left out.

    The percentage is deliberately **not** averaged: :class:`Row` recomputes it from
    the summed counts, which is the ratio of the totals rather than the mean of the
    ratios. Twenty years of 0% and six of 94% must not average to 22%.

    ``fsum`` rather than ``sum``: an exact accumulation over a handful of partial
    totals costs nothing and cannot drift from what one pass would have returned.
    """
    return Row(
        country=country,
        year=year,
        minimum=min(row.minimum for row in rows),
        maximum=max(row.maximum for row in rows),
        total=math.fsum(row.total for row in rows),
        fires=sum(row.fires for row in rows),
        located=sum(row.located for row in rows),
    )


def summarise(measured: list[Row]) -> list[Row]:
    """Build the report from the years measured: the years newest first, then the total.

    Parameters
    ----------
    measured : list of Row
        One row per year, as the statement returned them.

    Returns
    -------
    list of Row
        The years newest first with the summary row last. Empty if nothing was
        measured — a report of no fires has no total either.

    Notes
    -----
    No grouping by country and no lookup of one, unlike the four worldwide reports:
    there is exactly one country here and it is a constant, so ordering the countries
    would be sorting a list of one against the database's collation.
    """
    if not measured:
        return []
    rows = sorted(measured, key=lambda row: row.year, reverse=True)
    return rows + [combine(rows)]


def hectares(text: str) -> float:
    """Argparse type for ``--min-area``: a finite, non-negative number of hectares.

    Raises
    ------
    argparse.ArgumentTypeError
        If the text is not a number, or is negative, or is a non-finite float.

    Notes
    -----
    A bare ``type=float`` would accept ``-5``, ``nan`` and ``inf``. The first two are
    almost certainly a typo and would silently produce the unfiltered report —
    ``nan`` compares false against every area, so it would produce an empty one — and
    none of the three is a size a fire can have.
    """
    try:
        area = float(text)
    except ValueError:
        raise argparse.ArgumentTypeError(f"{text!r} is not a number of hectares")
    if not math.isfinite(area):
        raise argparse.ArgumentTypeError(f"{text!r} is not a finite number of hectares")
    if area < 0:
        raise argparse.ArgumentTypeError(
            f"a burnt area cannot be negative, and {area:g} ha is: pass 0 or more, or "
            f"leave --min-area out to count every fire"
        )
    return area


def parse_arguments(argv: list[str] | None = None) -> argparse.Namespace:
    """Parse the command line."""
    parser = argparse.ArgumentParser(
        description="Burnt-area statistics for the Greek Fire Service wildfire records, "
                    "with the number of fires that publish a coordinate.",
        epilog="Areas are the hectares the service reports, split eight ways by land "
               "cover — this dataset publishes no perimeter, so there is no "
               "--area-method. Every fire is Greek, so there is no --country and nothing "
               "is tested against a boundary. Database settings not given here are read "
               "from the environment (.env).",
    )
    parser.add_argument("-y", "--year", type=int,
                        help="restrict to one year, e.g. 2021; this is the year of the "
                             "published sheet, not of the start date")
    parser.add_argument("--surface", default=SURFACE_BURNT, choices=SURFACES,
                        help="which published burnt area to report: 'burnt' (default) is "
                             "all eight land covers summed, the nearest thing to the "
                             "total the service does not publish; 'forest-total' is the "
                             "three woodland classes summed, which is this report's "
                             "grouping and not a published figure; the rest are the "
                             "published columns one at a time")
    parser.add_argument("--min-area", type=hectares, default=None, metavar="HECTARES",
                        help="count only fires that burnt at least this many hectares of "
                             "the chosen surface; by default every fire counts")
    parser.add_argument("--include-false-alarms", action="store_true",
                        help=f"count the 2025 records categorised "
                             f"{greece_ffa.CATEGORY_FALSE_ALARM} — a call-out that found "
                             f"no fire. 1,255 of the 9,043 rows of 2025, and no earlier "
                             f"year publishes the category at all. Left out by default")

    # Accepted only so that they can be refused clearly. Anyone reaching for one has
    # copied a command line from one of the other reports, which is a reasonable
    # thing to have done, and argparse's own message would not say why this report is
    # different.
    parser.add_argument("--country", help=argparse.SUPPRESS)
    parser.add_argument("--country-source", help=argparse.SUPPRESS)
    parser.add_argument("--area-method", help=argparse.SUPPRESS)

    output = parser.add_argument_group("output", "at least one is required")
    output.add_argument("--csv", type=Path, help="write the report to this .csv")
    output.add_argument("--docx", type=Path, help="write the report to this .docx (MS Word)")

    common.add_database_arguments(parser)
    parser.add_argument("--log-level", default=os.getenv("GISFIRE_LOG_LEVEL", "INFO"),
                        choices=["CRITICAL", "ERROR", "WARNING", "INFO", "DEBUG"],
                        help="verbosity (env: GISFIRE_LOG_LEVEL, default INFO)")

    arguments = parser.parse_args(argv)
    if arguments.country is not None:
        parser.error(
            f"there is no --country here: the Fire Service publishes the fires of "
            f"{COUNTRY_NAME} and nothing else, so there is nothing to select between. "
            f"Every fire is counted and the Country column is {COUNTRY_NAME} on every "
            f"row."
        )
    if arguments.country_source is not None:
        parser.error(
            f"there is no --country-source here: these are {COUNTRY_NAME}'s own records "
            f"of its own fires, so nothing is tested against a boundary and there is "
            f"nothing for a containment test to find. Three quarters of the archive "
            f"publishes no coordinate to test in any case — see the Located column."
        )
    if arguments.area_method is not None:
        parser.error(
            "there is no --area-method here: this dataset publishes no perimeter, so "
            "nothing is measured. The hectares are the ones the service reports, and "
            "--surface picks which of the eight."
        )
    if arguments.csv is None and arguments.docx is None:
        parser.error("nothing to write: pass --csv, --docx, or both")
    return arguments


def compute(session: Session, year: int | None, logger: logging.Logger,
            surface: str = SURFACE_BURNT,
            min_area: float | None = None,
            include_false_alarms: bool = False) -> list[Row]:
    """Run the statement and return the report's rows in order.

    Notes
    -----
    One statement, under one spinner — see the module docstring. The ``Total`` row is
    arithmetic over its result, not a second query.

    The located fires are logged as well as reported, because the shape of that
    number is the first thing anyone should know about this dataset: a report whose
    ``Located`` column is zero throughout is not a broken import, it is twenty years
    of an archive that never published a coordinate.
    """
    with common.Spinner(f"Summing the burnt area ({surface}) of the {COUNTRY_NAME} fires",
                        logger):
        measured = [
            Row(country=record.country, year=record.year,
                minimum=float(record.minimum),
                maximum=float(record.maximum),
                total=float(record.total),
                fires=record.fires,
                located=record.located)
            for record in session.execute(
                statistics_query(surface, year, min_area, include_false_alarms))
        ]
        excluded = 0 if include_false_alarms else false_alarm_count(session, year)

    rows = summarise(measured)
    logger.info("Computed %d rows over %d year(s) (%s hectares as published, %s)",
                len(rows), len(measured), surface,
                "every fire" if min_area is None else f"fires of {min_area:g} ha or more")
    if excluded:
        logger.info("Excluded %d false alarm(s) (%s); pass --include-false-alarms to "
                    "count them", excluded, greece_ffa.CATEGORY_FALSE_ALARM)
    if rows:
        total = rows[-1]
        logger.info("%d of %d fire(s) publish a coordinate (%s%%)",
                    total.located, total.fires, share_label(total.located, total.fires))
        if not total.located:
            logger.warning(
                "No fire in scope publishes a coordinate, so every Located column is "
                "zero. That is the expected answer for any scope ending before %d — no "
                "earlier year publishes one — and is not a sign of a failed import",
                greece_ffa.FIRST_YEAR_WITH_COORDINATES)
    return rows


def write_csv(rows: list[Row], path: Path, logger: logging.Logger) -> None:
    """Write the report as CSV.

    The numbers go out unformatted apart from being rounded to two decimals — no
    thousands separators — because a CSV is read by another program far more often
    than by a person, and a separator would make every figure a string.
    """
    path.parent.mkdir(parents=True, exist_ok=True)
    with path.open("w", newline="", encoding="utf-8") as handle:
        writer = csv.writer(handle)
        writer.writerow(COLUMNS)
        for row in rows:
            writer.writerow(row.values)
    logger.info("Wrote %s", path)


def write_docx(rows: list[Row], path: Path, year: int | None,
               logger: logging.Logger,
               surface: str = SURFACE_BURNT,
               min_area: float | None = None,
               include_false_alarms: bool = False) -> None:
    """Write the report as a Word document.

    One table, with the summary row in bold. Numbers get thousands separators here —
    the opposite of the CSV, and for the opposite reason: this one is for reading.

    The opening paragraphs name the surface, the scope, what ``Located`` is and what
    happened to the false alarms, because a table of hectares that is in fact one
    land cover, a column of zeros that is a property of the archive rather than of
    the import, and a silently dropped 14% of 2025 are all things a reader should not
    have to remember.
    """
    # Imported here rather than at module scope so that --csv keeps working if
    # python-docx is not installed, which matters because it is the only dependency
    # this application adds.
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt

    document = Document()
    document.add_heading(f"Greek Fire Service wildfire burnt area ({COUNTRY_NAME})",
                         level=1)

    scope = [f"year: {year}" if year is not None else "all years"]
    if min_area is not None:
        scope.append(f"only fires of {min_area:g} ha or more")
    scope.append("false alarms counted" if include_false_alarms
                 else "false alarms excluded")
    document.add_paragraph(
        f"Areas in hectares as the service reports them — this dataset publishes no "
        f"perimeter, so nothing is measured. Surface: {SURFACE_PROSE[surface]}. The "
        f"service publishes στρέμματα, a tenth of a hectare each; the import converted "
        f"them. Years are the published sheet's, not the start date's. "
        f"Scope: {'; '.join(scope)}."
    )
    document.add_paragraph(
        f"The Country column is {COUNTRY_NAME} on every row and nothing is tested "
        f"against a boundary: the Fire Service publishes the fires of {COUNTRY_NAME} "
        f"and nothing else."
    )
    document.add_paragraph(
        f"Located counts the fires that publish a coordinate. It is a column and not a "
        f"filter — an unlocated fire still contributes its hectares — and it is zero "
        f"for every year before {greece_ffa.FIRST_YEAR_WITH_COORDINATES}, which is a "
        f"property of the archive and not of the import: no earlier year publishes one. "
        f"Any spatial analysis of this dataset is an analysis of its last six years."
    )
    document.add_paragraph(
        "The service publishes no cause for any fire in any year, so there is no "
        "counts-by-cause companion to this report as there is for Portugal and Spain."
    )

    table = document.add_table(rows=1, cols=len(COLUMNS))
    table.style = "Table Grid"
    for cell, heading in zip(table.rows[0].cells, COLUMNS):
        cell.text = heading
        cell.paragraphs[0].runs[0].bold = True

    for row in rows:
        cells = table.add_row().cells
        for index, (cell, value) in enumerate(zip(cells, row.readable_values)):
            cell.text = value
            paragraph = cell.paragraphs[0]
            if index >= FIRST_NUMERIC_COLUMN:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            for run in paragraph.runs:
                run.bold = row.is_total
                run.font.size = Pt(9)

    path.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(path))
    logger.info("Wrote %s", path)


def report(args: argparse.Namespace, engine: Engine, logger: logging.Logger) -> list[Row]:
    """Compute the statistics and write whichever outputs were asked for."""
    with Session(engine) as session:
        rows = compute(session, args.year, logger, args.surface, args.min_area,
                       args.include_false_alarms)

    if not rows:
        # An empty report is almost always a year with no data, and writing an empty
        # file would hide that. A threshold is named when there is one, because then
        # it is at least as likely to be the reason as the year is.
        threshold = "" if args.min_area is None else \
            f" No fire reached the --min-area of {args.min_area:g} ha."
        raise RuntimeError(
            f"No wildfires matched. Check --year, and that the {COUNTRY_NAME} fires are "
            f"imported — the published workbooks run from "
            f"{greece_ffa.FIRST_YEAR}." + threshold
        )

    if args.csv is not None:
        write_csv(rows, args.csv, logger)
    if args.docx is not None:
        write_docx(rows, args.docx, args.year, logger, args.surface, args.min_area,
                   args.include_false_alarms)
    return rows


def main(argv: list[str] | None = None) -> int:
    args = parse_arguments(argv)
    logging.basicConfig(level=args.log_level, format="%(asctime)s %(levelname)s %(message)s")
    logger = logging.getLogger("greece-ffa-statistics")

    try:
        settings = common.resolve_database_settings(args)
    except RuntimeError as error:
        logger.error("%s", error)
        return 1

    engine = create_engine(common.database_url(settings))
    try:
        report(args, engine, logger)
    except Exception as error:  # noqa: BLE001  (the CLI boundary: report, do not traceback)
        logger.error("%s", error)
        return 1
    finally:
        engine.dispose()
    return 0


if __name__ == "__main__":  # pragma nocover
    sys.exit(main())
