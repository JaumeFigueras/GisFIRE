#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Wildfire counts by cause for the Canadian NFDB agency fire reports.

Reports, per country and year, how many fires the agencies filed, how many have a
**determined** cause, how many of those were natural, and how many hectares those
natural fires burnt as the agencies reported them::

    Country  Year    Fires  Determined  Natural  Natural (%)  Natural (ha)  Natural (% of ha)
    Canada   2023     6830        6517     3825        58.69   16220878.79              97.03
    Canada   1995     8451        8250     3518        42.64    7002347.75              95.34
    Canada   Total  382047      374651   172430        46.02  129128595.02              90.70

(``--country-source filed``, which is the fast way to ask this dataset about causes;
the default tests every point against the country polygons and takes considerably
longer to reach the same conclusion about all but a handful of fires.)

Run it over everything, or narrow it to one year, one agency, or the human-caused
fires::

    python3 -m src.apps.statistics.wildfires.canada_nfdb.wildfire_causes --csv causes.csv
    python3 -m src.apps.statistics.wildfires.canada_nfdb.wildfire_causes \\
        --year 2023 --csv 2023.csv --docx 2023.docx
    python3 -m src.apps.statistics.wildfires.canada_nfdb.wildfire_causes \\
        --agency NT --csv northwest-territories.csv
    python3 -m src.apps.statistics.wildfires.canada_nfdb.wildfire_causes \\
        --cause human --csv human.csv

At least one of ``--csv`` and ``--docx`` is required.

The companion of
:mod:`~src.apps.statistics.wildfires.canada_nfdb.wildfire_statistics`, over the same
fires, the same years and the same scope — so the ``Country``, ``Year`` and ``Fires``
columns of the two reports agree row for row. What that one sums in hectares, this one
counts by cause. It is also the twin of
:mod:`the NBAC report <src.apps.statistics.wildfires.canada_nbac.wildfire_causes>`,
over the same fires seen from the other side, and the two are meant to be read
together.

There is no lightning in this dataset either
----------------------------------------------

``CAUSE`` takes three values — ``N``, ``H`` and ``U`` — and
:data:`~src.providers.canada_nfdb.CAUSE_NATURAL` is the nearest thing to a lightning
category. It is dominated by lightning in the Canadian boreal and it **is not defined
as lightning**, so the column is headed ``Natural`` and never ``Lightning``.

The heading is the word and not the published letter (:data:`CAUSE_LABELS`): a column
headed ``N`` says nothing to a reader, and the word is the one the NBAC report
publishes, which is what lets the two CSVs be concatenated and compared.

At 172,430 fires it is by a wide margin the largest natural-cause set in GisFIRE, and
unlike NBAC's every one of them has a point and a date. That is what this dataset is
here for.

Half the fires, nine tenths of the hectares
---------------------------------------------

Over the whole archive, **46.02% of the fires with a determined cause are natural and
they account for 90.70% of the reported area.**

The companion :mod:`perimeter report
<src.apps.statistics.wildfires.canada_nbac.wildfire_causes>` puts the same two figures
at **67.18% and 90.54%** — twenty-one points apart on the fires and two tenths of a
point apart on the area.

.. important::

   The two archives are not contradicting each other. They count different
   populations: NBAC maps what burnt, so it is dominated by large remote fires, and
   NFDB records every call-out thirteen agencies filed, which includes tens of
   thousands of small human-caused fires near roads and towns that were never worth
   mapping. Two thirds of this archive is under one hectare.

   That the two disagree so much about the *count* and so little about the *area* is
   the most useful thing either of them says. It means the answer to "how much of
   Canada's burnt area is natural-cause" does not depend on which archive you ask, and
   the answer is about nine hectares in ten.

Why the denominator is the **determined** fires
------------------------------------------------

Like NBAC, ``U`` is a published category and not a missing value. Unlike NBAC it is
small — 7,396 fires, under 2% — so the choice of denominator changes this report much
less than it changes that one.

It is made the same way regardless, for two reasons: so the two reports' percentages
are the same kind of number and can be compared at all, and because ``U`` is **not
evenly spread between the agencies**, which is where it bites here:

==========  ========  ==========================  ==============
Agency      Fires     Natural, % of determined    ``U``, % of all
==========  ========  ==========================  ==============
NT            13,646                        83.0             2.7
PC             3,660                        65.0             3.9
YT             6,713                        62.4             3.1
BC           109,377                        53.6             0.4
MB            23,634                        50.9             0.1
ON            65,650                        49.0             2.0
SK            26,990                        47.2             0.2
AB            60,483                        44.6             2.9
QC            43,886                        28.4             0.0
NB            12,181                        10.0            12.2
NL             4,695                         9.8             0.3
NS            11,078                         2.6            14.7
PE                54                         0.0             0.0
==========  ========  ==========================  ==============

Where nothing is determined there is no percentage to give, and the cell is left
**empty** rather than filled with a zero that would be a claim.

.. warning::

   **The national figure is a weighted average of thirteen different fire regimes, and
   the weights are reporting volumes rather than areas.** The Northwest Territories
   file 83% natural and Nova Scotia 2.6%; British Columbia alone contributes 109,377 of
   the 382,047 rows, so the national percentage is largely British Columbia's.

   ``--agency`` is how to get a number that means one thing. A trend across the whole
   archive is partly a trend in which agencies were filing.

Why ``unknown`` cannot be counted
-----------------------------------

``--cause`` takes ``natural`` or ``human`` and not the third value, for the reason the
NBAC report gives: ``U`` is the complement of the denominator, so its share of the
determined fires is zero by construction. ``Fires`` minus ``Determined`` is exactly the
unknown count, and the report logs it.

The hectares are the agencies' own
------------------------------------

:attr:`~src.providers.canada_nfdb.wildfire.NfdbWildfire.size_ha`, as filed, and **not**
NBAC's mapped area: this dataset publishes no perimeter. A reported zero counts as a
fire and contributes nothing to the hectares, which is what it is.

Which fires are counted
-----------------------

Exactly the rule its companion report uses, so the two agree: every imported fire, with
the declared prescribed burns excluded unless ``--include-prescribed`` says otherwise,
and — under the default ``--country-source geometry`` — only those whose published
point falls inside a country. A fire over the American border is counted under the
United States, as it is there.

One year at a time
------------------

One statement per year, as everywhere else, and here it is the point: under
``--country-source geometry`` each fire is a point-in-polygon test against country
polygons of millions of vertices, and the memory that goes into them is only released
when the statement ends.

Shared with the companion report
---------------------------------

The years query, the scope conditions, the country resolution, the agency lookup and
the country ordering are **imported from**
:mod:`~src.apps.statistics.wildfires.canada_nfdb.wildfire_statistics` rather than
copied. Two reports over one dataset that disagreed about which fires are in scope, or
about which country one is in, would be worse than one report.
"""

from __future__ import annotations

import argparse
import csv
import logging
import math
import os
import sys

from dataclasses import dataclass
from pathlib import Path

from sqlalchemy import Engine
from sqlalchemy import Select
from sqlalchemy import create_engine
from sqlalchemy import func
from sqlalchemy import select
from sqlalchemy.orm import Session

import src.settings  # noqa: F401  (imported for the side effect of loading .env)

from src.apps.imports import common
from src.apps.statistics.wildfires.canada_nfdb.wildfire_statistics import COUNTRY_NAME
from src.apps.statistics.wildfires.canada_nfdb.wildfire_statistics import COUNTRY_SOURCES
from src.apps.statistics.wildfires.canada_nfdb.wildfire_statistics import (
    COUNTRY_SOURCE_GEOMETRY,
)
from src.apps.statistics.wildfires.canada_nfdb.wildfire_statistics import (
    FIRST_NUMERIC_COLUMN,
)
from src.apps.statistics.wildfires.canada_nfdb.wildfire_statistics import PUBLISHED_YEAR
from src.apps.statistics.wildfires.canada_nfdb.wildfire_statistics import SIZE_HA
from src.apps.statistics.wildfires.canada_nfdb.wildfire_statistics import TOTAL_LABEL
from src.apps.statistics.wildfires.canada_nfdb.wildfire_statistics import country_columns
from src.apps.statistics.wildfires.canada_nfdb.wildfire_statistics import ordered_countries
from src.apps.statistics.wildfires.canada_nfdb.wildfire_statistics import resolve_agency
from src.apps.statistics.wildfires.canada_nfdb.wildfire_statistics import scope_conditions
from src.apps.statistics.wildfires.canada_nfdb.wildfire_statistics import years_query
from src.providers import canada_nfdb
from src.providers.canada_nfdb.wildfire import NfdbWildfire

#: The causes this report can count, keyed to the published ``CAUSE`` value.
#:
#: :data:`~src.providers.canada_nfdb.CAUSE_UNKNOWN` is deliberately not here: it is the
#: complement of this report's denominator rather than one of the things being compared
#: against it. See the module docstring.
COUNTABLE_CAUSES = {
    "natural": canada_nfdb.CAUSE_NATURAL,
    "human": canada_nfdb.CAUSE_HUMAN,
}

#: How each cause is written in a column heading.
#:
#: Separate from :data:`COUNTABLE_CAUSES` because the published value is a single
#: letter and a column headed ``N`` says nothing to a reader. The word is also the one
#: :mod:`the NBAC report <src.apps.statistics.wildfires.canada_nbac.wildfire_causes>`
#: publishes, which is what lets the two CSVs be concatenated and compared — the whole
#: point of the pair.
#:
#: ``Natural`` and never ``Lightning``: see the module docstring.
CAUSE_LABELS = {
    "natural": "Natural",
    "human": "Human",
}

#: The cause counted unless ``--cause`` says otherwise, and the nearest thing this
#: dataset publishes to a lightning fire. It is **not** a lightning category.
DEFAULT_CAUSE = "natural"

#: The causes that count towards the denominator: the ones somebody determined.
DETERMINED_CAUSES = (canada_nfdb.CAUSE_NATURAL, canada_nfdb.CAUSE_HUMAN)


def columns(cause: str = DEFAULT_CAUSE) -> tuple[str, ...]:
    """The report's columns, in order, for a given cause.

    A function rather than a constant, for the reason the NBAC report's is: four of the
    headings name the cause being counted, and a file of ``H`` counts under a heading
    saying ``N`` would be a trap. The first three are the companion report's first
    three, unchanged.
    """
    label = CAUSE_LABELS[cause]
    return ("Country", "Year", "Fires", "Determined", label, f"{label} (%)",
            f"{label} (ha)", f"{label} (% of ha)")


def counts_query(year: int, cause: str = DEFAULT_CAUSE,
                 country_source: str = COUNTRY_SOURCE_GEOMETRY,
                 include_prescribed: bool = False,
                 agency: str | None = None) -> Select:
    """Build the counting query for one year.

    Returns
    -------
    Select
        A query yielding ``country, fires, determined, matching, hectares,
        determined_hectares``: one row per country the year's fires turn out to be in.

    Notes
    -----
    Five aggregates in one pass: whatever the country source costs, it is paid once per
    fire and not once per count.

    The cause is what this report counts, so it is never a filter — only the prescribed
    and agency halves of the companion's scope are applied, which is what keeps the
    ``Fires`` column equal to that report's.

    Both hectare figures are filtered sums over the reported ``SIZE_HA``, and the
    denominator one is restricted to the determined fires for the same reason the count
    is: the unknown hectares belong to fires nobody classified.
    """
    if cause not in COUNTABLE_CAUSES:
        raise ValueError(
            f"unknown cause {cause!r}; expected one of {', '.join(COUNTABLE_CAUSES)}"
        )
    nfdb = NfdbWildfire.__table__
    country_name, joins = country_columns(country_source)
    determined = nfdb.c.fire_cause.in_(DETERMINED_CAUSES)
    matching = nfdb.c.fire_cause == COUNTABLE_CAUSES[cause]

    counts = (
        select(
            country_name.label("country"),
            func.count().label("fires"),
            func.count().filter(determined).label("determined"),
            func.count().filter(matching).label("matching"),
            func.coalesce(func.sum(SIZE_HA).filter(matching), 0.0).label("hectares"),
            func.coalesce(func.sum(SIZE_HA).filter(determined),
                          0.0).label("determined_hectares"),
        )
        .select_from(nfdb)
        .where(PUBLISHED_YEAR == year)
    )
    for target, condition in joins:
        counts = counts.join(target, condition)
    for condition in scope_conditions(include_prescribed, None, agency):
        counts = counts.where(condition)
    return counts.group_by(country_name)


@dataclass(frozen=True)
class Row:
    """One line of the report.

    Attributes
    ----------
    country : str
        The country the fire was reported in — ``Canada`` under ``--country-source
        filed``, and whichever country contains the point under ``geometry``.
    year : int or None
        The published year, or ``None`` for a summary row.
    fires : int
        Every fire counted, determined or not.
    determined : int
        How many have a cause somebody determined. The denominator; the difference from
        :attr:`fires` is the unknown count.
    matching : int
        How many carry the cause asked for.
    hectares, determined_hectares : float
        What those fires reported burning, and what the determined ones did.
    """

    country: str
    year: int | None
    fires: int
    determined: int
    matching: int
    hectares: float
    determined_hectares: float

    @property
    def is_total(self) -> bool:
        return self.year is None

    @property
    def year_label(self) -> str:
        return TOTAL_LABEL if self.is_total else str(self.year)

    @property
    def unknown(self) -> int:
        """How many fires nobody classified. Not a column; logged, and derivable."""
        return self.fires - self.determined

    @property
    def share(self) -> float | None:
        """``matching`` as a percentage of ``determined``, or ``None`` if none is."""
        if not self.determined:
            return None
        return 100.0 * self.matching / self.determined

    @property
    def share_label(self) -> str:
        return "" if self.share is None else f"{self.share:.2f}"

    @property
    def area_share(self) -> float | None:
        """The matching hectares as a percentage of the determined ones.

        The figure the NBAC report corroborates: 90.70% here against its 90.54%, where
        the two disagree by twenty-one points about the fires.
        """
        if not self.determined_hectares:
            return None
        return 100.0 * self.hectares / self.determined_hectares

    @property
    def area_share_label(self) -> str:
        return "" if self.area_share is None else f"{self.area_share:.2f}"

    @property
    def values(self) -> tuple[str, ...]:
        """The row as the CSV writes it, in :func:`columns` order."""
        return (self.country, self.year_label, str(self.fires), str(self.determined),
                str(self.matching), self.share_label,
                f"{self.hectares:.2f}", self.area_share_label)

    @property
    def readable_values(self) -> tuple[str, ...]:
        """The row as the Word document writes it: the numbers with separators."""
        return (self.country, self.year_label, f"{self.fires:,}", f"{self.determined:,}",
                f"{self.matching:,}", self.share_label,
                f"{self.hectares:,.2f}", self.area_share_label)


def combine(rows: list[Row], country: str = COUNTRY_NAME,
            year: int | None = None) -> Row:
    """One row summarising several: the counts and the hectares added up.

    Counts and sums decompose over a partition of the fires, so a summary row is what a
    single pass over the same fires would have returned. Both shares are recomputed
    from the summed counts rather than averaged — the ratio of the totals, not the mean
    of the ratios.
    """
    return Row(
        country=country,
        year=year,
        fires=sum(row.fires for row in rows),
        determined=sum(row.determined for row in rows),
        matching=sum(row.matching for row in rows),
        hectares=math.fsum(row.hectares for row in rows),
        determined_hectares=math.fsum(row.determined_hectares for row in rows),
    )


def summarise(measured: list[Row], countries: list[str]) -> list[Row]:
    """Build the report: each country, its years newest first and its summary row last.

    No World block, exactly as in the companion report: this is one country's archive
    plus whatever fell over its border.
    """
    report: list[Row] = []
    for name in countries:
        rows = [row for row in measured if row.country == name]
        if not rows:
            continue
        report += sorted(rows, key=lambda row: row.year, reverse=True)
        report.append(combine(rows, name, None))
    return report


def parse_arguments(argv: list[str] | None = None) -> argparse.Namespace:
    """Parse the command line."""
    parser = argparse.ArgumentParser(
        description="Wildfire counts by cause for the Canadian NFDB agency fire "
                    "reports, with the area those fires reported burning.",
        epilog="The NFDB publishes no lightning category; 'N' is the closest it has and "
               "is what is counted by default. The percentages are of the fires whose "
               "cause somebody determined. The natural share varies enormously between "
               "the thirteen agencies, so --agency is how to get a number that means "
               "one thing. Database settings not given here are read from the "
               "environment (.env).",
    )
    parser.add_argument("-y", "--year", type=int,
                        help="restrict to one year, e.g. 2023; this is the published "
                             "YEAR, not the year of the report date")
    parser.add_argument("--cause", default=DEFAULT_CAUSE, choices=sorted(COUNTABLE_CAUSES),
                        help="which published CAUSE to count: 'natural' (default) is the "
                             "nearest this dataset comes to a lightning fire, and 'human' "
                             "is the other determined cause. 'U' cannot be counted: it is "
                             "the complement of the denominator, so its share of the "
                             "determined fires is zero by construction")
    parser.add_argument("-a", "--agency", default=None, metavar="CODE",
                        help="restrict to one filing agency — 'BC', 'NT', 'NS', … Case "
                             "does not matter. Worth reaching for: the natural share runs "
                             "from 83%% in the Northwest Territories to 2.6%% in Nova "
                             "Scotia, so the national figure is a weighted average of "
                             "thirteen different fire regimes")
    parser.add_argument("--include-prescribed", action="store_true",
                        help="count the fires flagged PRESCRIBED, as the companion "
                             "report's option of the same name does. Left out by default")
    parser.add_argument("--country-source", default=COUNTRY_SOURCE_GEOMETRY,
                        choices=COUNTRY_SOURCES,
                        help="how to decide which country a fire counts towards: "
                             "'geometry' (default) tests the published point against the "
                             "real country polygons, so a coordinate in the sea drops out "
                             "and one over the American border is counted as the United "
                             "States; 'filed' takes the agency's word for it")

    # Accepted only so that it can be refused clearly, exactly as in the companion
    # report: without it argparse's prefix matching resolves "--country" to
    # "--country-source" and complains about an invalid choice instead.
    parser.add_argument("--country", help=argparse.SUPPRESS)

    output = parser.add_argument_group("output", "at least one is required")
    output.add_argument("--csv", type=Path, help="write the report to this .csv")
    output.add_argument("--docx", type=Path, help="write the report to this .docx (MS Word)")

    common.add_database_arguments(parser)
    parser.add_argument("--log-level", default=os.getenv("GISFIRE_LOG_LEVEL", "INFO"),
                        choices=["CRITICAL", "ERROR", "WARNING", "INFO", "DEBUG"],
                        help="verbosity (env: GISFIRE_LOG_LEVEL, default INFO)")

    arguments = parser.parse_args(argv)
    if arguments.country is not None:
        parser.error(
            f"there is no --country here: these are {COUNTRY_NAME}'s own agencies' "
            f"reports of their own fires, so there is nothing to select between. Every "
            f"fire is counted, and the Country column says which country each one's "
            f"point turned out to be in — see --country-source."
        )
    if arguments.csv is None and arguments.docx is None:
        parser.error("nothing to write: pass --csv, --docx, or both")
    return arguments


def compute(session: Session, year: int | None, logger: logging.Logger,
            cause: str = DEFAULT_CAUSE,
            country_source: str = COUNTRY_SOURCE_GEOMETRY,
            include_prescribed: bool = False,
            agency: str | None = None) -> list[Row]:
    """Count the fires a year at a time, returning the report's rows in order.

    One statement per year, each under a spinner of its own and all of them in
    ``session``'s transaction, so a report assembled from fifty-three queries is as
    consistent as one assembled from a single query.
    """
    if year is not None:
        years = [year]
    else:
        with common.Spinner("Finding the years the NFDB reports cover", logger):
            years = list(session.scalars(years_query(include_prescribed, None, agency)))

    scope = "every agency" if agency is None else agency
    measured: list[Row] = []
    for index, counting in enumerate(years, start=1):
        with common.Spinner(f"Counting the NFDB fires by cause "
                            f"({cause}, {scope}, {counting}: {index} of {len(years)})",
                            logger):
            measured += [
                Row(country=record.country, year=counting,
                    fires=record.fires,
                    determined=record.determined,
                    matching=record.matching,
                    hectares=float(record.hectares),
                    determined_hectares=float(record.determined_hectares))
                for record in session.execute(
                    counts_query(counting, cause, country_source, include_prescribed,
                                 agency))
            ]

    countries = ordered_countries(session, {row.country for row in measured})
    rows = summarise(measured, countries)
    logger.info("Counted %d rows over %d country/countries and %d year(s) "
                "(%s fires, country from %s, %s)",
                len(rows), len(countries), len({row.year for row in measured}),
                cause, country_source, scope)
    if rows:
        total = combine([row for row in measured], COUNTRY_NAME, None) if measured \
            else rows[-1]
        logger.info("%d of %d fire(s) have a determined cause; %d of those are %s (%s%%)",
                    total.determined, total.fires, total.matching, cause,
                    total.share_label)
        logger.info("They reported %.0f ha of the %.0f ha whose cause was determined "
                    "(%s%%)", total.hectares, total.determined_hectares,
                    total.area_share_label)
        if total.unknown:
            logger.info("%d fire(s) in scope are %s and are in neither figure",
                        total.unknown, canada_nfdb.CAUSE_UNKNOWN)
        if agency is None:
            # A plain '%' and not '%%': logging only applies %-formatting when there
            # are arguments to substitute, and there are none here.
            logger.info("This is a national figure over thirteen agencies whose natural "
                        "share runs from 83% to 3%; pass --agency for one that means "
                        "one thing")
    return rows


def write_csv(rows: list[Row], path: Path, logger: logging.Logger,
              cause: str = DEFAULT_CAUSE) -> None:
    """Write the report as CSV.

    The percentages are written bare and are **empty** where nothing was determined: an
    empty field reads as no answer, which is what it is, while a zero would read as one.
    """
    path.parent.mkdir(parents=True, exist_ok=True)
    with path.open("w", newline="", encoding="utf-8") as handle:
        writer = csv.writer(handle)
        writer.writerow(columns(cause))
        for row in rows:
            writer.writerow(row.values)
    logger.info("Wrote %s", path)


def write_docx(rows: list[Row], path: Path, year: int | None,
               logger: logging.Logger,
               cause: str = DEFAULT_CAUSE,
               country_source: str = COUNTRY_SOURCE_GEOMETRY,
               include_prescribed: bool = False,
               agency: str | None = None) -> None:
    """Write the report as a Word document.

    The opening paragraphs say that the NFDB publishes no lightning category, what the
    denominator is, that the national figure is a weighted average of thirteen very
    different agencies, and how the area share compares with NBAC's. A table of natural
    fires read without them would be misread several ways over.
    """
    # Imported here rather than at module scope so that --csv keeps working if
    # python-docx is not installed.
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt

    label = CAUSE_LABELS[cause]
    document = Document()
    where = COUNTRY_NAME if agency is None else f"{agency}, {COUNTRY_NAME}"
    document.add_heading(f"NFDB wildfires by cause ({where})", level=1)

    scope = [f"year: {year}" if year is not None else "all years"]
    if agency is not None:
        scope.append(f"only the fires filed by {agency}")
    scope.append("prescribed burns counted" if include_prescribed
                 else "declared prescribed burns excluded")
    if country_source == COUNTRY_SOURCE_GEOMETRY:
        scope.append("only fires whose published point falls inside a country")
    document.add_paragraph(
        f"Counts of the agency fire reports whose published CAUSE is {label}, and the "
        f"area they reported burning (SIZE_HA) — not a measured perimeter, which this "
        f"dataset does not publish. Years are the published YEAR. "
        f"Scope: {'; '.join(scope)}."
    )
    document.add_paragraph(
        f"The NFDB publishes no lightning category. {canada_nfdb.CAUSE_NATURAL} is the "
        f"nearest it comes, it is dominated by lightning in the Canadian boreal, and it "
        f"is not defined as lightning — which is why the column is headed "
        f"{canada_nfdb.CAUSE_NATURAL} and not Lightning."
    )
    document.add_paragraph(
        f"The percentages are of the fires whose cause somebody determined, not of all "
        f"of them. {canada_nfdb.CAUSE_UNKNOWN} is a published category rather than a "
        f"missing value; Fires minus Determined is the unknown count."
    )
    if agency is None:
        document.add_paragraph(
            "This is a national figure and it is a weighted average of thirteen "
            "different fire regimes, weighted by how much each agency files rather than "
            "by area: the Northwest Territories report 83% natural and Nova Scotia 2.6%, "
            "and British Columbia alone contributes 109,377 of the 382,047 rows. Use the "
            "agency option for a number that means one thing."
        )
    if cause == DEFAULT_CAUSE:
        document.add_paragraph(
            "Note the two percentages against each other. Over the archive 46.02% of the "
            "determined fires are natural and they account for 90.70% of the reported "
            "area. The companion NBAC perimeter report, built from imagery rather than "
            "from the agencies' records, puts the same two figures at 67.18% and 90.54% "
            "— so the two archives disagree by twenty-one points about the fires and "
            "agree to two tenths of a point about the area. They count different "
            "populations of fire and the same hectares."
        )

    table = document.add_table(rows=1, cols=len(columns(cause)))
    table.style = "Table Grid"
    for cell, heading in zip(table.rows[0].cells, columns(cause)):
        cell.text = heading
        cell.paragraphs[0].runs[0].bold = True

    for row in rows:
        cells = table.add_row().cells
        for index, (cell, value) in enumerate(zip(cells, row.readable_values)):
            cell.text = value
            paragraph = cell.paragraphs[0]
            if index >= FIRST_NUMERIC_COLUMN:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            for run in paragraph.runs:
                run.bold = row.is_total
                run.font.size = Pt(9)

    path.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(path))
    logger.info("Wrote %s", path)


def report(args: argparse.Namespace, engine: Engine, logger: logging.Logger) -> list[Row]:
    """Count the causes and write whichever outputs were asked for.

    The agency is resolved first, inside the same session, so that a code nobody
    recognises fails before any fire is counted — and against the database, so the
    error can list the agencies that really are imported.
    """
    with Session(engine) as session:
        agency = None if args.agency is None else resolve_agency(session, args.agency)
        rows = compute(session, args.year, logger, args.cause, args.country_source,
                       args.include_prescribed, agency)

    if not rows:
        extra = "" if agency is None else \
            f" The report is of {agency}'s fires alone, and the agencies do not all " \
            f"cover the same years."
        raise RuntimeError(
            f"No wildfires matched. Check --year, and that the NFDB fires are imported "
            f"— the import reads from {canada_nfdb.FIRST_YEAR} on." + extra
        )

    if args.csv is not None:
        write_csv(rows, args.csv, logger, args.cause)
    if args.docx is not None:
        write_docx(rows, args.docx, args.year, logger, args.cause, args.country_source,
                   args.include_prescribed, agency)
    return rows


def main(argv: list[str] | None = None) -> int:
    args = parse_arguments(argv)
    logging.basicConfig(level=args.log_level, format="%(asctime)s %(levelname)s %(message)s")
    logger = logging.getLogger("nfdb-causes")

    try:
        settings = common.resolve_database_settings(args)
    except RuntimeError as error:
        logger.error("%s", error)
        return 1

    engine = create_engine(common.database_url(settings))
    try:
        report(args, engine, logger)
    except Exception as error:  # noqa: BLE001  (the CLI boundary: report, do not traceback)
        logger.error("%s", error)
        return 1
    finally:
        engine.dispose()
    return 0


if __name__ == "__main__":  # pragma nocover
    sys.exit(main())
