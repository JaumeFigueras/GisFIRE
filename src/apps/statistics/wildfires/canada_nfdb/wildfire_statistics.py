#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Burnt-area statistics for the Canadian National Fire Database agency fire data.

Reports, per country and year, how many fires the agencies filed and the smallest,
largest and total area they reported, in hectares — and, this dataset's own column,
how many agencies filed anything at all that year::

    Country   Year   Fires   Minimum (ha)   Maximum (ha)    Total (ha)  Agencies
    Canada    2025    5731           0.00      404069.00   6132908.13        12
    Canada    2024    6019           0.00      369950.00   5313994.62        12
    ...
    Canada    1973    4392           0.00       89030.00    849382.53         9
    Canada    Total  381...          0.00     1050000.00 146...              13

Run it over everything, or narrow it to one year, to the fires above a size, to one
cause or to one filing agency::

    python3 -m src.apps.statistics.wildfires.canada_nfdb.wildfire_statistics --csv burnt.csv
    python3 -m src.apps.statistics.wildfires.canada_nfdb.wildfire_statistics \\
        --year 2023 --csv 2023.csv --docx 2023.docx
    python3 -m src.apps.statistics.wildfires.canada_nfdb.wildfire_statistics \\
        --min-area 200 --csv large-fires.csv
    python3 -m src.apps.statistics.wildfires.canada_nfdb.wildfire_statistics \\
        --cause natural --agency BC --csv bc-lightning.csv

At least one of ``--csv`` and ``--docx`` is required: an application that computed a
report and then printed nothing would be a strange thing.

The application only reads. Database settings come from the environment (``.env``,
see :mod:`src.settings`); every one of them can be overridden with a command-line
argument.

The ninth of the burnt-area reports and the companion of
:mod:`NBAC's <src.apps.statistics.wildfires.canada_nbac.wildfire_statistics>`: the
same country, the same fires, two different measurements of them. Its first six
columns are the other eight reports', in their order, so the CSVs can still be
concatenated on them; the seventh is this dataset's own.

The area is reported, not measured
-----------------------------------

:attr:`~src.providers.canada_nfdb.wildfire.NfdbWildfire.size_ha` is what the agency
filed, and this dataset publishes **no perimeter, ever** — see
:mod:`src.providers.canada_nfdb`. So there is no ``--area-method`` and no
``--surface``: nothing is projected, nothing is measured, and there is exactly one
burnt area per fire.

.. warning::

   **This report and the NBAC one do not agree, and neither is wrong.** Over the
   archive the agencies' reported sizes sum to about 166.5 million hectares against
   NBAC's 132.7 million of mapped burn. One is what somebody recorded at the time,
   the other is what a satellite could see afterwards. Do not quote one as a
   correction of the other, and do not add them: they are the same fires counted
   twice.

A reported **zero is counted** and is a real answer — tens of thousands of rows
carry it, and two thirds of the archive is under one hectare. That is why a
``Minimum (ha)`` of ``0.00`` here is normal and not a missing value, and it is what
``--min-area`` is for: the service's own *large fire* distribution keeps only the
fires of 200 ha or more, which are a twentieth of the rows and over 97% of the area.

Which country a fire counts towards, and the points outside Canada
-------------------------------------------------------------------

This is the half of the report that matters most, and the reason the default here is
the opposite of the Spanish one's.

**Every stored fire is a point**, not a polygon, and the points are agency reports:
the published summary says outright that *"locations are approximate"*. The import's
only geometric guard is a plausibility box round Canada
(:func:`~src.providers.canada_nfdb.is_located`), and a box round a country is not the
country — it contains a great deal of the United States, of Greenland and of three
oceans. A good many published coordinates land there, which is visible the moment
the layer is opened over a basemap.

So ``--country-source`` chooses what to do about it:

``geometry`` (the default)
    Ask the database which country actually contains the point, at report time,
    against the real OCHA polygons. A fire inside no country is dropped and a fire
    over the American border is reported as a **United States of America** row rather
    than folded into Canada's. This is the default because the alternative quietly
    credits Canada with fires whose coordinates say otherwise.
``filed``
    Take the agency's word for it: an NFDB report is a Canadian agency's report, so
    the fire is in Canada and the ``Country`` column is the constant
    :data:`COUNTRY_NAME` on every row. Faster — it does no geometry at all — and the
    honest choice when the question is *what did the agencies file* rather than
    *where are these fires*.

Unlike the :mod:`EGIF report <src.apps.statistics.wildfires.spain_egif.
wildfire_statistics>`, where ``geometry`` costs half the archive because half of it
publishes no coordinate, here it costs almost nothing: this dataset is points first,
and only a couple of hundred imported fires have no usable one. That is what makes
the cautious mode affordable as a default.

Whenever ``geometry`` is in force the log says how many fires were left out and why,
split between *no point published* and *point in no country* — see
:func:`location_audit`. The second number is the one to watch: it is the count of
coordinates that are somewhere no country is, which for a fire report means the sea.

.. note::

   A fire's whole reported area is attributed to one country. Nothing here splits a
   fire between countries, and nothing needs to: these are points, and a point is in
   one place.

Which year a fire counts towards
---------------------------------

The published ``YEAR``, which is
:attr:`~src.providers.canada_nfdb.wildfire.NfdbWildfire.year` — the filing, as in
every report in this project, and not the year of
:attr:`~src.data_model.wildfire.Wildfire.start_date_time`.

Here the two rarely disagree, ``REP_DATE`` being what the instant is resolved from,
but the published year is ``NOT NULL`` in practice, is indexed, and needs no timezone
applied to it — which the instant would, Canada spanning six zones.

.. note::

   The column is nullable on the model, for the 95 published rows carrying
   :data:`~src.providers.canada_nfdb.UNKNOWN_YEAR`. None of them can reach the
   database: the import's ``--from-year`` floor is a lower bound on the same column,
   so a fire with no year is never stored. The report checks anyway and warns rather
   than silently leaving such a fire out of every year.

Which fires are counted
------------------------

Every imported fire, which is the archive from :data:`~src.providers.canada_nfdb.
FIRST_YEAR` on that publishes a report date, an agency and one of the three causes.
The 1930-1972 points are published and deliberately not imported — see
:mod:`src.providers.canada_nfdb` — so a run of this report is a run over 1973
onwards whatever ``--year`` says.

.. warning::

   **A count over this archive is a count of what thirteen agencies chose to file.**
   Coverage, accuracy, vocabulary and start year all vary by agency: British Columbia
   files well over a hundred thousand of the rows and Prince Edward Island fifty-five.
   A trend across fifty years is partly a trend in reporting practice, and a
   comparison between two provinces is partly a comparison of two filing standards.

   The ``Agencies`` column is that caveat made visible: it counts the distinct
   :attr:`~src.providers.canada_nfdb.wildfire.NfdbWildfire.src_agency` values behind
   the row. A year reported by nine agencies is not comparable with a year reported
   by thirteen, however similar the hectares look. In the ``Total`` row it is the
   number of agencies over the whole period, not a sum of the years — see
   :func:`combine`.

``--agency`` narrows the report to one of them, which is the only way to make a
series across the years mean one thing. Under it the ``Agencies`` column is the
constant 1.

One cause, if you want one
---------------------------

``--cause`` narrows the report to ``natural``, ``human`` or ``unknown`` fires — the
three values of ``CAUSE``, which :mod:`src.providers.canada_nfdb` defines.

.. warning::

   ``--cause natural`` is **not** a lightning filter, exactly as in the NBAC report:
   see :data:`~src.providers.canada_nfdb.CAUSE_NATURAL`. It is the nearest thing this
   archive has to one, it is dominated by lightning in the Canadian boreal, and it is
   not defined as lightning. It is also, at roughly 195,000 fires, the largest
   natural-cause set in GisFIRE and the reason this dataset is here.

Prescribed burns are excluded by default
-----------------------------------------

:attr:`~src.providers.canada_nfdb.wildfire.NfdbWildfire.prescribed` is a deliberate
burn rather than a wildfire, so those fires are left out and the log says how many;
``--include-prescribed`` counts them.

The column is a weak one and the report says so rather than pretending otherwise:
most agencies do not publish it at all, and
:func:`~src.providers.canada_nfdb.is_prescribed` reads silence as *not prescribed*.
So this exclusion removes the prescribed burns that were **declared**, not the ones
that happened. :data:`~src.providers.canada_nfdb.CAUSE2_PRESCRIBED_BURN` on
``fire_cause_detail`` is the other, independent statement of the same thing, and it
is deliberately not folded in here: it is a *cause*, and mixing a cause into a flag
would make the exclusion two different questions at once.

One year at a time
-------------------

The report is not one statement. The years are found first, then each is measured by
a statement of its own and the summary rows are computed from their results — the
:mod:`GWIS <src.apps.statistics.wildfires.gwis.wildfire_statistics>` shape, for the
reason given there.

Under ``filed`` it would not be needed: that mode is an indexed aggregate over one
column of one table. Under ``geometry``, which is the default, each fire means a
point-in-polygon test against country polygons of millions of vertices, and the
memory that goes into them is only released when the statement ends. Half a million
points is well short of the twenty-million-perimeter case that met the OOM killer,
but the shape that survives it costs nothing here and the two modes are worth
keeping identical.

Nothing about the figures changes. ``count``, ``sum``, ``min`` and ``max`` all
decompose over a partition of the fires, and the agencies are unioned rather than
added, so the ``Total`` row is exactly the number one pass would have returned.
Every statement runs in one transaction and so against one snapshot.
"""

from __future__ import annotations

import argparse
import csv
import logging
import math
import os
import sys

from dataclasses import dataclass
from dataclasses import field
from pathlib import Path

from sqlalchemy import ColumnElement
from sqlalchemy import Engine
from sqlalchemy import Select
from sqlalchemy import create_engine
from sqlalchemy import func
from sqlalchemy import literal
from sqlalchemy import select
from sqlalchemy import true
from sqlalchemy.orm import Session

import src.settings  # noqa: F401  (imported for the side effect of loading .env)

from src.apps.imports import common
from src.data_model.geography.admin_boundary import AdminBoundary
from src.data_model.ignition import Ignition
from src.providers import canada_nfdb
from src.providers.canada_nfdb.wildfire import NfdbWildfire
from src.providers.ocha.admin_boundary import OchaAdminBoundary

#: Label used in the ``Year`` column for a summary row.
TOTAL_LABEL = "Total"

#: The six columns this report shares with the other eight, in their order, so the
#: CSVs can still be concatenated on them.
SHARED_COLUMNS = ("Country", "Year", "Fires", "Minimum (ha)", "Maximum (ha)", "Total (ha)")

#: The report's columns, in order, shared by both output formats so that a change to
#: one cannot silently leave the other behind. The last is this dataset's own: how
#: many agencies filed the fires behind the row.
COLUMNS = SHARED_COLUMNS + ("Agencies",)

#: Index of the first column that holds a number, and so is right-aligned in the Word
#: table.
FIRST_NUMERIC_COLUMN = 2

#: The country these reports are filed in, and the whole of the ``Country`` column
#: under ``--country-source filed``. Spelled as the OCHA boundaries spell it, so a row
#: of this report sorts and groups with the rows of the other eight.
COUNTRY_NAME = "Canada"

#: The two ways of deciding which country a fire counts towards. Unlike the Spanish
#: report the default **is** ``geometry``: see the module docstring — almost every
#: fire here has a point, and a good many of those points are not in Canada.
COUNTRY_SOURCE_GEOMETRY = "geometry"
COUNTRY_SOURCE_FILED = "filed"
COUNTRY_SOURCES = (COUNTRY_SOURCE_GEOMETRY, COUNTRY_SOURCE_FILED)

#: Administrative level of a country in ``admin_boundary``.
COUNTRY_LEVEL = 0

#: The three causes, as ``--cause`` accepts them, keyed to the values ``CAUSE``
#: publishes. Spelled out on the command line because ``--cause N`` is a letter
#: nobody will remember, and mapped rather than upper-cased so that the vocabulary
#: stays :mod:`src.providers.canada_nfdb`'s and not this module's.
CAUSES = {
    "natural": canada_nfdb.CAUSE_NATURAL,
    "human": canada_nfdb.CAUSE_HUMAN,
    "unknown": canada_nfdb.CAUSE_UNKNOWN,
}

#: The year a fire counts towards: the published ``YEAR``. See the module docstring —
#: the filing, not the clock.
PUBLISHED_YEAR = NfdbWildfire.__table__.c.year

#: Which agency filed the report, and what ``--agency`` selects on. ``NOT NULL`` on
#: every stored fire — the import refuses a row without one — which is why the filter
#: reaches the whole archive and needs no join.
AGENCY = NfdbWildfire.__table__.c.src_agency

#: The reported burnt area, in hectares, as filed. The only area this dataset has:
#: there is no perimeter to measure and no second column to choose between.
SIZE_HA = NfdbWildfire.__table__.c.size_ha


def country_columns(source: str) -> tuple[ColumnElement, list]:
    """Where a fire's country comes from, as ``(name, joins)``.

    Parameters
    ----------
    source : str
        One of :data:`COUNTRY_SOURCES`.

    Returns
    -------
    tuple
        The country-name expression, and the joins that have to be applied to a
        ``select`` over ``nfdb_wildfire`` for it to resolve.

    Raises
    ------
    ValueError
        If ``source`` is not one of :data:`COUNTRY_SOURCES`.

    Notes
    -----
    **``filed``** is a literal and needs no join at all. It is still grouped on, like
    a real column: by the time the outer aggregate sees it, it is a column of the
    subquery rather than a constant, and PostgreSQL requires every one of those in the
    ``GROUP BY``. One extra grouping key with a single distinct value costs nothing
    and keeps the two modes the same shape.

    **``geometry``** tests the published point. Both of its joins are inner, and each
    drops a different kind of fire: the join to ``ignition`` drops the fires whose
    published coordinate the import could not use, and the lateral drops the ones
    whose coordinate is inside no country — which for a fire report means the sea.
    :func:`location_audit` counts the two separately, because they mean entirely
    different things about the data.

    ``LATERAL ... LIMIT 1`` rather than a plain join: a point on a shared border can
    satisfy ``ST_Contains`` for two countries — and the Canada-United States border is
    nine thousand kilometres of exactly that — so one fire must not become two rows.
    The join to ``ocha_admin_boundary`` is what keeps the test against the OCHA
    country outlines rather than against every level-0 boundary any provider has ever
    loaded.
    """
    if source == COUNTRY_SOURCE_FILED:
        return literal(COUNTRY_NAME), []

    if source != COUNTRY_SOURCE_GEOMETRY:
        raise ValueError(
            f"unknown country source {source!r}; expected one of {', '.join(COUNTRY_SOURCES)}"
        )

    ignition = Ignition.__table__
    ocha_boundary = OchaAdminBoundary.__table__
    containing = (
        select(AdminBoundary.name.label("name"))
        .select_from(AdminBoundary)
        .join(ocha_boundary, ocha_boundary.c.id == AdminBoundary.id)
        .where(AdminBoundary.level == COUNTRY_LEVEL)
        .where(func.ST_Contains(AdminBoundary.geometry, ignition.c.geometry))
        .limit(1)
        .lateral("containing")
    )
    joins = [
        (ignition, ignition.c.id == NfdbWildfire.__table__.c.ignition_id),
        (containing, true()),
    ]
    return containing.c.name, joins


def is_a_wildfire(include_prescribed: bool = False) -> ColumnElement | None:
    """The condition excluding the declared prescribed burns, or ``None`` to keep them.

    Parameters
    ----------
    include_prescribed : bool
        ``True`` to count prescribed burns as wildfires. ``False``, the default,
        leaves them out.

    Returns
    -------
    ColumnElement or None
        ``None`` when nothing is to be filtered, so the caller can leave the ``WHERE``
        off altogether rather than adding a clause that is always true.

    Notes
    -----
    ``~prescribed`` and not ``IS NOT true``: the column is ``NOT NULL`` with a default
    of false, so this cannot silently drop a row whose flag was never set.

    What it cannot do is find the prescribed burns nobody declared. Most agencies do
    not publish ``PRESCRIBED`` at all and
    :func:`~src.providers.canada_nfdb.is_prescribed` reads silence as *no*, so this
    excludes a declaration rather than a kind of fire. See the module docstring.
    """
    if include_prescribed:
        return None
    return ~NfdbWildfire.__table__.c.prescribed


def cause_condition(cause: str | None) -> ColumnElement | None:
    """The condition selecting one ``CAUSE`` value, or ``None`` for every cause.

    Parameters
    ----------
    cause : str or None
        One of :data:`CAUSES`' keys, or ``None`` to count every fire.

    Raises
    ------
    ValueError
        If ``cause`` is neither ``None`` nor one of the three.

    Notes
    -----
    The column is ``NOT NULL`` and constrained to the three values, so this is a
    partition of the archive: the three possible runs of ``--cause`` account for every
    fire between them, with none counted twice.
    """
    if cause is None:
        return None
    if cause not in CAUSES:
        raise ValueError(
            f"unknown cause {cause!r}; expected one of {', '.join(CAUSES)}"
        )
    return NfdbWildfire.__table__.c.fire_cause == CAUSES[cause]


def agency_condition(agency: str | None) -> ColumnElement | None:
    """The condition selecting one filing agency, or ``None`` for all of them."""
    if agency is None:
        return None
    return AGENCY == agency


def scope_conditions(include_prescribed: bool = False,
                     cause: str | None = None,
                     agency: str | None = None) -> list[ColumnElement]:
    """Every condition narrowing which fires the report is of, as a list.

    Built once and applied to each year's statement, so that the three filters cannot
    drift apart between the query that finds the years and the queries that measure
    them — or between the report and the audit that accounts for what it dropped.
    """
    return [condition for condition in (is_a_wildfire(include_prescribed),
                                        cause_condition(cause),
                                        agency_condition(agency))
            if condition is not None]


def available_agencies(session: Session) -> list[str]:
    """The agency codes in the database, in order.

    Read from the database rather than written out here: the thirteen are what the
    service published, and a fourteenth appearing in a future distribution should
    show up in ``--agency``'s error message without this module being edited.
    """
    return list(session.scalars(
        select(AGENCY).distinct().order_by(AGENCY)
    ))


def resolve_agency(session: Session, wanted: str) -> str:
    """The agency code ``--agency`` names, as it is stored.

    Parameters
    ----------
    session : Session
        An open session.
    wanted : str
        What was passed to ``--agency``: the published code, in any case — ``BC``,
        ``bc``, ``Pc``.

    Returns
    -------
    str
        The code as stored, which is what the statements filter on.

    Raises
    ------
    RuntimeError
        If no agency is imported at all, or if nothing matches.

    Notes
    -----
    Matched exactly and case-insensitively, and never by prefix. These are two-letter
    codes: a prefix match would make ``N`` mean four different things, and there is no
    long form of ``NT`` for anyone to have typed instead.

    The error lists the codes that really are imported, because thirteen two-letter
    codes are not something anyone remembers and reading the source to find them out
    would be a poor answer.
    """
    agencies = available_agencies(session)
    if not agencies:
        raise RuntimeError(
            "No agency is imported, so --agency has nothing to select from: import the "
            "NFDB fires first, with "
            "src.apps.imports.wildfires.canada_nfdb.import_wildfires"
        )
    asked = wanted.strip().casefold()
    for agency in agencies:
        if agency.casefold() == asked:
            return agency
    raise RuntimeError(
        f"No agency matches {wanted!r}. The agencies imported are: "
        f"{', '.join(agencies)}"
    )


def years_query(include_prescribed: bool = False,
                cause: str | None = None,
                agency: str | None = None) -> Select:
    """The years the dataset holds fires in, newest first.

    Returns
    -------
    Select
        A query yielding one ``int`` per year.

    Notes
    -----
    Run before anything else, because each of those years is then measured by a
    statement of its own — see the module docstring.

    It carries the scope filters but resolves no country: that is the expensive half
    of this report under ``geometry`` and there is nothing to gain by paying for it
    twice. A year whose fires all turn out to be in no country simply measures to
    nothing and never reaches the report.

    ``NULL`` years are excluded explicitly. The import cannot store one — its
    ``--from-year`` floor is a lower bound on the same column — but the column is
    nullable, and a fire in no year would otherwise be dropped by every statement
    without anything saying so. :func:`unknown_year_count` is what says so.

    ``DISTINCT`` over the table rather than ``min(year)`` to ``max(year)``: a gap in
    the record is a gap in the report, and a range would fill it with statements that
    can only return nothing.
    """
    query = (
        select(PUBLISHED_YEAR.label("year"))
        .select_from(NfdbWildfire.__table__)
        .where(PUBLISHED_YEAR.is_not(None))
        .distinct()
        .order_by(PUBLISHED_YEAR.desc())
    )
    for condition in scope_conditions(include_prescribed, cause, agency):
        query = query.where(condition)
    return query


def unknown_year_count(session: Session,
                       include_prescribed: bool = False,
                       cause: str | None = None,
                       agency: str | None = None) -> int:
    """How many fires in scope carry no published year, and so are in no row.

    Zero on any database the import wrote, which is the point: a number that can only
    be zero is cheap to check and expensive to have been wrong about. A fire with no
    year cannot be reported under a year, and it must not be silently absent instead.
    """
    query = (
        select(func.count())
        .select_from(NfdbWildfire.__table__)
        .where(PUBLISHED_YEAR.is_(None))
    )
    for condition in scope_conditions(include_prescribed, cause, agency):
        query = query.where(condition)
    return session.scalar(query) or 0


def statistics_query(year: int,
                     min_area: float | None = None,
                     country_source: str = COUNTRY_SOURCE_GEOMETRY,
                     include_prescribed: bool = False,
                     cause: str | None = None,
                     agency: str | None = None) -> Select:
    """Build the statistics query for one year.

    Parameters
    ----------
    year : int
        The published year to measure.
    min_area : float, optional
        Count only fires of at least this many reported hectares.
    country_source : str
        One of :data:`COUNTRY_SOURCES`.
    include_prescribed : bool
        Count the declared prescribed burns as wildfires.
    cause : str, optional
        Restrict to one ``CAUSE`` value.
    agency : str, optional
        Restrict to one filing agency, as :func:`resolve_agency` returned it.

    Returns
    -------
    Select
        A query yielding ``country, minimum, maximum, total, fires, agencies``: one
        row per country that had a fire in ``year``, unordered. The summary rows and
        the report's order are :func:`summarise`'s work.

    Notes
    -----
    Built against the mapped classes rather than written as SQL text, so a column
    renamed on a model breaks this at import time rather than in front of a user. It
    also lets the filters be plain conditionals: written as text each would have to
    become an "unset, or matching" disjunction so that one statement could serve every
    combination, leaving branches in the SQL that are dead on every actual run.

    Under ``filed`` it is selected from ``nfdb_wildfire`` alone — the year, the size,
    the cause, the agency and the prescribed flag all live there — so neither the
    parent ``wildfire`` row nor the ignition is joined at all. By table and not through
    the mapped class, to keep SQLAlchemy from adding a polymorphic join of its own.

    The agencies come back as an array and become a set in Python, rather than as a
    count, because a count cannot be combined: the ``Total`` row's ``Agencies`` is the
    number of agencies over the whole period, which is the size of the union of the
    years' sets and not the sum of their counts. Thirteen strings a year is nothing to
    carry.

    ``min_area`` filters the subquery's column and is applied before the aggregates
    rather than as a ``HAVING``: the threshold selects the fires the figures are
    computed from, it does not discard years whose total came out small.

    A year with no fires in scope produces no rows at all, rather than one row of
    ``NULL`` aggregates over zero fires — which is what keeps ``--year 1950`` on an
    archive starting in 1973 an empty report and so the caller's "no wildfires
    matched" message.
    """
    nfdb = NfdbWildfire.__table__
    country_name, joins = country_columns(country_source)

    fires = (
        select(country_name.label("country"),
               SIZE_HA.label("hectares"),
               AGENCY.label("agency"))
        .select_from(nfdb)
        .where(SIZE_HA.is_not(None))
        .where(PUBLISHED_YEAR == year)
    )
    for target, condition in joins:
        fires = fires.join(target, condition)
    for condition in scope_conditions(include_prescribed, cause, agency):
        fires = fires.where(condition)

    fire = fires.subquery("fire")
    statistics = (
        select(
            fire.c.country,
            func.min(fire.c.hectares).label("minimum"),
            func.max(fire.c.hectares).label("maximum"),
            func.sum(fire.c.hectares).label("total"),
            func.count().label("fires"),
            func.array_agg(fire.c.agency.distinct()).label("agencies"),
        )
        .group_by(fire.c.country)
    )
    if min_area is not None:
        statistics = statistics.where(fire.c.hectares >= min_area)
    return statistics


def location_audit(year: int | None = None,
                   min_area: float | None = None,
                   include_prescribed: bool = False,
                   cause: str | None = None,
                   agency: str | None = None) -> Select:
    """Count the fires ``--country-source geometry`` leaves out, and why.

    Returns
    -------
    Select
        A query yielding one row, ``no_point, outside``: how many fires in scope have
        no usable published coordinate, and how many have one that is inside no
        country.

    Notes
    -----
    The two are worth separating and the report does not add them up. *No point* is a
    published coordinate the import could not use — a null, a zero, or projected
    metres that leaked into the degrees columns — and there are very few of them.
    *Outside* is a coordinate that resolves to nowhere, which for a fire report means
    the sea, and it is the number that says how much the plausibility box round Canada
    let through.

    A point over the American border is in **neither**: it is inside a country, so it
    is reported, as that country's row. That is the whole reason ``geometry`` is the
    default here, and it is why this audit does not try to be a count of "wrong"
    coordinates — it counts what the report could not place at all.

    Run under the same scope as the report it accompanies — same year, same threshold,
    same cause, same agency — so that its two numbers and the ``Fires`` column account
    for the same set of fires.

    One statement over the whole scope rather than one per year: it does no
    point-in-polygon test that the report has not already done, an ``EXISTS`` stopping
    at the first containing polygon, and a single extra pass is cheaper than fifty
    separate ones.
    """
    nfdb = NfdbWildfire.__table__
    ignition = Ignition.__table__
    ocha_boundary = OchaAdminBoundary.__table__

    located = (
        select(AdminBoundary.id)
        .select_from(AdminBoundary)
        .join(ocha_boundary, ocha_boundary.c.id == AdminBoundary.id)
        .where(AdminBoundary.level == COUNTRY_LEVEL)
        .where(func.ST_Contains(AdminBoundary.geometry, ignition.c.geometry))
        .exists()
    )

    fires = (
        select(nfdb.c.ignition_id.label("ignition_id"),
               SIZE_HA.label("hectares"),
               located.label("in_a_country"))
        .select_from(nfdb)
        .outerjoin(ignition, ignition.c.id == nfdb.c.ignition_id)
        .where(SIZE_HA.is_not(None))
    )
    if year is not None:
        fires = fires.where(PUBLISHED_YEAR == year)
    for condition in scope_conditions(include_prescribed, cause, agency):
        fires = fires.where(condition)

    fire = fires.subquery("fire")
    audit = select(
        func.count().filter(fire.c.ignition_id.is_(None)).label("no_point"),
        func.count().filter(
            fire.c.ignition_id.is_not(None) & ~fire.c.in_a_country
        ).label("outside"),
    ).select_from(fire)
    if min_area is not None:
        audit = audit.where(fire.c.hectares >= min_area)
    return audit


@dataclass(frozen=True)
class Row:
    """One line of the report.

    Attributes
    ----------
    country : str
        Name of the country the fire was reported in. :data:`COUNTRY_NAME` under
        ``--country-source filed``; under ``geometry`` it is whichever country
        contains the published point, which is normally but not always Canada.
    year : int or None
        The published year, or ``None`` for a summary row.
    minimum, maximum, total : float
        Smallest single fire, largest single fire and sum of every fire, in hectares
        as the agencies reported them.
    fires : int
        How many fires the three area figures were computed from.
    agencies : frozenset of str
        Which agencies filed them. Held as the set and not as a count so that a
        summary row can union rather than add — see :func:`combine`.
    """

    country: str
    year: int | None
    minimum: float
    maximum: float
    total: float
    fires: int
    agencies: frozenset[str] = field(default_factory=frozenset)

    @property
    def is_total(self) -> bool:
        """Whether this is a summary row rather than one of the years."""
        return self.year is None

    @property
    def year_label(self) -> str:
        return TOTAL_LABEL if self.is_total else str(self.year)

    @property
    def agency_count(self) -> int:
        """How many agencies filed the fires behind this row."""
        return len(self.agencies)

    @property
    def values(self) -> tuple[str, ...]:
        """The row as the CSV writes it, in :data:`COLUMNS` order."""
        return (self.country, self.year_label, str(self.fires),
                f"{self.minimum:.2f}", f"{self.maximum:.2f}", f"{self.total:.2f}",
                str(self.agency_count))

    @property
    def readable_values(self) -> tuple[str, ...]:
        """The row as the Word document writes it: the numbers with separators."""
        return (self.country, self.year_label, f"{self.fires:,}",
                f"{self.minimum:,.2f}", f"{self.maximum:,.2f}", f"{self.total:,.2f}",
                str(self.agency_count))


def combine(rows: list[Row], country: str = COUNTRY_NAME, year: int | None = None) -> Row:
    """One row summarising several: every figure taken over all of them.

    Notes
    -----
    This is what makes measuring a year at a time cost nothing. The four numbers
    decompose over a partition of the fires — a minimum of minima is a minimum, a sum
    of sums is a sum — so the summary rows are the numbers one pass over the same rows
    would have returned, and no fire is counted twice or left out.

    The agencies are **unioned and not added**, which is the whole reason the set is
    carried around. Thirteen agencies filing every year for fifty years are thirteen
    agencies, not six hundred and fifty.

    ``fsum`` rather than ``sum``: an exact accumulation over a handful of partial
    totals costs nothing and cannot drift from what one pass would have returned.
    """
    return Row(
        country=country,
        year=year,
        minimum=min(row.minimum for row in rows),
        maximum=max(row.maximum for row in rows),
        total=math.fsum(row.total for row in rows),
        fires=sum(row.fires for row in rows),
        agencies=frozenset().union(*(row.agencies for row in rows)),
    )


def ordered_countries(session: Session, names: set[str]) -> list[str]:
    """The countries of the report, in the order the database would have sorted them.

    Notes
    -----
    One name, ``Canada``, under ``--country-source filed``; under ``geometry`` there
    can be a second, which is a point over a border — in practice the United States.

    Sorted by PostgreSQL and not by Python because the two do not agree: Python
    compares code points, which puts every accented name after every unaccented one,
    while the database sorts by its collation.

    A name the query does not return is appended rather than dropped — under ``filed``
    the constant ``Canada`` is not a row of ``admin_boundary`` at all unless the OCHA
    boundaries happen to be imported, and a report that vanished for want of them
    would be a poor trade.
    """
    if not names:
        return []
    ordered = list(session.scalars(
        select(AdminBoundary.name)
        .distinct()
        .where(AdminBoundary.level == COUNTRY_LEVEL)
        .where(AdminBoundary.name.in_(names))
        .order_by(AdminBoundary.name)
    ))
    return ordered + sorted(names - set(ordered))


def summarise(measured: list[Row], countries: list[str]) -> list[Row]:
    """Build the report from the years measured: the summary rows, in order.

    Parameters
    ----------
    measured : list of Row
        One row per country and year, as the per-year statements returned them.
    countries : list of str
        The countries in scope, in the order they are to be reported.

    Returns
    -------
    list of Row
        Each country, its years newest first and its summary row last. Empty if
        nothing was measured — a report of no fires has no total either.

    Notes
    -----
    No World block, unlike the GWIS and GFA reports: this is one country's archive,
    and a block summarising Canada and the handful of border fires beside it would be
    a total of two things that are not comparable.
    """
    report: list[Row] = []
    for name in countries:
        rows = [row for row in measured if row.country == name]
        if not rows:
            continue
        report += sorted(rows, key=lambda row: row.year, reverse=True)
        report.append(combine(rows, name, None))
    return report


def hectares(text: str) -> float:
    """Argparse type for ``--min-area``: a finite, non-negative number of hectares.

    Raises
    ------
    argparse.ArgumentTypeError
        If the text is not a number, or is negative, or is a non-finite float.

    Notes
    -----
    A bare ``type=float`` would accept ``-5``, ``nan`` and ``inf``. The first two are
    almost certainly a typo and would silently produce the unfiltered report — ``nan``
    compares false against every area, so it would produce an empty one — and none of
    the three is a size a fire can have.
    """
    try:
        area = float(text)
    except ValueError:
        raise argparse.ArgumentTypeError(f"{text!r} is not a number of hectares")
    if not math.isfinite(area):
        raise argparse.ArgumentTypeError(f"{text!r} is not a finite number of hectares")
    if area < 0:
        raise argparse.ArgumentTypeError(
            f"a burnt area cannot be negative, and {area:g} ha is: pass 0 or more, or "
            f"leave --min-area out to count every fire"
        )
    return area


def parse_arguments(argv: list[str] | None = None) -> argparse.Namespace:
    """Parse the command line."""
    parser = argparse.ArgumentParser(
        description="Burnt-area statistics for the Canadian National Fire Database "
                    "agency fire data, with the number of agencies behind each row.",
        epilog="Areas are the hectares the agencies reported — this dataset publishes no "
               "perimeter, so there is no --area-method and no --surface. By default the "
               "published point is tested against the real country polygons, because a "
               "good many of them are not in Canada. Database settings not given here are "
               "read from the environment (.env).",
    )
    parser.add_argument("-y", "--year", type=int,
                        help="restrict to one year, e.g. 2023; this is the published "
                             "YEAR, not the year of the report date")
    parser.add_argument("--min-area", type=hectares, default=None, metavar="HECTARES",
                        help="count only fires that reported at least this many hectares; "
                             "by default every fire counts, including a reported zero. "
                             "200 is the service's own large-fire threshold, which keeps "
                             "a twentieth of the rows and over 97%% of the area")
    parser.add_argument("--cause", default=None, choices=sorted(CAUSES),
                        help="restrict to one published CAUSE. 'natural' is the nearest "
                             "thing this dataset has to a lightning category and is not "
                             "one, though it is the largest natural-cause set in GisFIRE")
    parser.add_argument("-a", "--agency", default=None, metavar="CODE",
                        help="restrict to one filing agency, by its published code — "
                             "'BC', 'AB', 'ON', 'PC', … Case does not matter, and an "
                             "unrecognised code is answered with the list of the ones "
                             "that are imported. Worth reaching for: coverage and "
                             "practice vary so much between agencies that a series over "
                             "all thirteen is partly a series about reporting")
    parser.add_argument("--include-prescribed", action="store_true",
                        help="count the fires flagged PRESCRIBED — a deliberate burn "
                             "rather than a wildfire. Left out by default. Most agencies "
                             "do not publish the column at all, so this excludes a "
                             "declaration and not a kind of fire")
    parser.add_argument("--country-source", default=COUNTRY_SOURCE_GEOMETRY,
                        choices=COUNTRY_SOURCES,
                        help="how to decide which country a fire counts towards: "
                             "'geometry' (default) tests the published point against the "
                             "real country polygons, so a coordinate in the sea drops out "
                             "and one over the American border is reported as the United "
                             "States; 'filed' takes the agency's word for it, every "
                             "report being a Canadian one. The log says how many fires "
                             "'geometry' left out and why")

    # Accepted only so that they can be refused clearly. Anyone reaching for one has
    # copied a command line from one of the other eight reports, which is a reasonable
    # thing to have done, and argparse's own message would not say why this report is
    # different.
    parser.add_argument("--country", help=argparse.SUPPRESS)
    parser.add_argument("--area-method", help=argparse.SUPPRESS)
    parser.add_argument("--surface", help=argparse.SUPPRESS)

    output = parser.add_argument_group("output", "at least one is required")
    output.add_argument("--csv", type=Path, help="write the report to this .csv")
    output.add_argument("--docx", type=Path, help="write the report to this .docx (MS Word)")

    common.add_database_arguments(parser)
    parser.add_argument("--log-level", default=os.getenv("GISFIRE_LOG_LEVEL", "INFO"),
                        choices=["CRITICAL", "ERROR", "WARNING", "INFO", "DEBUG"],
                        help="verbosity (env: GISFIRE_LOG_LEVEL, default INFO)")

    arguments = parser.parse_args(argv)
    if arguments.country is not None:
        parser.error(
            f"there is no --country here: these are {COUNTRY_NAME}'s own agencies' "
            f"reports of their own fires, so there is nothing to select between. Every "
            f"fire is counted, and the Country column says which country each one's "
            f"point turned out to be in — see --country-source."
        )
    if arguments.area_method is not None:
        parser.error(
            "there is no --area-method here: this dataset publishes no perimeter, so "
            "nothing is measured and no CRS is involved. The hectares are the ones the "
            "agency reported. The companion NBAC report measures polygons and has one."
        )
    if arguments.surface is not None:
        parser.error(
            "there is no --surface here: the agencies report one burnt area per fire, "
            "SIZE_HA, and there is nothing to choose between."
        )
    if arguments.csv is None and arguments.docx is None:
        parser.error("nothing to write: pass --csv, --docx, or both")
    return arguments


def compute(session: Session, year: int | None, logger: logging.Logger,
            min_area: float | None = None,
            country_source: str = COUNTRY_SOURCE_GEOMETRY,
            include_prescribed: bool = False,
            cause: str | None = None,
            agency: str | None = None) -> list[Row]:
    """Measure the fires a year at a time, returning the report's rows in order.

    Notes
    -----
    One statement per year, for the reason given in the module docstring, each under a
    spinner of its own so that a long run says which year it is on rather than only
    that it is alive.

    Every one of them runs in ``session``'s transaction, and so against a single
    snapshot: a report assembled from many queries is then exactly as consistent as
    one assembled from a single query, and an import running alongside it cannot have
    a fire counted in one year's statement and not in another's.

    Under ``geometry`` a second statement follows them: :func:`location_audit` counts
    what they dropped and why. A report that quietly left out a fire whose coordinate
    is in the sea would be worse than one that did not offer the option.
    """
    if year is not None:
        years = [year]
    else:
        with common.Spinner("Finding the years the NFDB reports cover", logger):
            years = list(session.scalars(years_query(include_prescribed, cause, agency)))

    scope = "every agency" if agency is None else agency
    if cause is not None:
        scope = f"{cause}, {scope}"
    measured: list[Row] = []
    for index, measuring in enumerate(years, start=1):
        with common.Spinner(f"Summing the reported burnt area of the NFDB fires "
                            f"({scope}, {measuring}: {index} of {len(years)})", logger):
            measured += [
                Row(country=record.country, year=measuring,
                    minimum=float(record.minimum),
                    maximum=float(record.maximum),
                    total=float(record.total),
                    fires=record.fires,
                    agencies=frozenset(record.agencies or ()))
                for record in session.execute(
                    statistics_query(measuring, min_area, country_source,
                                     include_prescribed, cause, agency))
            ]

    if country_source == COUNTRY_SOURCE_GEOMETRY:
        with common.Spinner("Counting the fires with no usable point", logger):
            audit = session.execute(
                location_audit(year, min_area, include_prescribed, cause, agency)).one()
        counted = sum(row.fires for row in measured)
        logger.info("Excluded %d of %d fire(s): %d have no usable published point, "
                    "%d have one that is inside no country",
                    audit.no_point + audit.outside,
                    counted + audit.no_point + audit.outside,
                    audit.no_point, audit.outside)
        if audit.outside:
            logger.warning(
                "%d fire(s) have a published coordinate that is inside no country — the "
                "import's only geometric guard is a plausibility box round Canada, and a "
                "box round a country contains a great deal of sea", audit.outside)

    orphaned = unknown_year_count(session, include_prescribed, cause, agency)
    if orphaned:
        logger.warning(
            "%d fire(s) in scope publish no year and so are in no row of this report. "
            "The import cannot store one, so this is a database written by something "
            "else", orphaned)

    countries = ordered_countries(session, {row.country for row in measured})
    rows = summarise(measured, countries)
    logger.info("Computed %d rows over %d country/countries and %d year(s) "
                "(reported hectares, country from %s, %s, %s)",
                len(rows), len(countries), len({row.year for row in measured}),
                country_source,
                "every fire" if min_area is None else f"fires of {min_area:g} ha or more",
                scope)
    if rows:
        agencies = frozenset().union(*(row.agencies for row in rows))
        logger.info("%d agency/agencies filed the fires reported: %s",
                    len(agencies), ", ".join(sorted(agencies)))
    return rows


def write_csv(rows: list[Row], path: Path, logger: logging.Logger) -> None:
    """Write the report as CSV.

    The numbers go out unformatted apart from being rounded to two decimals — no
    thousands separators — because a CSV is read by another program far more often
    than by a person, and a separator would make every figure a string.
    """
    path.parent.mkdir(parents=True, exist_ok=True)
    with path.open("w", newline="", encoding="utf-8") as handle:
        writer = csv.writer(handle)
        writer.writerow(COLUMNS)
        for row in rows:
            writer.writerow(row.values)
    logger.info("Wrote %s", path)


def write_docx(rows: list[Row], path: Path, year: int | None,
               logger: logging.Logger,
               min_area: float | None = None,
               country_source: str = COUNTRY_SOURCE_GEOMETRY,
               include_prescribed: bool = False,
               cause: str | None = None,
               agency: str | None = None) -> None:
    """Write the report as a Word document.

    One table, with each country's summary row in bold. Numbers get thousands
    separators here — the opposite of the CSV, and for the opposite reason: this one
    is for reading.

    The opening paragraphs say that the hectares are reported rather than measured,
    that they are not NBAC's, what ``Agencies`` is and what the point test did,
    because a table that looks exactly like the NBAC one and disagrees with it by 34
    million hectares is the single easiest mistake to make with this dataset.
    """
    # Imported here rather than at module scope so that --csv keeps working if
    # python-docx is not installed, which matters because it is the only dependency
    # this application adds.
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt

    document = Document()
    where = COUNTRY_NAME if agency is None else f"{agency}, {COUNTRY_NAME}"
    document.add_heading(f"NFDB wildfire reported burnt area ({where})", level=1)

    scope = [f"year: {year}" if year is not None else "all years"]
    if agency is not None:
        scope.append(f"only the fires filed by {agency}")
    if cause is not None:
        scope.append(f"only {cause} fires")
    if min_area is not None:
        scope.append(f"only fires of {min_area:g} ha or more")
    scope.append("prescribed burns counted" if include_prescribed
                 else "declared prescribed burns excluded")
    if country_source == COUNTRY_SOURCE_GEOMETRY:
        scope.append("only fires whose published point falls inside a country")
    document.add_paragraph(
        f"Areas in hectares as the agencies reported them (SIZE_HA) — this dataset "
        f"publishes no perimeter, so nothing is measured. A reported zero is counted. "
        f"Years are the published YEAR. Scope: {'; '.join(scope)}."
    )
    document.add_paragraph(
        "These are not the NBAC figures for the same fires and are not a correction of "
        "them: over the archive the agencies' reported sizes sum to about 166.5 million "
        "hectares against NBAC's 132.7 million of mapped burn. One is what somebody "
        "recorded at the time, the other what a satellite could see afterwards. Do not "
        "add them."
    )
    document.add_paragraph(
        "Agencies counts the fire management agencies behind the row — thirteen "
        "contribute, over wildly different periods and at wildly different volumes. A "
        "year reported by nine of them is not comparable with a year reported by "
        "thirteen, so a trend across the archive is partly a trend in reporting "
        "practice. In a summary row it is the number of agencies over the whole period, "
        "not a sum of the years."
    )
    if country_source == COUNTRY_SOURCE_GEOMETRY:
        document.add_paragraph(
            "Each fire's published point was tested against the real country polygons. "
            "A fire whose point is inside no country is not in this table, and one whose "
            "point is over a border is reported under that country: the import's only "
            "geometric guard is a plausibility box round Canada, and a box round a "
            "country contains a great deal of sea and a great deal of the United States."
        )
    if cause is not None:
        document.add_paragraph(
            f"These are the fires whose published CAUSE is {CAUSES[cause]}. Natural is "
            f"the nearest thing this dataset has to a lightning category and is not one: "
            f"it is dominated by lightning in the Canadian boreal but is not defined as "
            f"lightning."
        )

    table = document.add_table(rows=1, cols=len(COLUMNS))
    table.style = "Table Grid"
    for cell, heading in zip(table.rows[0].cells, COLUMNS):
        cell.text = heading
        cell.paragraphs[0].runs[0].bold = True

    for row in rows:
        cells = table.add_row().cells
        for index, (cell, value) in enumerate(zip(cells, row.readable_values)):
            cell.text = value
            paragraph = cell.paragraphs[0]
            if index >= FIRST_NUMERIC_COLUMN:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            for run in paragraph.runs:
                run.bold = row.is_total
                run.font.size = Pt(9)

    path.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(path))
    logger.info("Wrote %s", path)


def report(args: argparse.Namespace, engine: Engine, logger: logging.Logger) -> list[Row]:
    """Compute the statistics and write whichever outputs were asked for.

    The agency is resolved first, inside the same session, so that a code nobody
    recognises fails before any fire is counted — and against the database, so the
    error can list the agencies that really are imported.
    """
    with Session(engine) as session:
        agency = None if args.agency is None else resolve_agency(session, args.agency)
        rows = compute(session, args.year, logger, args.min_area, args.country_source,
                       args.include_prescribed, args.cause, agency)

    if not rows:
        # An empty report is almost always a year with no data, and writing an empty
        # file would hide that. The narrowing options are named when they are in force,
        # because then each is at least as likely to be the reason as the year is.
        extra = "" if args.min_area is None else \
            f" No fire reached the --min-area of {args.min_area:g} ha."
        if args.cause is not None:
            extra += f" The report is of {args.cause} fires alone."
        if agency is not None:
            extra += (f" The report is of {agency}'s fires alone, and the agencies do "
                      f"not all cover the same years.")
        if args.country_source == COUNTRY_SOURCE_GEOMETRY:
            extra += (" --country-source geometry also needs the OCHA country boundaries "
                      "imported, and counts only fires whose point is inside one.")
        raise RuntimeError(
            f"No wildfires matched. Check --year, and that the NFDB fires are imported "
            f"— the import reads from {canada_nfdb.FIRST_YEAR} on." + extra
        )

    if args.csv is not None:
        write_csv(rows, args.csv, logger)
    if args.docx is not None:
        write_docx(rows, args.docx, args.year, logger, args.min_area,
                   args.country_source, args.include_prescribed, args.cause, agency)
    return rows


def main(argv: list[str] | None = None) -> int:
    args = parse_arguments(argv)
    logging.basicConfig(level=args.log_level, format="%(asctime)s %(levelname)s %(message)s")
    logger = logging.getLogger("nfdb-statistics")

    try:
        settings = common.resolve_database_settings(args)
    except RuntimeError as error:
        logger.error("%s", error)
        return 1

    engine = create_engine(common.database_url(settings))
    try:
        report(args, engine, logger)
    except Exception as error:  # noqa: BLE001  (the CLI boundary: report, do not traceback)
        logger.error("%s", error)
        return 1
    finally:
        engine.dispose()
    return 0


if __name__ == "__main__":  # pragma nocover
    sys.exit(main())
