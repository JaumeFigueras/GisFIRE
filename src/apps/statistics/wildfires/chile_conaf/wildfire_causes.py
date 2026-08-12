#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Report why Chilean fires started, season by season, from CONAF's own classification.

One row per season and cause: how many fires CONAF filed under it and how many
hectares they burnt.

Usage
-----

.. code-block:: console

   $ python3 -m src.apps.statistics.wildfires.chile_conaf.wildfire_causes \\
         --csv /tmp/conaf-causes.csv
   $ python3 -m src.apps.statistics.wildfires.chile_conaf.wildfire_causes \\
         --cause "Incendios naturales" --docx /tmp/conaf-natural.docx
   $ python3 -m src.apps.statistics.wildfires.chile_conaf.wildfire_causes \\
         --bridge-schemes --csv /tmp/conaf-causes-bridged.csv

The taxonomy changed in 2023-2024, and this report will not hide it
---------------------------------------------------------------------

CONAF renumbered its cause classification in 2023-2024 **and reused the codes**:
``4.1`` is *incendios de causa desconocida* before the break and *faenas forestales*
after it. Ten categories were also renamed, and most of the renamings changed what
the category covers — *Accidentes eléctricos* became the narrower *Líneas
eléctricas*, *Quema de desechos* became the broader *Otras quemas*.

:mod:`src.providers.chile_conaf.fire_cause` keeps those ten pairs apart, so a
fifteen-season series of any one of them has **a break at 2023-2024**: the pre-2023
name stops and its successor starts. That is what the published data says, and this
report prints it that way by default.

``--bridge-schemes`` joins each of the ten pairs into one series, using
:data:`~src.providers.chile_conaf.fire_cause.SCHEME_SUCCESSORS`. It is off by default
and both outputs say when it was used, because a bridged series asserts a continuity
CONAF did not publish. Use it deliberately or not at all.

.. danger::

   Never group these fires on
   :attr:`~src.providers.chile_conaf.fire_cause.ConafFireCause.cause_code`. This
   report groups on
   :attr:`~src.providers.chile_conaf.fire_cause.ConafFireCause.cause_normalised`,
   which is the only column that means the same thing in both numberings.

The fires with no cause are a row, not a silence
--------------------------------------------------

1,012 fires publish neither a *causa general* nor a *causa específica*, 6,221 publish
only a specific cause, and a handful publish a *causa general* this project has no
canonical form for — four of them, all DBF corruption. Every one of those appears in
the report under :data:`NO_CAUSE_LABEL`, :data:`SPECIFIC_ONLY_LABEL` or
:data:`UNRECONCILED_LABEL` rather than being dropped.

A report whose rows do not add up to the season's fire count is a report that has
quietly decided something, and the reader cannot see what.
"""

from __future__ import annotations

import argparse
import csv
import logging
import os
import sys

from dataclasses import dataclass
from pathlib import Path

from sqlalchemy import Engine
from sqlalchemy import Select
from sqlalchemy import case
from sqlalchemy import create_engine
from sqlalchemy import func
from sqlalchemy import select
from sqlalchemy.orm import Session

import src.settings  # noqa: F401  (imported for the side effect of loading .env)

from src.apps.imports import common
from src.apps.statistics.wildfires.chile_conaf.wildfire_statistics import SEASON
from src.apps.statistics.wildfires.chile_conaf.wildfire_statistics import TOTAL_LABEL
from src.apps.statistics.wildfires.chile_conaf.wildfire_statistics import season_label
from src.apps.statistics.wildfires.chile_conaf.wildfire_statistics import surface_area
from src.providers import chile_conaf
from src.providers.chile_conaf.fire_cause import ConafFireCause
from src.providers.chile_conaf.fire_cause import SCHEME_SUCCESSORS
from src.providers.chile_conaf.wildfire import ConafWildfire

LOG_FORMAT = "%(asctime)s %(levelname)-8s %(message)s"

#: The report's columns, in order, shared by both output formats.
COLUMNS = ("Season", "Cause", "Cause (English)", "Fires", "Fires (%)", "Total (ha)",
           "Total (% of ha)")

#: Index of the first column holding a number, and so right-aligned in the Word table.
FIRST_NUMERIC_COLUMN = 3

#: How a fire publishing no classification at all appears in the report. 1,012 fires.
NO_CAUSE_LABEL = "(no cause published)"

#: How a fire publishing only a *causa específica* appears. 6,221 fires, almost all of
#: them in the seasons whose ``CAUSA_GENE`` column is empty.
#:
#: They are not folded into :data:`NO_CAUSE_LABEL`: a fire whose specific cause is
#: *1.7.1. Uso de fuego por transeúntes* is classified, just not at the level this
#: report groups by, and calling that "no cause published" would be false.
SPECIFIC_ONLY_LABEL = "(specific cause only)"

#: How a fire whose published *causa general* has no canonical form appears.
UNRECONCILED_LABEL = "(unreconciled cause)"

#: The three labels above, in the order they are printed after the real causes.
SYNTHETIC_LABELS = (SPECIFIC_ONLY_LABEL, UNRECONCILED_LABEL, NO_CAUSE_LABEL)


def counts_query(season: int) -> Select:
    """Build the counting query for one season.

    Returns
    -------
    Select
        A query yielding ``cause, cause_en, scheme, fires, hectares``: one row per
        canonical cause the season's fires were filed under, plus one row for each of
        the three unclassified kinds.

    Notes
    -----
    Grouped on ``cause_normalised``, never on the code — see the module docstring.

    The ``CASE`` that produces the three synthetic labels is in SQL rather than in
    Python so that every fire of the season is in exactly one group by construction,
    and the report's rows add up to the season's fire count without anything having to
    check that they do.

    ``scheme`` comes along because ``--bridge-schemes`` needs to know which side of the
    2023-2024 break a row is on, and because printing it is the cheapest way for a
    reader to see the break for themselves.
    """
    conaf = ConafWildfire.__table__
    cause = ConafFireCause.__table__
    surface = surface_area("total")

    # A CASE rather than a COALESCE chain: these are four distinct states of
    # knowledge, and a coalesce would read as though they were degrees of one.
    label = case(
        (cause.c.cause_normalised.isnot(None), cause.c.cause_normalised),
        (cause.c.cause.isnot(None), UNRECONCILED_LABEL),
        (cause.c.specific_cause.isnot(None), SPECIFIC_ONLY_LABEL),
        else_=NO_CAUSE_LABEL,
    )
    english = case(
        (cause.c.cause_normalised.isnot(None), cause.c.cause_en),
        else_=None,
    )

    return (
        select(
            label.label("cause"),
            func.min(english).label("cause_en"),
            func.min(cause.c.scheme).label("scheme"),
            func.count().label("fires"),
            func.coalesce(func.sum(surface), 0.0).label("hectares"),
        )
        .select_from(conaf)
        .join(cause, cause.c.id == conaf.c.cause_id, isouter=True)
        .where(SEASON == season)
        .group_by(label)
    )


def seasons_query() -> Select:
    """The seasons that have fires, in order."""
    return (select(SEASON).select_from(ConafWildfire.__table__)
            .distinct().order_by(SEASON))


@dataclass(frozen=True)
class Row:
    """One line of the report.

    Attributes
    ----------
    season : int or None
        The season's first year, or ``None`` on a summary row.
    cause : str
        The canonical Spanish, or one of :data:`SYNTHETIC_LABELS`.
    cause_en : str or None
        English, where there is a canonical cause to translate.
    fires : int
        Fires filed under it.
    hectares : float
        Their published area, summed.
    season_fires : int
        Every fire of the season, which is what the percentage is of.
    season_hectares : float
        Every hectare of the season.
    """

    season: int | None
    cause: str
    cause_en: str | None
    fires: int
    hectares: float
    season_fires: int
    season_hectares: float

    @property
    def is_total(self) -> bool:
        return self.season is None

    @property
    def season_label(self) -> str:
        return TOTAL_LABEL if self.season is None else season_label(self.season)

    @property
    def fires_percent(self) -> float | None:
        """Share of the season's fires, or ``None`` where the season has none."""
        return None if not self.season_fires else 100.0 * self.fires / self.season_fires

    @property
    def hectares_percent(self) -> float | None:
        """Share of the season's hectares, or ``None`` where the season burnt none."""
        return (None if not self.season_hectares
                else 100.0 * self.hectares / self.season_hectares)

    @property
    def values(self) -> list[str]:
        """The row's cells as text, for both writers, in :data:`COLUMNS` order."""
        return [
            self.season_label, self.cause, self.cause_en or "", f"{self.fires}",
            "" if self.fires_percent is None else f"{self.fires_percent:.2f}",
            f"{self.hectares:.2f}",
            "" if self.hectares_percent is None else f"{self.hectares_percent:.2f}",
        ]


def sort_key(row: Row) -> tuple:
    """Real causes by descending fire count, then the three synthetic labels last."""
    synthetic = row.cause in SYNTHETIC_LABELS
    order = SYNTHETIC_LABELS.index(row.cause) if synthetic else 0
    return (synthetic, order, -row.fires, row.cause)


def bridge(cause: str) -> str:
    """The post-2023 name a pre-2023 cause became, or the cause unchanged.

    Used only under ``--bridge-schemes``. The successor's name is the one kept,
    because it is the one CONAF is publishing now and a series that ends in the
    current vocabulary is easier to extend than one that ends in a retired one.
    """
    return SCHEME_SUCCESSORS.get(cause, cause)


def compute(session: Session, season: int | None, logger: logging.Logger,
            bridge_schemes: bool = False) -> list[Row]:
    """Count every season in scope, with a summary block over all of them."""
    seasons = ([season] if season is not None
               else list(session.scalars(seasons_query()).all()))
    if not seasons:
        logger.warning("No CONAF fire in scope. Import them with "
                       "src.apps.imports.wildfires.chile_conaf.import_wildfires")
        return []

    rows: list[Row] = []
    overall: dict[str, list] = {}
    overall_fires = overall_hectares = 0

    for one in seasons:
        counted = session.execute(counts_query(one)).all()
        if not counted:
            continue
        merged: dict[str, tuple[str | None, int, float]] = {}
        for record in counted:
            name = bridge(record.cause) if bridge_schemes else record.cause
            english, fires, hectares = merged.get(name, (record.cause_en, 0, 0.0))
            merged[name] = (english or record.cause_en,
                            fires + record.fires, hectares + float(record.hectares))

        season_fires = sum(entry[1] for entry in merged.values())
        season_hectares = sum(entry[2] for entry in merged.values())
        for name, (english, fires, hectares) in merged.items():
            rows.append(Row(season=one, cause=name, cause_en=english, fires=fires,
                            hectares=hectares, season_fires=season_fires,
                            season_hectares=season_hectares))
            previous = overall.get(name, [english, 0, 0.0])
            overall[name] = [previous[0] or english, previous[1] + fires,
                             previous[2] + hectares]
        overall_fires += season_fires
        overall_hectares += season_hectares

    report_rows: list[Row] = []
    for one in seasons:
        report_rows += sorted((row for row in rows if row.season == one), key=sort_key)
    if len(seasons) > 1:
        report_rows += sorted(
            (Row(season=None, cause=name, cause_en=english, fires=fires,
                 hectares=hectares, season_fires=overall_fires,
                 season_hectares=overall_hectares)
             for name, (english, fires, hectares) in overall.items()),
            key=sort_key)
    return report_rows


def broken_series(rows: list[Row]) -> list[str]:
    """The causes whose series stops or starts at the 2023-2024 break.

    Reported rather than repaired: a reader looking at a column of counts that goes to
    zero needs to know whether the fires stopped or the category did.
    """
    seasons = sorted({row.season for row in rows if row.season is not None})
    if not seasons:
        return []
    before = {row.cause for row in rows
              if row.season is not None and row.season < 2023 and row.fires}
    after = {row.cause for row in rows
             if row.season is not None and row.season >= 2023 and row.fires}
    renamed = set(SCHEME_SUCCESSORS) | set(SCHEME_SUCCESSORS.values())
    return sorted((before ^ after) & renamed)


def write_csv(rows: list[Row], path: Path, logger: logging.Logger) -> None:
    """Write the report as CSV."""
    path.parent.mkdir(parents=True, exist_ok=True)
    with path.open("w", newline="", encoding="utf-8") as handle:
        writer = csv.writer(handle)
        writer.writerow(COLUMNS)
        for row in rows:
            writer.writerow(row.values)
    logger.info("Wrote %s", path)


def write_docx(rows: list[Row], path: Path, season: int | None,
               logger: logging.Logger, bridge_schemes: bool = False) -> None:
    """Write the report as a Word document, with the summary block in bold."""
    # Imported here rather than at module scope so that --csv keeps working if
    # python-docx is not installed.
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt

    document = Document()
    document.add_heading("CONAF wildfire causes (Chile)", level=1)
    document.add_paragraph(
        f"Fires and reported hectares by cause, grouped on CONAF's classification. "
        f"Seasons run 1 July to 30 June. Scope: "
        f"{season_label(season) if season is not None else 'all seasons'}."
    )
    if bridge_schemes:
        document.add_paragraph(
            "Schemes bridged: each pre-2023 cause is counted under the post-2023 "
            "category that replaced it. This asserts a continuity CONAF did not "
            "publish — most of the ten renamings also changed what the category "
            "covers — and is only sound for a series read as approximate."
        )
    else:
        document.add_paragraph(
            "CONAF renumbered and renamed its cause classification in 2023-2024, so "
            "ten categories stop there and their successors start. That break is real "
            "and is printed as it is; --bridge-schemes joins the pairs deliberately."
        )

    table = document.add_table(rows=1, cols=len(COLUMNS))
    table.style = "Table Grid"
    for cell, heading in zip(table.rows[0].cells, COLUMNS):
        cell.text = heading
        cell.paragraphs[0].runs[0].bold = True

    for row in rows:
        cells = table.add_row().cells
        values = [row.season_label, row.cause, row.cause_en or "", f"{row.fires:,}",
                  "" if row.fires_percent is None else f"{row.fires_percent:.2f}",
                  f"{row.hectares:,.2f}",
                  "" if row.hectares_percent is None
                  else f"{row.hectares_percent:.2f}"]
        for index, (cell, value) in enumerate(zip(cells, values)):
            cell.text = value
            paragraph = cell.paragraphs[0]
            if index >= FIRST_NUMERIC_COLUMN:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            for run in paragraph.runs:
                run.bold = row.is_total
                run.font.size = Pt(9)

    path.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(path))
    logger.info("Wrote %s", path)


def parse_arguments(argv: list[str] | None = None) -> argparse.Namespace:
    """Parse the command line."""
    parser = argparse.ArgumentParser(
        description="Report Chilean wildfire causes from CONAF's own classification.",
        epilog="CONAF renumbered its taxonomy in 2023-2024 and reused the codes, so "
               "this report groups on the canonical name and never on the code. Import "
               "the fires first. Database settings not given here are read from the "
               "environment (.env).",
    )
    parser.add_argument("-y", "--season", type=int, metavar="YEAR",
                        help="report only this season, named by its first year")
    parser.add_argument("--cause", metavar="NAME",
                        help="report only this canonical cause, in Spanish "
                             "(for example 'Incendios naturales')")
    parser.add_argument("--bridge-schemes", action="store_true",
                        help="count each pre-2023 cause under the post-2023 category "
                             "that replaced it, joining the ten broken series. Asserts "
                             "a continuity CONAF did not publish; off by default")

    output = parser.add_argument_group("output")
    output.add_argument("--csv", type=Path, help="write the report to this .csv")
    output.add_argument("--docx", type=Path, help="write the report to this .docx")

    common.add_database_arguments(parser)
    parser.add_argument("--log-level", default=os.getenv("GISFIRE_LOG_LEVEL", "INFO"),
                        choices=["CRITICAL", "ERROR", "WARNING", "INFO", "DEBUG"],
                        help="verbosity (env: GISFIRE_LOG_LEVEL, default INFO)")

    args = parser.parse_args(argv)
    if args.csv is None and args.docx is None:
        parser.error("at least one of --csv and --docx is required")
    return args


def report(args: argparse.Namespace, engine: Engine,
           logger: logging.Logger) -> list[Row]:
    """Compute and write the report."""
    common.require_tables(engine, ["conaf_wildfire", "conaf_fire_cause"], logger)
    with Session(engine) as session:
        rows = compute(session, args.season, logger, args.bridge_schemes)
    if not rows:
        return rows

    if not args.bridge_schemes:
        broken = broken_series(rows)
        if broken:
            logger.warning(
                "%d cause series stop or start at the 2023-2024 renumbering (%s). The "
                "zeros on either side of it are the category changing, not the fires. "
                "Pass --bridge-schemes to join them deliberately",
                len(broken), ", ".join(broken[:6]))

    if args.cause is not None:
        wanted = chile_conaf.normalise(args.cause)
        rows = [row for row in rows if chile_conaf.normalise(row.cause) == wanted]
        if not rows:
            logger.warning("No fire is filed under %r. The canonical names are the "
                           "values of "
                           "src.providers.chile_conaf.fire_cause.CAUSE_NORMALISATIONS",
                           args.cause)
            return rows

    if args.csv:
        write_csv(rows, args.csv, logger)
    if args.docx:
        write_docx(rows, args.docx, args.season, logger, args.bridge_schemes)
    return rows


def main(argv: list[str] | None = None) -> int:
    args = parse_arguments(argv)
    logging.basicConfig(level=args.log_level, format=LOG_FORMAT)
    logger = logging.getLogger("conaf-causes")

    try:
        settings = common.resolve_database_settings(args)
    except RuntimeError as error:
        logger.error("%s", error)
        return 1

    engine = create_engine(common.database_url(settings))
    try:
        report(args, engine, logger)
    except Exception as error:  # noqa: BLE001  (the CLI boundary: report, do not traceback)
        logger.error("Report failed: %s", error)
        return 1
    finally:
        engine.dispose()
    return 0


if __name__ == "__main__":
    sys.exit(main())
