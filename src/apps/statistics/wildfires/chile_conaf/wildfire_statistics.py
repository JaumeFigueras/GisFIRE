#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Report burnt area from CONAF's Chilean seasonal fire reports.

One row per season: how many fires CONAF filed, and how many hectares they burnt,
from the office's own ``SUPERFICIE``.

Usage
-----

.. code-block:: console

   $ python3 -m src.apps.statistics.wildfires.chile_conaf.wildfire_statistics \\
         --csv /tmp/conaf-seasons.csv
   $ python3 -m src.apps.statistics.wildfires.chile_conaf.wildfire_statistics \\
         --surface plantation --min-area 5 --docx /tmp/conaf-plantations.docx
   $ python3 -m src.apps.statistics.wildfires.chile_conaf.wildfire_statistics \\
         --dated-only --reporter Conaf --csv /tmp/conaf-dated.csv

The area is reported, not measured
------------------------------------

Every other burnt-area report in GisFIRE measures a polygon —
:mod:`src.apps.statistics.wildfires.portugal_icnf.wildfire_statistics` offers a
choice of geodesic or equal-area because it has one to measure. This archive has no
perimeter at all: :attr:`~src.data_model.wildfire.Wildfire.perimeter` is ``NULL`` on
every row, and the hectares come from
:attr:`~src.providers.chile_conaf.wildfire.ConafWildfire.area_ha_total`, which an
office wrote down.

So there is no ``--area-method`` here, and there is a ``--surface`` instead: CONAF
reports the area by what burnt, in three subtotals and nine components, and which of
them a question is about is a real choice. See :data:`SURFACES`.

.. warning::

   The mapped areas of :mod:`src.providers.chile_conaf_magnitud` are a **different
   measurement of the same fires** — traced from imagery rather than filed by an
   office — and the two routinely disagree.
   :mod:`src.apps.statistics.wildfires.chile_conaf_magnitud.wildfire_statistics`
   reports those. Do not add the two together.

Half the archive has no date, and this report says so
--------------------------------------------------------

49,470 of the 95,865 fires have no published start and are dated to 1 July of their
season. That does not affect this report — a season is a season either way, and the
season comes from ``TEMPORADA`` and not from the placeholder instant — but it affects
everything a reader might do next, so the ``Dated`` column carries the count and both
outputs say what it means.

``--dated-only`` restricts the report to the fires that have a real start, which is
what a comparison against a dated archive needs. It roughly halves the counts for
2010-2011 to 2016-2017 and for 2018-2019, and leaves the other seasons alone.

The season is the unit, and it is not a year
----------------------------------------------

Chile's fire season runs 1 July to 30 June, so ``2016-2017`` is one season and not
two half-years. The ``Season`` column is that label, and the CSV sorts by its first
year. A report that grouped by calendar year would cut every southern summer in half.
"""

from __future__ import annotations

import argparse
import csv
import logging
import os
import sys

from dataclasses import dataclass
from pathlib import Path

from sqlalchemy import ColumnElement
from sqlalchemy import Engine
from sqlalchemy import Select
from sqlalchemy import create_engine
from sqlalchemy import func
from sqlalchemy import select
from sqlalchemy import true as sql_true
from sqlalchemy.orm import Session

import src.settings  # noqa: F401  (imported for the side effect of loading .env)

from src.apps.imports import common
from src.data_model.geography.admin_boundary import AdminBoundary
from src.data_model.ignition import Ignition
from src.data_model.wildfire import Wildfire
from src.providers import chile_conaf
from src.providers.chile_conaf.wildfire import ConafWildfire

LOG_FORMAT = "%(asctime)s %(levelname)-8s %(message)s"

#: Country this archive is about, for the row that is not a season.
COUNTRY_NAME = "Chile"

#: Label of the row summarising every season.
TOTAL_LABEL = "Total"

#: The report's columns, in order, shared by both output formats so that a change to
#: one cannot silently leave the other behind.
COLUMNS = ("Country", "Season", "Fires", "Dated", "Minimum (ha)", "Maximum (ha)",
           "Total (ha)")

#: Index of the first column holding a number, and so right-aligned in the Word table.
FIRST_NUMERIC_COLUMN = 2

#: Which of CONAF's published areas the report measures.
#:
#: The archive reports burnt area by what burnt, and the three subtotals are not
#: interchangeable: a question about the forestry industry is about ``plantation``, a
#: question about ecology is about ``vegetation``, and ``total`` is the only one that
#: matches the figure CONAF itself publishes in its annual statistics.
#:
#: The nine components are not offered individually. They are on the model and a
#: reader who wants *eucalyptus alone* can ask for it in SQL; putting nine more
#: choices on this command line would suggest they are all equally meaningful
#: questions, and they are not.
SURFACES = {
    "total": ConafWildfire.__table__.c.area_ha_total,
    "plantation": ConafWildfire.__table__.c.area_ha_plantation,
    "vegetation": ConafWildfire.__table__.c.area_ha_vegetation,
    "other": ConafWildfire.__table__.c.area_ha_other,
}

#: The surface measured unless ``--surface`` says otherwise.
DEFAULT_SURFACE = "total"

#: How each surface is described in the report's own prose.
SURFACE_LABELS = {
    "total": "the whole fire (SUPERFICIE)",
    "plantation": "plantations only (TOTAL_PLAN: pine, eucalyptus and other)",
    "vegetation": "natural vegetation only (TOTAL_VEG: native forest, scrub, grassland)",
    "other": "other surfaces only (TOTAL_OTRA: agricultural land and debris)",
}

#: The two ways of deciding which country a fire counts towards.
COUNTRY_SOURCE_GEOMETRY = "geometry"
COUNTRY_SOURCE_REPORTED = "reported"
COUNTRY_SOURCES = (COUNTRY_SOURCE_GEOMETRY, COUNTRY_SOURCE_REPORTED)

#: Administrative level of a country in ``admin_boundary``.
COUNTRY_LEVEL = 0

#: The season a fire counts towards.
#:
#: The published ``TEMPORADA``, never a year derived from ``start_date_time``: for
#: 51.6% of these fires that instant is a placeholder built *from* this column in the
#: first place, so deriving the season back out of it would be circular for half the
#: archive and wrong for none of it.
SEASON = ConafWildfire.__table__.c.season_start_year


def season_label(season: int) -> str:
    """``2016`` as ``"2016-2017"``, which is how CONAF writes it."""
    return f"{season}-{season + 1}"


def surface_area(surface: str) -> ColumnElement:
    """The hectares column for one of :data:`SURFACES`.

    Raises
    ------
    ValueError
        If ``surface`` is not one of :data:`SURFACES`.
    """
    if surface not in SURFACES:
        raise ValueError(
            f"unknown surface {surface!r}; expected one of {', '.join(SURFACES)}")
    return SURFACES[surface]


def country_columns(source: str) -> tuple[ColumnElement, list]:
    """Where a fire's country comes from, as ``(name, joins)``.

    Raises
    ------
    ValueError
        If ``source`` is not one of :data:`COUNTRY_SOURCES`.

    Notes
    -----
    **``geometry``** asks the database which country actually contains the fire's
    point, at report time, against the real polygons. It is the default, and for a
    single-country archive its job is not to choose between countries but to catch the
    fires that are in none — a point mis-keyed into the Pacific keeps its Chilean
    ``admin_boundary_id`` and is silently in the total under ``reported``.

    **``reported``** trusts what the import stored, which is far cheaper: a foreign
    key lookup instead of a point-in-polygon test per fire. For 95,865 fires that
    difference is worth having.

    The containment test is against the **ignition's point**, not the fire's
    perimeter, because there is no perimeter:
    :attr:`~src.data_model.wildfire.Wildfire.perimeter` is ``NULL`` on every row of
    this archive. Every other perimeter-bearing provider's version of this function
    tests ``ST_PointOnSurface`` of the polygon; here the point *is* the published
    location, so no surrogate is needed.

    Both joins are inner, which is what drops a fire attributable to no country. With
    no boundaries imported at all that is *every* fire, and the report says so rather
    than printing zeros.
    """
    if source == COUNTRY_SOURCE_REPORTED:
        return AdminBoundary.name, [
            (AdminBoundary, AdminBoundary.id == Wildfire.__table__.c.admin_boundary_id),
        ]
    if source != COUNTRY_SOURCE_GEOMETRY:
        raise ValueError(
            f"unknown country source {source!r}; expected one of "
            f"{', '.join(COUNTRY_SOURCES)}")

    # LATERAL ... LIMIT 1 rather than a plain join: a point on a shared border can
    # satisfy ST_Contains for two countries, and one fire must not become two rows.
    containing = (
        select(AdminBoundary.name.label("name"))
        .where(AdminBoundary.level == COUNTRY_LEVEL)
        .where(func.ST_Contains(AdminBoundary.geometry, Ignition.__table__.c.geometry))
        .limit(1)
        .lateral("containing_country")
    )
    return containing.c.name, [
        (Ignition.__table__,
         Ignition.__table__.c.id == ConafWildfire.__table__.c.ignition_id),
        (containing, sql_true()),
    ]


def scope_conditions(dated_only: bool, min_area: float | None,
                     reporter: str | None, surface: ColumnElement) -> list:
    """The ``WHERE`` clauses ``--dated-only``, ``--min-area`` and ``--reporter`` add.

    Notes
    -----
    ``--min-area`` is applied to **the surface being measured**, not always to the
    total. Asking for plantations of 5 hectares or more and getting every fire whose
    *whole* burn reached 5 hectares would be a different question with the same
    command line.
    """
    conditions = []
    if dated_only:
        conditions.append(
            ConafWildfire.__table__.c.date_time_precision != chile_conaf.PRECISION_SEASON)
    if min_area is not None:
        conditions.append(surface >= min_area)
    if reporter is not None:
        conditions.append(ConafWildfire.__table__.c.reporter == reporter)
    return conditions


def seasons_query(dated_only: bool, min_area: float | None,
                  reporter: str | None, surface: ColumnElement) -> Select:
    """The seasons that have fires in scope, in order."""
    query = select(SEASON).select_from(ConafWildfire.__table__).distinct()
    for condition in scope_conditions(dated_only, min_area, reporter, surface):
        query = query.where(condition)
    return query.order_by(SEASON)


def counts_query(season: int, surface: ColumnElement,
                 country_source: str = COUNTRY_SOURCE_GEOMETRY,
                 dated_only: bool = False, min_area: float | None = None,
                 reporter: str | None = None) -> Select:
    """Build the counting query for one season.

    Returns
    -------
    Select
        A query yielding ``country, fires, dated, minimum, maximum, total``: one row
        per country the season's fires turn out to be in.

    Notes
    -----
    Five aggregates in one pass, so whatever the country source costs is paid once per
    fire rather than once per figure.

    A ``NULL`` area is not a zero and is not counted in ``minimum``: the aggregate
    functions skip it, which is right — a fire whose ``SUPERFICIE`` CONAF never filled
    in has an unknown area, not an area of nought, and letting it set the minimum to 0
    would make every season's minimum 0.
    """
    conaf = ConafWildfire.__table__
    country_name, joins = country_columns(country_source)

    counts = (
        select(
            country_name.label("country"),
            func.count().label("fires"),
            func.count().filter(
                conaf.c.date_time_precision != chile_conaf.PRECISION_SEASON
            ).label("dated"),
            func.min(surface).label("minimum"),
            func.max(surface).label("maximum"),
            func.coalesce(func.sum(surface), 0.0).label("total"),
        )
        .select_from(conaf)
        .join(Wildfire.__table__, Wildfire.__table__.c.id == conaf.c.id)
        .where(SEASON == season)
    )
    for target, condition in joins:
        counts = counts.join(target, condition)
    for condition in scope_conditions(dated_only, min_area, reporter, surface):
        counts = counts.where(condition)
    return counts.group_by(country_name)


@dataclass(frozen=True)
class Row:
    """One line of the report.

    Attributes
    ----------
    country : str
        The country the fires were attributed to.
    season : int or None
        The season's first year, or ``None`` on a summary row.
    fires : int
        Fires in scope.
    dated : int
        Of those, the ones with a published start rather than a season placeholder.
    minimum, maximum : float or None
        Smallest and largest published area, or ``None`` where no fire in the row
        published one.
    total : float
        Published area summed.
    """

    country: str
    season: int | None
    fires: int
    dated: int
    minimum: float | None
    maximum: float | None
    total: float

    @property
    def is_total(self) -> bool:
        return self.season is None

    @property
    def season_label(self) -> str:
        return TOTAL_LABEL if self.season is None else season_label(self.season)

    @property
    def values(self) -> list[str]:
        """The row's cells as text, for both writers, in :data:`COLUMNS` order."""
        return [
            self.country, self.season_label, f"{self.fires}", f"{self.dated}",
            "" if self.minimum is None else f"{self.minimum:.2f}",
            "" if self.maximum is None else f"{self.maximum:.2f}",
            f"{self.total:.2f}",
        ]


def combine(rows: list[Row], country: str) -> Row:
    """The summary row for one country's seasons.

    The minimum and maximum are taken over the seasons rather than recomputed, which
    gives the same answer because both are order statistics over a partition of the
    same fires — and costs one pass instead of a fourteenth query.
    """
    minima = [row.minimum for row in rows if row.minimum is not None]
    maxima = [row.maximum for row in rows if row.maximum is not None]
    return Row(
        country=country, season=None,
        fires=sum(row.fires for row in rows),
        dated=sum(row.dated for row in rows),
        minimum=min(minima) if minima else None,
        maximum=max(maxima) if maxima else None,
        total=sum(row.total for row in rows),
    )


def summarise(measured: list[Row]) -> list[Row]:
    """Build the report: each country, its seasons oldest first, its summary row last.

    Oldest first, unlike the Canadian and Portuguese reports, because a Chilean season
    is written ``2010-2011`` and a column of those reads as a series. The CSV is
    sorted, so a reader who wants it the other way round has one sort to do.
    """
    report: list[Row] = []
    countries = sorted({row.country for row in measured})
    for country in countries:
        rows = sorted((row for row in measured if row.country == country),
                      key=lambda row: row.season)
        report += rows
        if len(rows) > 1:
            report.append(combine(rows, country))
    return report


def compute(session: Session, season: int | None, surface_name: str,
            logger: logging.Logger, country_source: str = COUNTRY_SOURCE_GEOMETRY,
            dated_only: bool = False, min_area: float | None = None,
            reporter: str | None = None) -> list[Row]:
    """Measure every season in scope."""
    surface = surface_area(surface_name)
    seasons = ([season] if season is not None
               else list(session.scalars(
                   seasons_query(dated_only, min_area, reporter, surface)).all()))
    if not seasons:
        logger.warning("No CONAF fire in scope. Import them with "
                       "src.apps.imports.wildfires.chile_conaf.import_wildfires")
        return []

    measured: list[Row] = []
    for one in seasons:
        rows = session.execute(counts_query(one, surface, country_source, dated_only,
                                            min_area, reporter)).all()
        if not rows:
            logger.debug("%s: no fire attributable to a country", season_label(one))
        for row in rows:
            measured.append(Row(
                country=row.country, season=one, fires=row.fires, dated=row.dated,
                minimum=None if row.minimum is None else float(row.minimum),
                maximum=None if row.maximum is None else float(row.maximum),
                total=float(row.total)))
    if not measured:
        logger.warning("No fire could be attributed to a country. Import the OCHA "
                       "boundaries with "
                       "src.apps.imports.admin_boundaries.ocha.import_admin_boundaries, "
                       "or pass --country-source reported")
    return summarise(measured)


def scope_sentence(season: int | None, surface_name: str, dated_only: bool,
                   min_area: float | None, reporter: str | None) -> str:
    """What the report covers, in one sentence, for the Word document."""
    scope = [season_label(season) if season is not None else "all seasons"]
    if dated_only:
        scope.append("only fires with a published start date")
    if min_area is not None:
        scope.append(f"only fires of {min_area:g} ha or more on the measured surface")
    if reporter is not None:
        scope.append(f"only fires reported by {reporter}")
    return "; ".join(scope)


def write_csv(rows: list[Row], path: Path, logger: logging.Logger) -> None:
    """Write the report as CSV, bare numbers and no thousands separators."""
    path.parent.mkdir(parents=True, exist_ok=True)
    with path.open("w", newline="", encoding="utf-8") as handle:
        writer = csv.writer(handle)
        writer.writerow(COLUMNS)
        for row in rows:
            writer.writerow(row.values)
    logger.info("Wrote %s", path)


def write_docx(rows: list[Row], path: Path, season: int | None, surface_name: str,
               logger: logging.Logger, dated_only: bool = False,
               min_area: float | None = None, reporter: str | None = None) -> None:
    """Write the report as a Word document.

    One table, with the summary row in bold. Numbers get thousands separators here —
    the opposite of the CSV, and for the opposite reason: this one is for reading.
    """
    # Imported here rather than at module scope so that --csv keeps working if
    # python-docx is not installed, which matters because it is the only dependency
    # this application adds.
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt

    document = Document()
    document.add_heading("CONAF wildfire burnt area (Chile)", level=1)
    document.add_paragraph(
        f"Areas in hectares, as CONAF reported them: {SURFACE_LABELS[surface_name]}. "
        f"These are filed figures and not measured polygons — this archive publishes "
        f"no perimeter. Seasons run 1 July to 30 June. Fires not attributable to a "
        f"country are excluded. Scope: "
        f"{scope_sentence(season, surface_name, dated_only, min_area, reporter)}."
    )
    document.add_paragraph(
        "The Dated column counts the fires with a published start date. The rest have "
        "none at all and are stored at 1 July of their season, so any statistic about "
        "months, hours or durations is computable only over that column's fires."
    )

    table = document.add_table(rows=1, cols=len(COLUMNS))
    table.style = "Table Grid"
    for cell, heading in zip(table.rows[0].cells, COLUMNS):
        cell.text = heading
        cell.paragraphs[0].runs[0].bold = True

    for row in rows:
        cells = table.add_row().cells
        values = [row.country, row.season_label, f"{row.fires:,}", f"{row.dated:,}",
                  "" if row.minimum is None else f"{row.minimum:,.2f}",
                  "" if row.maximum is None else f"{row.maximum:,.2f}",
                  f"{row.total:,.2f}"]
        for index, (cell, value) in enumerate(zip(cells, values)):
            cell.text = value
            paragraph = cell.paragraphs[0]
            if index >= FIRST_NUMERIC_COLUMN:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            for run in paragraph.runs:
                run.bold = row.is_total
                run.font.size = Pt(9)

    path.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(path))
    logger.info("Wrote %s", path)


def parse_arguments(argv: list[str] | None = None) -> argparse.Namespace:
    """Parse the command line."""
    parser = argparse.ArgumentParser(
        description="Report burnt area from CONAF's Chilean seasonal fire reports.",
        epilog="Areas are CONAF's own reported figures, not measured polygons: this "
               "archive publishes no perimeter. Import the fires and the OCHA "
               "boundaries first. Database settings not given here are read from the "
               "environment (.env).",
    )
    parser.add_argument("-y", "--season", type=int, metavar="YEAR",
                        help="report only this season, named by its first year (2016 "
                             "for 2016-2017)")
    parser.add_argument("--surface", choices=sorted(SURFACES), default=DEFAULT_SURFACE,
                        help=f"which published area to measure (default "
                             f"{DEFAULT_SURFACE})")
    parser.add_argument("--min-area", type=float, metavar="HA",
                        help="ignore fires smaller than this on the measured surface")
    parser.add_argument("--reporter", choices=list(chile_conaf.REPORTERS),
                        help="count only the fires reported by CONAF, or only those "
                             "reported by a forestry company")
    parser.add_argument("--dated-only", action="store_true",
                        help="count only the fires with a published start date, "
                             "excluding the 49,470 that carry a season placeholder")
    parser.add_argument("--country-source", choices=COUNTRY_SOURCES,
                        default=COUNTRY_SOURCE_GEOMETRY, help=argparse.SUPPRESS)

    output = parser.add_argument_group("output")
    output.add_argument("--csv", type=Path, help="write the report to this .csv")
    output.add_argument("--docx", type=Path, help="write the report to this .docx")

    common.add_database_arguments(parser)
    parser.add_argument("--log-level", default=os.getenv("GISFIRE_LOG_LEVEL", "INFO"),
                        choices=["CRITICAL", "ERROR", "WARNING", "INFO", "DEBUG"],
                        help="verbosity (env: GISFIRE_LOG_LEVEL, default INFO)")

    args = parser.parse_args(argv)
    if args.csv is None and args.docx is None:
        parser.error("at least one of --csv and --docx is required")
    return args


def report(args: argparse.Namespace, engine: Engine,
           logger: logging.Logger) -> list[Row]:
    """Compute and write the report."""
    common.require_tables(engine, ["wildfire", "conaf_wildfire", "admin_boundary"],
                          logger)
    with Session(engine) as session:
        rows = compute(session, args.season, args.surface, logger,
                       args.country_source, args.dated_only, args.min_area,
                       args.reporter)
    if not rows:
        return rows
    if args.csv:
        write_csv(rows, args.csv, logger)
    if args.docx:
        write_docx(rows, args.docx, args.season, args.surface, logger,
                   args.dated_only, args.min_area, args.reporter)
    return rows


def main(argv: list[str] | None = None) -> int:
    args = parse_arguments(argv)
    logging.basicConfig(level=args.log_level, format=LOG_FORMAT)
    logger = logging.getLogger("conaf-statistics")

    try:
        settings = common.resolve_database_settings(args)
    except RuntimeError as error:
        logger.error("%s", error)
        return 1

    engine = create_engine(common.database_url(settings))
    try:
        report(args, engine, logger)
    except Exception as error:  # noqa: BLE001  (the CLI boundary: report, do not traceback)
        logger.error("Report failed: %s", error)
        return 1
    finally:
        engine.dispose()
    return 0


if __name__ == "__main__":
    sys.exit(main())
