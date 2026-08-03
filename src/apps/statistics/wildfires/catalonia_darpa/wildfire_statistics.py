#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Burnt-area statistics for the Catalan DARPA burnt area cartography.

Reports, per year, how many fires there were, the smallest, largest and total area
burnt, in hectares, and **how many of those fires are bound to the EGIF *parte* for
the same fire**::

    Country   Year   Fires  Minimum  Maximum      Total  EGIF matched  EGIF matched (%)
    Spain     2024      23     4.79   431.29    1213.13             0              0.00
    Spain     2023      22     4.32   856.64    2406.44             0              0.00
    Spain     2022      43     4.31  2683.60    6487.73            40             93.02
    Spain     Total    860     2.42 22932.24  318497.58           778             90.47

Run it over everything, or narrow it to one year, to the fires above a size, or to
the bindings you are willing to trust::

    python3 -m src.apps.statistics.wildfires.catalonia_darpa.wildfire_statistics --csv burnt.csv
    python3 -m src.apps.statistics.wildfires.catalonia_darpa.wildfire_statistics \\
        --year 2013 --csv 2013.csv --docx 2013.docx
    python3 -m src.apps.statistics.wildfires.catalonia_darpa.wildfire_statistics \\
        --min-area 5 --csv over-5-ha.csv
    python3 -m src.apps.statistics.wildfires.catalonia_darpa.wildfire_statistics \\
        --min-confidence 0.9 --csv identifier-matches.csv

At least one of ``--csv`` and ``--docx`` is required: an application that computed a
report and then printed nothing would be a strange thing.

The application only reads. Database settings come from the environment
(``.env``, see :mod:`src.settings`); every one of them can be overridden with a
command-line argument.

The fifth of the burnt-area reports, alongside
:mod:`GWIS <src.apps.statistics.wildfires.gwis.wildfire_statistics>`,
:mod:`GFA <src.apps.statistics.wildfires.gfa.wildfire_statistics>`,
:mod:`ICNF <src.apps.statistics.wildfires.portugal_icnf.wildfire_statistics>` and
:mod:`EGIF <src.apps.statistics.wildfires.spain_egif.wildfire_statistics>`. Its first
six columns are theirs, in their order, so the CSVs can still be concatenated; the
two after them are this dataset's own and are described below.

Measured, and the complement of the EGIF report
-----------------------------------------------

These hectares are **measured from the published perimeter**, like the GWIS, GFA and
ICNF ones and unlike the EGIF ones. That is not an implementation detail, it is the
whole reason this dataset exists in GisFIRE: EGIF publishes a burnt area in hectares
and no polygon, DARPA publishes a polygon and **no hectares in any layer of any
year** (see :mod:`src.providers.catalonia_darpa`). The two sources are the two halves
of the same fire.

So a Catalan fire's area appears twice in this project and the two figures are
different quantities: what the department's cartography says burnt, and what the
*parte* said burnt. Neither is a correction of the other. The ``EGIF matched`` column
is what makes the comparison possible at all — see below.

How many matched the EGIF data
------------------------------

``EGIF matched`` counts the fires of the year that carry a link to an EGIF *parte*,
:attr:`~src.providers.catalonia_darpa.wildfire.DarpaWildfire.egif_wildfire_id`, and
``EGIF matched (%)`` is that as a share of the ``Fires`` beside it.

Three things about it:

* **It is a column, not a filter.** An unbound fire is still a fire and still
  contributes its hectares to the row it is in. This report is a report of the
  Catalan cartography that says how much of it can be joined to the national
  statistic — it is not a report of the joinable part.
* **It follows the scope.** Whatever ``--year`` and ``--min-area`` select, the
  matched count is counted over exactly those fires, so the percentage always has
  the ``Fires`` column as its denominator.
* **A zero is usually a fact about EGIF's coverage, not about the binding.** The
  binding can only reach the campaigns that have been imported, so a DARPA year the
  EGIF exports do not cover matches nothing at all, however good the rules are.

The links themselves are none of this application's business: they are written by
:mod:`~src.apps.bindings.wildfires.catalonia_darpa.bind_egif_wildfires`, which is
also where the rules, the cascade and their results on the real archive are set out.
On the published data 778 of the 860 perimeters are bound — 90.5% — and 45 of the 82
that are not are 2023 and 2024, which the EGIF exports do not reach.

.. note::

   If nothing has ever been bound the column is zero everywhere, which looks exactly
   like a dataset that failed to match. The log says which it is at ``WARNING``
   rather than letting a table of zeros pass for an answer.

Not every binding is the same claim
^^^^^^^^^^^^^^^^^^^^^^^^^^^^^^^^^^^

``--min-confidence`` counts only the bindings at or above a given
:attr:`~src.providers.catalonia_darpa.wildfire.DarpaWildfire.match_confidence`.
Without it every binding counts, whatever produced it.

It is worth reaching for because roughly three quarters of the links rest on an
**identifier** — from 1997 the published ``CODI_FINAL`` *is* the EGIF
``report_number``, and four of the older formats decode into one — and the rest on a
date narrowed by a province and a municipality name, which is a good rule and not a
certainty. ``--min-confidence 0.9`` is the boundary between the two kinds; see
:data:`~src.providers.catalonia_darpa.wildfire.MATCH_METHOD_CONFIDENCE`.

.. warning::

   The confidences are **an ordering, not probabilities**. Nothing has been
   calibrated against ground truth — there is no independent answer key for a 1989
   fire — so ``--min-confidence 0.75`` selects a class of matching rule and does not
   mean "matches that are 75% likely to be right".

Which year a fire counts towards
--------------------------------

:attr:`~src.providers.catalonia_darpa.wildfire.DarpaWildfire.year` — the year of the
published layer the fire was read from — exactly as the ICNF report uses the
published ``Ano`` and the EGIF one the filed ``Campania``.

Here the three candidate years agree by construction and it is checked rather than
assumed: the import verifies on every fire that the year of the published
``DATA_INCEN`` is the year of the layer it is in, and it is true of all 4,533 burnt
features. The column is also ``NOT NULL`` and indexed, and needs no timezone applied
to it, which
:attr:`~src.data_model.wildfire.Wildfire.start_date_time` would — a Catalan fire's
instant is local midnight, the dataset publishing no time of day anywhere.

Which fires are counted
-----------------------

Every imported Catalan fire with a perimeter, which on this dataset is every one of
them: a fire here *is* a perimeter, and the features with no geometry, no code and no
date are the raster background class the import drops before anything is stored
(:data:`~src.providers.catalonia_darpa.GRID_CODE_BURNT`).

``--min-area`` narrows it to the fires of at least that many hectares. By default
there is no threshold.

.. note::

   A fire is one published ``(code, date)`` and not one polygon. Three layers were
   vectorised from a raster and never dissolved — 4,533 burnt features are 860 fires,
   and one fire of 1994 is published as 1,309 separate polygons — so the ``Fires``
   column counts fires and the areas are of the dissolved perimeter. Anything
   counting features instead would report 1994 as five times the fire year it was.

No country test, and no ``--country``
-------------------------------------

Every fire in this dataset is Catalan, and therefore Spanish, because the
department publishes the fires of Catalonia and nothing else. The ``Country`` column
is the constant :data:`COUNTRY_NAME` on every row and **nothing is tested against a
boundary**, so there is no ``--country`` and no ``--country-source``: both are
refused with a message saying why rather than argparse resolving a prefix and
complaining about something else.

That is the EGIF report's ``filed`` mode and not the ICNF report's default, and it is
a deliberate difference. There the point of testing is to catch a perimeter digitised
into the sea; here the perimeters are the department's own cartography of its own
territory, published on its own grid, and there is nothing for a containment test to
find.

.. warning::

   The column says ``Spain`` and the report is **Catalonia's fires alone**. A total
   here is not a Spanish total and must not be read beside the EGIF report's as one:
   it is one autonomous community of seventeen. The region is in the title of the
   ``.docx`` and in this page, not in a column, so that the CSV keeps the shape the
   other four reports have.

How the area is measured
------------------------

The same two ways as the GFA and ICNF reports, chosen with ``--area-method``:
``geodesic`` on the WGS84 ellipsoid (the default) or ``equal-area`` in EPSG:6933.

Why not the CRS the department publishes in
^^^^^^^^^^^^^^^^^^^^^^^^^^^^^^^^^^^^^^^^^^^

The perimeters are also stored as published, in EPSG:25831 (ETRS89 / UTM zone 31N),
on
:attr:`~src.providers.catalonia_darpa.wildfire.DarpaWildfire.perimeter_etrs89_utm31n`
— a projected national grid in metres, and the one the department measures on. It is
**not** offered as an area method, for the ICNF report's reason: UTM is a transverse
Mercator, which is *conformal and not equal-area*.

The error is far smaller here than the Portuguese case — Catalonia is one zone wide,
with no islands — but it is not nothing, and it is systematic:

.. code-block:: text

   the same polygon, EPSG:25831 vs geodesic
     Tortosa           (0.5E)    +0.028%
     Lleida            (0.6E)    +0.016%
     Val d'Aran        (0.8E)    +0.001%
     Barcelona         (2.2E)    -0.068%
     Girona            (2.8E)    -0.079%
     Cap de Creus      (3.3E)    -0.078%

A tenth of a percent, and it **varies with longitude** — over-measuring in the west
and under-measuring in the east — so it does not cancel over a year whose fires are
not evenly spread across the country, which no fire year is. The two methods that are
offered agree with each other to within 0.003%, which is the difference between two
ways of measuring the same thing rather than a projection's distortion.

To reproduce a figure measured on the department's grid, measure it there
explicitly rather than asking this report for it:

.. code-block:: sql

   SELECT year, sum(ST_Area(perimeter_etrs89_utm31n) / 10000.0) AS grid_ha
   FROM darpa_wildfire GROUP BY year ORDER BY year DESC;

One statement
-------------

The GWIS, GFA and ICNF reports measure one year per statement, because the memory a
point-in-polygon test against a country polygon needs is only released when the
statement ends, and a single pass over twenty million perimeters took a 30 GB machine
to the OOM killer.

This report is one statement, like the EGIF one. It tests nothing against a boundary
at all, and the whole archive is **860 fires** — four orders of magnitude short of
the case that died. The ``Total`` row is arithmetic over the years, by
:func:`combine`, so the output is the same shape as the other four either way.
"""

from __future__ import annotations

import argparse
import csv
import logging
import math
import os
import sys

from dataclasses import dataclass
from pathlib import Path

from geoalchemy2 import Geography
from sqlalchemy import ColumnElement
from sqlalchemy import Engine
from sqlalchemy import Select
from sqlalchemy import cast
from sqlalchemy import create_engine
from sqlalchemy import func
from sqlalchemy import literal
from sqlalchemy import select
from sqlalchemy.orm import Session

import src.settings  # noqa: F401  (imported for the side effect of loading .env)

from src.apps.imports import common
from src.data_model.wildfire import Wildfire
from src.providers.catalonia_darpa.wildfire import DarpaWildfire

#: Label used in the ``Year`` column for the summary row.
TOTAL_LABEL = "Total"

#: The six columns this report shares with the GWIS, GFA, ICNF and EGIF ones, in
#: their order, so the CSVs can still be concatenated on them.
SHARED_COLUMNS = ("Country", "Year", "Fires", "Minimum (ha)", "Maximum (ha)", "Total (ha)")

#: The report's columns, in order, shared by both output formats so that a change to
#: one cannot silently leave the other behind. The last two are this dataset's own:
#: how many of the year's fires are bound to an EGIF *parte*, and that as a share of
#: the fires counted.
COLUMNS = SHARED_COLUMNS + ("EGIF matched", "EGIF matched (%)")

#: Index of the first column that holds a number, and so is right-aligned in the
#: Word table.
FIRST_NUMERIC_COLUMN = 2

#: The country every fire in this dataset is in, and the whole of the ``Country``
#: column. Spelled as the OCHA boundaries spell it, so a row of this report sorts and
#: groups with the rows of the other four. Nothing is tested against a boundary to
#: arrive at it — see the module docstring.
COUNTRY_NAME = "Spain"

#: The autonomous community the dataset covers. Named in the ``.docx`` and in the
#: log, and deliberately **not** a column: the CSV keeps the other four reports'
#: shape, and every row of this one would carry the same value.
REGION_NAME = "Catalonia"

#: Square metres in a hectare.
SQUARE_METRES_PER_HECTARE = 10_000.0

#: NSIDC EASE-Grid 2.0 Global — a cylindrical equal-area projection in metres,
#: defined for the whole world. The CRS behind ``--area-method equal-area``.
EQUAL_AREA_SRID = 6933

#: The two ways of turning a perimeter in degrees into hectares.
AREA_METHOD_GEODESIC = "geodesic"
AREA_METHOD_EQUAL_AREA = "equal-area"
AREA_METHODS = (AREA_METHOD_GEODESIC, AREA_METHOD_EQUAL_AREA)

#: The year a fire counts towards: the year of the published layer, which the import
#: has already checked is the year of the published date. See the module docstring.
PUBLISHED_YEAR = DarpaWildfire.__table__.c.year

#: The confidence at or above which a binding rests on an identifier rather than on a
#: name — the value ``--min-confidence`` is reached for most often.
#:
#: Not a default: by default every binding counts. It is here so the help text and the
#: documentation cannot drift from
#: :data:`~src.providers.catalonia_darpa.wildfire.MATCH_METHOD_CONFIDENCE`.
IDENTIFIER_CONFIDENCE = 0.9


def burnt_area(method: str) -> ColumnElement:
    """Burnt area of one fire in hectares, by whichever method was asked for.

    Parameters
    ----------
    method : str
        One of :data:`AREA_METHODS`.

    Returns
    -------
    ColumnElement
        The SQL expression yielding hectares.

    Raises
    ------
    ValueError
        If ``method`` is not one of :data:`AREA_METHODS`.

    Notes
    -----
    Always measured from the EPSG:4326 perimeter on the parent ``wildfire`` row,
    never from the published EPSG:25831 copy — see the module docstring for why the
    department's own grid is not an option here.

    The geodesic cast carries the geometry type and SRID rather than being a bare
    ``Geography()``: that renders as ``geography(GEOMETRY,-1)``, which PostGIS
    rejects, because -1 is not a SRID it knows.
    """
    if method == AREA_METHOD_GEODESIC:
        square_metres = func.ST_Area(
            cast(Wildfire.perimeter, Geography(geometry_type="MULTIPOLYGON", srid=4326))
        )
    elif method == AREA_METHOD_EQUAL_AREA:
        square_metres = func.ST_Area(func.ST_Transform(Wildfire.perimeter, EQUAL_AREA_SRID))
    else:
        raise ValueError(
            f"unknown area method {method!r}; expected one of {', '.join(AREA_METHODS)}"
        )
    return square_metres / SQUARE_METRES_PER_HECTARE


def is_matched(min_confidence: float | None = None) -> ColumnElement:
    """Whether a fire counts as bound to an EGIF *parte*.

    Parameters
    ----------
    min_confidence : float, optional
        Count only bindings of at least this confidence. ``None``, the default,
        counts every binding whatever rule produced it.

    Returns
    -------
    ColumnElement
        A boolean expression, ``True`` for a fire the ``EGIF matched`` column counts.

    Notes
    -----
    The link and not the method is what is tested, because the two cannot disagree: a
    check constraint on the model makes ``egif_wildfire_id`` and ``match_method``
    null together, so a row with one and not the other cannot exist to be counted
    wrongly.

    With a threshold the confidence is tested as well, and the two conditions are
    both stated rather than the second alone. ``match_confidence >= 0.5`` would in
    fact select the same rows today, every method having a confidence — but it would
    be relying on that, where this relies on the constraint.
    """
    darpa = DarpaWildfire.__table__
    bound = darpa.c.egif_wildfire_id.is_not(None)
    if min_confidence is None:
        return bound
    return bound & (darpa.c.match_confidence >= min_confidence)


def statistics_query(method: str = AREA_METHOD_GEODESIC,
                     year: int | None = None,
                     min_area: float | None = None,
                     min_confidence: float | None = None) -> Select:
    """Build the statistics query: one row per year, newest first.

    Parameters
    ----------
    method : str
        One of :data:`AREA_METHODS`.
    year : int, optional
        Restrict to one published year. ``None``, the default, reports every year.
    min_area : float, optional
        Count only fires of at least this many hectares. ``None``, the default,
        counts every fire.
    min_confidence : float, optional
        Count as matched only the bindings of at least this confidence. ``None``, the
        default, counts every binding.

    Returns
    -------
    Select
        A query yielding ``country, year, minimum, maximum, total, fires, matched``.
        The summary row is :func:`summarise`'s work.

    Notes
    -----
    Built against the mapped classes rather than written as SQL text, so a column
    renamed on a model breaks this at import time rather than in front of a user.

    ``darpa_wildfire`` is joined by table, to keep SQLAlchemy from adding a
    polymorphic join of its own, and it has to be joined in any case:
    :data:`PUBLISHED_YEAR` and the EGIF link both live on it.

    The inner query computes each area exactly once. Folded into the outer aggregate
    instead, the area expression would be evaluated three times per row — for the
    minimum, the maximum and the sum — and it is by far the most expensive thing
    here.

    The country is grouped on although it is a literal: by the time the outer
    aggregate sees it, it is a column of the subquery rather than a constant, and
    PostgreSQL requires every one of those in the ``GROUP BY``. One grouping key with
    a single distinct value costs nothing and keeps this report the same shape as the
    four it is read beside.

    ``min_area`` filters the subquery's column rather than repeating the area
    expression, and is applied before the aggregates rather than as a ``HAVING``: the
    threshold selects the fires the figures are computed from, it does not discard
    years whose total came out small. The matched count is aggregated over the same
    filtered rows, which is what keeps ``EGIF matched`` a share of the ``Fires``
    beside it and not of some other set of fires.
    """
    darpa = DarpaWildfire.__table__

    fires = (
        select(
            literal(COUNTRY_NAME).label("country"),
            PUBLISHED_YEAR.label("year"),
            burnt_area(method).label("hectares"),
            is_matched(min_confidence).label("matched"),
        )
        .select_from(Wildfire)
        .join(darpa, darpa.c.id == Wildfire.id)
        .where(Wildfire.perimeter.is_not(None))
    )
    if year is not None:
        fires = fires.where(PUBLISHED_YEAR == year)

    fire = fires.subquery("fire")
    statistics = (
        select(
            fire.c.country,
            fire.c.year,
            func.min(fire.c.hectares).label("minimum"),
            func.max(fire.c.hectares).label("maximum"),
            func.sum(fire.c.hectares).label("total"),
            func.count().label("fires"),
            func.count().filter(fire.c.matched).label("matched"),
        )
        .group_by(fire.c.country, fire.c.year)
        .order_by(fire.c.year.desc())
    )
    if min_area is not None:
        statistics = statistics.where(fire.c.hectares >= min_area)
    return statistics


def share(part: int, whole: int) -> float | None:
    """``part`` as a percentage of ``whole``, or ``None`` where there is no whole.

    ``None`` and not zero: a percentage of nothing is not zero percent, it is no
    answer, and the writers turn it into an empty cell. A row of this report always
    has fires in it, so the empty case is one a caller can produce and the report
    cannot — which is exactly when it is worth not returning a number.
    """
    if not whole:
        return None
    return 100.0 * part / whole


def share_label(part: int, whole: int) -> str:
    """A percentage as it is written out, empty where there is none."""
    value = share(part, whole)
    return "" if value is None else f"{value:.2f}"


@dataclass(frozen=True)
class Row:
    """One line of the report.

    Attributes
    ----------
    country : str
        Always :data:`COUNTRY_NAME`. Nothing is tested against a boundary to arrive
        at it — the department publishes Catalonia's fires and nothing else.
    year : int or None
        The published year, or ``None`` for the summary row.
    minimum, maximum, total : float
        Smallest single fire, largest single fire and sum of every fire, in hectares
        measured from the perimeter.
    fires : int
        How many fires the three area figures were computed from — a fire being one
        published ``(code, date)``, not one polygon.
    matched : int
        How many of those ``fires`` are bound to an EGIF *parte*, at or above
        ``--min-confidence`` where one was given. Always at most ``fires``: it is
        counted over the same rows.
    """

    country: str
    year: int | None
    minimum: float
    maximum: float
    total: float
    fires: int
    matched: int

    @property
    def is_total(self) -> bool:
        """Whether this is the summary row rather than one of the years."""
        return self.year is None

    @property
    def year_label(self) -> str:
        return TOTAL_LABEL if self.is_total else str(self.year)

    @property
    def matched_share(self) -> float | None:
        """``matched`` as a percentage of ``fires``: how much of the year joins."""
        return share(self.matched, self.fires)

    @property
    def values(self) -> tuple[str, ...]:
        """The row as the CSV writes it, in :data:`COLUMNS` order."""
        return (self.country, self.year_label, str(self.fires),
                f"{self.minimum:.2f}", f"{self.maximum:.2f}", f"{self.total:.2f}",
                str(self.matched), share_label(self.matched, self.fires))

    @property
    def readable_values(self) -> tuple[str, ...]:
        """The row as the Word document writes it: the numbers with separators."""
        return (self.country, self.year_label, f"{self.fires:,}",
                f"{self.minimum:,.2f}", f"{self.maximum:,.2f}", f"{self.total:,.2f}",
                f"{self.matched:,}", share_label(self.matched, self.fires))


def combine(rows: list[Row], country: str = COUNTRY_NAME, year: int | None = None) -> Row:
    """One row summarising several: every figure taken over all of them.

    Notes
    -----
    All five decompose over a partition of the fires — a minimum of minima is a
    minimum, a sum of sums is a sum, a count of counts is a count — so the ``Total``
    row is the number a second aggregate over the same rows would have returned, and
    no fire is counted twice or left out.

    The percentage is deliberately **not** averaged: :class:`Row` recomputes it from
    the summed counts, which is the ratio of the totals rather than the mean of the
    ratios. A year of four fires and a year of ninety must not weigh the same in the
    answer.

    ``fsum`` rather than ``sum``: an exact accumulation over a handful of partial
    totals costs nothing and cannot drift from what one pass would have returned.
    """
    return Row(
        country=country,
        year=year,
        minimum=min(row.minimum for row in rows),
        maximum=max(row.maximum for row in rows),
        total=math.fsum(row.total for row in rows),
        fires=sum(row.fires for row in rows),
        matched=sum(row.matched for row in rows),
    )


def summarise(measured: list[Row]) -> list[Row]:
    """Build the report from the years measured: the years newest first, then the total.

    Parameters
    ----------
    measured : list of Row
        One row per year, as the statement returned them.

    Returns
    -------
    list of Row
        The years newest first with the summary row last. Empty if nothing was
        measured — a report of no fires has no total either.

    Notes
    -----
    No grouping by country and no lookup of one, unlike the four reports this is read
    beside: there is exactly one country here and it is a constant, so ordering the
    countries would be sorting a list of one against the database's collation.
    """
    if not measured:
        return []
    rows = sorted(measured, key=lambda row: row.year, reverse=True)
    return rows + [combine(rows)]


def hectares(text: str) -> float:
    """Argparse type for ``--min-area``: a finite, non-negative number of hectares.

    Raises
    ------
    argparse.ArgumentTypeError
        If the text is not a number, or is negative, or is a non-finite float.

    Notes
    -----
    A bare ``type=float`` would accept ``-5``, ``nan`` and ``inf``. The first two are
    almost certainly a typo and would silently produce the unfiltered report —
    ``nan`` compares false against every area, so it would produce an empty one — and
    none of the three is a size a fire can have.
    """
    try:
        area = float(text)
    except ValueError:
        raise argparse.ArgumentTypeError(f"{text!r} is not a number of hectares")
    if not math.isfinite(area):
        raise argparse.ArgumentTypeError(f"{text!r} is not a finite number of hectares")
    if area < 0:
        raise argparse.ArgumentTypeError(
            f"a burnt area cannot be negative, and {area:g} ha is: pass 0 or more, or "
            f"leave --min-area out to count every fire"
        )
    return area


def confidence(text: str) -> float:
    """Argparse type for ``--min-confidence``: a number from 0 to 1 inclusive.

    Raises
    ------
    argparse.ArgumentTypeError
        If the text is not a number, or is not finite, or is outside ``[0, 1]``.

    Notes
    -----
    The range is closed at both ends on purpose: ``0`` counts every binding, which is
    the default said explicitly, and ``1`` counts only the exact-identifier ones.

    A value above 1 is refused rather than quietly counting nothing. Someone passing
    ``90`` means ninety percent, and a report that answered them with a column of
    zeros would be worse than one that stopped.
    """
    try:
        value = float(text)
    except ValueError:
        raise argparse.ArgumentTypeError(f"{text!r} is not a number")
    if not math.isfinite(value):
        raise argparse.ArgumentTypeError(f"{text!r} is not a finite confidence")
    if not 0.0 <= value <= 1.0:
        raise argparse.ArgumentTypeError(
            f"a match confidence is between 0 and 1, and {value:g} is not: "
            f"{IDENTIFIER_CONFIDENCE:g} is the boundary between the identifier matches "
            f"and the name matches"
        )
    return value


def parse_arguments(argv: list[str] | None = None) -> argparse.Namespace:
    """Parse the command line."""
    parser = argparse.ArgumentParser(
        description="Burnt-area statistics for the Catalan DARPA burnt area cartography, "
                    "with the number of fires bound to the Spanish EGIF statistics.",
        epilog="Areas are in hectares measured from the published perimeter, geodesic on "
               "the WGS84 ellipsoid by default — this dataset publishes no hectares of "
               "its own. Every fire is Catalan, so there is no --country and nothing is "
               "tested against a boundary. Database settings not given here are read from "
               "the environment (.env).",
    )
    parser.add_argument("-y", "--year", type=int,
                        help="restrict to one year, e.g. 2013; this is the year of the "
                             "published layer, which is also the year of the published "
                             "DATA_INCEN")
    parser.add_argument("--min-area", type=hectares, default=None, metavar="HECTARES",
                        help="count only fires that burnt at least this many hectares; by "
                             "default every fire counts. The threshold is applied to the "
                             "area this report measures (see --area-method)")
    parser.add_argument("--min-confidence", type=confidence, default=None,
                        metavar="CONFIDENCE",
                        help=f"count a fire as matched only when its binding to an EGIF "
                             f"parte has at least this confidence, between 0 and 1; by "
                             f"default every binding counts. "
                             f"{IDENTIFIER_CONFIDENCE:g} is the boundary between the "
                             f"matches resting on the published identifier and those "
                             f"resting on a date and a municipality name")
    parser.add_argument("--area-method", default=AREA_METHOD_GEODESIC, choices=AREA_METHODS,
                        help="how to turn the EPSG:4326 perimeter into hectares: "
                             "'geodesic' measures on the WGS84 ellipsoid (default); "
                             "'equal-area' projects to EPSG:6933 and measures there. They "
                             "agree to within 0.003%%. The published EPSG:25831 grid is "
                             "not offered: UTM is conformal, not equal-area, and is a "
                             "tenth of a percent out across Catalonia in a way that "
                             "varies with longitude")

    # Accepted only so that they can be refused clearly. Anyone reaching for either
    # has copied a command line from one of the other four reports, which is a
    # reasonable thing to have done, and argparse's own message would not say why
    # this report is different.
    parser.add_argument("--country", help=argparse.SUPPRESS)
    parser.add_argument("--country-source", help=argparse.SUPPRESS)

    output = parser.add_argument_group("output", "at least one is required")
    output.add_argument("--csv", type=Path, help="write the report to this .csv")
    output.add_argument("--docx", type=Path, help="write the report to this .docx (MS Word)")

    common.add_database_arguments(parser)
    parser.add_argument("--log-level", default=os.getenv("GISFIRE_LOG_LEVEL", "INFO"),
                        choices=["CRITICAL", "ERROR", "WARNING", "INFO", "DEBUG"],
                        help="verbosity (env: GISFIRE_LOG_LEVEL, default INFO)")

    arguments = parser.parse_args(argv)
    if arguments.country is not None:
        parser.error(
            f"there is no --country here: the department publishes the fires of "
            f"{REGION_NAME} and nothing else, so there is nothing to select between. "
            f"Every fire is counted and the Country column is {COUNTRY_NAME} on every "
            f"row."
        )
    if arguments.country_source is not None:
        parser.error(
            f"there is no --country-source here: these perimeters are "
            f"{REGION_NAME}'s own cartography of its own territory, so nothing is "
            f"tested against a boundary and there is nothing for a containment test to "
            f"find. The Country column is the constant {COUNTRY_NAME}."
        )
    if arguments.csv is None and arguments.docx is None:
        parser.error("nothing to write: pass --csv, --docx, or both")
    return arguments


def compute(session: Session, year: int | None, logger: logging.Logger,
            method: str = AREA_METHOD_GEODESIC,
            min_area: float | None = None,
            min_confidence: float | None = None) -> list[Row]:
    """Run the statement and return the report's rows in order.

    Notes
    -----
    One statement, under one spinner — see the module docstring for why this report
    does not need the year-at-a-time machinery three of the other four are built on.
    The ``Total`` row is arithmetic over its result, not a second query.

    The matched fires are logged as well as reported, and a report in which nothing
    at all is bound says so at ``WARNING``: a column of zeros is what an unrun
    binding application and a dataset that matched nothing look like alike, and they
    are not the same thing.
    """
    with common.Spinner(f"Measuring the burnt area of the {REGION_NAME} fires "
                        f"and their EGIF matches", logger):
        measured = [
            Row(country=record.country, year=record.year,
                minimum=float(record.minimum),
                maximum=float(record.maximum),
                total=float(record.total),
                fires=record.fires,
                matched=record.matched)
            for record in session.execute(
                statistics_query(method, year, min_area, min_confidence))
        ]

    rows = summarise(measured)
    logger.info("Computed %d rows over %d year(s) (%s areas, %s)",
                len(rows), len(measured), method,
                "every fire" if min_area is None else f"fires of {min_area:g} ha or more")
    if rows:
        total = rows[-1]
        logger.info("%d of %d fire(s) are bound to an EGIF parte (%s%%)%s",
                    total.matched, total.fires, share_label(total.matched, total.fires),
                    "" if min_confidence is None
                    else f", counting only bindings of confidence {min_confidence:g} or more")
        if not total.matched:
            logger.warning(
                "No fire in scope is bound to an EGIF parte, so every EGIF matched "
                "column is zero. Nothing here can tell an unrun binding from a real "
                "absence of matches — run "
                "src.apps.bindings.wildfires.catalonia_darpa.bind_egif_wildfires, and "
                "note that it can only reach the EGIF campaigns that are imported")
    return rows


def write_csv(rows: list[Row], path: Path, logger: logging.Logger) -> None:
    """Write the report as CSV.

    The numbers go out unformatted apart from being rounded to two decimals — no
    thousands separators — because a CSV is read by another program far more often
    than by a person, and a separator would make every figure a string.
    """
    path.parent.mkdir(parents=True, exist_ok=True)
    with path.open("w", newline="", encoding="utf-8") as handle:
        writer = csv.writer(handle)
        writer.writerow(COLUMNS)
        for row in rows:
            writer.writerow(row.values)
    logger.info("Wrote %s", path)


def write_docx(rows: list[Row], path: Path, year: int | None,
               logger: logging.Logger,
               method: str = AREA_METHOD_GEODESIC,
               min_area: float | None = None,
               min_confidence: float | None = None) -> None:
    """Write the report as a Word document.

    One table, with the summary row in bold. Numbers get thousands separators here —
    the opposite of the CSV, and for the opposite reason: this one is for reading.

    The opening paragraphs name the region, the scope and what the ``EGIF matched``
    column is, because a table headed ``Spain`` that is in fact one autonomous
    community, and a count of matches that is a count of *bindings* rather than of
    fires that exist in both archives, are both things a reader should not have to
    remember.
    """
    # Imported here rather than at module scope so that --csv keeps working if
    # python-docx is not installed, which matters because it is the only dependency
    # this application adds.
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt

    document = Document()
    document.add_heading(f"DARPA wildfire burnt area ({REGION_NAME})", level=1)

    measured = ("geodesically on the WGS84 ellipsoid" if method == AREA_METHOD_GEODESIC
                else f"in the equal-area projection EPSG:{EQUAL_AREA_SRID}")
    scope = [f"year: {year}" if year is not None else "all years"]
    if min_area is not None:
        scope.append(f"only fires of {min_area:g} ha or more")
    if min_confidence is not None:
        scope.append(f"only EGIF bindings of confidence {min_confidence:g} or more")
    document.add_paragraph(
        f"Areas in hectares, computed {measured} from the published perimeter — this "
        f"dataset publishes no burnt area of its own. Years are the published layer's. A "
        f"fire is one published (code, date), which in three years is many polygons. "
        f"Scope: {'; '.join(scope)}."
    )
    document.add_paragraph(
        f"The Country column is {COUNTRY_NAME} on every row and nothing is tested against "
        f"a boundary: the department publishes the fires of {REGION_NAME} and nothing "
        f"else. These totals are therefore one autonomous community's and are not a "
        f"Spanish total."
    )
    document.add_paragraph(
        "EGIF matched counts the fires linked to the Spanish parte for the same fire. It "
        "is a column and not a filter — an unbound fire still contributes its hectares — "
        "and a year the EGIF exports do not cover matches nothing however good the rules "
        "are."
    )

    table = document.add_table(rows=1, cols=len(COLUMNS))
    table.style = "Table Grid"
    for cell, heading in zip(table.rows[0].cells, COLUMNS):
        cell.text = heading
        cell.paragraphs[0].runs[0].bold = True

    for row in rows:
        cells = table.add_row().cells
        for index, (cell, value) in enumerate(zip(cells, row.readable_values)):
            cell.text = value
            paragraph = cell.paragraphs[0]
            if index >= FIRST_NUMERIC_COLUMN:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            for run in paragraph.runs:
                run.bold = row.is_total
                run.font.size = Pt(9)

    path.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(path))
    logger.info("Wrote %s", path)


def report(args: argparse.Namespace, engine: Engine, logger: logging.Logger) -> list[Row]:
    """Compute the statistics and write whichever outputs were asked for."""
    with Session(engine) as session:
        rows = compute(session, args.year, logger, args.area_method, args.min_area,
                       args.min_confidence)

    if not rows:
        # An empty report is almost always a year with no data, and writing an empty
        # file would hide that. A threshold is named when there is one, because then
        # it is at least as likely to be the reason as the year is.
        threshold = "" if args.min_area is None else \
            f" No fire reached the --min-area of {args.min_area:g} ha."
        raise RuntimeError(
            f"No wildfires matched. Check --year, and that the {REGION_NAME} fires are "
            f"imported — the published layers run from 1986." + threshold
        )

    if args.csv is not None:
        write_csv(rows, args.csv, logger)
    if args.docx is not None:
        write_docx(rows, args.docx, args.year, logger, args.area_method, args.min_area,
                   args.min_confidence)
    return rows


def main(argv: list[str] | None = None) -> int:
    args = parse_arguments(argv)
    logging.basicConfig(level=args.log_level, format="%(asctime)s %(levelname)s %(message)s")
    logger = logging.getLogger("darpa-statistics")

    try:
        settings = common.resolve_database_settings(args)
    except RuntimeError as error:
        logger.error("%s", error)
        return 1

    engine = create_engine(common.database_url(settings))
    try:
        report(args, engine, logger)
    except Exception as error:  # noqa: BLE001  (the CLI boundary: report, do not traceback)
        logger.error("%s", error)
        return 1
    finally:
        engine.dispose()
    return 0


if __name__ == "__main__":  # pragma nocover
    sys.exit(main())
