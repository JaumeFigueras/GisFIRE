#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Report burnt area from CONAF's Chilean *incendio de magnitud* perimeters.

One row per season: how many large fires CONAF mapped, and how many hectares those
polygons cover.

Usage
-----

.. code-block:: console

   $ python3 -m src.apps.statistics.wildfires.chile_conaf_magnitud.wildfire_statistics \\
         --csv /tmp/conaf-magnitud.csv
   $ python3 -m src.apps.statistics.wildfires.chile_conaf_magnitud.wildfire_statistics \\
         --area-method geodesic --min-area 1000 --docx /tmp/conaf-large.docx
   $ python3 -m src.apps.statistics.wildfires.chile_conaf_magnitud.wildfire_statistics \\
         --bound-only --csv /tmp/conaf-magnitud-bound.csv

This is a sample, and the report says of what
-----------------------------------------------

The archive maps the fires that reached about
:data:`~src.providers.chile_conaf_magnitud.MAGNITUD_THRESHOLD_HA` hectares, and it
does not map all of them: 2021-2022 has 97 reports of 200 ha or more and 62
perimeters here.

So a total from this report is **the area CONAF mapped**, not the area Chile burnt.
The ``Reported (ha)`` column is the same fires' own filed figures, resolved through
the binding, so the two measurements of one set of fires sit side by side — which is
the only honest way to show either.

.. warning::

   Never add this report's total to
   :mod:`src.apps.statistics.wildfires.chile_conaf.wildfire_statistics`'s. They are
   two measurements of overlapping sets of fires, and 706 of these 743 are in that
   report already.

Three ways to measure a polygon, and one of them is the published number
--------------------------------------------------------------------------

``--area-method`` chooses between them:

``published``
    :attr:`~src.providers.chile_conaf_magnitud.wildfire.ConafMagnitudWildfire.area_ha_mapped`,
    measured at import on the grid the polygon was published on — metres, so the
    figure means something without a projection argument. The default.
``geodesic``
    Measured now, on the WGS84 ellipsoid, from the EPSG:4326 perimeter. Comparable
    with every other country's report in GisFIRE, which is what it is for.
``equal-area``
    Measured now, in EPSG:6933.

The three agree to within a fraction of a percent for a Chilean fire — UTM 19S is a
good projection for a country this shape — so the choice matters less here than it
does for a country spanning many zones. It is offered because a report that cannot
say how it measured is a report that cannot be checked.
"""

from __future__ import annotations

import argparse
import csv
import logging
import os
import sys

from dataclasses import dataclass
from pathlib import Path

from geoalchemy2 import Geography
from sqlalchemy import ColumnElement
from sqlalchemy import Engine
from sqlalchemy import Select
from sqlalchemy import cast
from sqlalchemy import create_engine
from sqlalchemy import func
from sqlalchemy import select
from sqlalchemy import true as sql_true
from sqlalchemy.orm import Session

import src.settings  # noqa: F401  (imported for the side effect of loading .env)

from src.apps.imports import common
from src.apps.statistics.wildfires.chile_conaf.wildfire_statistics import COUNTRY_LEVEL
from src.apps.statistics.wildfires.chile_conaf.wildfire_statistics import COUNTRY_SOURCES
from src.apps.statistics.wildfires.chile_conaf.wildfire_statistics import (
    COUNTRY_SOURCE_GEOMETRY)
from src.apps.statistics.wildfires.chile_conaf.wildfire_statistics import (
    COUNTRY_SOURCE_REPORTED)
from src.apps.statistics.wildfires.chile_conaf.wildfire_statistics import TOTAL_LABEL
from src.apps.statistics.wildfires.chile_conaf.wildfire_statistics import season_label
from src.data_model.geography.admin_boundary import AdminBoundary
from src.data_model.wildfire import Wildfire
from src.providers import chile_conaf_magnitud
from src.providers.chile_conaf.wildfire import ConafWildfire
from src.providers.chile_conaf_magnitud.wildfire import ConafMagnitudWildfire

LOG_FORMAT = "%(asctime)s %(levelname)-8s %(message)s"

#: The report's columns, in order, shared by both output formats.
COLUMNS = ("Country", "Season", "Fires", "Bound", "Minimum (ha)", "Maximum (ha)",
           "Mapped (ha)", "Reported (ha)")

#: Index of the first column holding a number, and so right-aligned in the Word table.
FIRST_NUMERIC_COLUMN = 2

#: Square metres in a hectare.
SQUARE_METRES_PER_HECTARE = 10_000.0

#: NSIDC EASE-Grid 2.0 Global — a cylindrical equal-area projection in metres,
#: defined for the whole world. The CRS behind ``--area-method equal-area``.
EQUAL_AREA_SRID = 6933

#: The published figure, measured at import on the grid the polygon came on.
AREA_METHOD_PUBLISHED = "published"

#: Measured now on the WGS84 ellipsoid, from the EPSG:4326 perimeter.
AREA_METHOD_GEODESIC = "geodesic"

#: Measured now in :data:`EQUAL_AREA_SRID`.
AREA_METHOD_EQUAL_AREA = "equal-area"

#: Every way of turning a perimeter into hectares.
AREA_METHODS = (AREA_METHOD_PUBLISHED, AREA_METHOD_GEODESIC, AREA_METHOD_EQUAL_AREA)

#: How each method is described in the report's own prose.
AREA_METHOD_LABELS = {
    AREA_METHOD_PUBLISHED: "on the UTM grid the polygon was published on, at import",
    AREA_METHOD_GEODESIC: "geodesically on the WGS84 ellipsoid",
    AREA_METHOD_EQUAL_AREA: f"in the equal-area projection EPSG:{EQUAL_AREA_SRID}",
}

#: The season a fire counts towards: the published ``TEMPORADA``.
SEASON = ConafMagnitudWildfire.__table__.c.season_start_year


def burnt_area(method: str) -> ColumnElement:
    """Burnt area of one perimeter in hectares, by whichever method was asked for.

    Raises
    ------
    ValueError
        If ``method`` is not one of :data:`AREA_METHODS`.

    Notes
    -----
    The two measured methods work from the EPSG:4326 perimeter on the parent
    ``wildfire`` row rather than from the published grid copy, so that they mean the
    same thing for a mainland fire and for the Easter Island one — which are on
    different grids and could not otherwise be added together.

    The geodesic cast carries the geometry type and SRID rather than being a bare
    ``Geography()``: that renders as ``geography(GEOMETRY,-1)``, which PostGIS rejects,
    because -1 is not a SRID it knows.
    """
    if method == AREA_METHOD_PUBLISHED:
        return ConafMagnitudWildfire.__table__.c.area_ha_mapped
    if method == AREA_METHOD_GEODESIC:
        square_metres = func.ST_Area(
            cast(Wildfire.__table__.c.perimeter,
                 Geography(geometry_type="MULTIPOLYGON", srid=4326)))
    elif method == AREA_METHOD_EQUAL_AREA:
        square_metres = func.ST_Area(
            func.ST_Transform(Wildfire.__table__.c.perimeter, EQUAL_AREA_SRID))
    else:
        raise ValueError(
            f"unknown area method {method!r}; expected one of {', '.join(AREA_METHODS)}")
    return square_metres / SQUARE_METRES_PER_HECTARE


def country_columns(source: str) -> tuple[ColumnElement, list]:
    """Where a fire's country comes from, as ``(name, joins)``.

    Raises
    ------
    ValueError
        If ``source`` is not one of :data:`COUNTRY_SOURCES`.

    Notes
    -----
    Unlike its counterpart in
    :mod:`src.apps.statistics.wildfires.chile_conaf.wildfire_statistics`, the
    containment test here uses ``ST_PointOnSurface`` of the perimeter — this archive
    has one. Not ``ST_Centroid``: the centroid of a crescent or a ring-shaped burn can
    fall outside the polygon entirely, and would then be tested against a country the
    fire never reached, or against none.
    """
    if source == COUNTRY_SOURCE_REPORTED:
        return AdminBoundary.name, [
            (AdminBoundary, AdminBoundary.id == Wildfire.__table__.c.admin_boundary_id),
        ]
    if source != COUNTRY_SOURCE_GEOMETRY:
        raise ValueError(
            f"unknown country source {source!r}; expected one of "
            f"{', '.join(COUNTRY_SOURCES)}")

    containing = (
        select(AdminBoundary.name.label("name"))
        .where(AdminBoundary.level == COUNTRY_LEVEL)
        .where(func.ST_Contains(AdminBoundary.geometry,
                                func.ST_PointOnSurface(Wildfire.__table__.c.perimeter)))
        .limit(1)
        .lateral("containing_country")
    )
    return containing.c.name, [(containing, sql_true())]


def scope_conditions(min_area: float | None, bound_only: bool,
                     area: ColumnElement) -> list:
    """The ``WHERE`` clauses ``--min-area`` and ``--bound-only`` add."""
    conditions = []
    if min_area is not None:
        conditions.append(area >= min_area)
    if bound_only:
        conditions.append(
            ConafMagnitudWildfire.__table__.c.conaf_wildfire_id.isnot(None))
    return conditions


def seasons_query(min_area: float | None, bound_only: bool,
                  area: ColumnElement) -> Select:
    """The seasons that have perimeters in scope, in order."""
    query = (select(SEASON).select_from(ConafMagnitudWildfire.__table__)
             .join(Wildfire.__table__,
                   Wildfire.__table__.c.id == ConafMagnitudWildfire.__table__.c.id)
             .distinct())
    for condition in scope_conditions(min_area, bound_only, area):
        query = query.where(condition)
    return query.order_by(SEASON)


def counts_query(season: int, area: ColumnElement,
                 country_source: str = COUNTRY_SOURCE_GEOMETRY,
                 min_area: float | None = None,
                 bound_only: bool = False) -> Select:
    """Build the counting query for one season.

    Returns
    -------
    Select
        A query yielding ``country, fires, bound, minimum, maximum, mapped,
        reported``: one row per country the season's perimeters turn out to be in.

    Notes
    -----
    ``reported`` is a sum over the **bound report's** own ``SUPERFICIE``, reached
    through ``conaf_wildfire_id``, and is ``0`` for a perimeter that has not been
    bound. That is why ``Bound`` is a column: without it the reported total would look
    like a measurement of every fire in the row, and it is a measurement of the bound
    ones only.
    """
    magnitud = ConafMagnitudWildfire.__table__
    report_table = ConafWildfire.__table__
    country_name, joins = country_columns(country_source)

    counts = (
        select(
            country_name.label("country"),
            func.count().label("fires"),
            func.count().filter(magnitud.c.conaf_wildfire_id.isnot(None)).label("bound"),
            func.min(area).label("minimum"),
            func.max(area).label("maximum"),
            func.coalesce(func.sum(area), 0.0).label("mapped"),
            func.coalesce(func.sum(report_table.c.area_ha_total), 0.0).label("reported"),
        )
        .select_from(magnitud)
        .join(Wildfire.__table__, Wildfire.__table__.c.id == magnitud.c.id)
        .join(report_table, report_table.c.id == magnitud.c.conaf_wildfire_id,
              isouter=True)
        .where(SEASON == season)
    )
    for target, condition in joins:
        counts = counts.join(target, condition)
    for condition in scope_conditions(min_area, bound_only, area):
        counts = counts.where(condition)
    return counts.group_by(country_name)


@dataclass(frozen=True)
class Row:
    """One line of the report."""

    country: str
    season: int | None
    fires: int
    bound: int
    minimum: float | None
    maximum: float | None
    mapped: float
    reported: float

    @property
    def is_total(self) -> bool:
        return self.season is None

    @property
    def season_label(self) -> str:
        return TOTAL_LABEL if self.season is None else season_label(self.season)

    @property
    def values(self) -> list[str]:
        """The row's cells as text, for both writers, in :data:`COLUMNS` order."""
        return [
            self.country, self.season_label, f"{self.fires}", f"{self.bound}",
            "" if self.minimum is None else f"{self.minimum:.2f}",
            "" if self.maximum is None else f"{self.maximum:.2f}",
            f"{self.mapped:.2f}", f"{self.reported:.2f}",
        ]


def combine(rows: list[Row], country: str) -> Row:
    """The summary row for one country's seasons."""
    minima = [row.minimum for row in rows if row.minimum is not None]
    maxima = [row.maximum for row in rows if row.maximum is not None]
    return Row(country=country, season=None,
               fires=sum(row.fires for row in rows),
               bound=sum(row.bound for row in rows),
               minimum=min(minima) if minima else None,
               maximum=max(maxima) if maxima else None,
               mapped=sum(row.mapped for row in rows),
               reported=sum(row.reported for row in rows))


def summarise(measured: list[Row]) -> list[Row]:
    """Each country, its seasons oldest first, its summary row last."""
    report: list[Row] = []
    for country in sorted({row.country for row in measured}):
        rows = sorted((row for row in measured if row.country == country),
                      key=lambda row: row.season)
        report += rows
        if len(rows) > 1:
            report.append(combine(rows, country))
    return report


def compute(session: Session, season: int | None, method: str,
            logger: logging.Logger, country_source: str = COUNTRY_SOURCE_GEOMETRY,
            min_area: float | None = None, bound_only: bool = False) -> list[Row]:
    """Measure every season in scope."""
    area = burnt_area(method)
    seasons = ([season] if season is not None
               else list(session.scalars(
                   seasons_query(min_area, bound_only, area)).all()))
    if not seasons:
        logger.warning(
            "No CONAF perimeter in scope. Import them with "
            "src.apps.imports.wildfires.chile_conaf_magnitud.import_wildfires")
        return []

    measured: list[Row] = []
    for one in seasons:
        for row in session.execute(counts_query(one, area, country_source, min_area,
                                                bound_only)).all():
            measured.append(Row(
                country=row.country, season=one, fires=row.fires, bound=row.bound,
                minimum=None if row.minimum is None else float(row.minimum),
                maximum=None if row.maximum is None else float(row.maximum),
                mapped=float(row.mapped), reported=float(row.reported)))
    if not measured:
        logger.warning("No perimeter could be attributed to a country. Import the OCHA "
                       "boundaries, or pass --country-source reported")
        return []

    rows = summarise(measured)
    unbound = sum(row.fires - row.bound for row in rows if not row.is_total)
    if unbound:
        logger.info("%d perimeter(s) in scope are not bound to a report, so their "
                    "reported area is not in the Reported column. Bind them with "
                    "src.apps.bindings.wildfires.chile_conaf_magnitud."
                    "bind_conaf_wildfires", unbound)
    return rows


def write_csv(rows: list[Row], path: Path, logger: logging.Logger) -> None:
    """Write the report as CSV."""
    path.parent.mkdir(parents=True, exist_ok=True)
    with path.open("w", newline="", encoding="utf-8") as handle:
        writer = csv.writer(handle)
        writer.writerow(COLUMNS)
        for row in rows:
            writer.writerow(row.values)
    logger.info("Wrote %s", path)


def write_docx(rows: list[Row], path: Path, season: int | None, method: str,
               logger: logging.Logger, min_area: float | None = None,
               bound_only: bool = False) -> None:
    """Write the report as a Word document, with the summary row in bold."""
    # Imported here rather than at module scope so that --csv keeps working if
    # python-docx is not installed.
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt

    scope = [season_label(season) if season is not None else "all seasons"]
    if min_area is not None:
        scope.append(f"only fires of {min_area:g} ha or more")
    if bound_only:
        scope.append("only perimeters bound to a report")

    document = Document()
    document.add_heading("CONAF mapped burnt area, incendios de magnitud (Chile)",
                         level=1)
    document.add_paragraph(
        f"Areas in hectares, measured {AREA_METHOD_LABELS[method]}. Seasons run 1 July "
        f"to 30 June. Fires not attributable to a country are excluded. Scope: "
        f"{'; '.join(scope)}."
    )
    document.add_paragraph(
        f"CONAF maps the fires that reached about "
        f"{chile_conaf_magnitud.MAGNITUD_THRESHOLD_HA:g} hectares, and does not map "
        f"all of them, so these totals are the area mapped and not the area burnt. "
        f"The Reported column is the same fires' own filed figures, reached through "
        f"the binding, and covers the Bound column's fires only."
    )

    table = document.add_table(rows=1, cols=len(COLUMNS))
    table.style = "Table Grid"
    for cell, heading in zip(table.rows[0].cells, COLUMNS):
        cell.text = heading
        cell.paragraphs[0].runs[0].bold = True

    for row in rows:
        cells = table.add_row().cells
        values = [row.country, row.season_label, f"{row.fires:,}", f"{row.bound:,}",
                  "" if row.minimum is None else f"{row.minimum:,.2f}",
                  "" if row.maximum is None else f"{row.maximum:,.2f}",
                  f"{row.mapped:,.2f}", f"{row.reported:,.2f}"]
        for index, (cell, value) in enumerate(zip(cells, values)):
            cell.text = value
            paragraph = cell.paragraphs[0]
            if index >= FIRST_NUMERIC_COLUMN:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            for run in paragraph.runs:
                run.bold = row.is_total
                run.font.size = Pt(9)

    path.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(path))
    logger.info("Wrote %s", path)


def parse_arguments(argv: list[str] | None = None) -> argparse.Namespace:
    """Parse the command line."""
    parser = argparse.ArgumentParser(
        description="Report mapped burnt area from CONAF's Chilean magnitud perimeters.",
        epilog="This archive maps only the fires of about 200 ha and more, and not all "
               "of them, so its totals are the area mapped rather than the area burnt. "
               "Import the perimeters and bind them first. Database settings not given "
               "here are read from the environment (.env).",
    )
    parser.add_argument("-y", "--season", type=int, metavar="YEAR",
                        help="report only this season, named by its first year")
    parser.add_argument("--area-method", choices=AREA_METHODS,
                        default=AREA_METHOD_PUBLISHED,
                        help=f"how to turn a perimeter into hectares (default "
                             f"{AREA_METHOD_PUBLISHED})")
    parser.add_argument("--min-area", type=float, metavar="HA",
                        help="ignore fires smaller than this")
    parser.add_argument("--bound-only", action="store_true",
                        help="count only the perimeters bound to a seasonal report, "
                             "which is what makes the Mapped and Reported columns "
                             "cover the same fires")
    parser.add_argument("--country-source", choices=COUNTRY_SOURCES,
                        default=COUNTRY_SOURCE_GEOMETRY, help=argparse.SUPPRESS)

    output = parser.add_argument_group("output")
    output.add_argument("--csv", type=Path, help="write the report to this .csv")
    output.add_argument("--docx", type=Path, help="write the report to this .docx")

    common.add_database_arguments(parser)
    parser.add_argument("--log-level", default=os.getenv("GISFIRE_LOG_LEVEL", "INFO"),
                        choices=["CRITICAL", "ERROR", "WARNING", "INFO", "DEBUG"],
                        help="verbosity (env: GISFIRE_LOG_LEVEL, default INFO)")

    args = parser.parse_args(argv)
    if args.csv is None and args.docx is None:
        parser.error("at least one of --csv and --docx is required")
    return args


def report(args: argparse.Namespace, engine: Engine,
           logger: logging.Logger) -> list[Row]:
    """Compute and write the report."""
    common.require_tables(engine, ["wildfire", "conaf_magnitud_wildfire",
                                   "conaf_wildfire", "admin_boundary"], logger)
    with Session(engine) as session:
        rows = compute(session, args.season, args.area_method, logger,
                       args.country_source, args.min_area, args.bound_only)
    if not rows:
        return rows
    if args.csv:
        write_csv(rows, args.csv, logger)
    if args.docx:
        write_docx(rows, args.docx, args.season, args.area_method, logger,
                   args.min_area, args.bound_only)
    return rows


def main(argv: list[str] | None = None) -> int:
    args = parse_arguments(argv)
    logging.basicConfig(level=args.log_level, format=LOG_FORMAT)
    logger = logging.getLogger("conaf-magnitud-statistics")

    try:
        settings = common.resolve_database_settings(args)
    except RuntimeError as error:
        logger.error("%s", error)
        return 1

    engine = create_engine(common.database_url(settings))
    try:
        report(args, engine, logger)
    except Exception as error:  # noqa: BLE001  (the CLI boundary: report, do not traceback)
        logger.error("Report failed: %s", error)
        return 1
    finally:
        engine.dispose()
    return 0


if __name__ == "__main__":
    sys.exit(main())
